"""OpenEvolve Protocol Improver - A Streamlit application for protocol improvement using LLMs.

This module provides a comprehensive interface for improving protocols and standard operating
procedures (SOPs) through two main approaches:

1. Evolution-based Improvement: Uses a single LLM provider to iteratively refine protocols
2. Adversarial Testing: Employs multiple LLM providers in a red team/blue team approach
   to identify vulnerabilities and generate hardened protocols

Key Features:
- Support for 34+ LLM providers including OpenAI, Anthropic, Google Gemini, and more
- Real-time protocol evaluation and improvement
- Cost estimation and token tracking
- Thread-safe operations for concurrent model evaluation
- Comprehensive logging and result visualization
- Collaborative features and version control
- Advanced analytics and reporting

The application uses Streamlit for the web interface and provides both single-provider
evolution and multi-provider adversarial testing capabilities.
"""

import functools
import json
import math
import os
import random
import re
import threading
import time
import traceback
import hashlib
import uuid
from typing import List, Dict, Any, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime

try:
    import streamlit as st
except ImportError:
    print("streamlit package not found. Please install it with 'pip install streamlit'")
    st = None

try:
    import requests
except ImportError:
    if st is not None:
        st.error("requests package not found. Please install it with 'pip install requests'")
    else:
        print("requests package not found. Please install it with 'pip install requests'")
    requests = None

# Graceful fallback for streamlit_tags
try:
    from streamlit_tags import st_tags
    HAS_STREAMLIT_TAGS = True
except ImportError:
    HAS_STREAMLIT_TAGS = False
    st.error("streamlit_tags package not found. Please install it with 'pip install streamlit-tags' for full functionality.")

try:
    from fpdf import FPDF
    HAS_FPDF = True
except ImportError:
    HAS_FPDF = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ------------------------------------------------------------------
# 0. Streamlit page config
# ------------------------------------------------------------------

st.set_page_config(
    page_title="OpenEvolve Protocol Improver (34 providers + Adversarial Testing)",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Custom CSS for enhanced styling
st.markdown("""
<style>
/* Custom color scheme */
:root {
    --primary-color: #4a6fa5;
    --secondary-color: #6b8cbc;
    --accent-color: #ff6b6b;
    --success-color: #4caf50;
    --warning-color: #ff9800;
    --error-color: #f44336;
    --background-color: #f8f9fa;
    --card-background: #ffffff;
    --text-color: #333333;
    --border-radius: 8px;
    --box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    --box-shadow-hover: 0 6px 12px rgba(0, 0, 0, 0.15);
}

/* Dark mode variables */
[data-theme="dark"] {
    --primary-color: #6b8cbc;
    --secondary-color: #4a6fa5;
    --accent-color: #ff6b6b;
    --success-color: #4caf50;
    --warning-color: #ff9800;
    --error-color: #f44336;
    --background-color: #0e1117;
    --card-background: #1e2130;
    --text-color: #fafafa;
}

/* Main title styling */
h1 {
    color: var(--primary-color);
    text-align: center;
    padding: 1rem 0;
    border-bottom: 2px solid var(--primary-color);
    margin-bottom: 2rem;
    font-weight: 700;
    letter-spacing: -0.5px;
}

/* Tab styling */
.stTabs [data-baseweb="tab-list"] {
    gap: 24px;
    padding: 0.5rem;
    background-color: var(--card-background);
    border-radius: var(--border-radius);
    box-shadow: var(--box-shadow);
    margin-bottom: 1.5rem;
}

.stTabs [data-baseweb="tab"] {
    height: 50px;
    white-space: pre-wrap;
    background-color: transparent;
    border-radius: var(--border-radius);
    color: var(--text-color);
    font-weight: 600;
    transition: all 0.3s ease;
}

.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg, var(--primary-color), var(--secondary-color));
    color: white;
    box-shadow: var(--box-shadow);
}

/* Card styling for sections */
.stMarkdown, .stDataFrame, .stTable {
    background-color: var(--card-background);
    border-radius: var(--border-radius);
    padding: 1.5rem;
    box-shadow: var(--box-shadow);
    margin-bottom: 1.5rem;
    border: 1px solid rgba(0, 0, 0, 0.05);
    transition: all 0.3s ease;
}

.stMarkdown:hover, .stDataFrame:hover, .stTable:hover {
    box-shadow: var(--box-shadow-hover);
    transform: translateY(-2px);
}

/* Button styling */
.stButton > button {
    background: linear-gradient(135deg, var(--primary-color), var(--secondary-color));
    color: white;
    border: none;
    border-radius: var(--border-radius);
    padding: 0.75rem 1.5rem;
    font-weight: 600;
    transition: all 0.3s ease;
    box-shadow: var(--box-shadow);
}

.stButton > button:hover {
    transform: translateY(-2px);
    box-shadow: var(--box-shadow-hover);
}

.stButton > button:active {
    transform: translateY(0);
}

.stButton > button[kind="secondary"] {
    background-color: var(--card-background);
    color: var(--primary-color);
    border: 1px solid var(--primary-color);
}

/* Input styling */
.stSelectbox, .stTextInput, .stTextArea, .stNumberInput, .stSlider {
    background-color: var(--card-background);
    border-radius: var(--border-radius);
    border: 1px solid #e0e0e0;
    box-shadow: 0 1px 2px rgba(0, 0, 0, 0.05);
    transition: all 0.3s ease;
}

.stSelectbox:focus, .stTextInput:focus, .stTextArea:focus, .stNumberInput:focus {
    border-color: var(--primary-color);
    box-shadow: 0 0 0 2px rgba(74, 111, 165, 0.2);
}

[data-theme="dark"] .stSelectbox, 
[data-theme="dark"] .stTextInput, 
[data-theme="dark"] .stTextArea, 
[data-theme="dark"] .stNumberInput, 
[data-theme="dark"] .stSlider {
    border: 1px solid #4a4a4a;
}

/* Metric styling */
.stMetric {
    background-color: var(--card-background);
    border-radius: var(--border-radius);
    padding: 1.5rem;
    box-shadow: var(--box-shadow);
    text-align: center;
    transition: all 0.3s ease;
    border: 1px solid rgba(0, 0, 0, 0.05);
}

.stMetric:hover {
    transform: translateY(-2px);
    box-shadow: var(--box-shadow-hover);
}

.stMetric label {
    font-size: 1rem;
    color: var(--text-color);
    font-weight: 500;
}

.stMetric div[data-testid="stMetricValue"] {
    font-size: 2rem;
    font-weight: 700;
    color: var(--primary-color);
    margin-top: 0.5rem;
}

/* Status message styling */
.stStatus {
    background-color: var(--card-background);
    border-radius: var(--border-radius);
    padding: 1.5rem;
    box-shadow: var(--box-shadow);
    margin-bottom: 1.5rem;
    border-left: 4px solid var(--primary-color);
}

/* Expander styling */
.stExpander {
    background-color: var(--card-background);
    border-radius: var(--border-radius);
    box-shadow: var(--box-shadow);
    margin-bottom: 1.5rem;
    border: 1px solid rgba(0, 0, 0, 0.05);
    transition: all 0.3s ease;
}

.stExpander:hover {
    box-shadow: var(--box-shadow-hover);
}

.stExpander div[data-testid="stExpanderDetails"] {
    padding: 1.5rem;
    border-top: 1px solid #e0e0e0;
}

[data-theme="dark"] .stExpander div[data-testid="stExpanderDetails"] {
    border-top: 1px solid #4a4a4a;
}

/* Code block styling */
.stCodeBlock {
    background-color: var(--card-background);
    border-radius: var(--border-radius);
    box-shadow: var(--box-shadow);
    margin-bottom: 1.5rem;
}

/* Sidebar styling */
[data-testid="stSidebar"] {
    background-color: var(--card-background);
    border-right: 1px solid #e0e0e0;
    box-shadow: 2px 0 5px rgba(0, 0, 0, 0.05);
}

[data-theme="dark"] [data-testid="stSidebar"] {
    border-right: 1px solid #4a4a4a;
}

/* Progress bar styling */
.stProgress {
    background-color: var(--card-background);
    border-radius: var(--border-radius);
    padding: 1.5rem;
    box-shadow: var(--box-shadow);
    margin-bottom: 1.5rem;
}

/* Alert styling */
.stAlert {
    border-radius: var(--border-radius);
    box-shadow: var(--box-shadow);
    margin-bottom: 1.5rem;
}

/* Custom classes for specific elements */
.protocol-analysis-card {
    background: linear-gradient(135deg, var(--primary-color), var(--secondary-color));
    color: white;
    border-radius: var(--border-radius);
    padding: 2rem;
    margin-bottom: 1.5rem;
    box-shadow: var(--box-shadow);
    transition: all 0.3s ease;
}

.protocol-analysis-card:hover {
    transform: translateY(-2px);
    box-shadow: var(--box-shadow-hover);
}

.protocol-analysis-card h3 {
    color: white;
    margin-top: 0;
    font-weight: 600;
}

.team-badge {
    display: inline-block;
    padding: 0.5rem 1rem;
    border-radius: 20px;
    font-weight: 600;
    font-size: 0.9rem;
    margin: 0.25rem;
    box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    transition: all 0.3s ease;
}

.team-badge:hover {
    transform: translateY(-2px);
    box-shadow: 0 4px 8px rgba(0, 0, 0, 0.15);
}

.red-team {
    background-color: #ffebee;
    color: #c62828;
    border: 1px solid #ffcdd2;
}

[data-theme="dark"] .red-team {
    background-color: #3a1414;
    color: #ef9a9a;
    border: 1px solid #5c2323;
}

.blue-team {
    background-color: #e3f2fd;
    color: #1565c0;
    border: 1px solid #bbdefb;
}

[data-theme="dark"] .blue-team {
    background-color: #14233a;
    color: #90caf9;
    border: 1px solid #233d5c;
}

.model-performance-table {
    width: 100%;
    border-collapse: collapse;
    margin: 1rem 0;
}

.model-performance-table th,
.model-performance-table td {
    padding: 1rem;
    text-align: left;
    border-bottom: 1px solid #e0e0e0;
}

[data-theme="dark"] .model-performance-table th,
[data-theme="dark"] .model-performance-table td {
    border-bottom: 1px solid #4a4a4a;
}

.model-performance-table th {
    background: linear-gradient(135deg, var(--primary-color), var(--secondary-color));
    color: white;
    font-weight: 600;
}

.model-performance-table tr:hover {
    background-color: rgba(74, 111, 165, 0.1);
}

[data-theme="dark"] .model-performance-table tr:hover {
    background-color: rgba(74, 111, 165, 0.2);
}

/* New enhanced styles */
.header-badge {
    display: inline-block;
    padding: 0.25rem 0.75rem;
    border-radius: 20px;
    font-weight: 600;
    font-size: 0.8rem;
    margin-left: 0.5rem;
    vertical-align: middle;
}

.success-badge {
    background-color: #e8f5e9;
    color: #2e7d32;
}

.warning-badge {
    background-color: #fff8e1;
    color: #f57f17;
}

.error-badge {
    background-color: #ffebee;
    color: #c62828;
}

[data-theme="dark"] .success-badge {
    background-color: #1b5e20;
    color: #a5d6a7;
}

[data-theme="dark"] .warning-badge {
    background-color: #33691e;
    color: #f4ff81;
}

[data-theme="dark"] .error-badge {
    background-color: #b71c1c;
    color: #ffcdd2;
}

/* Loading spinner */
.loading-spinner {
    border: 4px solid rgba(74, 111, 165, 0.2);
    border-top: 4px solid var(--primary-color);
    border-radius: 50%;
    width: 30px;
    height: 30px;
    animation: spin 1s linear infinite;
    margin: 0 auto;
}

@keyframes spin {
    0% { transform: rotate(0deg); }
    100% { transform: rotate(360deg); }
}

/* Enhanced dashboard styling */
.dashboard-card {
    background: linear-gradient(135deg, var(--card-background), #f0f2f6);
    border-radius: var(--border-radius);
    padding: 1.5rem;
    box-shadow: var(--box-shadow);
    margin-bottom: 1.5rem;
    border: 1px solid rgba(0, 0, 0, 0.05);
    transition: all 0.3s ease;
}

.dashboard-card:hover {
    box-shadow: var(--box-shadow-hover);
    transform: translateY(-2px);
}

.dashboard-card h3 {
    color: var(--primary-color);
    margin-top: 0;
    font-weight: 600;
    border-bottom: 1px solid #e0e0e0;
    padding-bottom: 0.5rem;
}

/* Enhanced button styling */
.stButton > button.primary {
    background: linear-gradient(135deg, var(--primary-color), var(--secondary-color));
    color: white;
    border: none;
    border-radius: var(--border-radius);
    padding: 0.75rem 1.5rem;
    font-weight: 600;
    transition: all 0.3s ease;
    box-shadow: var(--box-shadow);
}

.stButton > button.secondary {
    background-color: var(--card-background);
    color: var(--primary-color);
    border: 1px solid var(--primary-color);
    border-radius: var(--border-radius);
    padding: 0.75rem 1.5rem;
    font-weight: 600;
    transition: all 0.3s ease;
    box-shadow: var(--box-shadow);
}

.stButton > button.success {
    background: linear-gradient(135deg, var(--success-color), #81c784);
    color: white;
    border: none;
    border-radius: var(--border-radius);
    padding: 0.75rem 1.5rem;
    font-weight: 600;
    transition: all 0.3s ease;
    box-shadow: var(--box-shadow);
}

.stButton > button.warning {
    background: linear-gradient(135deg, var(--warning-color), #ffb74d);
    color: white;
    border: none;
    border-radius: var(--border-radius);
    padding: 0.75rem 1.5rem;
    font-weight: 600;
    transition: all 0.3s ease;
    box-shadow: var(--box-shadow);
}

.stButton > button.error {
    background: linear-gradient(135deg, var(--error-color), #e57373);
    color: white;
    border: none;
    border-radius: var(--border-radius);
    padding: 0.75rem 1.5rem;
    font-weight: 600;
    transition: all 0.3s ease;
    box-shadow: var(--box-shadow);
}

/* Enhanced input styling */
.stSelectbox, .stTextInput, .stTextArea, .stNumberInput, .stSlider {
    background-color: var(--card-background);
    border-radius: var(--border-radius);
    border: 1px solid #e0e0e0;
    box-shadow: 0 1px 2px rgba(0, 0, 0, 0.05);
    transition: all 0.3s ease;
}

.stSelectbox:focus, .stTextInput:focus, .stTextArea:focus, .stNumberInput:focus {
    border-color: var(--primary-color);
    box-shadow: 0 0 0 2px rgba(74, 111, 165, 0.2);
}

[data-theme="dark"] .stSelectbox, 
[data-theme="dark"] .stTextInput, 
[data-theme="dark"] .stTextArea, 
[data-theme="dark"] .stNumberInput, 
[data-theme="dark"] .stSlider {
    border: 1px solid #4a4a4a;
}

/* Enhanced metric styling */
.stMetric {
    background-color: var(--card-background);
    border-radius: var(--border-radius);
    padding: 1.5rem;
    box-shadow: var(--box-shadow);
    text-align: center;
    transition: all 0.3s ease;
    border: 1px solid rgba(0, 0, 0, 0.05);
}

.stMetric:hover {
    transform: translateY(-2px);
    box-shadow: var(--box-shadow-hover);
}

.stMetric label {
    font-size: 1rem;
    color: var(--text-color);
    font-weight: 500;
}

.stMetric div[data-testid="stMetricValue"] {
    font-size: 2rem;
    font-weight: 700;
    color: var(--primary-color);
    margin-top: 0.5rem;
}

/* Enhanced status message styling */
.stStatus {
    background-color: var(--card-background);
    border-radius: var(--border-radius);
    padding: 1.5rem;
    box-shadow: var(--box-shadow);
    margin-bottom: 1.5rem;
    border-left: 4px solid var(--primary-color);
}

/* Enhanced expander styling */
.stExpander {
    background-color: var(--card-background);
    border-radius: var(--border-radius);
    box-shadow: var(--box-shadow);
    margin-bottom: 1.5rem;
    border: 1px solid rgba(0, 0, 0, 0.05);
    transition: all 0.3s ease;
}

.stExpander:hover {
    box-shadow: var(--box-shadow-hover);
}

.stExpander div[data-testid="stExpanderDetails"] {
    padding: 1.5rem;
    border-top: 1px solid #e0e0e0;
}

[data-theme="dark"] .stExpander div[data-testid="stExpanderDetails"] {
    border-top: 1px solid #4a4a4a;
}

/* Enhanced code block styling */
.stCodeBlock {
    background-color: var(--card-background);
    border-radius: var(--border-radius);
    box-shadow: var(--box-shadow);
    margin-bottom: 1.5rem;
}

/* Enhanced sidebar styling */
[data-testid="stSidebar"] {
    background-color: var(--card-background);
    border-right: 1px solid #e0e0e0;
    box-shadow: 2px 0 5px rgba(0, 0, 0, 0.05);
}

[data-theme="dark"] [data-testid="stSidebar"] {
    border-right: 1px solid #4a4a4a;
}

/* Enhanced progress bar styling */
.stProgress {
    background-color: var(--card-background);
    border-radius: var(--border-radius);
    padding: 1.5rem;
    box-shadow: var(--box-shadow);
    margin-bottom: 1.5rem;
}

/* Enhanced alert styling */
.stAlert {
    border-radius: var(--border-radius);
    box-shadow: var(--box-shadow);
    margin-bottom: 1.5rem;
}

/* Responsive adjustments */
@media (max-width: 768px) {
    .stTabs [data-baseweb="tab-list"] {
        gap: 12px;
    }
    
    .stMetric div[data-testid="stMetricValue"] {
        font-size: 1.5rem;
    }
    
    .protocol-analysis-card {
        padding: 1.5rem;
    }
    
    .team-badge {
        padding: 0.25rem 0.5rem;
        font-size: 0.8rem;
    }
    
    .dashboard-card {
        padding: 1rem;
    }
}

/* Animation for new elements */
@keyframes fadeIn {
    from { opacity: 0; transform: translateY(10px); }
    to { opacity: 1; transform: translateY(0); }
}

.fade-in {
    animation: fadeIn 0.5s ease-out;
}

/* Enhanced team badge styling */
.team-badge-lg {
    display: inline-block;
    padding: 0.75rem 1.5rem;
    border-radius: 30px;
    font-weight: 700;
    font-size: 1.1rem;
    margin: 0.5rem;
    box-shadow: 0 4px 8px rgba(0, 0, 0, 0.15);
    transition: all 0.3s ease;
}

.team-badge-lg:hover {
    transform: translateY(-3px);
    box-shadow: 0 6px 12px rgba(0, 0, 0, 0.2);
}

/* Enhanced model performance table */
.model-performance-table-enhanced {
    width: 100%;
    border-collapse: collapse;
    margin: 1rem 0;
    border-radius: var(--border-radius);
    overflow: hidden;
    box-shadow: var(--box-shadow);
}

.model-performance-table-enhanced th,
.model-performance-table-enhanced td {
    padding: 1.2rem;
    text-align: left;
    border-bottom: 1px solid #e0e0e0;
}

[data-theme="dark"] .model-performance-table-enhanced th,
[data-theme="dark"] .model-performance-table-enhanced td {
    border-bottom: 1px solid #4a4a4a;
}

.model-performance-table-enhanced th {
    background: linear-gradient(135deg, var(--primary-color), var(--secondary-color));
    color: white;
    font-weight: 600;
}

.model-performance-table-enhanced tr:hover {
    background-color: rgba(74, 111, 165, 0.1);
}

[data-theme="dark"] .model-performance-table-enhanced tr:hover {
    background-color: rgba(74, 111, 165, 0.2);
}

/* Enhanced protocol comparison styling */
.protocol-comparison {
    display: flex;
    gap: 1rem;
    margin-bottom: 1.5rem;
}

.protocol-comparison-column {
    flex: 1;
    background-color: var(--card-background);
    border-radius: var(--border-radius);
    padding: 1.5rem;
    box-shadow: var(--box-shadow);
    border: 1px solid rgba(0, 0, 0, 0.05);
}

.protocol-comparison-header {
    font-weight: 600;
    color: var(--primary-color);
    margin-bottom: 1rem;
    padding-bottom: 0.5rem;
    border-bottom: 1px solid #e0e0e0;
}

/* Enhanced visualization styling */
.chart-container {
    background-color: var(--card-background);
    border-radius: var(--border-radius);
    padding: 1.5rem;
    box-shadow: var(--box-shadow);
    margin-bottom: 1.5rem;
    border: 1px solid rgba(0, 0, 0, 0.05);
}

.chart-title {
    font-weight: 600;
    color: var(--primary-color);
    margin-bottom: 1rem;
}
</style>
""", unsafe_allow_html=True)

# ------------------------------------------------------------------
# Constants and Prompts
# ------------------------------------------------------------------

JSON_RE = re.compile(r"\{[\s\S]*\}")
FENCE_RE = re.compile(r"```(?:json)?\s*(\{[\s\S]*?\})\s*```", re.IGNORECASE)

def safe_int(x, d=0):
    return d if x is None or not isinstance(x, (int, float, str)) or str(x) == '' else int(float(x))

def safe_float(x, d=0.0):
    return d if x is None or not isinstance(x, (int, float, str)) or str(x) == '' else float(x)

# Define adversarial prompts as constants for clarity and maintainability.
APPROVAL_PROMPT = (
    "You are a senior technical reviewer. Your task is to evaluate the provided Standard Operating Procedure (SOP). "
    "Return a single, STRICT JSON object and nothing else. The JSON object must have the following keys:\n"
    '{"verdict": "APPROVED" | "REJECTED", "score": 0-100, "reasons": ["..."], "notes": "..."}\n'
    "- `verdict`: Your final judgment. 'APPROVED' if the SOP is robust, clear, and secure. 'REJECTED' otherwise.\n"
    "- `score`: An integer from 0 to 100 reflecting the overall quality, security, and robustness to abuse.\n"
    "- `reasons`: A list of concise strings explaining the key factors behind your verdict and score.\n"
    "- `notes`: Any additional commentary or observations.\n"
    "Do not include any text, explanation, or formatting outside of the single JSON object."
)

RED_TEAM_CRITIQUE_PROMPT = (
    "You are a critical technical reviewer conducting a peer review. Your goal is to find every possible flaw in the provided SOP. "
    "Analyze it for logical gaps, ambiguities, edge cases, undefined responsibilities, missing "
    "preconditions, unsafe defaults, and potential paths for errors or misuse. "
    "Also, check for compliance with the following requirements:\\n{compliance_requirements}\\n"
    "Return a single, STRICT JSON object and nothing else with the following structure:\\n"
    '{\"issues\": [{\"title\": \"...\", \"severity\": \"low|medium|high|critical\", \"category\": \"...\", \"detail\": \"...\", '
    '\"reproduction\": \"...\", \"exploit_paths\": [\"...\"], \"mitigations\": [\"...\"]}], '
    '\"summary\": \"...\", \"overall_risk\": \"low|medium|high|critical\"}'
)

BLUE_TEAM_PATCH_PROMPT = (
    "You are a technical improvement specialist. Your task is to address the issues identified in the "
    "provided critiques and produce an improved version of the SOP. Incorporate the feedback to make the protocol "
    "explicit, verifiable, and robust. Add preconditions, validation steps, error handling, "
    "monitoring, auditability, and documentation where applicable. "
    "Return a single, STRICT JSON object and nothing else with the following keys:\\n"
    '{\"sop\": \"<the complete improved SOP in Markdown>\", \"changelog\": [\"...\"], \"residual_risks\": [\"...\"], '
    '\"mitigation_matrix\": [{\"issue\": \"...\", \"fix\": \"...\", \"status\": \"resolved|mitigated|wontfix\"}]}'
)

# Specialized prompts for different review types
CODE_REVIEW_RED_TEAM_PROMPT = (
    "You are a senior software engineer conducting a code peer review. Your goal is to find every possible flaw in the provided code. "
    "Analyze it for bugs, security vulnerabilities, performance issues, maintainability problems, "
    "code smells, anti-patterns, and potential runtime errors. "
    "Also, check for compliance with the following requirements:\\n{compliance_requirements}\\n"
    "Return a single, STRICT JSON object and nothing else with the following structure:\\n"
    '{\"issues\": [{\"title\": \"...\", \"severity\": \"low|medium|high|critical\", \"category\": \"bug|security|performance|maintainability|style\", \"detail\": \"...\", \"line_number\": \"...\", \"suggested_fix\": \"...\"}], '
    '\"summary\": \"...\", \"overall_quality\": \"poor|fair|good|excellent\"}'
)

CODE_REVIEW_BLUE_TEAM_PROMPT = (
    "You are a code improvement specialist. Your task is to address the issues identified in the "
    "provided code review critiques and produce an improved version of the code. Incorporate the feedback to make the code "
    "more secure, efficient, maintainable, and readable. Add proper error handling, "
    "documentation, and follow best practices. "
    "Return a single, STRICT JSON object and nothing else with the following keys:\\n"
    '{\"code\": \"<the complete improved code>\", \"changelog\": [\"...\"], \"residual_issues\": [\"...\"], '
    '\"fix_matrix\": [{\"issue\": \"...\", \"fix\": \"...\", \"status\": \"resolved|mitigated|wontfix\"}]}'
)

PLAN_REVIEW_RED_TEAM_PROMPT = (
    "You are a strategic planning expert conducting a peer review of the provided plan. Your goal is to find every possible flaw in the plan. "
    "Analyze it for unrealistic assumptions, missing steps, unclear objectives, resource constraints, "
    "timeline issues, risk factors, and potential failure points. "
    "Also, check for compliance with the following requirements:\\n{compliance_requirements}\\n"
    "Return a single, STRICT JSON object and nothing else with the following structure:\\n"
    '{\"issues\": [{\"title\": \"...\", \"severity\": \"low|medium|high|critical\", \"category\": \"assumptions|resources|timeline|risks|clarity\", \"detail\": \"...\", \"suggested_improvement\": \"...\"}], '
    '\"summary\": \"...\", \"overall_feasibility\": \"low|medium|high\"}'
)

PLAN_REVIEW_BLUE_TEAM_PROMPT = (
    "You are a planning improvement specialist. Your task is to address the issues identified in the "
    "provided plan review critiques and produce an improved version of the plan. Incorporate the feedback to make the plan "
    "more realistic, actionable, and robust. Add missing details, clarify objectives, "
    "adjust timelines, and improve resource allocation. "
    "Return a single, STRICT JSON object and nothing else with the following keys:\\n"
    '{\"plan\": \"<the complete improved plan in Markdown>\", \"changelog\": [\"...\"], \"residual_concerns\": [\"...\"], '
    '\"improvement_matrix\": [{\"issue\": \"...\", \"improvement\": \"...\", \"status\": \"resolved|mitigated|wontfix\"}]}'
)

# ------------------------------------------------------------------
# Utilities
# ------------------------------------------------------------------

def _now_ms() -> int:
    """Get current time in milliseconds since epoch.
    
    Returns:
        int: Current timestamp in milliseconds
    """
    return int(time.time() * 1000)

def calculate_protocol_complexity(protocol_text: str) -> Dict[str, Any]:
    """Calculate various complexity metrics for a protocol.
    
    Args:
        protocol_text (str): The protocol text to analyze
        
    Returns:
        Dict[str, Any]: Dictionary containing complexity metrics
    """
    if not protocol_text:
        return {
            "word_count": 0,
            "sentence_count": 0,
            "paragraph_count": 0,
            "avg_sentence_length": 0,
            "avg_paragraph_length": 0,
            "unique_words": 0,
            "complexity_score": 0
        }
    
    # Basic text statistics
    words = protocol_text.split()
    word_count = len(words)
    
    sentences = re.split(r'[.!?]+', protocol_text)
    sentences = [s.strip() for s in sentences if s.strip()]
    sentence_count = len(sentences)
    
    paragraphs = [p.strip() for p in protocol_text.split('\n') if p.strip()]
    paragraph_count = len(paragraphs)
    
    # Average lengths
    avg_sentence_length = word_count / max(1, sentence_count)
    avg_paragraph_length = sentence_count / max(1, paragraph_count)
    
    # Unique words
    unique_words = len(set(words))
    
    # Complexity score (weighted combination of metrics)
    # Higher scores indicate more complex protocols
    complexity_score = (
        (word_count / 100) * 0.3 +  # Normalize word count
        (avg_sentence_length / 20) * 0.3 +  # Longer sentences = more complex
        (avg_paragraph_length / 10) * 0.2 +  # Longer paragraphs = more complex
        (1 - (unique_words / max(1, word_count))) * 0.2 * 100  # Repetition = more complex
    )
    
    return {
        "word_count": word_count,
        "sentence_count": sentence_count,
        "paragraph_count": paragraph_count,
        "avg_sentence_length": round(avg_sentence_length, 2),
        "avg_paragraph_length": round(avg_paragraph_length, 2),
        "unique_words": unique_words,
        "complexity_score": round(complexity_score, 2)
    }

def extract_protocol_structure(protocol_text: str) -> Dict[str, Any]:
    """Extract structural elements from a protocol.
    
    Args:
        protocol_text (str): The protocol text to analyze
        
    Returns:
        Dict[str, Any]: Dictionary containing structural elements
    """
    if not protocol_text:
        return {
            "has_numbered_steps": False,
            "has_bullet_points": False,
            "has_headers": False,
            "section_count": 0,
            "has_preconditions": False,
            "has_postconditions": False,
            "has_error_handling": False
        }
    
    lines = protocol_text.split('\n')
    
    # Check for numbered steps (e.g., "1.", "2.", "1.1", etc.)
    numbered_pattern = re.compile(r'^\s*\d+\.')
    has_numbered_steps = any(numbered_pattern.match(line) for line in lines)
    
    # Check for bullet points
    bullet_pattern = re.compile(r'^\s*[*\-+]')
    has_bullet_points = any(bullet_pattern.match(line) for line in lines)
    
    # Check for headers (markdown style ## or === underlines)
    header_pattern = re.compile(r'^#{1,6}\s+|.*\n[=]{3,}|.*\n[-]{3,}')
    has_headers = bool(header_pattern.search(protocol_text))
    
    # Count sections (headers)
    section_count = len(header_pattern.findall(protocol_text))
    
    # Check for common protocol elements
    lower_text = protocol_text.lower()
    has_preconditions = 'precondition' in lower_text or 'prerequisite' in lower_text
    has_postconditions = 'postcondition' in lower_text
    has_error_handling = 'error' in lower_text or 'exception' in lower_text or 'failure' in lower_text
    
    return {
        "has_numbered_steps": has_numbered_steps,
        "has_bullet_points": has_bullet_points,
        "has_headers": has_headers,
        "section_count": section_count,
        "has_preconditions": has_preconditions,
        "has_postconditions": has_postconditions,
        "has_error_handling": has_error_handling
    }

# ------------------------------------------------------------------
# AI-Powered Recommendation Functions
# ------------------------------------------------------------------

def generate_protocol_recommendations(protocol_text: str) -> List[str]:
    """Generate AI-powered recommendations for improving a protocol.
    
    Args:
        protocol_text (str): The protocol text to analyze
        
    Returns:
        List[str]: List of recommendations
    """
    if not protocol_text:
        return ["Please enter a protocol to analyze."]
    
    recommendations = []
    complexity = calculate_protocol_complexity(protocol_text)
    structure = extract_protocol_structure(protocol_text)
    
    # Complexity-based recommendations
    if complexity["complexity_score"] > 50:
        recommendations.append("üìù Protocol is quite complex. Consider breaking it into smaller, more manageable sections.")
    
    if complexity["avg_sentence_length"] > 25:
        recommendations.append("üîó Average sentence length is high. Try to use shorter, clearer sentences.")
    
    if complexity["unique_words"] / max(1, complexity["word_count"]) < 0.4:
        recommendations.append("üîÑ High word repetition detected. Consider using more varied vocabulary for clarity.")
    
    # Structure-based recommendations
    if not structure["has_headers"]:
        recommendations.append("üìå Add headers to organize your protocol into clear sections.")
    
    if not structure["has_numbered_steps"] and not structure["has_bullet_points"]:
        recommendations.append("üî¢ Use numbered steps or bullet points to make instructions clearer.")
    
    if not structure["has_preconditions"]:
        recommendations.append("üîí Add preconditions to specify what must be true before executing the protocol.")
    
    if not structure["has_postconditions"]:
        recommendations.append("‚úÖ Add postconditions to specify what should be true after executing the protocol.")
    
    if not structure["has_error_handling"]:
        recommendations.append("‚ö†Ô∏è Include error handling procedures for when things go wrong.")
    
    # Length-based recommendations
    if complexity["word_count"] < 100:
        recommendations.append("üìÑ Protocol is quite short. Ensure all necessary details are included.")
    
    if structure["section_count"] == 0 and complexity["word_count"] > 300:
        recommendations.append("üìÇ Long protocol without sections. Consider organizing it with headers for better readability.")
    
    # Add some general recommendations if none were triggered
    if not recommendations:
        recommendations.append("üëç Your protocol looks well-structured! Consider running adversarial testing for additional hardening.")
    
    return recommendations

def suggest_protocol_template(protocol_text: str) -> str:
    """Suggest the most appropriate protocol template based on content.
    
    Args:
        protocol_text (str): The protocol text to analyze
        
    Returns:
        str: Suggested template name
    """
    if not protocol_text:
        return "Standard Operating Procedure"
    
    lower_text = protocol_text.lower()
    
    # Check for keywords that indicate specific template types
    if any(keyword in lower_text for keyword in ["security", "vulnerability", "attack", "threat", "penetration"]):
        return "Security Policy"
    elif any(keyword in lower_text for keyword in ["incident", "response", "emergency", "crisis"]):
        return "Incident Response Plan"
    elif any(keyword in lower_text for keyword in ["compliance", "regulation", "audit", "policy"]):
        return "Compliance Policy"
    elif any(keyword in lower_text for keyword in ["development", "code", "software", "programming"]):
        return "Development Process"
    elif any(keyword in lower_text for keyword in ["testing", "test", "qa", "quality"]):
        return "Testing Procedure"
    else:
        return "Standard Operating Procedure"  # Default suggestion

def compare_protocols(protocol_a: str, protocol_b: str) -> Dict[str, Any]:
    """Compare two protocols and highlight differences.
    
    Args:
        protocol_a (str): First protocol text
        protocol_b (str): Second protocol text
        
    Returns:
        Dict[str, Any]: Comparison results including similarity metrics and differences
    """
    if not protocol_a or not protocol_b:
        return {
            "similarity": 0.0,
            "length_difference": 0,
            "added_sections": [],
            "removed_sections": [],
            "complexity_change": 0.0,
            "word_count_change": 0,
            "sentence_count_change": 0,
            "unique_words_change": 0,
            "structure_changes": {}
        }
    
    # Calculate similarity using a simple approach
    # In a real implementation, we might use more sophisticated algorithms
    set_a = set(protocol_a.split())
    set_b = set(protocol_b.split())
    
    # Jaccard similarity
    intersection = len(set_a.intersection(set_b))
    union = len(set_a.union(set_b))
    similarity = intersection / max(1, union)
    
    # Length difference
    length_difference = len(protocol_b) - len(protocol_a)
    
    # Complexity metrics
    complexity_a = calculate_protocol_complexity(protocol_a)
    complexity_b = calculate_protocol_complexity(protocol_b)
    complexity_change = complexity_b["complexity_score"] - complexity_a["complexity_score"]
    
    # Detailed metrics changes
    word_count_change = complexity_b["word_count"] - complexity_a["word_count"]
    sentence_count_change = complexity_b["sentence_count"] - complexity_a["sentence_count"]
    unique_words_change = complexity_b["unique_words"] - complexity_a["unique_words"]
    
    # Simple section difference detection
    sections_a = re.findall(r'^#{1,6}\s+.*$|.*\n[=]{3,}|.*\n[-]{3,}', protocol_a, re.MULTILINE)
    sections_b = re.findall(r'^#{1,6}\s+.*$|.*\n[=]{3,}|.*\n[-]{3,}', protocol_b, re.MULTILINE)
    
    added_sections = [s for s in sections_b if s not in sections_a]
    removed_sections = [s for s in sections_a if s not in sections_b]
    
    # Structure analysis
    structure_a = extract_protocol_structure(protocol_a)
    structure_b = extract_protocol_structure(protocol_b)
    
    structure_changes = {}
    for key in structure_a.keys():
        if structure_a[key] != structure_b[key]:
            structure_changes[key] = {
                "before": structure_a[key],
                "after": structure_b[key]
            }
    
    return {
        "similarity": round(similarity * 100, 2),
        "length_difference": length_difference,
        "added_sections": added_sections,
        "removed_sections": removed_sections,
        "complexity_change": round(complexity_change, 2),
        "word_count_change": word_count_change,
        "sentence_count_change": sentence_count_change,
        "unique_words_change": unique_words_change,
        "structure_changes": structure_changes,
        "improvement": "increased" if complexity_change < 0 else "decreased" if complexity_change > 0 else "unchanged"
    }


def render_protocol_comparison(protocol_a: str, protocol_b: str, name_a: str = "Original", name_b: str = "Modified") -> str:
    """Render a visual comparison of two protocols.
    
    Args:
        protocol_a (str): First protocol text
        protocol_b (str): Second protocol text
        name_a (str): Name for first protocol
        name_b (str): Name for second protocol
        
    Returns:
        str: HTML formatted comparison
    """
    if not protocol_a or not protocol_b:
        return "<p>Cannot compare empty protocols</p>"
    
    comparison = compare_protocols(protocol_a, protocol_b)
    
    # Split protocols into lines for line-by-line comparison
    lines_a = protocol_a.split('\n')
    lines_b = protocol_b.split('\n')
    
    # Simple diff implementation
    html = f"""
    <div style="display: flex; gap: 20px; margin-bottom: 20px;">
        <div style="flex: 1;">
            <h3 style="color: #4a6fa5; border-bottom: 2px solid #4a6fa5; padding-bottom: 5px;">{name_a}</h3>
        </div>
        <div style="flex: 1;">
            <h3 style="color: #4a6fa5; border-bottom: 2px solid #4a6fa5; padding-bottom: 5px;">{name_b}</h3>
        </div>
    </div>
    """
    
    # Add summary metrics
    html += f"""
    <div style="background-color: #f8f9fa; padding: 15px; border-radius: 8px; margin-bottom: 20px;">
        <h4>üìä Comparison Summary</h4>
        <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px;">
            <div style="background-color: white; padding: 10px; border-radius: 5px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                <strong>Similarity</strong><br>
                <span style="font-size: 1.2em; color: {'#4caf50' if comparison['similarity'] > 80 else '#ff9800' if comparison['similarity'] > 50 else '#f44336'}">
                    {comparison['similarity']}%
                </span>
            </div>
            <div style="background-color: white; padding: 10px; border-radius: 5px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                <strong>Length Change</strong><br>
                <span style="font-size: 1.2em; color: {'#4caf50' if comparison['length_difference'] > 0 else '#f44336' if comparison['length_difference'] < 0 else '#666'}">
                    {comparison['length_difference']:+d} chars
                </span>
            </div>
            <div style="background-color: white; padding: 10px; border-radius: 5px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                <strong>Complexity Change</strong><br>
                <span style="font-size: 1.2em; color: {'#4caf50' if comparison['complexity_change'] < 0 else '#f44336' if comparison['complexity_change'] > 0 else '#666'}">
                    {comparison['complexity_change']:+.2f}
                </span>
            </div>
        </div>
    </div>
    """
    
    # Add structural changes
    if comparison["structure_changes"]:
        html += "<div style='background-color: #e3f2fd; padding: 15px; border-radius: 8px; margin-bottom: 20px;'>"
        html += "<h4>üß© Structural Changes</h4><ul>"
        for key, change in comparison["structure_changes"].items():
            html += f"<li><strong>{key.replace('_', ' ').title()}:</strong> "
            html += f"{'‚úÖ' if not change['before'] and change['after'] else '‚ùå' if change['before'] and not change['after'] else 'üîÑ'} "
            html += f"Changed from {change['before']} to {change['after']}</li>"
        html += "</ul></div>"
    
    # Add section changes
    if comparison["added_sections"] or comparison["removed_sections"]:
        html += "<div style='background-color: #fff3e0; padding: 15px; border-radius: 8px; margin-bottom: 20px;'>"
        html += "<h4>üìã Section Changes</h4>"
        if comparison["added_sections"]:
            html += "<p><strong>Added Sections:</strong></p><ul>"
            for section in comparison["added_sections"][:5]:  # Limit to first 5
                html += f"<li>‚ûï {section[:50]}{'...' if len(section) > 50 else ''}</li>"
            if len(comparison["added_sections"]) > 5:
                html += f"<li>... and {len(comparison['added_sections']) - 5} more</li>"
            html += "</ul>"
        if comparison["removed_sections"]:
            html += "<p><strong>Removed Sections:</strong></p><ul>"
            for section in comparison["removed_sections"][:5]:  # Limit to first 5
                html += f"<li>‚ûñ {section[:50]}{'...' if len(section) > 50 else ''}</li>"
            if len(comparison["removed_sections"]) > 5:
                html += f"<li>... and {len(comparison['removed_sections']) - 5} more</li>"
            html += "</ul>"
        html += "</div>"
    
    return html

def _rand_jitter_ms(base: int = 250, spread: int = 500) -> float:
    """Generate random jitter time in seconds for retry backoff.
    
    Args:
        base: Base milliseconds value
        spread: Random spread in milliseconds
        
    Returns:
        float: Random jitter time in seconds
    """
    return (base + random.randint(0, spread)) / 1000.0

def _hash_text(s: str) -> str:
    """Generate a short hash of text for comparison purposes.
    
    Args:
        s: Text to hash
        
    Returns:
        str: First 16 characters of SHA256 hash
    """
    return hashlib.sha256(s.encode("utf-8")).hexdigest()[:16]

def _approx_tokens(txt: str) -> int:
    """Approximate token count using character length.
    
    Args:
        txt: Text to estimate tokens for
        
    Returns:
        int: Estimated token count (minimum 1)
    """
    # A conservative approximation: 1 token ~ 4 chars in English. Clamp to a minimum of 1.
    return max(1, math.ceil(len(txt) / 4))

def _safe_json_loads(s: str) -> Optional[dict]:
    """Safely parse JSON string, returning None on failure.
    
    Args:
        s: JSON string to parse
        
    Returns:
        Optional[dict]: Parsed JSON object or None if invalid
    """
    if not isinstance(s, str):
        return None
    try:
        return json.loads(s)
    except (json.JSONDecodeError, TypeError):
        return None

def _safe_list(d: dict, key: str) -> list:
    """Safely retrieve a list from a dictionary, returning an empty list if not found or not a list.
    
    Args:
        d: Dictionary to retrieve from
        key: Key to look up
        
    Returns:
        list: The value at the key if it's a list, otherwise an empty list
    """
    if not isinstance(d, dict):
        return []
    value = d.get(key, [])
    return value if isinstance(value, list) else []

def _extract_json_block(txt: str) -> Optional[dict]:
    """Extract JSON object from text, handling fenced code blocks.
    
    Args:
        txt: Text containing potential JSON
        
    Returns:
        Optional[dict]: Extracted JSON object or None if not found
    """
    if not txt or not isinstance(txt, str):
        return None
    # Try to find a JSON block fenced with ```json
    m = FENCE_RE.search(txt)
    if m:
        cand = _safe_json_loads(m.group(1).strip())
        if cand is not None:
            return cand
    # If not found, try to find any substring that looks like a JSON object
    m2 = JSON_RE.search(txt)
    if m2:
        cand = _safe_json_loads(m2.group(0))
        if cand is not None:
            return cand
    return None

def _clamp(n: float, lo: float, hi: float) -> float:
    return max(lo, min(hi, n))

def _parse_price_per_million(v: Any) -> Optional[float]:
    """
    Safely parses OpenRouter pricing fields (e.g., "0.15", 0.15, "N/A")
    into a float representing price per 1M tokens, or None if invalid.
    """
    if v is None or isinstance(v, (dict, list)):
        return None
    try:
        return float(v)
    except (ValueError, TypeError):
        return None

def _cost_estimate(prompt_toks: int, completion_toks: int, ppm_prompt: Optional[float], ppm_comp: Optional[float]) -> float:
    """Calculates the estimated cost for a given number of tokens and prices per million."""
    cost = 0.0
    if ppm_prompt is not None and isinstance(ppm_prompt, (int, float)):
        cost += (prompt_toks / 1_000_000.0) * ppm_prompt
    if ppm_comp is not None and isinstance(ppm_comp, (int, float)):
        cost += (completion_toks / 1_000_000.0) * ppm_comp
    return cost

# ------------------------------------------------------------------
# 1. Generic helper ‚Äì fetch model lists with tiny caching layer
# ------------------------------------------------------------------

@functools.lru_cache(maxsize=128)
def _cached_get(url: str, bearer: str | None = None, timeout: int = 10) -> dict | list:
    headers = {}
    if bearer:
        headers["Authorization"] = f"Bearer {bearer}"
    r = requests.get(url, headers=headers, timeout=timeout)
    r.raise_for_status()
    return r.json()

# ------------------------------------------------------------------
# 2. Loader helpers for each provider that exposes an endpoint
# ------------------------------------------------------------------

def _openai_style_loader(url: str, api_key: str | None = None) -> List[str]:
    try:
        data = _cached_get(url, bearer=api_key)
        # Handle cases where data is a list (e.g., Together) or a dict with a 'data' key
        # (e.g., OpenAI)
        models_list = data if isinstance(data, list) else data.get("data", [])
        return sorted([m["id"] for m in models_list
                      if isinstance(m, dict) and "id" in m])
    except Exception as e:
        st.error(f"Error fetching models from {url}: {e}")
        return []

def _together_loader(api_key: str | None = None) -> List[str]:
    try:
        data = _cached_get("https://api.together.xyz/v1/models", bearer=api_key)
        return sorted([m["id"] for m in data
                      if isinstance(m, dict) and m.get("display_type") != "image"
                      and "id" in m])
    except Exception as e:
        st.error(f"Error fetching Together AI models: {e}")
        return []

def _fireworks_loader(api_key: str | None = None) -> List[str]:
    return _openai_style_loader("https://api.fireworks.ai/inference/v1/models", api_key)

def _groq_loader(api_key: str | None = None) -> List[str]:
    return _openai_style_loader("https://api.groq.com/openai/v1/models", api_key)

def _deepseek_loader(api_key: str | None = None) -> List[str]:
    return _openai_style_loader("https://api.deepseek.com/v1/models", api_key)

def _moonshot_loader(api_key: str | None = None) -> List[str]:
    return _openai_style_loader("https://api.moonshot.cn/v1/models", api_key)

def _baichuan_loader(api_key: str | None = None) -> List[str]:
    return _openai_style_loader("https://api.baichuan-ai.com/v1/models", api_key)

def _zhipu_loader(api_key: str | None = None) -> List[str]:
    # Zhipu's API might return models without an 'id', so we filter defensively
    try:
        data = _cached_get("https://open.bigmodel.cn/api/paas/v4/models",
                          bearer=api_key)
        return sorted([m["id"] for m in data.get("data", [])
                      if isinstance(m, dict) and "id" in m])
    except Exception as e:
        st.error(f"Error fetching Zhipu AI models: {e}")
        return []

def _minimax_loader(api_key: str | None = None) -> List[str]:
    return _openai_style_loader("https://api.minimax.chat/v1/models", api_key)

def _yi_loader(api_key: str | None = None) -> List[str]:
    return _openai_style_loader("https://api.lingyiwanwu.com/v1/models", api_key)


# ------------------------------------------------------------------
# 3. Central provider catalogue
# ------------------------------------------------------------------

PROVIDERS: dict[str, dict] = {
    # OpenAI official ------------------------------------------------
    "OpenAI": {
        "base": "https://api.openai.com/v1",
        "model": "gpt-4o-mini",
        "env": "OPENAI_API_KEY",
        "loader": lambda api_key=None: _openai_style_loader("https://api.openai.com/v1/models", api_key),
    },
    # Azure
    "Azure-OpenAI": {
        "base": "https://<your-resource>.openai.azure.com/openai/deployments/<deployment-name>",
        "model": "gpt-4o-mini", # This is often ignored as it's part of the deployment
        "env": "AZURE_OPENAI_API_KEY",
        "omit_model_in_payload": True, # Azure includes model in URL, not body
    },
    # Anthropic
    "Anthropic": {
        "base": "https://api.anthropic.com/v1", # Note: Anthropic has a non-OpenAI-compatible API
        "model": "claude-3-haiku-20240307",
        "env": "ANTHROPIC_API_KEY",
    },
    # Google Gemini
    "Google (Gemini)": {
        "base": "https://generativelanguage.googleapis.com/v1beta", # Note: Non-OpenAI-compatible API
        "model": "gemini-1.5-flash",
        "env": "GOOGLE_API_KEY",
    },
    # Mistral
    "Mistral": {
        "base": "https://api.mistral.ai/v1",
        "model": "mistral-small-latest",
        "env": "MISTRAL_API_KEY",
        "loader": lambda api_key=None: _openai_style_loader("https://api.mistral.ai/v1/models", api_key),
    },
    # Cohere
    "Cohere": {
        "base": "https://api.cohere.ai/v1", # Note: Non-OpenAI-compatible API
        "model": "command-r-plus",
        "env": "COHERE_API_KEY",
    },
    # Perplexity
    "Perplexity": {
        "base": "https://api.perplexity.ai",
        "model": "llama-3.1-sonar-small-128k-online",
        "env": "PERPLEXITY_API_KEY",
        "loader": lambda api_key=None: _openai_style_loader("https://api.perplexity.ai/models", api_key),
    },
    # Groq
    "Groq": {
        "base": "https://api.groq.com/openai/v1",
        "model": "llama-3.1-8b-instant",
        "env": "GROQ_API_KEY",
        "loader": _groq_loader,
    },
    # Databricks
    "Databricks": {
        "base": "https://<workspace>.cloud.databricks.com/serving-endpoints", # Note: Custom API structure
        "model": "databricks-dbrx-instruct",
        "env": "DATABRICKS_TOKEN",
    },
    # Together AI
    "Together": {
        "base": "https://api.together.xyz/v1",
        "model": "meta-llama/Llama-3-8b-chat-hf",
        "env": "TOGETHER_API_KEY",
        "loader": _together_loader,
    },
    # Fireworks
    "Fireworks": {
        "base": "https://api.fireworks.ai/inference/v1",
        "model": "accounts/fireworks/models/llama-v3-8b-instruct",
        "env": "FIREWORKS_API_KEY",
        "loader": _fireworks_loader,
    },
    # Replicate
    "Replicate": {
        "base": "https://api.replicate.com/v1", # Note: Non-OpenAI-compatible API
        "model": "meta/meta-llama-3-8b-instruct",
        "env": "REPLICATE_API_TOKEN",
    },
    # Anyscale Endpoints
    "Anyscale": {
        "base": "https://api.endpoints.anyscale.com/v1",
        "model": "meta-llama/Llama-3.1-8B-Instruct",
        "env": "ANYSCALE_API_KEY",
        "loader": lambda api_key=None: _openai_style_loader("https://api.endpoints.anyscale.com/v1/models", api_key),
    },
    # OpenRouter
    "OpenRouter": {
        "base": "https://openrouter.ai/api/v1",
        "model": "openai/gpt-4o-mini",
        "env": "OPENROUTER_API_KEY",
        "loader": lambda api_key=None: _openai_style_loader("https://openrouter.ai/api/v1/models", api_key),
    },
    # DeepInfra
    "DeepInfra": {
        "base": "https://api.deepinfra.com/v1/openai",
        "model": "meta-llama/Llama-3.1-8B-Instruct",
        "env": "DEEPINFRA_API_KEY",
        "loader": lambda api_key=None: _openai_style_loader("https://api.deepinfra.com/v1/models", api_key),
    },
    # OctoAI
    "OctoAI": {
        "base": "https://text.octoai.run/v1",
        "model": "meta-llama-3-8b-instruct",
        "env": "OCTOAI_TOKEN",
    },
    # AI21
    "AI21": {
        "base": "https://api.ai21.com/studio/v1", # Note: Non-OpenAI-compatible API
        "model": "jamba-instruct",
        "env": "AI21_API_KEY",
    },
    # Aleph-Alpha
    "AlephAlpha": {
        "base": "https://api.aleph-alpha.com", # Note: Non-OpenAI-compatible API
        "model": "luminous-supreme-control",
        "env": "ALEPH_ALPHA_API_KEY",
    },
    # Bedrock variants
    "Bedrock-Claude": {
        "base": "https://bedrock-runtime.<region>.amazonaws.com", # Note: AWS Signature v4 auth needed
        "model": "anthropic.claude-3-haiku-20240307-v1:0",
        "env": "AWS_SECRET_ACCESS_KEY",
    },
    "Bedrock-Titan": {
        "base": "https://bedrock-runtime.<region>.amazonaws.com",
        "model": "amazon.titan-text-express-v1",
        "env": "AWS_SECRET_ACCESS_KEY",
    },
    "Bedrock-Cohere": {
        "base": "https://bedrock-runtime.<region>.amazonaws.com",
        "model": "cohere.command-text-v14",
        "env": "AWS_SECRET_ACCESS_KEY",
    },
    "Bedrock-Jurassic": {
        "base": "https://bedrock-runtime.<region>.amazonaws.com",
        "model": "ai21.j2-ultra-v1",
        "env": "AWS_SECRET_ACCESS_KEY",
    },
    "Bedrock-Llama": {
        "base": "https://bedrock-runtime.<region>.amazonaws.com",
        "model": "meta.llama3-1-8b-instruct-v1:0",
        "env": "AWS_SECRET_ACCESS_KEY",
    },
    # Hugging Face Inference
    "HuggingFace": {
        "base": "https://api-inference.huggingface.co/models", # Note: Non-OpenAI-compatible API
        "model": "microsoft/DialoGPT-medium",
        "env": "HF_API_KEY",
    },
    # Local / self-hosted
    "Ollama": {
        "base": "http://localhost:11434/v1",
        "model": "llama3.1",
        "env": None,
    },
    "LM-Studio": {
        "base": "http://localhost:1234/v1",
        "model": "local-model",
        "env": None,
    },
    "vLLM": {
        "base": "http://localhost:8000/v1",
        "model": "meta-llama/Llama-3.1-8B-Instruct",
        "env": None,
    },
    # SageMaker
    "SageMaker": {
        "base": "https://runtime.sagemaker.<region>.amazonaws.com/endpoints/<endpoint>/invocations", # AWS Sig v4
        "model": "jumpstart-dft-meta-textgeneration-llama-3-1-8b",
        "env": "AWS_SECRET_ACCESS_KEY",
    },
    # Cloudflare Workers AI
    "Cloudflare": {
        "base": "https://api.cloudflare.com/client/v4/accounts/<account>/ai/run", # Non-OpenAI format
        "model": "@cf/meta/llama-3.1-8b-instruct-awq",
        "env": "CLOUDFLARE_API_TOKEN",
    },
    # Vertex AI
    "VertexAI": {
        "base": "https://<region>-aiplatform.googleapis.com/v1/projects/<project>/locations/<region>/publishers/google/models", # Non-OpenAI
        "model": "gemini-1.5-flash",
        "env": "GOOGLE_APPLICATION_CREDENTIALS",
    },
    # Chinese providers
    "Moonshot": {
        "base": "https://api.moonshot.cn/v1",
        "model": "moonshot-v1-8k",
        "env": "MOONSHOT_API_KEY",
        "loader": _moonshot_loader,
    },
    "Baichuan": {
        "base": "https://api.baichuan-ai.com/v1",
        "model": "Baichuan3-Turbo",
        "env": "BAICHUAN_API_KEY",
        "loader": _baichuan_loader,
    },
    "Zhipu": {
        "base": "https://open.bigmodel.cn/api/paas/v4",
        "model": "glm-4",
        "env": "ZHIPU_API_KEY",
        "loader": _zhipu_loader,
    },
    "MiniMax": {
        "base": "https://api.minimax.chat/v1",
        "model": "abab6.5-chat",
        "env": "MINIMAX_API_KEY",
        "loader": _minimax_loader,
    },
    "Yi": {
        "base": "https://api.lingyiwanwu.com/v1",
        "model": "yi-large",
        "env": "YI_API_KEY",
        "loader": _yi_loader,
    },
    "DeepSeek": {
        "base": "https://api.deepseek.com/v1",
        "model": "deepseek-chat",
        "env": "DEEPSEEK_API_KEY",
        "loader": _deepseek_loader,
    },
    # Bring-your-own
    "Custom": {"base": "", "model": "", "env": None},
}


# ------------------------------------------------------------------
# 4. Session-state helpers
# ------------------------------------------------------------------

# Thread lock for safely updating shared session state from background threads.
# Use a lock to ensure thread-safe initialization of the session state lock
if "thread_lock" not in st.session_state:
    with threading.Lock():
        # Double-checked locking pattern to ensure thread safety
        if "thread_lock" not in st.session_state:
            st.session_state.thread_lock = threading.Lock()

# Real-time collaboration state management
if "collaboration_session" not in st.session_state:
    st.session_state.collaboration_session = {
        "active_users": [],
        "last_activity": _now_ms(),
        "chat_messages": [],
        "notifications": [],
        "shared_cursor_position": 0,
        "edit_locks": {}
    }

# Configuration profiles
CONFIG_PROFILES = {
    "Security Hardening": {
        "system_prompt": "You are a security expert focused on making protocols robust against attacks. Focus on identifying and closing security gaps, enforcing least privilege, and adding comprehensive error handling.",
        "evaluator_system_prompt": "You are a security auditor evaluating the protocol for vulnerabilities and weaknesses.",
        "temperature": 0.8,
        "top_p": 0.9,
        "max_iterations": 15,
        "adversarial_confidence": 95,
        "adversarial_min_iter": 5,
        "adversarial_max_iter": 20
    },
    "Compliance Focus": {
        "system_prompt": "You are a compliance expert ensuring protocols meet regulatory requirements. Focus on completeness, auditability, and regulatory alignment.",
        "evaluator_system_prompt": "You are a compliance auditor checking if the protocol meets all necessary regulatory requirements.",
        "temperature": 0.5,
        "top_p": 0.8,
        "max_iterations": 10,
        "adversarial_confidence": 90,
        "adversarial_min_iter": 3,
        "adversarial_max_iter": 15
    },
    "Operational Efficiency": {
        "system_prompt": "You are an operations expert focused on making protocols efficient and practical. Focus on streamlining processes while maintaining effectiveness.",
        "evaluator_system_prompt": "You are an operations expert evaluating the protocol for practicality and efficiency.",
        "temperature": 0.6,
        "top_p": 0.85,
        "max_iterations": 12,
        "adversarial_confidence": 85,
        "adversarial_min_iter": 3,
        "adversarial_max_iter": 12
    },
    "Beginner-Friendly": {
        "system_prompt": "You are helping a beginner write clear, understandable protocols. Focus on clarity, simplicity, and completeness.",
        "evaluator_system_prompt": "You are evaluating if the protocol is clear and understandable for beginners.",
        "temperature": 0.7,
        "top_p": 1.0,
        "max_iterations": 8,
        "adversarial_confidence": 80,
        "adversarial_min_iter": 2,
        "adversarial_max_iter": 10
    }
}

# Protocol templates
PROTOCOL_TEMPLATES = {
    "Security Policy": """# Security Policy Template

## Overview
[Brief description of the policy's purpose and scope]

## Scope
[Define what systems, processes, and personnel are covered by this policy]

## Policy Statements
[Specific security requirements and guidelines]

## Roles and Responsibilities
[Define who is responsible for what aspects of the policy]

## Compliance
[How compliance will be measured and enforced]

## Exceptions
[Process for requesting policy exceptions]

## Review and Updates
[How often the policy will be reviewed and updated]""",
    
    "Standard Operating Procedure": """# Standard Operating Procedure (SOP) Template

## Title
[Name of the procedure]

## Purpose
[Why this procedure exists]

## Scope
[What this procedure covers and who it applies to]

## Responsibilities
[Who is responsible for each step]

## Procedure
1. [First step]
   - [Detailed instructions]
   - [Expected outcomes]
2. [Second step]
   - [Detailed instructions]
   - [Expected outcomes]

## Safety Considerations
[Any safety risks and how to mitigate them]

## Quality Control
[How to ensure quality and consistency]

## Documentation
[What records need to be maintained]

## Revision History
[Track changes to the procedure]""",
    
    "Incident Response Plan": """# Incident Response Plan Template

## Overview
[Brief description of the plan's purpose]

## Incident Classification
[Types of incidents and severity levels]

## Response Team
[Key personnel and their roles]

## Detection and Reporting
[How incidents are detected and reported]

## Containment
[Immediate actions to limit impact]

## Eradication
[Steps to remove the threat]

## Recovery
[How to restore normal operations]

## Post-Incident Activities
[Lessons learned and plan updates]

## Communication Plan
[Who to notify and when]

## Contact Information
[Key contacts and their availability]""",
    
    "Software Development Process": """# Software Development Process Template

## Overview
[Brief description of the development process]

## Scope
[What types of projects this process applies to]

## Roles and Responsibilities
- Project Manager: [Responsibilities]
- Developers: [Responsibilities]
- QA Engineers: [Responsibilities]
- DevOps Engineers: [Responsibilities]

## Development Lifecycle
### 1. Requirements Gathering
- [Process for collecting requirements]
- [Stakeholder involvement]

### 2. Design
- [System architecture design]
- [UI/UX design]
- [Database design]

### 3. Implementation
- [Coding standards]
- [Version control practices]
- [Code review process]

### 4. Testing
- [Unit testing]
- [Integration testing]
- [System testing]
- [User acceptance testing]

### 5. Deployment
- [Deployment process]
- [Rollback procedures]
- [Monitoring]

## Quality Assurance
[QA processes and standards]

## Documentation
[Required documentation at each stage]

## Tools and Technologies
[List of tools used in the process]

## Metrics and KPIs
[Key performance indicators to track]

## Review and Improvement
[Process for continuous improvement]""",
    
    "Data Privacy Policy": """# Data Privacy Policy Template

## Overview
[Statement of commitment to data privacy]

## Scope
[What data and processes this policy covers]

## Legal Compliance
[List of applicable regulations (GDPR, CCPA, etc.)]

## Data Collection
[What data is collected and why]

## Data Usage
[How collected data is used]

## Data Storage
[Where and how data is stored]

## Data Sharing
[When and with whom data may be shared]

## Data Retention
[How long data is retained]

## Individual Rights
- Right to Access
- Right to Rectification
- Right to Eradication
- Right to Restrict Processing
- Right to Data Portability
- Right to Object

## Security Measures
[Technical and organizational measures to protect data]

## Breach Notification
[Process for reporting data breaches]

## Training and Awareness
[Employee training requirements]

## Policy Enforcement
[Consequences for policy violations]

## Review and Updates
[How often the policy is reviewed]""",
    
    "Business Continuity Plan": """# Business Continuity Plan Template

## Overview
[Purpose and scope of the business continuity plan]

## Risk Assessment
[Identified risks and their potential impact]

## Business Impact Analysis
[Critical business functions and maximum tolerable downtime]

## Recovery Strategies
[Strategies for recovering critical functions]

## Emergency Response
### 1. Incident Declaration
[Criteria for declaring an emergency]

### 2. Emergency Response Team
- Team Members: [List]
- Contact Information: [Details]
- Roles and Responsibilities: [Details]

### 3. Communication Plan
[Internal and external communication procedures]

## Recovery Procedures
### Critical Function 1
- Recovery Steps: [Detailed steps]
- Resources Required: [List]
- Recovery Time Objective: [Timeframe]

### Critical Function 2
- Recovery Steps: [Detailed steps]
- Resources Required: [List]
- Recovery Time Objective: [Timeframe]

## Plan Testing and Maintenance
[Testing schedule and procedures]

## Training and Awareness
[Training requirements for personnel]

## Plan Distribution
[List of plan recipients]

## Plan Activation and Deactivation
[Criteria and procedures for plan activation and deactivation]""",
    
    "API Security Review Checklist": """# API Security Review Checklist Template

## Overview
[Description of the API and its purpose]

## Authentication
- [ ] Authentication mechanism implemented
- [ ] Strong password policies enforced
- [ ] Multi-factor authentication supported
- [ ] Session management secure

## Authorization
- [ ] Role-based access control implemented
- [ ] Permissions properly configured
- [ ] Least privilege principle applied
- [ ] Access controls tested

## Input Validation
- [ ] All inputs validated
- [ ] SQL injection protection implemented
- [ ] Cross-site scripting (XSS) prevention
- [ ] File upload restrictions in place

## Data Protection
- [ ] Data encryption in transit (TLS)
- [ ] Data encryption at rest
- [ ] Sensitive data masked in logs
- [ ] Personal data handling compliant

## Error Handling
- [ ] Descriptive error messages suppressed
- [ ] Error logging implemented
- [ ] Exception handling in place
- [ ] Stack traces not exposed

## Rate Limiting
- [ ] Rate limiting implemented
- [ ] Throttling configured
- [ ] Brute force protection
- [ ] DDoS protection measures

## Security Headers
- [ ] Content Security Policy (CSP) implemented
- [ ] X-Frame-Options set
- [ ] X-Content-Type-Options set
- [ ] Strict-Transport-Security configured

## API Gateway Security
- [ ] API gateway configured
- [ ] Traffic monitoring enabled
- [ ] Threat detection implemented
- [ ] Request/response filtering

## Third-Party Dependencies
- [ ] Dependencies regularly updated
- [ ] Vulnerability scanning performed
- [ ] Security patches applied
- [ ] Dependency security monitoring

## Logging and Monitoring
- [ ] Security events logged
- [ ] Audit trail maintained
- [ ] Anomaly detection configured
- [ ] Alerting mechanisms in place

## Compliance
- [ ] GDPR compliance (if applicable)
- [ ] HIPAA compliance (if applicable)
- [ ] PCI DSS compliance (if applicable)
- [ ] Industry-specific regulations met

## Review and Approval
- Security Reviewer: [Name]
- Review Date: [Date]
- Approval Status: [Approved/Rejected/Pending]
- Notes: [Additional comments]""",
    
    "DevOps Workflow": """# DevOps Workflow Template

## Overview
[Brief description of the DevOps workflow and its objectives]

## Scope
[What systems, applications, and environments this workflow covers]

## Roles and Responsibilities
- DevOps Engineer: [Responsibilities]
- Developers: [Responsibilities]
- QA Engineers: [Responsibilities]
- Security Team: [Responsibilities]

## CI/CD Pipeline
### 1. Code Commit
- Branching strategy: [e.g., GitFlow, GitHub Flow]
- Code review process: [Description]
- Static code analysis: [Tools and criteria]

### 2. Continuous Integration
- Automated build process: [Description]
- Unit test execution: [Process]
- Integration test execution: [Process]
- Security scanning: [Tools and criteria]

### 3. Continuous Deployment
- Deployment environments: [List]
- Deployment approval process: [Description]
- Rollback procedures: [Process]
- Monitoring setup: [Tools and metrics]

## Infrastructure as Code
- Tools used: [e.g., Terraform, CloudFormation]
- Version control: [Repository structure]
- Review process: [Approval workflow]
- Testing strategy: [How infrastructure changes are tested]

## Monitoring and Observability
- Metrics collection: [Tools and what is measured]
- Log aggregation: [Tools and retention policy]
- Alerting thresholds: [What triggers alerts]
- Incident response: [Process for handling alerts]

## Security Practices
- Vulnerability scanning: [Schedule and tools]
- Compliance checks: [Process and tools]
- Secret management: [How secrets are handled]
- Access control: [How access is managed]

## Backup and Recovery
- Backup strategy: [What is backed up and how often]
- Recovery time objectives: [RTO targets]
- Recovery point objectives: [RPO targets]
- Testing schedule: [How often recovery is tested]

## Documentation
- Runbooks: [Location and update process]
- Architecture diagrams: [Location and update process]
- Onboarding guides: [For new team members]

## Review and Improvement
- Retrospectives: [Schedule and process]
- KPI tracking: [Metrics monitored]
- Continuous improvement: [Process for implementing changes]""",
    
    "Risk Assessment Framework": """# Risk Assessment Framework Template

## Overview
[Purpose and scope of the risk assessment framework]

## Risk Categories
- Operational Risks: [Description]
- Security Risks: [Description]
- Compliance Risks: [Description]
- Financial Risks: [Description]
- Reputational Risks: [Description]

## Risk Assessment Process
### 1. Risk Identification
- Methods: [Brainstorming, historical data, expert interviews, etc.]
- Participants: [Who is involved]
- Frequency: [How often assessments are conducted]

### 2. Risk Analysis
- Qualitative analysis: [Method and criteria]
- Quantitative analysis: [Method and criteria]
- Risk owners: [Who is responsible for each risk]

### 3. Risk Evaluation
- Risk appetite: [Organization's tolerance for risk]
- Risk criteria: [How risks are prioritized]
- Risk matrix: [Likelihood vs Impact matrix]

### 4. Risk Treatment
- Avoidance: [When and how risks are avoided]
- Mitigation: [How risks are reduced]
- Transfer: [How risks are transferred]
- Acceptance: [How risks are accepted]

## Risk Monitoring
- Key risk indicators: [Metrics tracked]
- Reporting frequency: [How often reports are generated]
- Escalation procedures: [When and how risks are escalated]

## Roles and Responsibilities
- Risk Manager: [Responsibilities]
- Risk Owners: [Responsibilities]
- Senior Management: [Responsibilities]

## Documentation
- Risk register: [Format and maintenance]
- Assessment reports: [Template and distribution]
- Action plans: [Format and tracking]

## Review and Updates
- Framework review: [Frequency and process]
- Lessons learned: [How insights are captured]
- Continuous improvement: [Process for enhancing the framework]""",
    
    "Disaster Recovery Plan": """# Disaster Recovery Plan Template

## Overview
[Brief description of the disaster recovery plan's purpose]

## Scope
[What systems, applications, and data are covered by this plan]

## Recovery Objectives
- Recovery Time Objective (RTO): [Maximum acceptable downtime]
- Recovery Point Objective (RPO): [Maximum acceptable data loss]

## Critical Systems and Applications
[Identify and prioritize critical systems and applications]

## Disaster Recovery Team
- Team Members: [List]
- Contact Information: [Details]
- Roles and Responsibilities: [Details]

## Disaster Recovery Sites
- Primary Site: [Location and details]
- Backup Site: [Location and details]
- Hot/Warm/Cold Site: [Specify type and capabilities]

## Data Backup and Recovery
- Backup Schedule: [Frequency and methods]
- Backup Storage: [Locations and security]
- Recovery Procedures: [Step-by-step instructions]

## Communication Plan
- Internal Communication: [How to notify employees]
- External Communication: [How to notify customers, vendors, etc.]
- Media Relations: [How to handle media inquiries]

## Recovery Procedures
### System Restoration
- [Step-by-step instructions for restoring systems]
- [Required resources and personnel]

### Data Recovery
- [Step-by-step instructions for recovering data]
- [Validation procedures to ensure data integrity]

### Application Recovery
- [Step-by-step instructions for recovering applications]
- [Testing procedures to ensure functionality]

## Plan Testing and Maintenance
- Testing Schedule: [How often the plan will be tested]
- Testing Procedures: [Methods for testing the plan]
- Update Procedures: [How and when the plan will be updated]

## Training and Awareness
- Training Schedule: [How often personnel will be trained]
- Training Materials: [Resources for training personnel]
- Awareness Program: [How to keep personnel informed]

## Plan Distribution
[List of plan recipients and distribution methods]

## Plan Activation and Deactivation
- Activation Criteria: [When to activate the plan]
- Deactivation Criteria: [When to deactivate the plan]
- Activation Procedures: [How to activate the plan]
- Deactivation Procedures: [How to deactivate the plan]""",
    
    "Change Management Process": """# Change Management Process Template

## Overview
[Brief description of the change management process]

## Scope
[What types of changes are covered by this process]

## Change Categories
- Emergency Changes: [Description and criteria]
- Standard Changes: [Description and criteria]
- Normal Changes: [Description and criteria]
- Major Changes: [Description and criteria]

## Change Management Team
- Change Manager: [Name and contact information]
- Change Advisory Board (CAB): [Members and roles]
- Change Requester: [Role and responsibilities]
- Change Implementer: [Role and responsibilities]

## Change Request Process
### 1. Change Request Submission
- Request Form: [Template and submission process]
- Required Information: [What information must be provided]

### 2. Change Request Review
- Initial Review: [Who conducts the review and criteria]
- Detailed Assessment: [How the change is assessed]
- Risk Assessment: [How risks are identified and evaluated]

### 3. Change Approval
- Approval Authority: [Who has approval authority]
- Approval Criteria: [What criteria are used for approval]
- Approval Process: [How approvals are obtained]

### 4. Change Implementation
- Implementation Plan: [How the change is implemented]
- Implementation Schedule: [When the change is implemented]
- Implementation Team: [Who implements the change]

### 5. Change Closure
- Validation: [How the change is validated]
- Documentation: [How the change is documented]
- Communication: [How stakeholders are informed]

## Change Management Database
- Database Structure: [How change information is organized]
- Data Retention: [How long change information is retained]
- Reporting: [What reports are generated]

## Training and Awareness
- Training Program: [How personnel are trained]
- Awareness Program: [How personnel are kept informed]

## Plan Review and Improvement
- Review Schedule: [How often the process is reviewed]
- Improvement Process: [How improvements are made]""",
    
    "Business Impact Analysis": """# Business Impact Analysis Template

## Overview
[Brief description of the business impact analysis]

## Scope
[What business functions and processes are covered]

## Business Functions
[List of critical business functions]

## Dependencies
[Dependencies between business functions]

## Impact Criteria
- Financial Impact: [How financial impact is measured]
- Operational Impact: [How operational impact is measured]
- Legal/Regulatory Impact: [How legal/regulatory impact is measured]
- Reputational Impact: [How reputational impact is measured]

## Recovery Time Objectives
[Recovery time objectives for each business function]

## Recovery Point Objectives
[Recovery point objectives for each business function]

## Minimum Business Continuity Objective
[Minimum business continuity objective for each business function]

## Resource Requirements
[Resources required to recover each business function]

## Interim Recovery Strategies
[Interim recovery strategies for each business function]

## Long-Term Recovery Strategies
[Long-term recovery strategies for each business function]

## Plan Testing and Maintenance
[How the business impact analysis is tested and maintained]

## Review and Approval
[How the business impact analysis is reviewed and approved]"""
}

# Adversarial Testing Presets
ADVERSARIAL_PRESETS = {
    "Security Hardening": {
        "name": "üîê Security Hardening",
        "description": "Focus on identifying and closing security gaps, enforcing least privilege, and adding comprehensive error handling.",
        "red_team_models": ["openai/gpt-4o-mini", "anthropic/claude-3-haiku", "google/gemini-1.5-flash"],
        "blue_team_models": ["openai/gpt-4o", "anthropic/claude-3-sonnet", "google/gemini-1.5-pro"],
        "min_iter": 5,
        "max_iter": 15,
        "confidence_threshold": 95,
        "review_type": "General SOP",
        "compliance_requirements": "Security best practices, OWASP guidelines, least privilege principle",
        "advanced_settings": {
            "critique_depth": 8,
            "patch_quality": 9,
            "detailed_tracking": True,
            "early_stopping": True
        }
    },
    "Compliance Focus": {
        "name": "‚öñÔ∏è Compliance Focus",
        "description": "Ensure protocols meet regulatory requirements with comprehensive auditability.",
        "red_team_models": ["openai/gpt-4o-mini", "mistral/mistral-small-latest"],
        "blue_team_models": ["openai/gpt-4o", "mistral/mistral-medium-latest"],
        "min_iter": 3,
        "max_iter": 10,
        "confidence_threshold": 90,
        "review_type": "General SOP",
        "compliance_requirements": "GDPR, ISO 27001, SOC 2, industry-specific regulations",
        "advanced_settings": {
            "critique_depth": 7,
            "patch_quality": 8,
            "detailed_tracking": True,
            "performance_analytics": True
        }
    },
    "Operational Efficiency": {
        "name": "‚ö° Operational Efficiency",
        "description": "Streamline processes while maintaining effectiveness and clarity.",
        "red_team_models": ["openai/gpt-4o-mini", "meta-llama/llama-3-8b-instruct"],
        "blue_team_models": ["openai/gpt-4o", "meta-llama/llama-3-70b-instruct"],
        "min_iter": 3,
        "max_iter": 12,
        "confidence_threshold": 85,
        "review_type": "General SOP",
        "compliance_requirements": "Process optimization, resource efficiency, clarity standards",
        "advanced_settings": {
            "critique_depth": 6,
            "patch_quality": 7,
            "early_stopping": True,
            "target_complexity": 50
        }
    },
    "Beginner-Friendly": {
        "name": "üë∂ Beginner-Friendly",
        "description": "Focus on clarity, simplicity, and completeness for newcomers.",
        "red_team_models": ["openai/gpt-4o-mini", "google/gemini-1.5-flash"],
        "blue_team_models": ["openai/gpt-4o", "google/gemini-1.5-pro"],
        "min_iter": 2,
        "max_iter": 8,
        "confidence_threshold": 80,
        "review_type": "General SOP",
        "compliance_requirements": "Clear language, simple concepts, comprehensive examples",
        "advanced_settings": {
            "critique_depth": 5,
            "patch_quality": 8,
            "target_complexity": 30,
            "target_length": 500
        }
    },
    "Code Review": {
        "name": "üíª Code Review",
        "description": "Specialized testing for software development protocols and code reviews.",
        "red_team_models": ["openai/gpt-4o", "anthropic/claude-3-opus", "codellama/codellama-70b-instruct"],
        "blue_team_models": ["openai/gpt-4o", "anthropic/claude-3-sonnet", "codellama/codellama-70b-instruct"],
        "min_iter": 3,
        "max_iter": 10,
        "confidence_threshold": 90,
        "review_type": "Code Review",
        "compliance_requirements": "Clean code principles, security best practices, performance optimization",
        "advanced_settings": {
            "critique_depth": 9,
            "patch_quality": 9,
            "detailed_tracking": True,
            "performance_analytics": True
        }
    },
    "Mission Critical": {
        "name": "üî• Mission Critical",
        "description": "Maximum rigor for high-stakes protocols requiring the highest assurance.",
        "red_team_models": ["openai/gpt-4o", "anthropic/claude-3-opus", "google/gemini-1.5-pro"],
        "blue_team_models": ["openai/gpt-4o", "anthropic/claude-3-sonnet", "google/gemini-1.5-pro"],
        "min_iter": 10,
        "max_iter": 25,
        "confidence_threshold": 98,
        "review_type": "General SOP",
        "compliance_requirements": "Highest security standards, fault tolerance, disaster recovery",
        "advanced_settings": {
            "critique_depth": 10,
            "patch_quality": 10,
            "detailed_tracking": True,
            "performance_analytics": True,
            "early_stopping": False
        }
    },
    "AI Safety Review": {
        "name": "üõ°Ô∏è AI Safety Review",
        "description": "Specialized testing for AI safety considerations, bias detection, and ethical alignment.",
        "red_team_models": ["openai/gpt-4o", "anthropic/claude-3-opus", "google/gemini-1.5-pro"],
        "blue_team_models": ["openai/gpt-4o", "anthropic/claude-3-sonnet", "meta-llama/llama-3-70b-instruct"],
        "min_iter": 5,
        "max_iter": 15,
        "confidence_threshold": 92,
        "review_type": "General SOP",
        "compliance_requirements": "AI safety guidelines, fairness principles, explainability requirements",
        "advanced_settings": {
            "critique_depth": 9,
            "patch_quality": 9,
            "detailed_tracking": True,
            "performance_analytics": True,
            "bias_detection": True,
            "explainability_focus": True
        }
    },
    "Privacy Protection": {
        "name": "üîí Privacy Protection",
        "description": "Focus on data privacy protection, consent mechanisms, and regulatory compliance.",
        "red_team_models": ["openai/gpt-4o-mini", "anthropic/claude-3-haiku", "google/gemini-1.5-flash"],
        "blue_team_models": ["openai/gpt-4o", "anthropic/claude-3-sonnet", "google/gemini-1.5-pro"],
        "min_iter": 4,
        "max_iter": 12,
        "confidence_threshold": 93,
        "review_type": "General SOP",
        "compliance_requirements": "GDPR, CCPA, PIPEDA, data minimization principles",
        "advanced_settings": {
            "critique_depth": 8,
            "patch_quality": 8,
            "detailed_tracking": True,
            "privacy_by_design": True,
            "consent_mechanisms": True
        }
    }
}

DEFAULTS = {
    "provider": "OpenAI",
    "api_key": "",
    "base_url": PROVIDERS["OpenAI"]["base"],
    "model": PROVIDERS["OpenAI"]["model"],
    "extra_headers": "{}",
    "max_tokens": 4096,
    "temperature": 0.7,
    "top_p": 1.0,
    "frequency_penalty": 0.0,
    "presence_penalty": 0.0,
    "seed": "",
    "protocol_text": "",
    "system_prompt": "You are an assistant that makes the draft airtight, precise, and foolproof.",
    "evaluator_system_prompt": "You are a trivial evaluator that accepts everything.",
    "max_iterations": 20,
    "population_size": 1,
    "num_islands": 1,
    "checkpoint_interval": 5,
    "elite_ratio": 1.0,
    "exploration_ratio": 0.0,
    "exploitation_ratio": 0.0,
    "archive_size": 0,
    "evolution_running": False,
    "evolution_log": [],
    "evolution_current_best": "",
    "evolution_stop_flag": False,
    "openrouter_key": "",
    "red_team_models": [],
    "blue_team_models": [],
    "adversarial_running": False,
    "adversarial_results": {},
    "adversarial_status_message": "Idle.",
    "adversarial_log": [],
    "adversarial_stop_flag": False,
    "adversarial_cost_estimate_usd": 0.0,
    "adversarial_total_tokens_prompt": 0,
    "adversarial_total_tokens_completion": 0,
    "adversarial_min_iter": 3,
    "adversarial_max_iter": 10,
    "adversarial_confidence": 95,
    "adversarial_max_tokens": 8000,
    "adversarial_max_workers": 6,
    "adversarial_force_json": True,
    "adversarial_seed": "",
    "adversarial_rotation_strategy": "None",
    "adversarial_red_team_sample_size": 3,
    "adversarial_blue_team_sample_size": 3,
    "adversarial_model_performance": {},
    "adversarial_confidence_history": [],
    "adversarial_staged_rotation_config": "",
    "compliance_requirements": "",
    # Collaborative features
    "project_name": "Untitled Project",
    "collaborators": [],
    "comments": [],
    "protocol_versions": [],
    "current_version_id": "",
    "tags": [],
    "project_description": "",
    # Tutorial and onboarding
    "tutorial_completed": False,
    "current_tutorial_step": 0,
}

for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

def reset_defaults():
    p = st.session_state.provider
    if p in PROVIDERS:
        st.session_state.base_url = PROVIDERS[p].get("base", "")
        st.session_state.model = PROVIDERS[p].get("model") or ""
    st.session_state.api_key = ""
    st.session_state.extra_headers = "{}"

def save_config_profile(profile_name: str) -> bool:
    """Save current configuration as a profile.
    
    Args:
        profile_name (str): Name for the profile
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        profile_data = {
            "provider": st.session_state.provider,
            "base_url": st.session_state.base_url,
            "model": st.session_state.model,
            "temperature": st.session_state.temperature,
            "top_p": st.session_state.top_p,
            "frequency_penalty": st.session_state.frequency_penalty,
            "presence_penalty": st.session_state.presence_penalty,
            "max_tokens": st.session_state.max_tokens,
            "max_iterations": st.session_state.max_iterations,
            "system_prompt": st.session_state.system_prompt,
            "adversarial_confidence": st.session_state.adversarial_confidence,
            "adversarial_min_iter": st.session_state.adversarial_min_iter,
            "adversarial_max_iter": st.session_state.adversarial_max_iter,
            "adversarial_max_tokens": st.session_state.adversarial_max_tokens,
        }
        
        # Save to session state
        if "config_profiles" not in st.session_state:
            st.session_state.config_profiles = {}
        st.session_state.config_profiles[profile_name] = profile_data
        return True
    except Exception as e:
        st.error(f"Error saving profile: {e}")
        return False

def load_config_profile(profile_name: str) -> bool:
    """Load a configuration profile.
    
    Args:
        profile_name (str): Name of the profile to load
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        if "config_profiles" not in st.session_state:
            st.session_state.config_profiles = {}
            
        if profile_name not in st.session_state.config_profiles:
            # Check if it's a built-in profile
            if profile_name in CONFIG_PROFILES:
                profile_data = CONFIG_PROFILES[profile_name]
            else:
                st.error(f"Profile '{profile_name}' not found.")
                return False
        else:
            profile_data = st.session_state.config_profiles[profile_name]
            
        # Apply profile data to session state
        for key, value in profile_data.items():
            if key in st.session_state:
                st.session_state[key] = value
                
        return True
    except Exception as e:
        st.error(f"Error loading profile: {e}")
        return False

def list_config_profiles() -> List[str]:
    """List all available configuration profiles.
    
    Returns:
        List[str]: List of profile names
    """
    profiles = list(CONFIG_PROFILES.keys())
    if "config_profiles" in st.session_state:
        profiles.extend(list(st.session_state.config_profiles.keys()))
    return sorted(list(set(profiles)))

def list_protocol_templates() -> List[str]:
    """List all available protocol templates.
    
    Returns:
        List[str]: List of template names
    """
    return list(PROTOCOL_TEMPLATES.keys())

def load_protocol_template(template_name: str) -> str:
    """Load a protocol template.
    
    Args:
        template_name (str): Name of the template to load
        
    Returns:
        str: Template content
    """
    return PROTOCOL_TEMPLATES.get(template_name, "")

def list_adversarial_presets() -> List[str]:
    """List all available adversarial testing presets.
    
    Returns:
        List[str]: List of preset names
    """
    return list(ADVERSARIAL_PRESETS.keys())

def load_adversarial_preset(preset_name: str) -> Dict:
    """Load an adversarial testing preset.
    
    Args:
        preset_name (str): Name of the preset to load
        
    Returns:
        Dict: Preset configuration
    """
    return ADVERSARIAL_PRESETS.get(preset_name, {})

def apply_adversarial_preset(preset_name: str) -> bool:
    """Apply an adversarial testing preset to the current session state.
    
    Args:
        preset_name (str): Name of the preset to apply
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        preset = load_adversarial_preset(preset_name)
        if not preset:
            return False
            
        # Apply preset configuration to session state
        st.session_state.red_team_models = preset.get("red_team_models", [])
        st.session_state.blue_team_models = preset.get("blue_team_models", [])
        st.session_state.adversarial_min_iter = preset.get("min_iter", 3)
        st.session_state.adversarial_max_iter = preset.get("max_iter", 10)
        st.session_state.adversarial_confidence = preset.get("confidence_threshold", 85)
        st.session_state.adversarial_review_type = preset.get("review_type", "General SOP")
        st.session_state.compliance_requirements = preset.get("compliance_requirements", "")
        
        # Apply advanced settings if present
        advanced_settings = preset.get("advanced_settings", {})
        if advanced_settings:
            st.session_state.adversarial_critique_depth = advanced_settings.get("critique_depth", 5)
            st.session_state.adversarial_patch_quality = advanced_settings.get("patch_quality", 5)
            st.session_state.adversarial_detailed_tracking = advanced_settings.get("detailed_tracking", False)
            st.session_state.adversarial_performance_analytics = advanced_settings.get("performance_analytics", False)
            st.session_state.adversarial_early_stopping = advanced_settings.get("early_stopping", False)
            if "target_complexity" in advanced_settings:
                st.session_state.adversarial_target_complexity = advanced_settings.get("target_complexity", 0)
            if "target_length" in advanced_settings:
                st.session_state.adversarial_target_length = advanced_settings.get("target_length", 0)
            
        return True
    except Exception as e:
        st.error(f"Error applying preset: {e}")
        return False

# ------------------------------------------------------------------
# Version Control and Collaboration Functions
# ------------------------------------------------------------------

def create_new_version(protocol_text: str, version_name: str = "", comment: str = "") -> str:
    """Create a new version of the protocol.
    
    Args:
        protocol_text (str): The protocol text to save
        version_name (str): Optional name for the version
        comment (str): Optional comment about the changes
        
    Returns:
        str: Version ID of the created version
    """
    version_id = str(uuid.uuid4())
    timestamp = datetime.now().isoformat()
    
    version = {
        "id": version_id,
        "name": version_name or f"Version {len(st.session_state.protocol_versions) + 1}",
        "timestamp": timestamp,
        "protocol_text": protocol_text,
        "comment": comment,
        "author": "Current User",  # In a real implementation, this would be the actual user
        "complexity_metrics": calculate_protocol_complexity(protocol_text),
        "structure_analysis": extract_protocol_structure(protocol_text)
    }
    
    with st.session_state.thread_lock:
        st.session_state.protocol_versions.append(version)
        st.session_state.current_version_id = version_id
    
    return version_id

def load_version(version_id: str) -> bool:
    """Load a specific version of the protocol.
    
    Args:
        version_id (str): ID of the version to load
        
    Returns:
        bool: True if successful, False otherwise
    """
    with st.session_state.thread_lock:
        for version in st.session_state.protocol_versions:
            if version["id"] == version_id:
                st.session_state.protocol_text = version["protocol_text"]
                st.session_state.current_version_id = version_id
                return True
    return False

def get_version_history() -> List[Dict]:
    """Get the version history.
    
    Returns:
        List[Dict]: List of versions
    """
    with st.session_state.thread_lock:
        return st.session_state.protocol_versions.copy()

def add_comment(comment_text: str, version_id: str = None) -> str:
    """Add a comment to a version or the current protocol.
    
    Args:
        comment_text (str): The comment text
        version_id (str): Optional version ID to comment on
        
    Returns:
        str: Comment ID
    """
    comment_id = str(uuid.uuid4())
    timestamp = datetime.now().isoformat()
    
    comment = {
        "id": comment_id,
        "text": comment_text,
        "timestamp": timestamp,
        "author": "Current User",  # In a real implementation, this would be the actual user
        "version_id": version_id or st.session_state.current_version_id
    }
    
    with st.session_state.thread_lock:
        st.session_state.comments.append(comment)
    
    return comment_id

def get_comments(version_id: str = None) -> List[Dict]:
    """Get comments for a specific version or all comments.
    
    Args:
        version_id (str): Optional version ID to get comments for
        
    Returns:
        List[Dict]: List of comments
    """
    with st.session_state.thread_lock:
        if version_id:
            return [c for c in st.session_state.comments if c["version_id"] == version_id]
        return st.session_state.comments.copy()

def export_project() -> Dict:
    """Export the entire project including versions and comments.
    
    Returns:
        Dict: Project data
    """
    with st.session_state.thread_lock:
        return {
            "project_name": st.session_state.project_name,
            "project_description": st.session_state.project_description,
            "versions": st.session_state.protocol_versions,
            "comments": st.session_state.comments,
            "collaborators": st.session_state.collaborators,
            "tags": st.session_state.tags,
            "export_timestamp": datetime.now().isoformat()
        }

def export_project_detailed() -> Dict:
    """Export the entire project with detailed analytics and history.
    
    Returns:
        Dict: Detailed project data
    """
    with st.session_state.thread_lock:
        # Get analytics if adversarial testing was run
        analytics = {}
        if st.session_state.adversarial_results:
            analytics = generate_advanced_analytics(st.session_state.adversarial_results)
        
        return {
            "project_name": st.session_state.project_name,
            "project_description": st.session_state.project_description,
            "versions": st.session_state.protocol_versions,
            "comments": st.session_state.comments,
            "collaborators": st.session_state.collaborators,
            "tags": st.session_state.tags,
            "export_timestamp": datetime.now().isoformat(),
            "analytics": analytics,
            "adversarial_results": st.session_state.adversarial_results,
            "evolution_history": st.session_state.evolution_log,
            "model_performance": st.session_state.adversarial_model_performance
        }


def generate_shareable_link(project_data: Dict) -> str:
    """Generate a shareable link for the project.
    
    Args:
        project_data (Dict): Project data to share
        
    Returns:
        str: Shareable link
    """
    # In a real implementation, this would generate a real shareable link
    # For now, we'll simulate it
    project_id = hashlib.md5(json.dumps(project_data, sort_keys=True).encode()).hexdigest()[:16]
    return f"https://open-evolve.app/shared/{project_id}"


def get_version_by_name(version_name: str) -> Optional[Dict]:
    """Get a specific version by name.
    
    Args:
        version_name (str): Name of the version to retrieve
        
    Returns:
        Optional[Dict]: Version data or None if not found
    """
    with st.session_state.thread_lock:
        for version in st.session_state.protocol_versions:
            if version["name"] == version_name:
                return version
    return None


def get_version_by_timestamp(timestamp: str) -> Optional[Dict]:
    """Get a specific version by timestamp.
    
    Args:
        timestamp (str): Timestamp of the version to retrieve
        
    Returns:
        Optional[Dict]: Version data or None if not found
    """
    with st.session_state.thread_lock:
        for version in st.session_state.protocol_versions:
            if version["timestamp"].startswith(timestamp):
                return version
    return None


def delete_version(version_id: str) -> bool:
    """Delete a specific version.
    
    Args:
        version_id (str): ID of the version to delete
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        with st.session_state.thread_lock:
            st.session_state.protocol_versions = [
                v for v in st.session_state.protocol_versions if v["id"] != version_id
            ]
            # If we deleted the current version, set to the latest remaining version
            if st.session_state.current_version_id == version_id and st.session_state.protocol_versions:
                latest_version = st.session_state.protocol_versions[-1]
                st.session_state.protocol_text = latest_version["protocol_text"]
                st.session_state.current_version_id = latest_version["id"]
        return True
    except Exception as e:
        st.error(f"Error deleting version: {e}")
        return False


def branch_version(version_id: str, new_version_name: str) -> Optional[str]:
    """Create a new branch from an existing version.
    
    Args:
        version_id (str): ID of the version to branch from
        new_version_name (str): Name for the new branched version
        
    Returns:
        Optional[str]: ID of the new version or None if failed
    """
    version = None
    with st.session_state.thread_lock:
        for v in st.session_state.protocol_versions:
            if v["id"] == version_id:
                version = v
                break
    
    if not version:
        st.error("Version not found")
        return None
    
    # Create new version with branched content
    new_version_id = create_new_version(
        version["protocol_text"], 
        new_version_name, 
        f"Branched from {version['name']}"
    )
    
    # Add branch metadata
    with st.session_state.thread_lock:
        for v in st.session_state.protocol_versions:
            if v["id"] == new_version_id:
                v["branch_from"] = version_id
                v["branch_name"] = new_version_name
                break
    
    return new_version_id


def get_version_timeline() -> List[Dict]:
    """Get a chronological timeline of all versions.
    
    Returns:
        List[Dict]: Sorted list of versions by timestamp
    """
    with st.session_state.thread_lock:
        versions = st.session_state.protocol_versions.copy()
    
    # Sort by timestamp
    versions.sort(key=lambda x: x["timestamp"])
    return versions


def render_version_timeline() -> str:
    """Render a visual timeline of versions.
    
    Returns:
        str: HTML formatted timeline
    """
    versions = get_version_timeline()
    
    if not versions:
        return "<p>No version history available</p>"
    
    html = """
    <div style="background-color: #f8f9fa; padding: 20px; border-radius: 10px;">
        <h3 style="color: #4a6fa5; margin-top: 0;">üïí Version Timeline</h3>
        <div style="position: relative; padding-left: 30px;">
            <div style="position: absolute; left: 15px; top: 0; bottom: 0; width: 2px; background-color: #4a6fa5;"></div>
    """
    
    for i, version in enumerate(versions):
        is_current = version["id"] == st.session_state.get("current_version_id", "")
        timestamp = version["timestamp"][:16].replace("T", " ")
        
        html += f"""
        <div style="position: relative; margin-bottom: 20px;">
            <div style="position: absolute; left: -20px; top: 5px; width: 12px; height: 12px; border-radius: 50%; background-color: {'#4a6fa5' if is_current else '#6b8cbc'}; border: 2px solid white;"></div>
            <div style="background-color: {'#e3f2fd' if is_current else 'white'}; padding: 15px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); border-left: 4px solid {'#4a6fa5' if is_current else '#6b8cbc'};">
                <div style="display: flex; justify-content: space-between; align-items: center;">
                    <h4 style="margin: 0; color: #4a6fa5;">{version['name']}</h4>
                    <span style="font-size: 0.9em; color: #666;">{timestamp}</span>
                </div>
                <p style="margin: 5px 0 0 0; color: #666;">{version.get('comment', 'No comment')}</p>
                <div style="margin-top: 10px; display: flex; gap: 10px;">
        """
        
        # Add action buttons
        html += f"""
                    <button onclick="loadVersion('{version['id']}')" style="background-color: #4a6fa5; color: white; border: none; padding: 5px 10px; border-radius: 4px; cursor: pointer; font-size: 0.8em;">Load</button>
        """
        
        if not is_current:
            html += f"""
                    <button onclick="branchVersion('{version['id']}', 'Branch of {version['name']}')" style="background-color: #6b8cbc; color: white; border: none; padding: 5px 10px; border-radius: 4px; cursor: pointer; font-size: 0.8em;">Branch</button>
            """
        
        html += """
                </div>
            </div>
        </div>
        """
    
    html += """
        </div>
    </div>
    <script>
    function loadVersion(versionId) {
        // In a real implementation, this would trigger a reload with the version
        alert('Loading version: ' + versionId);
    }
    
    function branchVersion(versionId, branchName) {
        // In a real implementation, this would create a new branch
        alert('Branching from version: ' + versionId + ' as ' + branchName);
    }
    </script>
    """
    
    return html
    """Export protocol as a reusable template.
    
    Args:
        protocol_text (str): Protocol text to export as template
        template_name (str): Name for the template
        
    Returns:
        Dict: Template data
    """
    return {
        "name": template_name,
        "content": protocol_text,
        "created_at": datetime.now().isoformat(),
        "complexity_metrics": calculate_protocol_complexity(protocol_text),
        "structure_analysis": extract_protocol_structure(protocol_text),
        "tags": []
    }


def generate_ai_insights(protocol_text: str) -> Dict[str, Any]:
    """Generate AI-powered insights about the protocol.
    
    Args:
        protocol_text (str): Protocol text to analyze
        
    Returns:
        Dict[str, Any]: AI insights and recommendations
    """
    if not protocol_text:
        return {
            "overall_score": 0,
            "strengths": [],
            "weaknesses": [],
            "opportunities": [],
            "threats": [],
            "recommendations": [],
            "complexity_analysis": {},
            "readability_score": 0,
            "compliance_risk": "low"
        }
    
    # Calculate metrics
    complexity = calculate_protocol_complexity(protocol_text)
    structure = extract_protocol_structure(protocol_text)
    
    # Overall score calculation (weighted)
    structure_score = (
        (1 if structure["has_headers"] else 0) * 0.2 +
        (1 if structure["has_numbered_steps"] or structure["has_bullet_points"] else 0) * 0.2 +
        (1 if structure["has_preconditions"] else 0) * 0.15 +
        (1 if structure["has_postconditions"] else 0) * 0.15 +
        (1 if structure["has_error_handling"] else 0) * 0.15 +
        min(structure["section_count"] / 10, 1) * 0.15
    ) * 100
    
    complexity_score = max(0, 100 - complexity["complexity_score"])  # Invert complexity
    
    overall_score = (structure_score * 0.6 + complexity_score * 0.4)
    
    # Strengths
    strengths = []
    if structure["has_headers"]:
        strengths.append("‚úÖ Well-structured with clear headers")
    if structure["has_numbered_steps"] or structure["has_bullet_points"]:
        strengths.append("‚úÖ Uses lists or numbered steps for clarity")
    if structure["has_preconditions"]:
        strengths.append("‚úÖ Defines clear preconditions")
    if structure["has_postconditions"]:
        strengths.append("‚úÖ Specifies expected outcomes")
    if structure["has_error_handling"]:
        strengths.append("‚úÖ Includes error handling procedures")
    if complexity["unique_words"] / max(1, complexity["word_count"]) > 0.6:
        strengths.append("‚úÖ Good vocabulary diversity")
    
    # Weaknesses
    weaknesses = []
    if not structure["has_headers"]:
        weaknesses.append("‚ùå Lacks clear section headers")
    if not structure["has_numbered_steps"] and not structure["has_bullet_points"]:
        weaknesses.append("‚ùå Could use lists or numbered steps for better readability")
    if not structure["has_preconditions"]:
        weaknesses.append("‚ùå Missing preconditions specification")
    if not structure["has_postconditions"]:
        weaknesses.append("‚ùå No defined postconditions or expected outcomes")
    if not structure["has_error_handling"]:
        weaknesses.append("‚ùå Lacks error handling procedures")
    if complexity["avg_sentence_length"] > 25:
        weaknesses.append("‚ùå Sentences are quite long (hard to read)")
    if complexity["complexity_score"] > 60:
        weaknesses.append("‚ùå Protocol is quite complex")
    
    # Opportunities
    opportunities = []
    if complexity["word_count"] < 500:
        opportunities.append("‚ú® Protocol is brief - opportunity to add more detail")
    if structure["section_count"] == 0 and complexity["word_count"] > 300:
        opportunities.append("‚ú® Can improve organization with section headers")
    if not structure["has_preconditions"]:
        opportunities.append("‚ú® Add preconditions to clarify requirements")
    if not structure["has_postconditions"]:
        opportunities.append("‚ú® Define postconditions to specify expected outcomes")
    if not structure["has_error_handling"]:
        opportunities.append("‚ú® Include error handling for robustness")
    
    # Threats (potential issues)
    threats = []
    if complexity["complexity_score"] > 70:
        threats.append("‚ö†Ô∏è High complexity may lead to misinterpretation")
    if complexity["avg_sentence_length"] > 30:
        threats.append("‚ö†Ô∏è Long sentences may reduce clarity")
    if structure["section_count"] == 0 and complexity["word_count"] > 500:
        threats.append("‚ö†Ô∏è Lack of sections makes long protocols hard to navigate")
    
    # Recommendations
    recommendations = generate_protocol_recommendations(protocol_text)
    
    # Readability score
    readability_score = 100 - (complexity["avg_sentence_length"] / 50 * 100)
    readability_score = max(0, min(100, readability_score))
    
    # Compliance risk assessment
    compliance_risk = "low"
    if complexity["complexity_score"] > 70:
        compliance_risk = "high"
    elif complexity["complexity_score"] > 50:
        compliance_risk = "medium"
    
    return {
        "overall_score": round(overall_score, 1),
        "strengths": strengths,
        "weaknesses": weaknesses,
        "opportunities": opportunities,
        "threats": threats,
        "recommendations": recommendations,
        "complexity_analysis": complexity,
        "structure_analysis": structure,
        "readability_score": round(readability_score, 1),
        "compliance_risk": compliance_risk
    }


def render_ai_insights_dashboard(protocol_text: str) -> str:
    """Render an AI insights dashboard for the protocol.
    
    Args:
        protocol_text (str): Protocol text to analyze
        
    Returns:
        str: HTML formatted dashboard
    """
    insights = generate_ai_insights(protocol_text)
    
    # Create a visual dashboard
    html = """
    <div style="background: linear-gradient(135deg, #4a6fa5, #6b8cbc); color: white; padding: 20px; border-radius: 10px; margin-bottom: 20px;">
        <h2 style="margin-top: 0; text-align: center;">ü§ñ AI Insights Dashboard</h2>
        <div style="display: flex; justify-content: center; align-items: center;">
            <div style="background: white; color: #4a6fa5; border-radius: 50%; width: 100px; height: 100px; display: flex; justify-content: center; align-items: center; font-size: 2em; font-weight: bold;">
                """ + str(insights["overall_score"]) + """%
            </div>
        </div>
        <p style="text-align: center; margin-top: 10px;">Overall Protocol Quality Score</p>
    </div>
    """
    
    # Add metrics cards
    html += """
    <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px; margin-bottom: 20px;">
        <div style="background-color: #e8f5e9; padding: 15px; border-radius: 8px; border-left: 4px solid #4caf50;">
            <h4 style="margin-top: 0; color: #2e7d32;">üìä Readability</h4>
            <p style="font-size: 1.5em; font-weight: bold; margin: 0;">""" + str(insights["readability_score"]) + """%</p>
        </div>
        <div style="background-color: #fff8e1; padding: 15px; border-radius: 8px; border-left: 4px solid #ff9800;">
            <h4 style="margin-top: 0; color: #f57f17;">üìã Structure</h4>
            <p style="font-size: 1.5em; font-weight: bold; margin: 0;">""" + str(len([s for s in insights["structure_analysis"].values() if s])) + """/7</p>
        </div>
        <div style="background-color: #ffebee; padding: 15px; border-radius: 8px; border-left: 4px solid #f44336;">
            <h4 style="margin-top: 0; color: #c62828;">‚ö†Ô∏è Compliance Risk</h4>
            <p style="font-size: 1.5em; font-weight: bold; margin: 0; text-transform: capitalize;">""" + insights["compliance_risk"] + """</p>
        </div>
    </div>
    """
    
    # Add insights sections
    html += """
    <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)); gap: 20px;">
    """
    
    # Strengths
    if insights["strengths"]:
        html += """
        <div style="background-color: #e8f5e9; padding: 15px; border-radius: 8px;">
            <h3 style="color: #2e7d32; margin-top: 0;">‚úÖ Strengths</h3>
            <ul style="padding-left: 20px;">
        """
        for strength in insights["strengths"][:5]:  # Limit to first 5
            html += f"<li>{strength}</li>"
        html += """
            </ul>
        </div>
        """
    
    # Weaknesses
    if insights["weaknesses"]:
        html += """
        <div style="background-color: #ffebee; padding: 15px; border-radius: 8px;">
            <h3 style="color: #c62828; margin-top: 0;">‚ùå Areas for Improvement</h3>
            <ul style="padding-left: 20px;">
        """
        for weakness in insights["weaknesses"][:5]:  # Limit to first 5
            html += f"<li>{weakness}</li>"
        html += """
            </ul>
        </div>
        """
    
    # Opportunities
    if insights["opportunities"]:
        html += """
        <div style="background-color: #e3f2fd; padding: 15px; border-radius: 8px;">
            <h3 style="color: #1565c0; margin-top: 0;">‚ú® Opportunities</h3>
            <ul style="padding-left: 20px;">
        """
        for opportunity in insights["opportunities"][:5]:  # Limit to first 5
            html += f"<li>{opportunity}</li>"
        html += """
            </ul>
        </div>
        """
    
    # Threats
    if insights["threats"]:
        html += """
        <div style="background-color: #fff3e0; padding: 15px; border-radius: 8px;">
            <h3 style="color: #ef6c00; margin-top: 0;">‚ö†Ô∏è Potential Threats</h3>
            <ul style="padding-left: 20px;">
        """
        for threat in insights["threats"][:5]:  # Limit to first 5
            html += f"<li>{threat}</li>"
        html += """
            </ul>
        </div>
        """
    
    html += "</div>"
    
    # Add recommendations
    if insights["recommendations"]:
        html += """
        <div style="background-color: #f3e5f5; padding: 15px; border-radius: 8px; margin-top: 20px;">
            <h3 style="color: #6a1b9a; margin-top: 0;">üí° AI Recommendations</h3>
            <ul style="padding-left: 20px;">
        """
        for recommendation in insights["recommendations"]:
            html += f"<li>{recommendation}</li>"
        html += """
            </ul>
        </div>
        """
    
    return html


def import_project(project_data: Dict) -> bool:
    """Import a project including versions and comments.
    
    Args:
        project_data (Dict): Project data to import
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        with st.session_state.thread_lock:
            st.session_state.project_name = project_data.get("project_name", "Imported Project")
            st.session_state.project_description = project_data.get("project_description", "")
            st.session_state.protocol_versions = project_data.get("versions", [])
            st.session_state.comments = project_data.get("comments", [])
            st.session_state.collaborators = project_data.get("collaborators", [])
            st.session_state.tags = project_data.get("tags", [])
            
            # Set current version to the latest one
            if st.session_state.protocol_versions:
                latest_version = st.session_state.protocol_versions[-1]
                st.session_state.protocol_text = latest_version["protocol_text"]
                st.session_state.current_version_id = latest_version["id"]
        return True
    except Exception as e:
        st.error(f"Error importing project: {e}")
        return False

# ------------------------------------------------------------------
# Advanced Analytics and Reporting Functions
# ------------------------------------------------------------------

def generate_advanced_analytics(results: Dict) -> Dict:
    """Generate advanced analytics from adversarial testing results.
    
    Args:
        results (Dict): Adversarial testing results
        
    Returns:
        Dict: Advanced analytics data
    """
    analytics = {
        "total_iterations": len(results.get("iterations", [])),
        "final_approval_rate": results.get("final_approval_rate", 0),
        "total_cost_usd": results.get("cost_estimate_usd", 0),
        "total_tokens": results.get("tokens", {}).get("prompt", 0) + results.get("tokens", {}).get("completion", 0),
        "confidence_trend": [],
        "issue_resolution_rate": 0,
        "model_performance": {},
        "efficiency_score": 0,
        "security_strength": 0,
        "compliance_coverage": 0,
        "clarity_score": 0,
        "completeness_score": 0
    }
    
    # Calculate confidence trend
    if results.get("iterations"):
        confidence_history = [iter.get("approval_check", {}).get("approval_rate", 0) 
                            for iter in results.get('iterations', [])]
        analytics["confidence_trend"] = confidence_history
    
    # Calculate issue resolution rate
    if results.get("iterations"):
        total_issues_found = 0
        total_issues_resolved = 0
        severity_weights = {"low": 1, "medium": 3, "high": 6, "critical": 12}
        
        for iteration in results.get("iterations", []):
            critiques = iteration.get("critiques", [])
            for critique in critiques:
                critique_json = critique.get("critique_json", {})
                issues = critique_json.get("issues", [])
                total_issues_found += len(issues)
                
                # Count resolved issues with severity weighting
                patches = iteration.get("patches", [])
                resolved_weighted = 0
                total_weighted = 0
                
                for issue in issues:
                    severity = issue.get("severity", "low").lower()
                    weight = severity_weights.get(severity, 1)
                    total_weighted += weight
                    
                for patch in patches:
                    patch_json = patch.get("patch_json", {})
                    mitigation_matrix = patch_json.get("mitigation_matrix", [])
                    for mitigation in mitigation_matrix:
                        if str(mitigation.get("status", "")).lower() in ["resolved", "mitigated"]:
                            # Approximate the severity of resolved issues
                            resolved_weighted += 3  # Average weight
                
                total_issues_resolved += min(len(issues), len(patches))
        
        if total_issues_found > 0:
            analytics["issue_resolution_rate"] = (total_issues_resolved / total_issues_found) * 100
    
    # Calculate efficiency score
    efficiency = 100
    if analytics["total_cost_usd"] > 0:
        # Lower cost = higher efficiency
        efficiency -= min(50, analytics["total_cost_usd"] * 10)
    if analytics["total_iterations"] > 10:
        # More iterations = lower efficiency
        efficiency -= min(30, (analytics["total_iterations"] - 10) * 2)
    analytics["efficiency_score"] = max(0, efficiency)
    
    # Calculate security strength based on resolved critical/high issues
    if results.get("iterations"):
        critical_high_resolved = 0
        total_critical_high = 0
        
        for iteration in results.get("iterations", []):
            critiques = iteration.get("critiques", [])
            for critique in critiques:
                critique_json = critique.get("critique_json", {})
                issues = critique_json.get("issues", [])
                for issue in issues:
                    severity = issue.get("severity", "low").lower()
                    if severity in ["critical", "high"]:
                        total_critical_high += 1
                        # Check if this issue was resolved in any patch
                        for patch in iteration.get("patches", []):
                            patch_json = patch.get("patch_json", {})
                            mitigation_matrix = patch_json.get("mitigation_matrix", [])
                            for mitigation in mitigation_matrix:
                                if (mitigation.get("issue") == issue.get("title") and 
                                    str(mitigation.get("status", "")).lower() in ["resolved", "mitigated"]):
                                    critical_high_resolved += 1
                                    break
        
        if total_critical_high > 0:
            analytics["security_strength"] = (critical_high_resolved / total_critical_high) * 100
    
    # Calculate compliance coverage
    if results.get("compliance_requirements"):
        # Simple check for compliance mentions in final protocol
        final_sop = results.get("final_sop", "")
        compliance_reqs = results.get("compliance_requirements", "")
        
        # Count how many compliance requirements are addressed
        reqs_addressed = 0
        total_reqs = 0
        
        for req in compliance_reqs.split(","):
            req = req.strip().lower()
            if req:
                total_reqs += 1
                if req in final_sop.lower():
                    reqs_addressed += 1
        
        if total_reqs > 0:
            analytics["compliance_coverage"] = (reqs_addressed / total_reqs) * 100
    
    # Calculate clarity score based on protocol structure
    final_sop = results.get("final_sop", "")
    if final_sop:
        structure = extract_protocol_structure(final_sop)
        complexity = calculate_protocol_complexity(final_sop)
        
        # Clarity score based on structure elements
        clarity_score = 0
        if structure["has_headers"]:
            clarity_score += 25
        if structure["has_numbered_steps"] or structure["has_bullet_points"]:
            clarity_score += 25
        if structure["has_preconditions"]:
            clarity_score += 15
        if structure["has_postconditions"]:
            clarity_score += 15
        if structure["has_error_handling"]:
            clarity_score += 20
            
        analytics["clarity_score"] = clarity_score
        
        # Completeness score based on structure and complexity
        completeness_score = min(100, (
            structure["section_count"] * 5 +  # Sections contribute to completeness
            complexity["unique_words"] / max(1, complexity["word_count"]) * 100 * 0.3 +  # Vocabulary diversity
            (1 - complexity["avg_sentence_length"] / 50) * 100 * 0.7  # Sentence complexity balance
        ))
        analytics["completeness_score"] = completeness_score
    
    return analytics

def create_performance_comparison_chart(results: Dict) -> str:
    """Create a performance comparison chart for models.
    
    Args:
        results (Dict): Adversarial testing results
        
    Returns:
        str: HTML chart code
    """
    # Simplified implementation - in a real app, this would generate actual charts
    return """
    <div style="background-color: #f0f2f6; padding: 20px; border-radius: 10px;">
        <h3>üìä Model Performance Comparison</h3>
        <p>Chart would display here showing model performance metrics</p>
    </div>
    """

def generate_executive_summary(results: Dict) -> str:
    """Generate an executive summary of adversarial testing results.
    
    Args:
        results (Dict): Adversarial testing results
        
    Returns:
        str: Executive summary in markdown format
    """
    analytics = generate_advanced_analytics(results)
    
    summary = f"""# Executive Summary

## üìä Key Metrics
- **Final Approval Rate**: {analytics['final_approval_rate']:.1f}%
- **Iterations Completed**: {analytics['total_iterations']}
- **Total Cost**: ${analytics['total_cost_usd']:.4f}
- **Issue Resolution Rate**: {analytics['issue_resolution_rate']:.1f}%
- **Efficiency Score**: {analytics['efficiency_score']:.1f}/100

## üìà Performance Insights
- **Confidence Improvement**: The protocol's approval confidence improved from {analytics['confidence_trend'][0] if analytics['confidence_trend'] else 'N/A'}% to {analytics['final_approval_rate']:.1f}%
- **Cost Efficiency**: Process completed within budget constraints
- **Resolution Effectiveness**: Issues were effectively identified and addressed

## üèÜ Recommendations
1. Continue monitoring for emerging threats
2. Periodically re-evaluate with updated models
3. Consider expanding the red team for broader vulnerability coverage
4. Implement continuous integration for automated protocol hardening

## üìÖ Next Steps
- Review and implement outstanding recommendations
- Schedule periodic adversarial testing cycles
- Share results with stakeholders for feedback
- Update documentation with hardened protocol
"""
    
    return summary

# ------------------------------------------------------------------
# 5. Adversarial Testing Functions (robust)
# ------------------------------------------------------------------

MODEL_META_BY_ID: Dict[str, Dict[str, Any]] = {}
MODEL_META_LOCK = threading.Lock()

@st.cache_data(ttl=600)
def get_openrouter_models(api_key: str) -> List[Dict]:
    """Fetch available models from OpenRouter (cached)."""
    if not api_key:
        return []
    try:
        response = requests.get(
            "https://openrouter.ai/api/v1/models",
            headers={"Authorization": f"Bearer {api_key}"},
            timeout=15,
        )
        response.raise_for_status()
        data = response.json()
        models = data.get("data", []) if isinstance(data, dict) else []
        return models
    except Exception as e:
        st.warning(f"Could not fetch OpenRouter models: {e}")
        return []

def _compose_messages(system_prompt: str, user_prompt: str) -> List[Dict[str, str]]:
    messages = []
    if system_prompt and system_prompt.strip():
        messages.append({"role": "system", "content": system_prompt})
    messages.append({"role": "user", "content": user_prompt})
    return messages

def _request_openrouter_chat(
    api_key: str,
    model_id: str,
    messages: List[Dict[str, str]],
    temperature: float,
    top_p: float,
    frequency_penalty: float,
    presence_penalty: float,
    max_tokens: int,
    force_json: bool = False,
    seed: Optional[int] = None,
    req_timeout: int = 60,
    max_retries: int = 5,
) -> Tuple[str, int, int, float]:
    """
    Robust OpenRouter chat call with exponential backoff, jitter, and cost/token estimation.
    Returns: (content, prompt_tokens, completion_tokens, cost)
    """
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://github.com/google/gemini-pro-builder", # Recommended by OpenRouter
        "X-Title": "OpenEvolve Protocol Improver", # Recommended by OpenRouter
    }
    payload: Dict[str, Any] = {
        "model": model_id,
        "messages": messages,
        "temperature": _clamp(temperature, 0.0, 2.0),
        "max_tokens": max_tokens,
        "top_p": _clamp(top_p, 0.0, 1.0),
        "frequency_penalty": _clamp(frequency_penalty, -2.0, 2.0),
        "presence_penalty": _clamp(presence_penalty, -2.0, 2.0),
    }
    if seed is not None:
        payload["seed"] = int(seed)
    if force_json:
        payload["response_format"] = {"type": "json_object"}

    last_err = None
    for attempt in range(max_retries):
        try:
            r = requests.post(url, headers=headers, json=payload, timeout=req_timeout)
            if r.status_code == 400:
                # HTTP 400 Bad Request - client error, don't retry
                last_err = Exception(f"HTTP 400 Bad Request: {r.text[:200]}...")
                break  # Break out of retry loop for client errors
            if r.status_code in {429, 500, 502, 503, 504}:
                sleep_s = (2 ** attempt) + _rand_jitter_ms()
                time.sleep(sleep_s)
                last_err = Exception(f"Transient error {r.status_code}: Retrying...")
                continue
            r.raise_for_status()
            data = r.json()
            
            # Safely access the first choice to prevent IndexError if "choices" is an empty list.
            choices = data.get("choices", [])
            if choices:
                choice = choices[0]
                content = choice.get("message", {}).get("content", "")
            else:
                content = ""

            usage = data.get("usage", {})
            p_tok = safe_int(usage.get("prompt_tokens"), _approx_tokens(json.dumps(messages)))
            c_tok = safe_int(usage.get("completion_tokens"), _approx_tokens(content or ""))
            cost = _cost_estimate(p_tok, c_tok, None, None)  # Simplified cost calculation
            return content or "", p_tok, c_tok, cost
        except Exception as e:
            last_err = e
            sleep_s = (2 ** attempt) + _rand_jitter_ms()
            time.sleep(sleep_s)
    raise RuntimeError(f"Request failed for {model_id} after {max_retries} attempts: {last_err}")

def _request_anthropic_chat(
    api_key: str, base_url: str, model: str, messages: List, extra_headers: Dict,
    temperature: float, top_p: float, max_tokens: int, seed: Optional[int],
    frequency_penalty: float = 0.0, presence_penalty: float = 0.0,
    req_timeout: int = 60, max_retries: int = 5
) -> str:
    url = base_url.rstrip('/') + "/messages"
    headers = {
        "x-api-key": api_key,
        "anthropic-version": "2023-06-01",
        "Content-Type": "application/json",
        **extra_headers
    }
    
    # Separate system prompt from messages
    system_prompt = ""
    user_messages = []
    for msg in messages:
        if msg['role'] == 'system':
            system_prompt = msg['content']
        else:
            user_messages.append(msg)

    payload = {
        "model": model,
        "messages": user_messages,
        "max_tokens": max_tokens,
        "temperature": _clamp(temperature, 0.0, 1.0),
        "top_p": _clamp(top_p, 0.0, 1.0),
    }
    if system_prompt:
        payload['system'] = system_prompt

    last_err = None
    for attempt in range(max_retries):
        try:
            r = requests.post(url, headers=headers, json=payload, timeout=req_timeout)
            if r.status_code in {429, 500, 502, 503, 504}:
                sleep_s = (2 ** attempt) + _rand_jitter_ms()
                time.sleep(sleep_s)
                last_err = Exception(f"HTTP {r.status_code}: {r.text}")
                continue
            r.raise_for_status()
            data = r.json()
            if data.get('content') and isinstance(data['content'], list):
                return data['content'][0].get('text', '')
            else:
                last_err = Exception("No content in response")
        except Exception as e:
            last_err = e
            sleep_s = (2 ** attempt) + _rand_jitter_ms()
            time.sleep(sleep_s)
    raise RuntimeError(f"Request failed after {max_retries} attempts for model {model}: {last_err}")

def _request_google_gemini_chat(
    api_key: str, base_url: str, model: str, messages: List, extra_headers: Dict,
    temperature: float, top_p: float, max_tokens: int, seed: Optional[int],
    frequency_penalty: float = 0.0, presence_penalty: float = 0.0,
    req_timeout: int = 60, max_retries: int = 5
) -> str:
    url = f"{base_url.rstrip('/')}/models/{model}:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json", **extra_headers}
    
    # Gemini uses a different message format
    contents = []
    for msg in messages:
        contents.append({"role": msg['role'], "parts": [{"text": msg['content']}]})

    payload = {
        "contents": contents,
        "generationConfig": {
            "temperature": _clamp(temperature, 0.0, 1.0),
            "topP": _clamp(top_p, 0.0, 1.0),
            "maxOutputTokens": max_tokens,
        }
    }

    last_err = None
    for attempt in range(max_retries):
        try:
            r = requests.post(url, headers=headers, json=payload, timeout=req_timeout)
            if r.status_code in {429, 500, 502, 503, 504}:
                sleep_s = (2 ** attempt) + _rand_jitter_ms()
                time.sleep(sleep_s)
                last_err = Exception(f"HTTP {r.status_code}: {r.text}")
                continue
            r.raise_for_status()
            data = r.json()
            if data.get('candidates') and isinstance(data['candidates'], list):
                return data['candidates'][0]['content']['parts'][0]['text']
            else:
                last_err = Exception("No content in response")
        except Exception as e:
            last_err = e
            sleep_s = (2 ** attempt) + _rand_jitter_ms()
            time.sleep(sleep_s)
    raise RuntimeError(f"Request failed after {max_retries} attempts for model {model}: {last_err}")

def _request_cohere_chat(
    api_key: str, base_url: str, model: str, messages: List, extra_headers: Dict,
    temperature: float, top_p: float, max_tokens: int, seed: Optional[int],
    frequency_penalty: float = 0.0, presence_penalty: float = 0.0,
    req_timeout: int = 60, max_retries: int = 5
) -> str:
    url = base_url.rstrip('/') + "/chat"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        **extra_headers
    }
    
    # Separate system prompt and history from the last user message
    system_prompt = ""
    chat_history = []
    for msg in messages[:-1]:
        if msg['role'] == 'system':
            system_prompt = msg['content']
        else:
            chat_history.append({"role": msg['role'].upper(), "message": msg['content']})
    
    user_message = messages[-1]['content']

    payload = {
        "model": model,
        "message": user_message,
        "chat_history": chat_history,
        "max_tokens": max_tokens,
        "temperature": _clamp(temperature, 0.0, 5.0),
        "p": _clamp(top_p, 0.0, 1.0),
    }
    if system_prompt:
        payload['preamble'] = system_prompt

    last_err = None
    for attempt in range(max_retries):
        try:
            r = requests.post(url, headers=headers, json=payload, timeout=req_timeout)
            if r.status_code in {429, 500, 502, 503, 504}:
                sleep_s = (2 ** attempt) + _rand_jitter_ms()
                time.sleep(sleep_s)
                last_err = Exception(f"HTTP {r.status_code}: {r.text}")
                continue
            r.raise_for_status()
            data = r.json()
            if data.get('text'):
                return data['text']
            else:
                last_err = Exception("No text in response")
        except Exception as e:
            last_err = e
            sleep_s = (2 ** attempt) + _rand_jitter_ms()
            time.sleep(sleep_s)
    raise RuntimeError(f"Request failed after {max_retries} attempts for model {model}: {last_err}")

def analyze_with_model(
    api_key: str,
    model_id: str,
    sop: str,
    config: Dict,
    system_prompt: str,
    user_suffix: str = "",
    force_json: bool = False,
    seed: Optional[int] = None,
    compliance_requirements: str = "",
) -> Dict[str, Any]:
    """
    Analyzes an SOP with a specific model, handling context limits and returning structured results.
    """
    try:
        if compliance_requirements:
            system_prompt = system_prompt.format(compliance_requirements=compliance_requirements)
        max_tokens = safe_int(config.get("max_tokens"), 8000)
        user_prompt = f"Here is the Standard Operating Procedure (SOP):\n\n---\n\n{sop}\n\n---\n\n{user_suffix}"
        full_prompt_text = system_prompt + user_prompt

        # Simplified context length estimation
        context_len = 8192
        prompt_toks_est = _approx_tokens(full_prompt_text)

        if prompt_toks_est + max_tokens >= context_len:
            err_msg = (f"ERROR[{model_id}]: Estimated prompt tokens ({prompt_toks_est}) + max_tokens ({max_tokens}) "
                       f"exceeds context window ({context_len}). Skipping.")
            return {"ok": False, "text": err_msg, "json": None, "ptoks": 0, "ctoks": 0, "cost": 0.0, "model_id": model_id}

        content, p_tok, c_tok, cost = _request_openrouter_chat(
            api_key=api_key, model_id=model_id,
            messages=_compose_messages(system_prompt, user_prompt),
            temperature=safe_float(config.get("temperature"), 0.7),
            top_p=safe_float(config.get("top_p"), 1.0),
            frequency_penalty=safe_float(config.get("frequency_penalty"), 0.0),
            presence_penalty=safe_float(config.get("presence_penalty"), 0.0),
            max_tokens=max_tokens, force_json=force_json, seed=seed,
        )
        json_content = _extract_json_block(content)
        return {"ok": True, "text": content, "json": json_content, "ptoks": p_tok, "ctoks": c_tok, "cost": cost, "model_id": model_id}
    except Exception as e:
        return {"ok": False, "text": f"ERROR[{model_id}]: {e}", "json": None, "ptoks": 0, "ctoks": 0, "cost": 0.0, "model_id": model_id}

def determine_review_type(content: str) -> str:
    """Determine the appropriate review type based on content analysis.
    
    Args:
        content (str): The content to analyze
        
    Returns:
        str: Review type ('general', 'code', 'plan')
    """
    if not content:
        return "general"
    
    # Convert to lowercase for analysis
    lower_content = content.lower()
    
    # Check for code indicators
    code_indicators = [
        'function ', 'def ', 'class ', 'import ', 'require(', 'var ', 'let ', 'const ',
        'public ', 'private ', 'protected ', 'static ', 'void ', 'int ', 'string ',
        '<html', '<?php', '<script', 'console.', 'print(', 'printf(', 'scanf(',
        'if(', 'for(', 'while(', 'switch(', 'try{', 'catch(', 'finally{'
    ]
    
    # Check for plan indicators
    plan_indicators = [
        'objective', 'goal', 'milestone', 'deliverable', 'resource', 'budget',
        'timeline', 'schedule', 'risk', 'dependency', 'assumption',
        'stakeholder', 'communication', 'review', 'approval'
    ]
    
    # Count matches
    code_matches = sum(1 for indicator in code_indicators if indicator in lower_content)
    plan_matches = sum(1 for indicator in plan_indicators if indicator in lower_content)
    
    # Determine review type
    if code_matches > plan_matches and code_matches > 2:
        return "code"
    elif plan_matches > code_matches and plan_matches > 2:
        return "plan"
    else:
        return "general"

def get_appropriate_prompts(review_type: str) -> Tuple[str, str]:
    """Get the appropriate prompts based on review type.
    
    Args:
        review_type (str): Type of review ('general', 'code', 'plan')
        
    Returns:
        Tuple[str, str]: Red team and blue team prompts
    """
    if review_type == "code":
        return CODE_REVIEW_RED_TEAM_PROMPT, CODE_REVIEW_BLUE_TEAM_PROMPT
    elif review_type == "plan":
        return PLAN_REVIEW_RED_TEAM_PROMPT, PLAN_REVIEW_BLUE_TEAM_PROMPT
    else:
        return RED_TEAM_CRITIQUE_PROMPT, BLUE_TEAM_PATCH_PROMPT

def _severity_rank(sev: str) -> int:
    order = {"low": 0, "medium": 1, "high": 2, "critical": 3}
    return order.get(str(sev).lower(), 0)

def _merge_consensus_sop(base_sop: str, blue_patches: List[dict], critiques: List[dict]) -> Tuple[str, dict]:
    """
    Selects the best patch from the blue team based on coverage, resolution, and quality.
    """
    valid_patches = [p for p in blue_patches if p and (p.get("patch_json") or {}).get("sop", "").strip()]
    if not valid_patches:
        return base_sop, {"reason": "no_valid_patches_received", "score": -1, "resolution_by_severity": {}, "resolution_by_category": {}}

    # Create a lookup for issue severity and category
    issue_details = {}
    for critique in critiques:
        if critique and critique.get("critique_json"):
            for issue in _safe_list(critique["critique_json"], "issues"):
                issue_details[issue.get("title")] = {
                    "severity": issue.get("severity", "low"),
                    "category": issue.get("category", "uncategorized")
                }

    scored = []
    for patch in valid_patches:
        patch_json = patch.get("patch_json", {})
        sop_text = patch_json.get("sop", "").strip()

        mm = _safe_list(patch_json, "mitigation_matrix")
        residual = _safe_list(patch_json, "residual_risks")

        resolved = sum(1 for r in mm if str(r.get("status", "")).lower() == "resolved")
        mitigated = sum(1 for r in mm if str(r.get("status", "")).lower() == "mitigated")

        # Score based on resolved issues, then mitigated, penalize for residuals, and use length as tie-breaker
        coverage_score = (resolved * 2) + mitigated
        final_score = coverage_score - (len(residual) * 2)

        # Track resolution by severity and category
        resolution_by_severity = {}
        resolution_by_category = {}
        for r in mm:
            issue_title = r.get("issue")
            if issue_title in issue_details:
                details = issue_details[issue_title]
                severity = details["severity"]
                category = details["category"]
                status = str(r.get("status", "")).lower()

                if status in ["resolved", "mitigated"]:
                    resolution_by_severity[severity] = resolution_by_severity.get(severity, 0) + 1
                    resolution_by_category[category] = resolution_by_category.get(category, 0) + 1

        scored.append((final_score, resolved, len(sop_text), sop_text, patch.get("model"), resolution_by_severity, resolution_by_category))

    if not scored:
        return base_sop, {"reason": "all_patches_were_empty_or_invalid", "score": -1, "resolution_by_severity": {}, "resolution_by_category": {}}

    # Sort by score, then resolved count, then SOP length
    scored.sort(key=lambda x: (-x[0], -x[1], x[2]))
    best_score, best_resolved, _, best_sop, best_model, best_res_sev, best_res_cat = scored[0]
    diagnostics = {"reason": "best_patch_selected", "score": best_score, "resolved": best_resolved, "model": best_model, "resolution_by_severity": best_res_sev, "resolution_by_category": best_res_cat}
    return best_sop, diagnostics

def _aggregate_red_risk(critiques: List[dict]) -> Dict[str, Any]:
    """Computes an aggregate risk score from all red-team critiques."""
    sev_weight = {"low": 1, "medium": 3, "high": 6, "critical": 12}
    total_weight, issue_count = 0, 0
    categories = {}
    severities = {}

    valid_critiques = [c.get("critique_json") for c in critiques if c and c.get("critique_json")]

    for critique in valid_critiques:
        for issue in _safe_list(critique, "issues"):
            sev = str(issue.get("severity", "low")).lower()
            weight = sev_weight.get(sev, 1)
            total_weight += weight
            issue_count += 1
            cat = str(issue.get("category", "uncategorized")).lower()
            categories[cat] = categories.get(cat, 0) + weight
            severities[sev] = severities.get(sev, 0) + 1

    avg_weight = (total_weight / max(1, issue_count)) if issue_count > 0 else 0
    return {"total_weight": total_weight, "avg_issue_weight": avg_weight, "categories": categories, "severities": severities, "count": issue_count}

def _update_model_performance(critiques: List[dict]):
    """Updates the performance scores of models based on the critiques they generated."""
    with st.session_state.thread_lock:
        if "adversarial_model_performance" not in st.session_state:
            st.session_state.adversarial_model_performance = {}

        sev_weight = {"low": 1, "medium": 3, "high": 6, "critical": 12}
        for critique in critiques:
            model_id = critique.get("model")
            if not model_id:
                continue

            if model_id not in st.session_state.adversarial_model_performance:
                st.session_state.adversarial_model_performance[model_id] = {"score": 0, "issues_found": 0}

            critique_json = critique.get("critique_json")
            if critique_json and isinstance(critique_json.get("issues"), list):
                for issue in critique_json["issues"]:
                    sev = str(issue.get("severity", "low")).lower()
                    st.session_state.adversarial_model_performance[model_id]["score"] += sev_weight.get(sev, 1)
                    st.session_state.adversarial_model_performance[model_id]["issues_found"] += 1

def _collect_model_configs(model_ids: List[str], max_tokens: int) -> Dict[str, Dict[str, Any]]:
    return {
        model_id: {
            "temperature": st.session_state.get(f"temp_{model_id}", 0.7),
            "top_p": st.session_state.get(f"topp_{model_id}", 1.0),
            "frequency_penalty": st.session_state.get(f"freqpen_{model_id}", 0.0),
            "presence_penalty": st.session_state.get(f"prespen_{model_id}", 0.0),
            "max_tokens": max_tokens,
        } for model_id in model_ids
    }

def _update_adv_log_and_status(msg: str):
    """Thread-safe way to update logs and status message."""
    with st.session_state.thread_lock:
        st.session_state.adversarial_log.append(f"[{time.strftime('%H:%M:%S')}] {msg}")
        st.session_state.adversarial_status_message = msg

def _update_adv_counters(ptoks: int, ctoks: int, cost: float):
    """Thread-safe way to update token and cost counters."""
    with st.session_state.thread_lock:
        st.session_state.adversarial_total_tokens_prompt += ptoks
        st.session_state.adversarial_total_tokens_completion += ctoks
        st.session_state.adversarial_cost_estimate_usd += cost

def check_approval_rate(
    api_key: str, red_team_models: List[str], sop_markdown: str, model_configs: Dict,
    seed: Optional[int], max_workers: int, approval_prompt: str = APPROVAL_PROMPT
) -> Dict[str, Any]:
    """Asks all red-team models for a final verdict on the SOP."""
    votes, scores, approved = [], [], 0
    total_ptoks, total_ctoks, total_cost = 0, 0, 0.0

    with ThreadPoolExecutor(max_workers=max_workers) as ex:
        future_to_model = {
            ex.submit(
                analyze_with_model, api_key, model_id, sop_markdown,
                model_configs.get(model_id, {}), approval_prompt, force_json=True, seed=seed
            ): model_id for model_id in red_team_models
        }
        for future in as_completed(future_to_model):
            model_id = future_to_model[future]
            res = future.result()
            total_ptoks += res["ptoks"]; total_ctoks += res["ctoks"]; total_cost += res["cost"]
            if res.get("ok") and res.get("json"):
                j = res["json"]
                verdict = str(j.get("verdict", "REJECTED")).upper()
                score = _clamp(safe_int(j.get("score"), 0), 0, 100)
                if verdict == "APPROVED":
                    approved += 1
                scores.append(score)
                votes.append({"model": model_id, "verdict": verdict, "score": score, "reasons": _safe_list(j, "reasons")})
            else:
                votes.append({"model": model_id, "verdict": "ERROR", "score": 0, "reasons": [res.get("text")]})

    rate = (approved / max(1, len(red_team_models))) * 100.0
    avg_score = (sum(scores) / max(1, len(scores))) if scores else 0

    # Calculate agreement
    if not votes:
        agreement = 0.0
    else:
        verdicts = [v["verdict"] for v in votes]
        approved_count = verdicts.count("APPROVED")
        rejected_count = verdicts.count("REJECTED")
        agreement = max(approved_count, rejected_count) / len(verdicts) * 100.0

    return {"approval_rate": rate, "avg_score": avg_score, "votes": votes, "prompt_tokens": total_ptoks, "completion_tokens": total_ctoks, "cost": total_cost, "agreement": agreement}

def generate_docx_report(results: dict) -> bytes:
    """Generates a DOCX report from the adversarial testing results."""
    document = docx.Document()
    document.add_heading('Adversarial Testing Report', 0)

    document.add_heading('Summary', level=1)
    document.add_paragraph(
        f"Final Approval Rate: {results.get('final_approval_rate', 0.0):.1f}%\n"
        f"Total Iterations: {len(results.get('iterations', []))}\n"
        f"Total Cost (USD): ${results.get('cost_estimate_usd', 0.0):,.4f}\n"
        f"Total Prompt Tokens: {results.get('tokens', {}).get('prompt', 0):,}\n"
        f"Total Completion Tokens: {results.get('tokens', {}).get('completion', 0):,}"
    )

    document.add_heading('Final Hardened SOP', level=1)
    document.add_paragraph(results.get("final_sop", ""))

    document.add_heading('Issues Found', level=1)
    for i, iteration in enumerate(results.get("iterations", [])):
        document.add_heading(f"Iteration {i+1}", level=2)
        for critique in iteration.get("critiques", []):
            if critique.get("critique_json"):
                for issue in _safe_list(critique["critique_json"], "issues"):
                    document.add_paragraph(f"- {issue.get('title')} ({issue.get('severity')})", style='List Bullet')

    document.add_heading('Final Votes', level=1)
    if results.get("iterations"):
        for vote in results["iterations"][-1].get("approval_check", {}).get("votes", []):
            document.add_paragraph(f"- {vote.get('model')}: {vote.get('verdict')} ({vote.get('score')})", style='List Bullet')

    document.add_heading('Audit Trail', level=1)
    for log_entry in results.get("log", []):
        document.add_paragraph(log_entry)

    from io import BytesIO
    bio = BytesIO()
    document.save(bio)
    return bio.getvalue()

def generate_pdf_report(results: dict) -> bytes:
    """Generates a PDF report from the adversarial testing results."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    pdf.cell(200, 10, txt="Adversarial Testing Report", ln=True, align='C')

    pdf.ln(10)

    pdf.set_font("Arial", 'B', size=12)
    pdf.cell(200, 10, txt="Summary", ln=True)
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, f"Final Approval Rate: {results.get('final_approval_rate', 0.0):.1f}%\n"
                         f"Total Iterations: {len(results.get('iterations', []))}\n"
                         f"Total Cost (USD): ${results.get('cost_estimate_usd', 0.0):,.4f}\n"
                         f"Total Prompt Tokens: {results.get('tokens', {}).get('prompt', 0):,}\n"
                         f"Total Completion Tokens: {results.get('tokens', {}).get('completion', 0):,}")

    pdf.ln(10)

    pdf.set_font("Arial", 'B', size=12)
    pdf.cell(200, 10, txt="Final Hardened SOP", ln=True)
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, results.get("final_sop", ""))

    pdf.ln(10)

    pdf.set_font("Arial", 'B', size=12)
    pdf.cell(200, 10, txt="Issues Found", ln=True)
    pdf.set_font("Arial", size=12)
    for i, iteration in enumerate(results.get("iterations", [])):
        pdf.set_font("Arial", 'B', size=10)
        pdf.cell(200, 10, txt=f"Iteration {i+1}", ln=True)
        pdf.set_font("Arial", size=10)
        for critique in iteration.get("critiques", []):
            if critique.get("critique_json"):
                for issue in _safe_list(critique["critique_json"], "issues"):
                    pdf.multi_cell(0, 10, f"- {issue.get('title')} ({issue.get('severity')})")

    pdf.ln(10)

    pdf.set_font("Arial", 'B', size=12)
    pdf.cell(200, 10, txt="Final Votes", ln=True)
    pdf.set_font("Arial", size=10)
    if results.get("iterations"):
        for vote in results["iterations"][-1].get("approval_check", {}).get("votes", []):
            pdf.multi_cell(0, 10, f"- {vote.get('model')}: {vote.get('verdict')} ({vote.get('score')})")

    pdf.ln(10)

    pdf.set_font("Arial", 'B', size=12)
    pdf.cell(200, 10, txt="Audit Trail", ln=True)
    pdf.set_font("Arial", size=8)
    for log_entry in results.get("log", []):
        pdf.multi_cell(0, 5, log_entry)

    return pdf.output(dest='S').encode('latin-1')

def generate_html_report(results: dict) -> str:
    """Generates an HTML report from the adversarial testing results."""
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Adversarial Testing Report</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; background-color: #f8f9fa; }}
            h1, h2, h3 {{ color: #4a6fa5; }}
            .summary {{ background-color: #ffffff; padding: 15px; border-radius: 8px; box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1); margin-bottom: 20px; }}
            .section {{ margin: 20px 0; background-color: #ffffff; padding: 20px; border-radius: 8px; box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1); }}
            .log {{ font-family: monospace; font-size: 0.9em; background-color: #f9f9f9; padding: 10px; border-radius: 4px; }}
            table {{ border-collapse: collapse; width: 100%; }}
            th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
            th {{ background-color: #4a6fa5; color: white; }}
            .metric {{ text-align: center; padding: 10px; background-color: #e9ecef; border-radius: 4px; margin: 5px; }}
            .improvement {{ color: #4caf50; font-weight: bold; }}
            .regression {{ color: #f44336; font-weight: bold; }}
        </style>
    </head>
    <body>
        <h1>Adversarial Testing Report</h1>
        
        <div class="summary">
            <h2>Summary</h2>
            <p><strong>Final Approval Rate:</strong> {results.get('final_approval_rate', 0.0):.1f}%</p>
            <p><strong>Total Iterations:</strong> {len(results.get('iterations', []))}</p>
            <p><strong>Total Cost (USD):</strong> ${results.get('cost_estimate_usd', 0.0):,.4f}</p>
            <p><strong>Total Prompt Tokens:</strong> {results.get('tokens', {}).get('prompt', 0):,}</p>
            <p><strong>Total Completion Tokens:</strong> {results.get('tokens', {}).get('completion', 0):,}</p>
        </div>
        
        <div class="section">
            <h2>Final Hardened SOP</h2>
            <pre style="white-space: pre-wrap; background-color: #f9f9f9; padding: 15px; border-radius: 4px;">{results.get("final_sop", "")}</pre>
        </div>
    """
    
    if results.get("iterations"):
        html += """
        <div class="section">
            <h2>Issues Found</h2>
        """
        for i, iteration in enumerate(results.get("iterations", [])):
            html += f"<h3>Iteration {i+1}</h3><ul>"
            for critique in iteration.get("critiques", []):
                if critique.get("critique_json"):
                    for issue in _safe_list(critique["critique_json"], "issues"):
                        severity = issue.get('severity', 'low')
                        severity_color = {
                            'low': '#4caf50',
                            'medium': '#ff9800',
                            'high': '#f44336',
                            'critical': '#9c27b0'
                        }.get(severity, '#000000')
                        html += f"<li><span style='color: {severity_color}; font-weight: bold;'>{severity.upper()}</span>: {issue.get('title')}</li>"
            html += "</ul>"
        html += "</div>"
        
        html += """
        <div class="section">
            <h2>Final Votes</h2>
            <table>
                <tr><th>Model</th><th>Verdict</th><th>Score</th></tr>
        """
        for vote in results["iterations"][-1].get("approval_check", {}).get("votes", []):
            verdict = vote.get('verdict', '')
            verdict_color = '#4caf50' if verdict.upper() == 'APPROVED' else '#f44336'
            html += f"<tr><td>{vote.get('model')}</td><td style='color: {verdict_color}; font-weight: bold;'>{verdict}</td><td>{vote.get('score')}</td></tr>"
        html += "</table></div>"
        
        # Add performance chart data
        html += """
        <div class="section">
            <h2>Performance Metrics</h2>
            <div style="display: flex; justify-content: space-around; flex-wrap: wrap;">
        """
        
        # Approval rate chart
        approval_rates = [iter["approval_check"].get("approval_rate", 0) for iter in results.get("iterations", [])]
        if approval_rates:
            avg_approval = sum(approval_rates) / len(approval_rates)
            html += f"""
            <div class="metric">
                <div>Avg Approval Rate</div>
                <div style="font-size: 24px;">{avg_approval:.1f}%</div>
            </div>
            """
        
        # Issue count chart
        issue_counts = [iter["agg_risk"].get("count", 0) for iter in results.get("iterations", [])]
        if issue_counts:
            total_issues = sum(issue_counts)
            html += f"""
            <div class="metric">
                <div>Total Issues Found</div>
                <div style="font-size: 24px;">{total_issues}</div>
            </div>
            """
        
        html += "</div></div>"
    
    html += """
        <div class="section">
            <h2>Audit Trail</h2>
            <div class="log">
    """
    for log_entry in results.get("log", []):
        html += f"<div>{log_entry}</div>"
    html += """
            </div>
        </div>
    </body>
    </html>
    """
    
    return html

# ------------------------------------------------------------------
# Performance Optimization Functions
# ------------------------------------------------------------------

def optimize_model_selection(red_team_models: List[str], blue_team_models: List[str], 
                            protocol_complexity: int, budget_limit: float = 0.0) -> Dict[str, List[str]]:
    """Optimize model selection based on protocol complexity and budget.
    
    Args:
        red_team_models (List[str]): Available red team models
        blue_team_models (List[str]): Available blue team models
        protocol_complexity (int): Complexity score of the protocol (0-100)
        budget_limit (float): Maximum budget in USD (0 = no limit)
        
    Returns:
        Dict[str, List[str]]: Optimized model selections
    """
    # Enhanced optimization logic
    optimized = {
        "red_team": [],
        "blue_team": []
    }
    
    # For complex protocols, use more capable models
    if protocol_complexity > 70:
        # Use premium models for complex protocols
        optimized["red_team"] = [m for m in red_team_models if "gpt-4" in m or "claude-3-opus" in m or "gemini-1.5-pro" in m][:3]
        optimized["blue_team"] = [m for m in blue_team_models if "gpt-4" in m or "claude-3-sonnet" in m or "gemini-1.5-pro" in m][:3]
    elif protocol_complexity > 40:
        # Use balanced models for medium complexity
        optimized["red_team"] = [m for m in red_team_models if "gpt-4" in m or "claude-3-haiku" in m or "gemini-1.5-flash" in m][:3]
        optimized["blue_team"] = [m for m in blue_team_models if "gpt-4" in m or "claude-3-sonnet" in m or "gemini-1.5-flash" in m][:3]
    else:
        # Use cost-effective models for simple protocols
        optimized["red_team"] = [m for m in red_team_models if "gpt-4o-mini" in m or "claude-3-haiku" in m or "llama-3-8b" in m][:3]
        optimized["blue_team"] = [m for m in blue_team_models if "gpt-4o" in m or "claude-3-sonnet" in m or "llama-3-70b" in m][:3]
    
    # Budget-based optimization
    if budget_limit > 0:
        # Filter out expensive models if budget is constrained
        expensive_models = ["gpt-4", "claude-3-opus", "gemini-1.5-pro"]
        if budget_limit < 0.1:  # Very low budget
            optimized["red_team"] = [m for m in optimized["red_team"] if not any(exp in m for exp in expensive_models)]
            optimized["blue_team"] = [m for m in optimized["blue_team"] if not any(exp in m for exp in expensive_models)]
        elif budget_limit < 0.5:  # Moderate budget
            # Reduce count of expensive models
            expensive_red = [m for m in optimized["red_team"] if any(exp in m for exp in expensive_models)]
            if len(expensive_red) > 1:
                optimized["red_team"] = [m for m in optimized["red_team"] if m not in expensive_red[1:]]
                
            expensive_blue = [m for m in optimized["blue_team"] if any(exp in m for exp in expensive_models)]
            if len(expensive_blue) > 1:
                optimized["blue_team"] = [m for m in optimized["blue_team"] if m not in expensive_blue[1:]]
    
    # If no models matched criteria, use defaults
    if not optimized["red_team"]:
        optimized["red_team"] = red_team_models[:min(3, len(red_team_models))]
    if not optimized["blue_team"]:
        optimized["blue_team"] = blue_team_models[:min(3, len(blue_team_models))]
    
    # Ensure diversity in model selection
    if len(optimized["red_team"]) < 3 and len(red_team_models) >= 3:
        # Add models from different providers
        providers = set()
        selected_models = []
        for model in red_team_models:
            provider = model.split("/")[0] if "/" in model else model.split("-")[0]
            if provider not in providers or len(selected_models) < 3:
                selected_models.append(model)
                providers.add(provider)
        optimized["red_team"] = selected_models[:3]
        
    if len(optimized["blue_team"]) < 3 and len(blue_team_models) >= 3:
        # Add models from different providers
        providers = set()
        selected_models = []
        for model in blue_team_models:
            provider = model.split("/")[0] if "/" in model else model.split("-")[0]
            if provider not in providers or len(selected_models) < 3:
                selected_models.append(model)
                providers.add(provider)
        optimized["blue_team"] = selected_models[:3]
    
    return optimized

def estimate_testing_time_and_cost(red_team_models: List[str], blue_team_models: List[str], 
                                  iterations: int, protocol_length: int) -> Dict[str, Any]:
    """Estimate testing time and cost based on configuration.
    
    Args:
        red_team_models (List[str]): Selected red team models
        blue_team_models (List[str]): Selected blue team models
        iterations (int): Number of iterations
        protocol_length (int): Length of protocol in words
        
    Returns:
        Dict[str, Any]: Time and cost estimates
    """
    # Simplified estimation logic
    # Base estimates per model per iteration
    avg_response_time = 5  # seconds
    avg_cost_per_1000_tokens = 0.002  # USD
    
    # Calculate total operations
    total_red_operations = len(red_team_models) * iterations
    total_blue_operations = len(blue_team_models) * iterations
    
    # Estimate time (parallel processing assumed)
    max_parallel_workers = min(6, len(red_team_models) + len(blue_team_models))
    estimated_time_seconds = ((total_red_operations + total_blue_operations) / max_parallel_workers) * avg_response_time
    
    # Estimate cost (simplified token estimation)
    avg_tokens_per_operation = protocol_length * 3  # Rough estimate
    total_tokens = (total_red_operations + total_blue_operations) * avg_tokens_per_operation
    estimated_cost = (total_tokens / 1000) * avg_cost_per_1000_tokens
    
    return {
        "estimated_time_minutes": round(estimated_time_seconds / 60, 1),
        "estimated_cost_usd": round(estimated_cost, 4),
        "total_operations": total_red_operations + total_blue_operations,
        "total_tokens_estimated": total_tokens
    }

def suggest_performance_improvements(current_config: Dict) -> List[str]:
    """Suggest performance improvements for the current configuration.
    
    Args:
        current_config (Dict): Current adversarial testing configuration
        
    Returns:
        List[str]: List of suggested improvements
    """
    suggestions = []
    
    red_models = current_config.get("red_team_models", [])
    blue_models = current_config.get("blue_team_models", [])
    iterations = current_config.get("adversarial_max_iter", 10)
    protocol_text = current_config.get("protocol_text", "")
    
    # Check for common performance issues
    if len(red_models) > 5:
        suggestions.append("üî¥ Reduce red team models to 3-5 for better performance and cost control")
    
    if len(blue_models) > 5:
        suggestions.append("üîµ Reduce blue team models to 3-5 for better performance and cost control")
    
    if iterations > 20:
        suggestions.append("üîÑ Consider reducing max iterations to 15-20 for faster results")
    
    if len(protocol_text.split()) > 5000:
        suggestions.append("üìÑ Your protocol is quite long (>5000 words). Consider breaking it into smaller sections")
    
    # Check for model diversity
    all_models = red_models + blue_models
    if len(set(all_models)) < len(all_models) * 0.7:
        suggestions.append("üîÄ Increase model diversity by selecting models from different providers")
    
    # Check for expensive model combinations
    expensive_models = [m for m in all_models if "gpt-4" in m or "claude-3-opus" in m]
    if len(expensive_models) > 3:
        suggestions.append("üí∞ You're using many expensive models. Consider mixing in some cost-effective models")
    
    # If no suggestions, provide positive feedback
    if not suggestions:
        suggestions.append("‚úÖ Your configuration looks well-balanced for optimal performance!")
    
    return suggestions

# ------------------------------------------------------------------
# Advanced Testing Strategies
# ------------------------------------------------------------------

def adaptive_testing_strategy(results_history: List[Dict], current_config: Dict) -> Dict[str, Any]:
    """Adapt testing strategy based on historical results.
    
    Args:
        results_history (List[Dict]): History of previous testing results
        current_config (Dict): Current testing configuration
        
    Returns:
        Dict[str, Any]: Adapted strategy recommendations
    """
    strategy = {
        "recommended_models": {"red_team": [], "blue_team": []},
        "iteration_adjustments": {},
        "focus_areas": [],
        "confidence_threshold": current_config.get("adversarial_confidence", 85)
    }
    
    if not results_history:
        # First run - use balanced approach
        strategy["recommended_models"]["red_team"] = current_config.get("red_team_models", [])[:3]
        strategy["recommended_models"]["blue_team"] = current_config.get("blue_team_models", [])[:3]
        strategy["iteration_adjustments"] = {"min_iter": 3, "max_iter": 10}
        return strategy
    
    # Analyze recent results
    recent_results = results_history[-3:]  # Last 3 iterations
    avg_confidence = sum(r.get("approval_check", {}).get("approval_rate", 0) for r in recent_results) / len(recent_results)
    avg_issue_count = sum(len(r.get("agg_risk", {}).get("issues", [])) for r in recent_results) / len(recent_results)
    
    # Adjust based on performance
    if avg_confidence > 90:
        # High confidence - focus on efficiency
        strategy["recommended_models"]["red_team"] = current_config.get("red_team_models", [])[:2]
        strategy["recommended_models"]["blue_team"] = current_config.get("blue_team_models", [])[:2]
        strategy["iteration_adjustments"] = {"min_iter": 2, "max_iter": 8}
        strategy["focus_areas"] = ["efficiency", "cost_reduction"]
    elif avg_confidence < 70:
        # Low confidence - increase intensity
        strategy["recommended_models"]["red_team"] = current_config.get("red_team_models", [])[:5]
        strategy["recommended_models"]["blue_team"] = current_config.get("blue_team_models", [])[:5]
        strategy["iteration_adjustments"] = {"min_iter": 5, "max_iter": 15}
        strategy["confidence_threshold"] = min(95, strategy["confidence_threshold"] + 5)
        strategy["focus_areas"] = ["thoroughness", "coverage"]
    else:
        # Balanced approach
        strategy["recommended_models"]["red_team"] = current_config.get("red_team_models", [])[:3]
        strategy["recommended_models"]["blue_team"] = current_config.get("blue_team_models", [])[:3]
        strategy["iteration_adjustments"] = {"min_iter": 3, "max_iter": 12}
        strategy["focus_areas"] = ["balanced_approach"]
    
    return strategy

def category_focused_testing(issues_by_category: Dict[str, int], current_config: Dict) -> Dict[str, Any]:
    """Focus testing on specific issue categories.
    
    Args:
        issues_by_category (Dict[str, int]): Count of issues by category
        current_config (Dict): Current testing configuration
        
    Returns:
        Dict[str, Any]: Category-focused testing recommendations
    """
    if not issues_by_category:
        return {"focus_category": None, "recommended_models": {"red_team": [], "blue_team": []}}
    
    # Find category with most issues
    focus_category = max(issues_by_category.items(), key=lambda x: x[1])[0]
    
    # Recommend models based on category
    category_experts = {
        "security": ["openai/gpt-4o", "anthropic/claude-3-opus", "google/gemini-1.5-pro"],
        "compliance": ["openai/gpt-4o", "mistral/mistral-medium-latest"],
        "clarity": ["openai/gpt-4o-mini", "google/gemini-1.5-flash"],
        "completeness": ["anthropic/claude-3-sonnet", "meta-llama/llama-3-70b-instruct"],
        "efficiency": ["openai/gpt-4o", "meta-llama/llama-3-70b-instruct"]
    }
    
    recommended_models = category_experts.get(focus_category, current_config.get("red_team_models", [])[:3])
    
    return {
        "focus_category": focus_category,
        "recommended_models": {
            "red_team": recommended_models,
            "blue_team": current_config.get("blue_team_models", [])[:3]
        }
    }

def performance_based_model_rotation(model_performance: Dict[str, Dict], 
                                    current_red_team: List[str], 
                                    current_blue_team: List[str]) -> Dict[str, List[str]]:
    """Rotate models based on performance metrics.
    
    Args:
        model_performance (Dict[str, Dict]): Performance data for each model
        current_red_team (List[str]): Current red team models
        current_blue_team (List[str]): Current blue team models
        
    Returns:
        Dict[str, List[str]]: Updated model selections
    """
    # Sort models by performance score
    sorted_models = sorted(model_performance.items(), key=lambda x: x[1].get("score", 0), reverse=True)
    
    # Select top performers for red team (critics)
    top_red_models = [model_id for model_id, _ in sorted_models[:3] if model_id in current_red_team]
    if not top_red_models:
        top_red_models = current_red_team[:min(3, len(current_red_team))]
    
    # Select diverse models for blue team (fixers)
    top_blue_models = [model_id for model_id, _ in sorted_models[:3] if model_id in current_blue_team]
    if not top_blue_models:
        top_blue_models = current_blue_team[:min(3, len(current_blue_team))]
    
    return {
        "red_team": top_red_models,
        "blue_team": top_blue_models
    }

# ------------------------------------------------------------------
# Advanced Analytics Functions
# ------------------------------------------------------------------

def analyze_code_quality(code_text: str) -> Dict[str, Any]:
    """Analyze code quality metrics.
    
    Args:
        code_text (str): The code to analyze
        
    Returns:
        Dict[str, Any]: Code quality metrics
    """
    if not code_text:
        return {
            "lines_of_code": 0,
            "functions": 0,
            "classes": 0,
            "comments": 0,
            "complexity": 0,
            "duplicate_lines": 0,
            "quality_score": 0
        }
    
    lines = code_text.split('\n')
    lines_of_code = len([line for line in lines if line.strip()])
    
    # Count functions (simplified)
    function_patterns = [r'\bdef\s+\w+\s*\(', r'\bfunction\s+\w+\s*\(', r'\w+\s*\([^)]*\)\s*{']
    functions = 0
    for pattern in function_patterns:
        functions += len(re.findall(pattern, code_text))
    
    # Count classes (simplified)
    class_patterns = [r'\bclass\s+\w+', r'\bstruct\s+\w+']
    classes = 0
    for pattern in class_patterns:
        classes += len(re.findall(pattern, code_text))
    
    # Count comments (simplified)
    comment_patterns = [r'#.*', r'//.*', r'/\*.*?\*/', r'<!--.*?-->']
    comments = 0
    for pattern in comment_patterns:
        comments += len(re.findall(pattern, code_text, re.DOTALL))
    
    # Simplified complexity calculation
    complexity_keywords = ['if', 'for', 'while', 'switch', 'try', 'catch', '&&', '||']
    complexity = 0
    for keyword in complexity_keywords:
        complexity += code_text.lower().count(keyword)
    
    # Estimate duplicate lines (very simplified)
    unique_lines = len(set(lines))
    duplicate_lines = lines_of_code - unique_lines if lines_of_code > 0 else 0
    
    # Calculate quality score (simplified)
    quality_score = 100
    if lines_of_code > 0:
        # Deduct points for low comment ratio
        comment_ratio = comments / lines_of_code if lines_of_code > 0 else 0
        if comment_ratio < 0.1:
            quality_score -= (0.1 - comment_ratio) * 100 * 2  # Up to 20 points deduction
        
        # Deduct points for high complexity
        complexity_ratio = complexity / lines_of_code if lines_of_code > 0 else 0
        if complexity_ratio > 0.3:
            quality_score -= (complexity_ratio - 0.3) * 100 * 1.5  # Up to 15 points deduction
        
        # Deduct points for duplicate lines
        duplicate_ratio = duplicate_lines / lines_of_code if lines_of_code > 0 else 0
        if duplicate_ratio > 0.05:
            quality_score -= (duplicate_ratio - 0.05) * 100 * 3  # Up to 30 points deduction
    
    quality_score = max(0, min(100, quality_score))  # Clamp to 0-100
    
    return {
        "lines_of_code": lines_of_code,
        "functions": functions,
        "classes": classes,
        "comments": comments,
        "complexity": complexity,
        "duplicate_lines": duplicate_lines,
        "quality_score": round(quality_score, 2)
    }

def analyze_plan_quality(plan_text: str) -> Dict[str, Any]:
    """Analyze plan quality metrics.
    
    Args:
        plan_text (str): The plan to analyze
        
    Returns:
        Dict[str, Any]: Plan quality metrics
    """
    if not plan_text:
        return {
            "sections": 0,
            "objectives": 0,
            "milestones": 0,
            "resources": 0,
            "risks": 0,
            "dependencies": 0,
            "timeline_elements": 0,
            "quality_score": 0
        }
    
    # Count sections (headers)
    sections = len(re.findall(r'^#{1,6}\s+|.*\n[=]{3,}|.*\n[-]{3,}', plan_text, re.MULTILINE))
    
    # Count objectives (look for objective-related terms)
    objective_patterns = [r'\bobjectives?\b', r'\bgoals?\b', r'\bpurpose\b', r'\baim\b']
    objectives = 0
    for pattern in objective_patterns:
        objectives += len(re.findall(pattern, plan_text, re.IGNORECASE))
    
    # Count milestones (look for milestone-related terms)
    milestone_patterns = [r'\bmilestones?\b', r'\bdeadlines?\b', r'\btimelines?\b', r'\bschedule\b']
    milestones = 0
    for pattern in milestone_patterns:
        milestones += len(re.findall(pattern, plan_text, re.IGNORECASE))
    
    # Count resources (look for resource-related terms)
    resource_patterns = [r'\bresources?\b', r'\bbudget\b', r'\bcosts?\b', r'\bmaterials?\b']
    resources = 0
    for pattern in resource_patterns:
        resources += len(re.findall(pattern, plan_text, re.IGNORECASE))
    
    # Count risks (look for risk-related terms)
    risk_patterns = [r'\brisks?\b', r'\bthreats?\b', r'\bvulnerabilit(?:y|ies)\b', r'\bhazards?\b']
    risks = 0
    for pattern in risk_patterns:
        risks += len(re.findall(pattern, plan_text, re.IGNORECASE))
    
    # Count dependencies (look for dependency-related terms)
    dependency_patterns = [r'\bdependenc(?:y|ies)\b', r'\bprerequisites?\b', r'\brequires?\b', r'\bneeds?\b']
    dependencies = 0
    for pattern in dependency_patterns:
        dependencies += len(re.findall(pattern, plan_text, re.IGNORECASE))
    
    # Count timeline elements (dates, time-related terms)
    timeline_patterns = [r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}', r'\d{4}[/-]\d{1,2}[/-]\d{1,2}', 
                       r'\bweeks?\b', r'\bmonths?\b', r'\byears?\b', r'\bdays?\b']
    timeline_elements = 0
    for pattern in timeline_patterns:
        timeline_elements += len(re.findall(pattern, plan_text, re.IGNORECASE))
    
    # Calculate quality score (simplified)
    quality_score = 50  # Start with baseline
    
    # Add points for completeness
    quality_score += min(20, sections * 2)  # Up to 20 points for sections
    quality_score += min(15, objectives * 3)  # Up to 15 points for objectives
    quality_score += min(10, milestones * 2)  # Up to 10 points for milestones
    quality_score += min(10, resources * 2)  # Up to 10 points for resources
    quality_score += min(10, risks * 2)  # Up to 10 points for risks
    quality_score += min(10, dependencies * 2)  # Up to 10 points for dependencies
    quality_score += min(10, timeline_elements * 2)  # Up to 10 points for timeline
    
    quality_score = max(0, min(100, quality_score))  # Clamp to 0-100
    
    return {
        "sections": sections,
        "objectives": objectives,
        "milestones": milestones,
        "resources": resources,
        "risks": risks,
        "dependencies": dependencies,
        "timeline_elements": timeline_elements,
        "quality_score": round(quality_score, 2)
    }

def run_adversarial_testing():
    """Main logic for the adversarial testing loop, designed to be run in a background thread."""
    try:
        # --- Initialization ---
        api_key = st.session_state.openrouter_key
        red_team_base = list(st.session_state.red_team_models or [])
        blue_team_base = list(st.session_state.blue_team_models or [])
        min_iter, max_iter = st.session_state.adversarial_min_iter, st.session_state.adversarial_max_iter
        confidence = st.session_state.adversarial_confidence
        max_tokens = st.session_state.adversarial_max_tokens
        json_mode = st.session_state.adversarial_force_json
        max_workers = st.session_state.adversarial_max_workers
        rotation_strategy = st.session_state.adversarial_rotation_strategy
        seed_str = str(st.session_state.adversarial_seed or "").strip()
        seed = None
        if seed_str:
            try:
                seed = int(float(seed_str))  # Handle floats by truncating to int
            except (ValueError, TypeError):
                pass  # Invalid input, keep seed as None

        # Validation
        if not api_key:
            _update_adv_log_and_status("‚ùå Error: OpenRouter API key is required for adversarial testing.")
            with st.session_state.thread_lock:
                st.session_state.adversarial_running = False
            return
            
        if not red_team_base or not blue_team_base:
            _update_adv_log_and_status("‚ùå Error: Please select at least one model for both red and blue teams.")
            with st.session_state.thread_lock:
                st.session_state.adversarial_running = False
            return
            
        if not st.session_state.protocol_text.strip():
            _update_adv_log_and_status("‚ùå Error: Please enter a protocol to test.")
            with st.session_state.thread_lock:
                st.session_state.adversarial_running = False
            return

        with st.session_state.thread_lock:
            st.session_state.adversarial_log = []
            st.session_state.adversarial_stop_flag = False
            st.session_state.adversarial_total_tokens_prompt = 0
            st.session_state.adversarial_total_tokens_completion = 0
            st.session_state.adversarial_cost_estimate_usd = 0.0

        model_configs = _collect_model_configs(red_team_base + blue_team_base, max_tokens)
        current_sop = st.session_state.protocol_text
        base_hash = _hash_text(current_sop)
        results, sop_hashes = [], [base_hash]
        iteration, approval_rate = 0, 0.0

        # Determine review type and get appropriate prompts
        if st.session_state.get("adversarial_custom_mode", False):
            # Use custom prompts when custom mode is enabled
            red_team_prompt = st.session_state.get("adversarial_custom_red_prompt", RED_TEAM_CRITIQUE_PROMPT)
            blue_team_prompt = st.session_state.get("adversarial_custom_blue_prompt", BLUE_TEAM_PATCH_PROMPT)
            review_type = "custom"
        else:
            # Use standard prompts based on review type
            if st.session_state.adversarial_review_type == "Auto-Detect":
                review_type = determine_review_type(current_sop)
            elif st.session_state.adversarial_review_type == "Code Review":
                review_type = "code"
            elif st.session_state.adversarial_review_type == "Plan Review":
                review_type = "plan"
            else:
                review_type = "general"
                
            red_team_prompt, blue_team_prompt = get_appropriate_prompts(review_type)
        
        _update_adv_log_and_status(f"üöÄ Start: {len(red_team_base)} red / {len(blue_team_base)} blue | seed={seed} | base_hash={base_hash} | rotation={rotation_strategy} | review_type={review_type}")

        # --- Main Loop ---
        while iteration < max_iter and not st.session_state.adversarial_stop_flag:
            iteration += 1
            
            # --- Team Rotation Logic ---
            if rotation_strategy == "Round Robin":
                red_team = [red_team_base[(iteration - 1 + i) % len(red_team_base)] for i in range(len(red_team_base))]
                blue_team = [blue_team_base[(iteration - 1 + i) % len(blue_team_base)] for i in range(len(blue_team_base))]
                _update_adv_log_and_status(f"üîÑ Iteration {iteration}/{max_iter}: Rotated teams (Round Robin). Red: {len(red_team)}, Blue: {len(blue_team)}")
            elif rotation_strategy == "Staged":
                try:
                    stages = json.loads(st.session_state.adversarial_staged_rotation_config)
                    if isinstance(stages, list) and len(stages) > 0:
                        stage_index = (iteration - 1) % len(stages)
                        stage = stages[stage_index]
                        red_team = stage.get("red", red_team_base)
                        blue_team = stage.get("blue", blue_team_base)
                        _update_adv_log_and_status(f"üîÑ Iteration {iteration}/{max_iter}: Rotated teams (Staged - Stage {stage_index + 1}). Red: {len(red_team)}, Blue: {len(blue_team)}")
                    else:
                        red_team = red_team_base
                        blue_team = blue_team_base
                        _update_adv_log_and_status(f"‚ö†Ô∏è Invalid Staged Rotation Config. Using base teams.")
                except json.JSONDecodeError:
                    red_team = red_team_base
                    blue_team = blue_team_base
                    _update_adv_log_and_status(f"‚ö†Ô∏è Invalid JSON in Staged Rotation Config. Using base teams.")
            elif rotation_strategy == "Performance-Based":
                model_performance = st.session_state.adversarial_model_performance
                red_team_weights = [model_performance.get(m, {"score": 1})["score"] for m in red_team_base]
                red_team_sample_size = min(st.session_state.adversarial_red_team_sample_size, len(red_team_base))
                if sum(red_team_weights) == 0:
                    red_team = random.sample(red_team_base, k=red_team_sample_size)
                else:
                    red_team = random.choices(red_team_base, weights=red_team_weights, k=red_team_sample_size)
                
                blue_team_sample_size = min(st.session_state.adversarial_blue_team_sample_size, len(blue_team_base))
                blue_team = random.sample(blue_team_base, k=blue_team_sample_size)
                _update_adv_log_and_status(f"üîÑ Iteration {iteration}/{max_iter}: Rotated teams (Performance-Based). Red: {len(red_team)}, Blue: {len(blue_team)}")
            elif rotation_strategy == "Adaptive":
                # Adaptive strategy based on previous iteration performance
                if iteration > 1 and len(results) > 0:
                    last_iteration = results[-1]
                    # If approval rate is low, use more diverse models
                    if last_iteration.get("approval_check", {}).get("approval_rate", 100) < 70:
                        red_team = random.sample(red_team_base, min(len(red_team_base), st.session_state.adversarial_red_team_sample_size + 1))
                        blue_team = random.sample(blue_team_base, min(len(blue_team_base), st.session_state.adversarial_blue_team_sample_size + 1))
                        _update_adv_log_and_status(f"üîÑ Iteration {iteration}/{max_iter}: Adaptive rotation - increasing diversity. Red: {len(red_team)}, Blue: {len(blue_team)}")
                    # If approval rate is high, focus on specialized models
                    elif last_iteration.get("approval_check", {}).get("approval_rate", 0) > 90:
                        # Use top performing models
                        top_red_models = sorted(st.session_state.adversarial_model_performance.items(), 
                                              key=lambda x: x[1].get("score", 0), reverse=True)[:3]
                        top_red_model_ids = [m[0] for m in top_red_models if m[0] in red_team_base]
                        
                        top_blue_models = sorted(st.session_state.adversarial_model_performance.items(), 
                                               key=lambda x: x[1].get("score", 0), reverse=True)[:3]
                        top_blue_model_ids = [m[0] for m in top_blue_models if m[0] in blue_team_base]
                        
                        if top_red_model_ids:
                            red_team = top_red_model_ids
                        else:
                            red_team = red_team_base[:min(3, len(red_team_base))]
                            
                        if top_blue_model_ids:
                            blue_team = top_blue_model_ids
                        else:
                            blue_team = blue_team_base[:min(3, len(blue_team_base))]
                        _update_adv_log_and_status(f"üîÑ Iteration {iteration}/{max_iter}: Adaptive rotation - focusing on top models. Red: {len(red_team)}, Blue: {len(blue_team)}")
                    else:
                        red_team = red_team_base
                        blue_team = blue_team_base
                else:
                    red_team = red_team_base
                    blue_team = blue_team_base
                _update_adv_log_and_status(f"üîÑ Iteration {iteration}/{max_iter}: Adaptive team selection. Red: {len(red_team)}, Blue: {len(blue_team)}")
            
            # Advanced Testing Strategies
            elif "Adaptive Testing" in st.session_state.get("advanced_testing_strategies", []):
                # Use adaptive testing strategy
                strategy = adaptive_testing_strategy(results, {
                    "red_team_models": red_team_base,
                    "blue_team_models": blue_team_base,
                    "adversarial_confidence": confidence
                })
                red_team = strategy["recommended_models"]["red_team"]
                blue_team = strategy["recommended_models"]["blue_team"]
                _update_adv_log_and_status(f"üîÑ Iteration {iteration}/{max_iter}: Adaptive testing strategy applied. Red: {len(red_team)}, Blue: {len(blue_team)}")
            
            elif "Category-Focused Testing" in st.session_state.get("advanced_testing_strategies", []):
                # Focus on specific issue categories
                if results and "agg_risk" in results[-1]:
                    categories = results[-1]["agg_risk"].get("categories", {})
                    if categories:
                        focus_recommendation = category_focused_testing(categories, {
                            "red_team_models": red_team_base,
                            "blue_team_models": blue_team_base
                        })
                        red_team = focus_recommendation["recommended_models"]["red_team"]
                        blue_team = focus_recommendation["recommended_models"]["blue_team"]
                        focus_category = focus_recommendation["focus_category"]
                        _update_adv_log_and_status(f"üîÑ Iteration {iteration}/{max_iter}: Category-focused testing on '{focus_category}'. Red: {len(red_team)}, Blue: {len(blue_team)}")
                    else:
                        red_team = red_team_base
                        blue_team = blue_team_base
                else:
                    red_team = red_team_base
                    blue_team = blue_team_base
            
            elif "Performance-Based Rotation" in st.session_state.get("advanced_testing_strategies", []):
                # Rotate models based on performance
                if st.session_state.get("adversarial_model_performance"):
                    rotated_teams = performance_based_model_rotation(
                        st.session_state.adversarial_model_performance,
                        red_team_base,
                        blue_team_base
                    )
                    red_team = rotated_teams["red_team"]
                    blue_team = rotated_teams["blue_team"]
                    _update_adv_log_and_status(f"üîÑ Iteration {iteration}/{max_iter}: Performance-based rotation. Red: {len(red_team)}, Blue: {len(blue_team)}")
                else:
                    red_team = red_team_base
                    blue_team = blue_team_base
            
            else: # "None" or any other case
                red_team = red_team_base
                blue_team = blue_team_base

            _update_adv_log_and_status(f"üîÑ Iteration {iteration}/{max_iter}: Starting red team analysis.")

            # --- RED TEAM: CRITIQUES ---
            critiques_raw = []
            with ThreadPoolExecutor(max_workers=max_workers) as ex:
                futures = {ex.submit(analyze_with_model, api_key, m, current_sop,
                                    model_configs.get(m,{}), red_team_prompt,
                                    force_json=json_mode, seed=seed, compliance_requirements=st.session_state.compliance_requirements): m for m in red_team}
                for fut in as_completed(futures):
                    res = fut.result()
                    _update_adv_counters(res['ptoks'], res['ctoks'], res['cost'])
                    if not res.get("ok") or not res.get("json"):
                        _update_adv_log_and_status(f"üî¥ {res['model_id']}: Invalid response. Details: {res.get('text', 'N/A')}")
                    critiques_raw.append({"model": res['model_id'], "critique_json": res.get("json"), "raw_text": res.get("text")})

            _update_model_performance(critiques_raw)
            agg_risk = _aggregate_red_risk(critiques_raw)
            if agg_risk['count'] == 0:
                _update_adv_log_and_status(f"üîÑ Iteration {iteration}: Red team found no exploitable issues. Checking for approval.")
            else:
                _update_adv_log_and_status(f"üîÑ Iteration {iteration}: Red team found {agg_risk['count']} issues. Starting blue team patching.")

            # --- BLUE TEAM: PATCHING ---
            blue_patches_raw = []
            valid_critiques_json = [c['critique_json'] for c in critiques_raw if c.get('critique_json')]
            critique_block = json.dumps({"critiques": valid_critiques_json}, ensure_ascii=False, indent=2)

            with ThreadPoolExecutor(max_workers=max_workers) as ex:
                futures = {ex.submit(analyze_with_model, api_key, m, current_sop,
                                    model_configs.get(m,{}), blue_team_prompt,
                                    user_suffix="\n\nCRITIQUES TO ADDRESS:\n" + critique_block,
                                    force_json=True, seed=seed): m for m in blue_team}
                for fut in as_completed(futures):
                    res = fut.result()
                    _update_adv_counters(res['ptoks'], res['ctoks'], res['cost'])
                    if not res.get("ok") or not res.get("json") or not res.get("json", {}).get("sop","").strip():
                         _update_adv_log_and_status(f"üîµ {res['model_id']}: Invalid or empty patch received. Details: {res.get('text', 'N/A')}")
                    blue_patches_raw.append({"model": res['model_id'], "patch_json": res.get("json"), "raw_text": res.get("text")})

            next_sop, consensus_diag = _merge_consensus_sop(current_sop, blue_patches_raw, critiques_raw)
            _update_adv_log_and_status(f"üîÑ Iteration {iteration}: Consensus SOP generated (Best patch from '{consensus_diag.get('model', 'N/A')}'). Starting approval check.")

            # --- APPROVAL CHECK ---
            # Use custom approval prompt when in custom mode
            if st.session_state.get("adversarial_custom_mode", False):
                approval_prompt = st.session_state.get("adversarial_custom_approval_prompt", APPROVAL_PROMPT)
            else:
                approval_prompt = APPROVAL_PROMPT
                
            eval_res = check_approval_rate(api_key, red_team, next_sop, model_configs, seed, max_workers, approval_prompt)
            approval_rate = eval_res["approval_rate"]
            _update_adv_counters(eval_res['prompt_tokens'], eval_res['completion_tokens'], eval_res['cost'])
            _update_adv_log_and_status(f"üîÑ Iteration {iteration}: Approval rate: {approval_rate:.1f}%, Avg Score: {eval_res['avg_score']:.1f}")

            results.append({
                "iteration": iteration, "critiques": critiques_raw, "patches": blue_patches_raw,
                "current_sop": next_sop, "approval_check": eval_res, "agg_risk": agg_risk, "consensus": consensus_diag,
                "cost_effectiveness": (agg_risk['count'] / st.session_state.adversarial_cost_estimate_usd) if st.session_state.adversarial_cost_estimate_usd > 0 else 0
            })
            current_sop = next_sop

            # --- Confidence Plateau and Critical Issue Triggers ---
            with st.session_state.thread_lock:
                st.session_state.adversarial_confidence_history.append(approval_rate)
                history = st.session_state.adversarial_confidence_history
                if len(history) > 3 and history[-1] == history[-2] and history[-2] == history[-3]:
                    _update_adv_log_and_status("‚ö†Ô∏è Confidence plateau detected: Confidence has not changed for 3 iterations.")
            
            if agg_risk["total_weight"] > 0 and any(issue.get("severity") == "critical" for critique in critiques_raw if critique.get("critique_json") for issue in critique["critique_json"].get("issues", [])):
                _update_adv_log_and_status("üö® Critical issue found! Activating specialist security models (simulation).")

            # --- Stagnation Check ---
            current_hash = _hash_text(current_sop)
            if len(sop_hashes) > 1 and current_hash == sop_hashes[-1] and current_hash == sop_hashes[-2]:
                _update_adv_log_and_status("‚ö†Ô∏è Stagnation detected: SOP has not changed for 2 iterations. Consider adjusting models or temperature.")
            sop_hashes.append(current_hash)

            if iteration >= min_iter and approval_rate >= confidence:
                _update_adv_log_and_status(f"‚úÖ Success! Confidence threshold of {confidence}% reached after {iteration} iterations.")
                break
        # --- End of Loop ---
        if st.session_state.adversarial_stop_flag:
            _update_adv_log_and_status("‚èπÔ∏è Process stopped by user.")
        elif iteration >= max_iter:
            _update_adv_log_and_status(f"üèÅ Reached max iterations ({max_iter}). Final approval rate: {approval_rate:.1f}%")

        with st.session_state.thread_lock:
            st.session_state.adversarial_results = {
                "final_sop": current_sop, "iterations": results, "final_approval_rate": approval_rate,
                "cost_estimate_usd": st.session_state.adversarial_cost_estimate_usd,
                "tokens": {"prompt": st.session_state.adversarial_total_tokens_prompt, "completion": st.session_state.adversarial_total_tokens_completion},
                "log": list(st.session_state.adversarial_log), "seed": seed, "base_hash": base_hash,
                "review_type": review_type
            }
            st.session_state.protocol_text = current_sop
            st.session_state.adversarial_running = False

    except Exception as e:
        # --- Global Error Handler ---
        tb_str = traceback.format_exc()
        error_message = f"üí• A critical error occurred: {e}\n{tb_str}"
        _update_adv_log_and_status(error_message)
        with st.session_state.thread_lock:
            st.session_state.adversarial_running = False
            if 'adversarial_results' not in st.session_state or not st.session_state.adversarial_results:
                st.session_state.adversarial_results = {}
            st.session_state.adversarial_results["critical_error"] = error_message
            # Ensure error is visible in UI by storing a simplified message
            st.session_state.adversarial_status_message = f"Error: {str(e)[:100]}..."

# ------------------------------------------------------------------
# 6. Sidebar ‚Äì every provider + every knob
# ------------------------------------------------------------------

with st.sidebar:
    st.title("‚öôÔ∏è Provider Configuration")
    st.caption("Controls the 'Evolution' tab. Adversarial Testing always uses OpenRouter.")
    
    # Add a visual separator
    st.markdown("---")
    
    # Provider selection section
    st.subheader("üåê Provider Selection")
    
    @st.cache_data(ttl=300)
    def _query_models(provider_key: str, api_key: str = ""):
        if provider_key not in PROVIDERS:
            return []
        loader = PROVIDERS[provider_key].get("loader")
        if loader:
            with st.spinner(f"Fetching models for {provider_key}..."):
                models = loader(api_key or None)
                if models:
                    return models
        # Fallback to the default model for the provider if loader fails or doesn't exist
        default_model = PROVIDERS[provider_key].get("model")
        return [default_model] if default_model else []

    provider = st.selectbox(
        "Provider", list(PROVIDERS.keys()), key="provider", on_change=reset_defaults,
        help="Choose the backend for the Evolution tab. Adversarial Testing always uses OpenRouter."
    )

    api_key_for_loader = st.session_state.openrouter_key if provider == "OpenRouter" else st.session_state.api_key
    model_options = _query_models(provider, api_key_for_loader)

    model_idx = 0
    if st.session_state.model in model_options:
        model_idx = model_options.index(st.session_state.model)
    elif model_options:
        st.session_state.model = model_options[0]  # Default to first in list if previous selection is invalid

    st.selectbox("Model", model_options, index=model_idx, key="model")
    st.text_input("API Key", type="password", key="api_key",
                 help=f"Leave empty to use env var: {PROVIDERS.get(provider, {}).get('env', 'Not specified')}")
    st.text_input("Base URL", key="base_url", help="Endpoint for chat/completions.")
    
    # Model parameters section
    st.markdown("---")
    st.subheader("‚öôÔ∏è Model Parameters")
    st.number_input("Max tokens", min_value=1, max_value=128_000, step=1, key="max_tokens")
    st.slider("Temperature", 0.0, 2.0, key="temperature")
    st.slider("Top-p", 0.0, 1.0, key="top_p")
    st.slider("Frequency penalty", -2.0, 2.0, key="frequency_penalty")
    st.slider("Presence penalty", -2.0, 2.0, key="presence_penalty")
    st.text_input("Seed (optional)", key="seed", help="Integer for deterministic sampling.")
    st.text_area("Extra headers (JSON dict)", height=80, key="extra_headers")
    
    # Evolution parameters section
    st.markdown("---")
    st.subheader("üîÑ Evolution Parameters")
    st.number_input("Max iterations", 1, 1000, key="max_iterations")
    st.number_input("Checkpoint interval", 1, 100, key="checkpoint_interval")
    # Disabled params from original code, kept for UI consistency
    st.number_input("Population size", 1, 100, key="population_size", disabled=True)
    st.number_input("Num islands", 1, 10, key="num_islands", disabled=True)
    st.number_input("Elite ratio", 0.0, 1.0, key="elite_ratio", disabled=True)
    st.number_input("Exploration ratio", 0.0, 1.0, key="exploration_ratio", disabled=True)
    st.number_input("Exploitation ratio", 0.0, 1.0, key="exploitation_ratio", disabled=True)
    st.number_input("Archive size", 0, 1000, key="archive_size", disabled=True)
    
    # Prompts section
    st.markdown("---")
    st.subheader("üí¨ Prompts")
    st.text_area("System prompt", height=120, key="system_prompt")
    st.text_area("Evaluator system prompt", height=120, key="evaluator_system_prompt", disabled=True)
    
    # Configuration profiles section
    st.markdown("---")
    st.subheader("üìã Configuration Profiles")
    
    # Load profile
    profiles = list_config_profiles()
    if profiles:
        selected_profile = st.selectbox("Load Profile", [""] + profiles, key="load_profile_select")
        if selected_profile and st.button("Load Selected Profile", key="load_profile_btn"):
            if load_config_profile(selected_profile):
                st.success(f"Loaded profile: {selected_profile}")
                st.rerun()
    
    # Save profile
    profile_name = st.text_input("Save Current Config As", key="save_profile_name")
    if profile_name and st.button("Save Profile", key="save_profile_btn"):
        if save_config_profile(profile_name):
            st.success(f"Saved profile: {profile_name}")
            st.rerun()
    
    # Tutorial section
    st.markdown("---")
    st.subheader("üéì Tutorial")
    if not st.session_state.tutorial_completed:
        if st.button("Start Tutorial"):
            st.session_state.current_tutorial_step = 1
            st.rerun()
        
        if st.session_state.current_tutorial_step > 0:
            tutorial_steps = [
                "Welcome to OpenEvolve! This tutorial will guide you through the main features.",
                "Step 1: Enter your protocol in the text area in the Evolution tab.",
                "Step 2: Configure your LLM provider and parameters in the sidebar.",
                "Step 3: Click 'Start Evolution' to begin improving your protocol.",
                "Step 4: Try the Adversarial Testing tab for advanced security hardening.",
                "Step 5: Use version control to track changes and collaborate with others.",
                "Tutorial completed! You're ready to use OpenEvolve."
            ]
            
            if st.session_state.current_tutorial_step <= len(tutorial_steps):
                st.info(tutorial_steps[st.session_state.current_tutorial_step - 1])
                col1, col2 = st.columns(2)
                if st.session_state.current_tutorial_step > 1:
                    if col1.button("Previous"):
                        st.session_state.current_tutorial_step -= 1
                        st.rerun()
                if st.session_state.current_tutorial_step < len(tutorial_steps):
                    if col2.button("Next"):
                        st.session_state.current_tutorial_step += 1
                        st.rerun()
                else:
                    if col2.button("Finish Tutorial"):
                        st.session_state.tutorial_completed = True
                        st.session_state.current_tutorial_step = 0
                        st.rerun()
    else:
        if st.button("Restart Tutorial"):
            st.session_state.tutorial_completed = False
            st.session_state.current_tutorial_step = 1
            st.rerun()
    
    # Reset button
    st.markdown("---")
    st.button("üîÑ Reset to provider defaults", on_click=reset_defaults, use_container_width=True)

# ------------------------------------------------------------------
# 7. Main layout with tabs
# ------------------------------------------------------------------

st.markdown('<h1 style="text-align: center; color: #4a6fa5;">üß¨ OpenEvolve Protocol Improver</h1>', unsafe_allow_html=True)
st.markdown('<p style="text-align: center; font-size: 1.2rem;">AI-Powered Protocol Hardening with Multi-LLM Consensus</p>', unsafe_allow_html=True)
st.markdown("---")

# Project information with enhanced UI
col1, col2, col3 = st.columns([3, 1, 1])
with col1:
    st.markdown("## üî¥üîµ Adversarial Testing & Evolution-based Protocol Improvement")
with col2:
    st.markdown('<span class="team-badge-lg red-team">Red Team</span><span class="team-badge-lg blue-team">Blue Team</span>', unsafe_allow_html=True)
with col3:
    # Add a quick action button with enhanced styling
    if st.button("üìã Quick Guide", key="quick_guide_btn", use_container_width=True):
        st.session_state.show_quick_guide = not st.session_state.get("show_quick_guide", False)

# Show quick guide if requested with enhanced UI
if st.session_state.get("show_quick_guide", False):
    with st.expander("üìò Quick Guide", expanded=True):
        st.markdown("""
        ### üöÄ Getting Started
        
        1. **Choose Your Approach**:
           - **Evolution Tab**: Iteratively improve a single protocol using one AI model
           - **Adversarial Testing Tab**: Harden protocols using multiple AI models in red team/blue team approach
        
        2. **Configure Your Models**:
           - Select a provider and model in the sidebar (Evolution tab)
           - Enter your OpenRouter API key for Adversarial Testing
           - Choose models for red team (critics) and blue team (fixers)
        
        3. **Input Your Protocol**:
           - Paste your existing protocol or load a template
           - Add compliance requirements if needed
        
        4. **Run the Process**:
           - Adjust parameters as needed
           - Click "Start" and monitor progress
           - Review results and save improved versions
        
        5. **Collaborate & Share**:
           - Add collaborators to your project
           - Save versions and track changes
           - Export results in multiple formats
        """)
        if st.button("Close Guide"):
            st.session_state.show_quick_guide = False
            st.rerun()

# Project info in sidebar
with st.sidebar:
    st.markdown("---")
    st.subheader("üìÅ Project Information")
    st.text_input("Project Name", key="project_name")
    st.text_area("Project Description", key="project_description", height=100)
    
    # Tags
    if HAS_STREAMLIT_TAGS:
        st.multiselect("Tags", st.session_state.tags, key="tags")
    
    # Collaborators
    st.multiselect("Collaborators", st.session_state.collaborators, key="collaborators")
    
    # Version control
    st.markdown("---")
    st.subheader("üîÑ Version Control")
    if st.button("üíæ Save Current Version"):
        if st.session_state.protocol_text.strip():
            version_name = st.text_input("Version Name", f"Version {len(st.session_state.protocol_versions) + 1}")
            comment = st.text_area("Comment", height=50)
            if st.button("‚úÖ Confirm Save", key="confirm_save_version"):
                version_id = create_new_version(st.session_state.protocol_text, version_name, comment)
                st.success(f"‚úÖ Version saved! ID: {version_id[:8]}")
                st.rerun()
    
    # Show version history with enhanced UI
    versions = get_version_history()
    if versions:
        st.write("### üìö Version History")
        # Add version comparison feature
        version_options = [f"{v['name']} ({v['timestamp'][:10]})" for v in reversed(versions[-10:])]
        selected_versions = st.multiselect("Select versions to compare", version_options, key="version_comparison")
        
        # Show timeline view toggle
        show_timeline = st.toggle("Show Timeline View", key="show_timeline")
        
        if show_timeline:
            # Show visual timeline
            st.markdown(render_version_timeline(), unsafe_allow_html=True)
        else:
            # Show traditional list view with enhanced features
            for version in reversed(versions[-5:]):  # Show last 5 versions
                col1, col2, col3, col4 = st.columns([2, 1, 1, 1])
                with col1:
                    if st.button(f"{version['name']}", key=f"load_version_{version['id']}", help=f"Load this version"):
                        load_version(version['id'])
                        st.success(f"‚úÖ Loaded version: {version['name']}")
                        st.rerun()
                with col2:
                    st.caption(f"{version['timestamp'][:10]}")
                with col3:
                    st.caption(f"v{version['id'][:8]}")
                with col4:
                    # Add branch button for non-current versions
                    if version['id'] != st.session_state.get('current_version_id', ''):
                        if st.button("üåø", key=f"branch_icon_{version['id']}", help="Create a branch from this version"):
                            branch_name = st.text_input("Branch Name", f"Branch of {version['name']}", key=f"branch_name_{version['id']}")
                            if st.button("‚úÖ Create Branch", key=f"create_branch_{version['id']}"):
                                new_branch_id = branch_version(version['id'], branch_name)
                                if new_branch_id:
                                    st.success(f"‚úÖ Branched to: {branch_name}")
                                    st.rerun()
    
    # Add collaborative workspace features
    st.markdown("---")
    st.subheader("üë• Collaborative Workspace")
    
    # Real-time collaboration status
    st.markdown("### üü¢ Collaboration Status")
    if st.session_state.collaborators:
        st.success(f"üë• {len(st.session_state.collaborators)} collaborators")
        for collaborator in st.session_state.collaborators[:3]:  # Show first 3
            st.caption(f"‚Ä¢ {collaborator}")
        if len(st.session_state.collaborators) > 3:
            st.caption(f"... and {len(st.session_state.collaborators) - 3} more")
    else:
        st.info("‚ÑπÔ∏è Add collaborators to enable real-time collaboration")
    
    # Activity feed
    st.markdown("### üìù Recent Activity")
    comments = get_comments()
    if comments:
        for comment in comments[-3:]:  # Show last 3 comments
            st.caption(f"üí¨ {comment['author']}: {comment['text'][:50]}{'...' if len(comment['text']) > 50 else ''}")
            st.caption(f"üïí {comment['timestamp'][:16]}")
    else:
        st.caption("No recent activity")
    
    # Notifications
    st.markdown("### üîî Notifications")
    if st.session_state.get("adversarial_running"):
        st.info("üîÑ Adversarial testing in progress")
    if st.session_state.get("evolution_running"):
        st.info("üîÑ Evolution process in progress")
    if st.session_state.protocol_versions and len(st.session_state.protocol_versions) > len(st.session_state.get("notified_versions", [])):
        st.success("‚úÖ New version saved")
        # Update notified versions
        st.session_state.notified_versions = st.session_state.protocol_versions.copy()

tab1, tab2 = st.tabs(["üîÑ Evolution", "‚öîÔ∏è Adversarial Testing"])

with tab1:
    # Protocol input section
    st.subheader("üìù Protocol Input")
    st.text_area("Paste your draft protocol / procedure here:", height=300, key="protocol_text", disabled=st.session_state.adversarial_running)
    
    # Protocol Templates
    templates = list_protocol_templates()
    if templates:
        col1, col2 = st.columns([3, 1])
        with col1:
            selected_template = st.selectbox("Load Template", [""] + templates, key="load_template_select")
        with col2:
            if selected_template and st.button("Load Selected Template", key="load_template_btn", use_container_width=True):
                template_content = load_protocol_template(selected_template)
                st.session_state.protocol_text = template_content
                st.success(f"Loaded template: {selected_template}")
                st.rerun()
    
    # AI Recommendations
    if st.session_state.protocol_text.strip():
        with st.expander("ü§ñ AI Recommendations", expanded=False):
            recommendations = generate_protocol_recommendations(st.session_state.protocol_text)
            suggested_template = suggest_protocol_template(st.session_state.protocol_text)
            
            st.markdown("### üí° Improvement Suggestions")
            for i, rec in enumerate(recommendations, 1):
                st.markdown(f"{i}. {rec}")
            
            st.markdown(f"### üìã Suggested Template: **{suggested_template}**")
            if st.button("Load Suggested Template"):
                template_content = load_protocol_template(suggested_template)
                if template_content:
                    st.session_state.protocol_text = template_content
                    st.success(f"Loaded template: {suggested_template}")
                    st.rerun()
    
    # Action buttons
    st.markdown("---")
    c1, c2 = st.columns(2)
    run_button = c1.button("üöÄ Start Evolution", type="primary", disabled=st.session_state.evolution_running, use_container_width=True)
    stop_button = c2.button("‚èπÔ∏è Stop Evolution", disabled=not st.session_state.evolution_running, use_container_width=True)

    # Results section
    st.markdown("---")
    left, right = st.columns(2)
    with left:
        st.subheader("üìÑ Current Best Protocol")
        proto_out = st.empty()
        
        # Protocol Analysis
        if st.session_state.evolution_current_best or st.session_state.protocol_text:
            current_protocol = st.session_state.evolution_current_best or st.session_state.protocol_text
            with st.expander("üîç Protocol Analysis", expanded=False):
                complexity = calculate_protocol_complexity(current_protocol)
                structure = extract_protocol_structure(current_protocol)
                
                # Use the new CSS class for the analysis card
                st.markdown('<div class="protocol-analysis-card">', unsafe_allow_html=True)
                st.markdown("### üìä Protocol Metrics")
                
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("_WORDS", complexity["word_count"])
                    st.metric("_SENTENCES", complexity["sentence_count"])
                    st.metric("_COMPLEXITY", complexity["complexity_score"])
                    
                with col2:
                    st.metric("PARAGRAPHS", complexity["paragraph_count"])
                    st.metric("UNIQUE WORDS", complexity["unique_words"])
                    st.metric("SECTIONS", structure["section_count"])
                st.markdown('</div>', unsafe_allow_html=True)
                
                st.markdown("### üß© Structure Analysis")
                col3, col4 = st.columns(2)
                with col3:
                    st.write("Numbered Steps:", "‚úÖ" if structure["has_numbered_steps"] else "‚ùå")
                    st.write("Bullet Points:", "‚úÖ" if structure["has_bullet_points"] else "‚ùå")
                    st.write("Headers:", "‚úÖ" if structure["has_headers"] else "‚ùå")
                with col4:
                    st.write("Preconditions:", "‚úÖ" if structure["has_preconditions"] else "‚ùå")
                    st.write("Postconditions:", "‚úÖ" if structure["has_postconditions"] else "‚ùå")
                    st.write("Error Handling:", "‚úÖ" if structure["has_error_handling"] else "‚ùå")
    with right:
        st.subheader("üîç Logs")
        log_out = st.empty()
        
        # Comments section
        with st.expander("üí¨ Comments", expanded=False):
            comments = get_comments()
            if comments:
                for comment in comments:
                    st.markdown(f"**{comment['author']}** ({comment['timestamp'][:16]})")
                    st.markdown(f"> {comment['text']}")
                    st.markdown("---")
            
            new_comment = st.text_area("Add a comment", key="new_comment")
            if st.button("Post Comment"):
                if new_comment.strip():
                    add_comment(new_comment)
                    st.success("Comment added!")
                    st.rerun()

    # Display the current state from the session state
    with st.session_state.thread_lock:
        current_log = "\n".join(st.session_state.evolution_log)
        current_protocol = st.session_state.evolution_current_best or st.session_state.protocol_text

    log_out.code(current_log, language="text")
    proto_out.code(current_protocol, language="markdown")

    # Enhanced visualization for evolution process
    if st.session_state.evolution_running or st.session_state.evolution_current_best:
        st.markdown("---")
        st.subheader("üìä Evolution Progress")
        
        # Progress metrics
        if st.session_state.evolution_current_best:
            original_complexity = calculate_protocol_complexity(st.session_state.protocol_text)
            current_complexity = calculate_protocol_complexity(st.session_state.evolution_current_best)
            
            progress_col1, progress_col2, progress_col3 = st.columns(3)
            progress_col1.metric("üìù Original Length", f"{original_complexity['word_count']} words")
            progress_col2.metric("üìù Current Length", f"{current_complexity['word_count']} words")
            progress_col3.metric("üìà Improvement", 
                                f"{current_complexity['word_count'] - original_complexity['word_count']} words", 
                                f"{((current_complexity['word_count'] / max(1, original_complexity['word_count'])) - 1) * 100:.1f}%")

    # If evolution is running, sleep for 1 second and then rerun to update the UI
    if st.session_state.evolution_running:
        time.sleep(1)
        st.rerun()

def render_adversarial_testing_tab():
    st.header("üî¥üîµ Adversarial Testing with Multi-LLM Consensus")
    
    # Add a brief introduction
    st.markdown("""
    > **How it works:** Adversarial Testing uses two teams of AI models to improve your protocols:
    > - **üî¥ Red Team** finds flaws and vulnerabilities
    > - **üîµ Blue Team** fixes the identified issues
    > The process repeats until your protocol reaches the desired confidence level.
    """)
    
    # Project Information Section
    st.markdown("---")
    st.subheader("üìÅ Project Information")
    col1, col2 = st.columns([3, 1])
    with col1:
        st.text_input("Project Name", key="project_name")
        st.text_area("Project Description", key="project_description", height=100)
    with col2:
        # Version control
        if st.button("üíæ Save Version"):
            if st.session_state.protocol_text.strip():
                version_name = st.text_input("Version Name", f"Version {len(st.session_state.protocol_versions) + 1}")
                comment = st.text_area("Comment", height=100, key="version_comment")
                if st.button("‚úÖ Confirm Save"):
                    version_id = create_new_version(st.session_state.protocol_text, version_name, comment)
                    st.success(f"‚úÖ Version saved! ID: {version_id[:8]}")
                    st.rerun()
        
        # Show version history
        versions = get_version_history()
        if versions:
            st.markdown("### üìö Versions")
            for version in reversed(versions[-5:]):  # Show last 5 versions
                col1, col2 = st.columns([3, 1])
                with col1:
                    if st.button(f"{version['name']} ({version['timestamp'][:10]})", key=f"load_version_{version['id']}"):
                        load_version(version['id'])
                        st.success(f"‚úÖ Loaded version: {version['name']}")
                        st.rerun()
                with col2:
                    st.caption(f"v{version['id'][:8]}")
    
    # Collaborative Features
    with st.expander("üë• Collaborative Features", expanded=False):
        st.markdown("### ü§ù Team Collaboration")
        collaborators = st.multiselect("Add Collaborators (email addresses)", 
                                      st.session_state.collaborators,
                                      key="collaborators")
        
        st.markdown("### üí¨ Comments & Discussions")
        comments = get_comments()
        if comments:
            for comment in comments:
                st.markdown(f"**{comment['author']}** ({comment['timestamp'][:16]})")
                st.markdown(f"> {comment['text']}")
                st.markdown("---")
        
        new_comment = st.text_area("Add a comment", key="new_comment")
        if st.button("üì§ Post Comment"):
            if new_comment.strip():
                add_comment(new_comment)
                st.success("‚úÖ Comment added!")
                st.rerun()
        
        st.markdown("### üè∑Ô∏è Tags")
        tags = st.multiselect("Add tags to organize this project", 
                             st.session_state.tags,
                             key="tags")
    
    # Quick Start Wizard
    with st.expander("‚ö° Quick Start Wizard", expanded=True):
        st.markdown("### üöÄ Get Started in 3 Easy Steps")
        
        # Step 1: Configure API Key
        st.markdown("#### 1Ô∏è‚É£ Configure OpenRouter API Key")
        openrouter_key = st.text_input("üîë Enter your OpenRouter API Key", type="password", key="wizard_openrouter_key")
        if openrouter_key:
            st.session_state.openrouter_key = openrouter_key
            st.success("‚úÖ API key saved!")
        else:
            st.info("‚ÑπÔ∏è Need an API key? Get one at [OpenRouter.ai](https://openrouter.ai/keys)")
        
        # Step 2: Select Models
        st.markdown("#### 2Ô∏è‚É£ Select AI Models")
        if openrouter_key:
            models = get_openrouter_models(openrouter_key)
            if models:
                model_names = [m['id'] for m in models if isinstance(m, dict) and 'id' in m][:10]  # Top 10 models
                
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("**üî¥ Red Team (Critics)**")
                    red_models = st.multiselect(
                        "Select 2-3 models for finding flaws",
                        options=model_names,
                        default=model_names[:2] if len(model_names) >= 2 else model_names,
                        key="wizard_red_models"
                    )
                    st.session_state.red_team_models = red_models
                    
                with col2:
                    st.markdown("**üîµ Blue Team (Fixers)**")
                    blue_models = st.multiselect(
                        "Select 2-3 models for fixing issues",
                        options=model_names,
                        default=model_names[2:4] if len(model_names) >= 4 else model_names[-2:],
                        key="wizard_blue_models"
                    )
                    st.session_state.blue_team_models = blue_models
                
                if red_models and blue_models:
                    st.success(f"‚úÖ {len(red_models)} red team and {len(blue_models)} blue team models selected!")
                else:
                    st.info("‚ÑπÔ∏è Please select at least one model for each team")
            else:
                st.warning("‚ö†Ô∏è Unable to fetch models. Please check your API key.")
        else:
            st.info("‚ÑπÔ∏è Please enter your OpenRouter API key to select models")
        
        # Step 3: Load or Create Protocol
        st.markdown("#### 3Ô∏è‚É£ Load or Create Your Protocol")
        templates = list_protocol_templates()
        if templates:
            selected_template = st.selectbox("üìù Choose a template or start from scratch", [""] + templates, key="wizard_template")
            if selected_template:
                template_content = load_protocol_template(selected_template)
                st.session_state.protocol_text = template_content
                st.success(f"‚úÖ Loaded {selected_template} template!")
        
        # Quick protocol editor
        protocol_content = st.text_area("‚úèÔ∏è Or paste/write your protocol here", 
                                       value=st.session_state.protocol_text,
                                       height=200,
                                       key="wizard_protocol_text")
        if protocol_content != st.session_state.protocol_text:
            st.session_state.protocol_text = protocol_content
        
        # Quick start button
        if st.button("üöÄ Quick Start Adversarial Testing", 
                    disabled=not (openrouter_key and st.session_state.red_team_models and st.session_state.blue_team_models and protocol_content.strip()),
                    type="primary",
                    use_container_width=True):
            st.success("üéâ Ready to go! Scroll down to configure advanced settings or click 'Start Adversarial Testing' below.")
            st.rerun()
    
    # Project controls
    col1, col2, col3, col4 = st.columns([2, 2, 1, 1])
    with col1:
        if st.button("üíæ Save Project"):
            project_data = export_project()
            st.download_button(
                label="üì• Download Project",
                data=json.dumps(project_data, indent=2),
                file_name=f"{st.session_state.project_name.replace(' ', '_')}_project.json",
                mime="application/json",
                use_container_width=True,
            )
    with col2:
        uploaded_file = st.file_uploader("üìÅ Import Project", type=["json"])
        if uploaded_file:
            try:
                project_data = json.load(uploaded_file)
                if import_project(project_data):
                    st.success("Project imported successfully!")
                    st.rerun()
            except Exception as e:
                st.error(f"Error importing project: {e}")
    with col3:
        if st.button("üìã Export Report"):
            if st.session_state.adversarial_results:
                # Create tabs for different export formats
                export_format = st.radio("Export Format", ["HTML", "PDF", "DOCX"], horizontal=True)
                
                if export_format == "HTML":
                    html_content = generate_html_report(st.session_state.adversarial_results)
                    st.download_button(
                        label="üì• Download HTML Report",
                        data=html_content,
                        file_name="adversarial_testing_report.html",
                        mime="text/html"
                    )
                elif export_format == "PDF" and HAS_FPDF:
                    pdf_content = generate_pdf_report(st.session_state.adversarial_results)
                    st.download_button(
                        label="üì• Download PDF Report",
                        data=pdf_content,
                        file_name="adversarial_testing_report.pdf",
                        mime="application/pdf"
                    )
                elif export_format == "DOCX" and HAS_DOCX:
                    docx_content = generate_docx_report(st.session_state.adversarial_results)
                    st.download_button(
                        label="üì• Download DOCX Report",
                        data=docx_content,
                        file_name="adversarial_testing_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                elif export_format == "PDF" and not HAS_FPDF:
                    st.error("FPDF library not installed. Please install it with 'pip install fpdf' to export PDF reports.")
                elif export_format == "DOCX" and not HAS_DOCX:
                    st.error("python-docx library not installed. Please install it with 'pip install python-docx' to export DOCX reports.")
    with col4:
        if st.button("‚ùì Tutorial"):
            st.session_state.show_adversarial_tutorial = True
    
    # Sharing and Collaboration Controls
    with st.expander("üîó Share & Collaborate", expanded=False):
        st.markdown("### üåê Public Sharing")
        share_publicly = st.toggle("Share publicly", key="share_publicly")
        if share_publicly:
            st.info("üîí Your project will be accessible via a public link. Only people with the link can view it.")
            if st.button("üîó Generate Shareable Link"):
                # In a real implementation, this would generate a real shareable link
                share_link = f"https://open-evolve.app/shared/{uuid.uuid4()}"
                st.code(share_link, language="markdown")
                st.info("üìã Copy this link to share your project with others.")
        
        st.markdown("### üìß Invite Collaborators")
        collaborator_emails = st.text_area("Enter email addresses (one per line)", 
                                          key="collaborator_emails",
                                          height=100)
        if st.button("‚úâÔ∏è Send Invitations"):
            if collaborator_emails.strip():
                emails = [email.strip() for email in collaborator_emails.split("\n") if email.strip()]
                st.success(f"üìß Sent invitations to {len(emails)} collaborators!")
            else:
                st.warning("üìß Please enter at least one email address.")
        
        st.markdown("### üì§ Export Options")
        export_options = st.multiselect(
            "Select what to export",
            ["Protocol Versions", "Adversarial Results", "Comments", "Analytics", "Full Project"],
            default=["Full Project"]
        )
        
        if "Full Project" in export_options:
            st.download_button(
                label="üì¶ Export Full Project (.json)",
                data=json.dumps(export_project(), indent=2),
                file_name=f"{st.session_state.project_name.replace(' ', '_')}_full_export.json",
                mime="application/json"
            )
        else:
            # Custom export
            if st.button("‚öôÔ∏è Generate Custom Export"):
                custom_export = {}
                if "Protocol Versions" in export_options:
                    custom_export["versions"] = st.session_state.protocol_versions
                if "Adversarial Results" in export_options:
                    custom_export["results"] = st.session_state.adversarial_results
                if "Comments" in export_options:
                    custom_export["comments"] = st.session_state.comments
                if "Analytics" in export_options:
                    if st.session_state.adversarial_results:
                        custom_export["analytics"] = generate_advanced_analytics(st.session_state.adversarial_results)
                
                st.download_button(
                    label="üì• Download Custom Export (.json)",
                    data=json.dumps(custom_export, indent=2),
                    file_name=f"{st.session_state.project_name.replace(' ', '_')}_custom_export.json",
                    mime="application/json"
                )
        
        # Export to different formats
        st.markdown("### üìÑ Format Export")
        format_options = st.selectbox("Export Format", 
                                      ["Markdown", "PDF", "Word Document", "HTML", "LaTeX", "Plain Text"])
        if st.button(f"üñ®Ô∏è Export as {format_options}"):
            if format_options == "Markdown":
                st.download_button(
                    label="üì• Download Markdown (.md)",
                    data=st.session_state.protocol_text,
                    file_name=f"{st.session_state.project_name.replace(' ', '_')}.md",
                    mime="text/markdown"
                )
            elif format_options == "PDF":
                st.info("Generating PDF... (This would generate a formatted PDF in a real implementation)")
            elif format_options == "Word Document":
                st.info("Generating Word document... (This would generate a .docx file in a real implementation)")
            elif format_options == "HTML":
                st.download_button(
                    label="üì• Download HTML (.html)",
                    data=f"<html><body><h1>{st.session_state.project_name}</h1><pre>{st.session_state.protocol_text}</pre></body></html>",
                    file_name=f"{st.session_state.project_name.replace(' ', '_')}.html",
                    mime="text/html"
                )
            elif format_options == "LaTeX":
                st.info("Generating LaTeX... (This would generate a .tex file in a real implementation)")
            elif format_options == "Plain Text":
                st.download_button(
                    label="üì• Download Plain Text (.txt)",
                    data=st.session_state.protocol_text,
                    file_name=f"{st.session_state.project_name.replace(' ', '_')}.txt",
                    mime="text/plain"
                )
    
    # Show tutorial modal if requested
    if st.session_state.get("show_adversarial_tutorial", False):
        with st.expander("üìò Adversarial Testing Tutorial", expanded=True):
            st.markdown("""
            ### üéì Adversarial Testing Guide
            
            Adversarial Testing uses two teams of AI models to improve your protocols:
            
            #### üî¥ Red Team (Critics)
            - Finds flaws, vulnerabilities, and weaknesses in your protocol
            - Analyzes for logical gaps, ambiguities, and potential misuse
            
            #### üîµ Blue Team (Fixers)
            - Addresses the issues identified by the Red Team
            - Produces improved versions of the protocol
            
            #### üîÑ Process
            1. Red Team critiques the protocol
            2. Blue Team patches the identified issues
            3. Consensus mechanism selects the best patch
            4. Approval check validates the improvements
            5. Process repeats until confidence threshold is reached
            
            #### ‚öôÔ∏è Key Parameters
            - **Confidence Threshold**: Percentage of Red Team that must approve the protocol
            - **Iterations**: Number of improvement cycles to run
            - **Rotation Strategy**: How to select models for each iteration
            - **Custom Mode**: Use your own prompts for testing
            
            #### üéØ Tips for Best Results
            - Use diverse models for both teams
            - Set appropriate confidence thresholds (80-95%)
            - Provide clear compliance requirements
            - Use custom prompts for domain-specific testing
            """)
            
            if st.button("Close Tutorial"):
                st.session_state.show_adversarial_tutorial = False
                st.rerun()
    
    # Quick Start Guide
    with st.expander("‚ö° Quick Start Guide", expanded=False):
        st.markdown("""
        ### üöÄ Getting Started in 3 Steps:
        
        1. **üîë Configure OpenRouter**
           - Enter your OpenRouter API key
           - Select models for Red and Blue teams
        
        2. **üìù Input Your Protocol**
           - Paste your protocol or load a template
           - Add compliance requirements if needed
        
        3. **‚ñ∂Ô∏è Run Adversarial Testing**
           - Adjust parameters as needed
           - Click "Start Adversarial Testing"
           - Monitor progress in real-time
        """)
        
        if st.button("üìã Load Sample Protocol"):
            sample_protocol = """# Sample Security Policy

## Overview
This policy defines security requirements for accessing company systems.

## Scope
Applies to all employees, contractors, and vendors with system access.

## Policy Statements
1. All users must use strong passwords
2. Multi-factor authentication is required for sensitive systems
3. Regular security training is mandatory
4. Incident reporting must occur within 24 hours

## Roles and Responsibilities
- IT Security Team: Enforces policy and monitors compliance
- Employees: Follow security practices and report incidents
- Managers: Ensure team compliance and provide resources

## Compliance
- Audits conducted quarterly
- Violations result in disciplinary action
- Continuous monitoring through SIEM tools

## Exceptions
- Emergency access requests require manager approval
- Temporary exceptions require security team approval

## Review and Updates
- Policy reviewed annually
- Updates approved by CISO"""
            st.session_state.protocol_text = sample_protocol
            st.success("Sample protocol loaded! You can now start adversarial testing.")
            st.rerun()
    
    # OpenRouter Configuration
    st.subheader("üîë OpenRouter Configuration")
    openrouter_key = st.text_input("OpenRouter API Key", type="password", key="openrouter_key")
    if not openrouter_key:
        st.info("Enter your OpenRouter API key to enable model selection and testing.")
        return

    models = get_openrouter_models(openrouter_key)
    # Update global model metadata with thread safety
    for m in models:
        if isinstance(m, dict) and (mid := m.get("id")):
            with MODEL_META_LOCK:
                MODEL_META_BY_ID[mid] = m
    if not models:
        st.error("No models fetched. Check your OpenRouter key and connection.")
        return

    model_options = sorted([
        f"{m['id']} (Ctx: {m.get('context_length', 'N/A')}, "
        f"In: {_parse_price_per_million(m.get('pricing', {}).get('prompt')) or 'N/A'}/M, "
        f"Out: {_parse_price_per_million(m.get('pricing', {}).get('completion')) or 'N/A'}/M)"
        for m in models if isinstance(m, dict) and "id" in m
    ])

    # Protocol Templates
    st.markdown("---")
    st.subheader("üìù Protocol Input")
    
    # Add protocol input guidance
    st.info("üí° **Tip:** Start with a clear, well-structured protocol. The better your starting point, the better the results.")
    
    # Protocol editor with enhanced features
    protocol_col1, protocol_col2 = st.columns([3, 1])
    with protocol_col1:
        protocol_text = st.text_area("‚úèÔ∏è Enter or paste your protocol/procedure:", 
                                    value=st.session_state.protocol_text,
                                    height=300,
                                    key="protocol_text",
                                    placeholder="Paste your draft protocol/procedure here...\n\nExample:\n# Security Policy\n\n## Overview\nThis policy defines requirements for secure system access.\n\n## Scope\nApplies to all employees and contractors.\n\n## Policy Statements\n1. All users must use strong passwords\n2. Multi-factor authentication is required for sensitive systems\n3. Regular security training is mandatory\n\n## Compliance\nViolations result in disciplinary action.")
    
    with protocol_col2:
        st.markdown("**üìã Quick Actions**")
        
        # Template loading
        templates = list_protocol_templates()
        if templates:
            selected_template = st.selectbox("Load Template", [""] + templates, key="adv_load_template_select")
            if selected_template and st.button("üì• Load Template", use_container_width=True):
                template_content = load_protocol_template(selected_template)
                st.session_state.protocol_text = template_content
                st.success(f"Loaded: {selected_template}")
                st.rerun()
        
        # Sample protocol
        if st.button("üß™ Load Sample", use_container_width=True):
            sample_protocol = """# Sample Security Policy

## Overview
This policy defines security requirements for accessing company systems.

## Scope
Applies to all employees, contractors, and vendors with system access.

## Policy Statements
1. All users must use strong passwords
2. Multi-factor authentication is required for sensitive systems
3. Regular security training is mandatory
4. Incident reporting must occur within 24 hours

## Roles and Responsibilities
- IT Security Team: Enforces policy and monitors compliance
- Employees: Follow security practices and report incidents
- Managers: Ensure team compliance and provide resources

## Compliance
- Audits conducted quarterly
- Violations result in disciplinary action
- Continuous monitoring through SIEM tools

## Exceptions
- Emergency access requests require manager approval
- Temporary exceptions require security team approval

## Review and Updates
- Policy reviewed annually
- Updates approved by CISO"""
            st.session_state.protocol_text = sample_protocol
            st.success("Sample protocol loaded!")
            st.rerun()
        
        # Clear button
        if st.session_state.protocol_text.strip() and st.button("üóëÔ∏è Clear", use_container_width=True):
            st.session_state.protocol_text = ""
            st.rerun()
        
        # Protocol analysis
        if st.session_state.protocol_text.strip():
            complexity = calculate_protocol_complexity(st.session_state.protocol_text)
            structure = extract_protocol_structure(st.session_state.protocol_text)
            
            st.markdown("**üìä Quick Stats**")
            st.metric("_WORDS", complexity["word_count"])
            st.metric("_SENTENCES", complexity["sentence_count"])
            st.metric("_COMPLEXITY", complexity["complexity_score"])
            
            # Structure indicators
            structure_icons = []
            if structure["has_numbered_steps"]:
                structure_icons.append("üî¢")
            if structure["has_bullet_points"]:
                structure_icons.append("‚Ä¢")
            if structure["has_headers"]:
                structure_icons.append("#")
            if structure["has_preconditions"]:
                structure_icons.append("üîí")
            if structure["has_postconditions"]:
                structure_icons.append("‚úÖ")
            if structure["has_error_handling"]:
                structure_icons.append("‚ö†Ô∏è")
            
            if structure_icons:
                st.markdown(" ".join(structure_icons))

    # AI Recommendations
    if st.session_state.protocol_text.strip():
        with st.expander("ü§ñ AI Recommendations", expanded=False):
            recommendations = generate_protocol_recommendations(st.session_state.protocol_text)
            suggested_template = suggest_protocol_template(st.session_state.protocol_text)
            
            st.markdown("### üí° Improvement Suggestions")
            for i, rec in enumerate(recommendations, 1):
                st.markdown(f"{i}. {rec}")
            
            st.markdown(f"### üìã Suggested Template: **{suggested_template}**")
            if st.button("Load Suggested Template", key="adv_load_suggested_template"):
                template_content = load_protocol_template(suggested_template)
                if template_content:
                    st.session_state.protocol_text = template_content
                    st.success(f"Loaded template: {suggested_template}")
                    st.rerun()

    # Model Selection
    st.markdown("---")
    st.subheader("ü§ñ Model Selection")
    
    # Add model selection guidance
    st.info("üí° **Tip:** Select 3-5 diverse models for each team for best results. Mix small and large models for cost-effectiveness.")
    
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("#### üî¥ Red Team (Critics)")
        st.caption("Models that find flaws and vulnerabilities in your protocol")
        
        # Quick selection buttons
        if model_options:
            quick_red_models = st.multiselect(
                "Quick Select Red Team Models", 
                options=[opt.split(" (")[0] for opt in model_options[:10]],  # First 10 models
                default=st.session_state.red_team_models[:3] if st.session_state.red_team_models else [],
                key="quick_red_select"
            )
            if quick_red_models:
                st.session_state.red_team_models = quick_red_models
        
        if HAS_STREAMLIT_TAGS:
            red_team_selected_full = st_tags(
                label="Search and select models:", 
                text="Type to search models...", 
                value=st.session_state.red_team_models,
                suggestions=model_options, 
                key="red_team_select"
            )
            # Robust model ID extraction from descriptive string
            red_team_models = []
            for m in red_team_selected_full:
                # Extract model ID by splitting on first occurrence of " (" or using entire string
                # if not found
                if " (" in m:
                    model_id = m.split(" (")[0].strip()
                else:
                    model_id = m.strip()
                if model_id:
                    red_team_models.append(model_id)
            st.session_state.red_team_models = sorted(list(set(red_team_models)))
        else:
            st.warning("streamlit_tags not available. Using text input for model selection.")
            red_team_input = st.text_input("Enter Red Team models (comma-separated):", value=",".join(st.session_state.red_team_models))
            st.session_state.red_team_models = sorted(list(set([model.strip() for model in red_team_input.split(",") if model.strip()])))
            
        # Model count indicator
        st.caption(f"Selected: {len(st.session_state.red_team_models)} models")
        
    with col2:
        st.markdown("#### üîµ Blue Team (Fixers)")
        st.caption("Models that patch the identified flaws and improve the protocol")
        
        # Quick selection buttons
        if model_options:
            quick_blue_models = st.multiselect(
                "Quick Select Blue Team Models", 
                options=[opt.split(" (")[0] for opt in model_options[:10]],  # First 10 models
                default=st.session_state.blue_team_models[:3] if st.session_state.blue_team_models else [],
                key="quick_blue_select"
            )
            if quick_blue_models:
                st.session_state.blue_team_models = quick_blue_models
        
        if HAS_STREAMLIT_TAGS:
            blue_team_selected_full = st_tags(
                label="Search and select models:", 
                text="Type to search models...", 
                value=st.session_state.blue_team_models,
                suggestions=model_options, 
                key="blue_team_select"
            )
            # Robust model ID extraction from descriptive string
            blue_team_models = []
            for m in blue_team_selected_full:
                # Extract model ID by splitting on first occurrence of " (" or using entire string
                # if not found
                if " (" in m:
                    model_id = m.split(" (")[0].strip()
                else:
                    model_id = m.strip()
                if model_id:
                    blue_team_models.append(model_id)
            st.session_state.blue_team_models = sorted(list(set(blue_team_models)))
        else:
            st.warning("streamlit_tags not available. Using text input for model selection.")
            blue_team_input = st.text_input("Enter Blue Team models (comma-separated):", value=",".join(st.session_state.blue_team_models))
            st.session_state.blue_team_models = sorted(list(set([model.strip() for model in blue_team_input.split(",") if model.strip()])))
            
        # Model count indicator
        st.caption(f"Selected: {len(st.session_state.blue_team_models)} models")
    
    # Model selection validation
    if st.session_state.red_team_models and st.session_state.blue_team_models:
        total_models = len(st.session_state.red_team_models) + len(st.session_state.blue_team_models)
        if total_models > 10:
            st.warning(f"‚ö†Ô∏è You have selected {total_models} models. Consider reducing the number to control costs and processing time.")
        else:
            st.success(f"‚úÖ Ready! {len(st.session_state.red_team_models)} red team and {len(st.session_state.blue_team_models)} blue team models selected.")
    elif not st.session_state.red_team_models or not st.session_state.blue_team_models:
        st.info("‚ÑπÔ∏è Please select at least one model for each team to proceed.")
    
    # Testing Parameters
    st.markdown("---")
    st.subheader("üß™ Testing Parameters")
    
    # Preset Selector
    with st.expander("üéØ Presets", expanded=True):
        st.markdown("### üöÄ Quick Start with Presets")
        st.info("üí° **Tip:** Use presets to quickly configure adversarial testing for common scenarios.")
        
        preset_names = list_adversarial_presets()
        if preset_names:
            col1, col2 = st.columns([3, 1])
            with col1:
                selected_preset = st.selectbox("Choose a preset configuration", [""] + preset_names, key="preset_selector")
            with col2:
                if st.button("Apply Preset", key="apply_preset_btn", use_container_width=True):
                    if selected_preset and apply_adversarial_preset(selected_preset):
                        st.success(f"‚úÖ Applied {selected_preset} preset!")
                        st.rerun()
                    elif selected_preset:
                        st.error("‚ùå Failed to apply preset.")
            
            # Show preset details
            if selected_preset:
                preset = load_adversarial_preset(selected_preset)
                if preset:
                    st.markdown(f"**{preset['name']}**")
                    st.caption(preset['description'])
                    col1, col2 = st.columns(2)
                    with col1:
                        st.write("**üî¥ Red Team Models:**")
                        for model in preset.get("red_team_models", []):
                            st.code(model, language="markdown")
                    with col2:
                        st.write("**üîµ Blue Team Models:**")
                        for model in preset.get("blue_team_models", []):
                            st.code(model, language="markdown")
                    st.write("**‚öôÔ∏è Settings:**")
                    st.write(f"- Iterations: {preset.get('min_iter', 3)}-{preset.get('max_iter', 10)}")
                    st.write(f"- Confidence Threshold: {preset.get('confidence_threshold', 85)}%")
                    st.write(f"- Review Type: {preset.get('review_type', 'General SOP')}")
        
        # Advanced Testing Strategies
        st.markdown("### üß† Advanced Testing Strategies")
        strategy_options = st.multiselect(
            "Select testing strategies to enable:",
            ["Adaptive Testing", "Category-Focused Testing", "Performance-Based Rotation", "Continuous Learning"],
            default=[],
            key="advanced_testing_strategies"
        )
        
        if "Adaptive Testing" in strategy_options:
            st.info("üîÑ **Adaptive Testing**: Automatically adjusts testing intensity based on results.")
        
        if "Category-Focused Testing" in strategy_options:
            focus_category = st.selectbox(
                "Focus on specific issue category:",
                ["", "Security", "Compliance", "Clarity", "Completeness", "Efficiency"],
                key="category_focus"
            )
            if focus_category:
                st.info(f"üéØ **Category Focus**: Testing will emphasize {focus_category.lower()} issues.")
        
        if "Performance-Based Rotation" in strategy_options:
            st.info("‚ö° **Performance-Based Rotation**: Automatically rotates models based on performance metrics.")
        
        if "Continuous Learning" in strategy_options:
            st.info("üìö **Continuous Learning**: Uses historical results to improve future testing runs.")
    
    # Custom Mode Toggle
    use_custom_mode = st.toggle("üîß Use Custom Mode", key="adversarial_custom_mode", 
                               help="Enable custom prompts and configurations for adversarial testing")
    
    if use_custom_mode:
        with st.expander("üîß Custom Prompts", expanded=True):
            st.text_area("Red Team Prompt (Critique)", 
                        value=RED_TEAM_CRITIQUE_PROMPT,
                        key="adversarial_custom_red_prompt", 
                        height=200,
                        help="Custom prompt for the red team to find flaws in the protocol")
            
            st.text_area("Blue Team Prompt (Patch)", 
                        value=BLUE_TEAM_PATCH_PROMPT,
                        key="adversarial_custom_blue_prompt", 
                        height=200,
                        help="Custom prompt for the blue team to patch the identified flaws")
            
            st.text_area("Approval Prompt", 
                        value=APPROVAL_PROMPT,
                        key="adversarial_custom_approval_prompt", 
                        height=150,
                        help="Custom prompt for final approval checking")
    
    # Review Type Selection
    review_types = ["Auto-Detect", "General SOP", "Code Review", "Plan Review"]
    st.selectbox("Review Type", review_types, key="adversarial_review_type", 
                 help="Select the type of review to perform. Auto-Detect will analyze the content and choose the appropriate review type.")
    
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.number_input("Min iterations", 1, 50, key="adversarial_min_iter")
        st.number_input("Max iterations", 1, 200, key="adversarial_max_iter")
    with c2:
        st.slider("Confidence threshold (%)", 50, 100, key="adversarial_confidence", help="Stop if this % of Red Team approves the SOP.")
    with c3:
        st.number_input("Max tokens per model", 1000, 100000, key="adversarial_max_tokens")
        st.number_input("Max parallel workers", 1, 24, key="adversarial_max_workers")
    with c4:
        st.toggle("Force JSON mode", key="adversarial_force_json", help="Use model's built-in JSON mode if available. Increases reliability.")
        st.text_input("Deterministic seed", key="adversarial_seed", help="Integer for reproducible runs.")
        st.selectbox("Rotation Strategy", ["None", "Round Robin", "Random Sampling", "Performance-Based", "Staged", "Adaptive", "Focus-Category"], key="adversarial_rotation_strategy")
        if st.session_state.adversarial_rotation_strategy == "Staged":
            st.text_area("Staged Rotation Config (JSON)", key="adversarial_staged_rotation_config", height=150, help="""
[{"red": ["model1", "model2"], "blue": ["model3"]},
 {"red": ["model4"], "blue": ["model5", "model6"]}]
""")
        st.number_input("Red Team Sample Size", 1, 100, key="adversarial_red_team_sample_size")
        st.number_input("Blue Team Sample Size", 1, 100, key="adversarial_blue_team_sample_size")

    st.text_area("Compliance Requirements", key="compliance_requirements", height=150, help="Enter any compliance requirements that the red team should check for.")
    
    # Advanced customization options
    with st.expander("‚öôÔ∏è Advanced Customization", expanded=False):
        st.markdown("### üéØ Target Metrics")
        col1, col2 = st.columns(2)
        with col1:
            st.number_input("Target Complexity Score", 0, 100, key="adversarial_target_complexity", 
                           help="Target complexity score for the final protocol (0-100)")
        with col2:
            st.number_input("Target Length (words)", 0, 10000, key="adversarial_target_length", 
                           help="Target length for the final protocol in words (0 = no limit)")
        
        st.markdown("### üß† Intelligence Settings")
        st.slider("Critique Depth", 1, 10, key="adversarial_critique_depth", 
                 help="How deeply the red team should analyze the protocol (1-10)")
        st.slider("Patch Quality", 1, 10, key="adversarial_patch_quality", 
                 help="Quality level for blue team patches (1-10)")
        
        st.markdown("### üìä Evaluation Settings")
        st.toggle("Detailed Issue Tracking", key="adversarial_detailed_tracking", 
                 help="Track issues by category and severity in detail")
        st.toggle("Performance Analytics", key="adversarial_performance_analytics", 
                 help="Show detailed model performance analytics")
        
        st.markdown("### üîÑ Iteration Controls")
        st.toggle("Early Stopping", key="adversarial_early_stopping", 
                 help="Stop early if no improvement is detected")
        st.number_input("Early Stopping Patience", 1, 10, key="adversarial_early_stopping_patience", 
                       help="Number of iterations to wait before early stopping")
        
        st.markdown("### üé® Style Customization")
        st.selectbox("Writing Style", ["Professional", "Concise", "Detailed", "Casual", "Technical", "Executive"], 
                    key="adversarial_writing_style", 
                    help="Preferred writing style for the final protocol")
        st.selectbox("Tone", ["Neutral", "Authoritative", "Friendly", "Strict", "Persuasive"], 
                    key="adversarial_tone", 
                    help="Desired tone for the protocol")
        st.text_input("Custom Style Instructions", 
                     key="adversarial_custom_style", 
                     help="Additional style instructions for the protocol writer")
        
        st.markdown("### üõ°Ô∏è Security Settings")
        st.toggle("Include Security Headers", key="adversarial_include_security_headers", 
                 help="Add security-focused headers to the protocol")
        st.toggle("Include Compliance Checks", key="adversarial_include_compliance_checks", 
                 help="Automatically add compliance-related sections")
        st.text_area("Custom Security Requirements", 
                    key="adversarial_custom_security", 
                    height=100,
                    help="Additional security requirements to enforce")
        
        st.markdown("### üì¶ Format Options")
        st.selectbox("Output Format", ["Markdown", "Plain Text", "HTML", "LaTeX"], 
                    key="adversarial_output_format", 
                    help="Desired output format for the final protocol")
        st.toggle("Include Table of Contents", key="adversarial_include_toc", 
                 help="Add automatically generated table of contents")
        st.toggle("Include Revision History", key="adversarial_include_revision_history", 
                 help="Track changes with revision history section")
        
        st.markdown("### üß™ Experimental Features")
        st.toggle("Use Chain-of-Thought Reasoning", key="adversarial_use_cot", 
                 help="Enable chain-of-thought reasoning for deeper analysis")
        st.toggle("Include Confidence Intervals", key="adversarial_include_confidence", 
                 help="Add confidence intervals to issue severity ratings")
        st.toggle("Enable Self-Critique", key="adversarial_enable_self_critique", 
                 help="Have models critique their own suggestions before finalizing")
        
        st.markdown("### ‚ö° Performance Optimization")
        st.toggle("Auto-Optimize Model Selection", key="adversarial_auto_optimize_models", 
                 help="Automatically select optimal models based on protocol complexity and budget")
        budget_limit = st.number_input("Budget Limit (USD)", 0.0, 100.0, 0.0, 0.1,
                                      key="adversarial_budget_limit",
                                      help="Maximum budget for this testing session (0 = no limit)")
        
        # Performance suggestions button
        if st.button("üí° Get Performance Suggestions"):
            current_config = {
                "red_team_models": st.session_state.red_team_models,
                "blue_team_models": st.session_state.blue_team_models,
                "adversarial_max_iter": st.session_state.adversarial_max_iter,
                "protocol_text": st.session_state.protocol_text
            }
            suggestions = suggest_performance_improvements(current_config)
            st.markdown("### üöÄ Performance Suggestions")
            for suggestion in suggestions:
                st.write(suggestion)
        
        # Time and cost estimation
        if st.button("‚è±Ô∏è Estimate Time & Cost"):
            protocol_length = len(st.session_state.protocol_text.split())
            estimate = estimate_testing_time_and_cost(
                st.session_state.red_team_models,
                st.session_state.blue_team_models,
                st.session_state.adversarial_max_iter,
                protocol_length
            )
            st.markdown("### üìä Time & Cost Estimate")
            col1, col2, col3, col4 = st.columns(4)
            col1.metric("‚è∞ Est. Time", f"{estimate['estimated_time_minutes']} min")
            col2.metric("üí∞ Est. Cost", f"${estimate['estimated_cost_usd']:.4f}")
            col3.metric("üîÑ Operations", f"{estimate['total_operations']:,}")
            col4.metric("üî§ Tokens", f"{estimate['total_tokens_estimated']:,}")

    all_models = sorted(list(set(st.session_state.red_team_models + st.session_state.blue_team_models)))
    if all_models:
        with st.expander("üîß Per-Model Configuration", expanded=False):
            for model_id in all_models:
                st.markdown(f"**{model_id}**")
                cc1, cc2, cc3, cc4 = st.columns(4)
                cc1.slider(f"Temp##{model_id}", 0.0, 2.0, 0.7, 0.1, key=f"temp_{model_id}")
                cc2.slider(f"Top-P##{model_id}", 0.0, 1.0, 1.0, 0.1, key=f"topp_{model_id}")
                cc3.slider(f"Freq Pen##{model_id}", -2.0, 2.0, 0.0, 0.1, key=f"freqpen_{model_id}")
                cc4.slider(f"Pres Pen##{model_id}", -2.0, 2.0, 0.0, 0.1, key=f"prespen_{model_id}")

    # Metrics Dashboard Preview
    with st.expander("üìä Metrics Dashboard Preview", expanded=False):
        st.markdown("### üìä Real-time Metrics (During Testing)")
        st.info("These metrics will be updated in real-time during adversarial testing:")
        col1, col2 = st.columns(2)
        with col1:
            st.metric("üìà Current Confidence", "0.0%")
            st.metric("üí∞ Est. Cost (USD)", "$0.0000")
            st.metric("üî§ Prompt Tokens", "0")
        with col2:
            st.metric("üîÑ Iterations", "0/0")
            st.metric("üìù Completion Tokens", "0")
            st.metric("‚ö° Avg Response Time", "0ms")
        
        st.markdown("### üìà Confidence Trend")
        st.line_chart([0, 0, 0, 0, 0])  # Placeholder chart
        
        st.markdown("### üèÜ Top Performing Models")
        st.write("Model performance rankings will appear here during testing.")

    # Start/Stop buttons for adversarial testing
    st.markdown("---")
    col1, col2, col3 = st.columns([2, 2, 1])
    start_button = col1.button("üöÄ Start Adversarial Testing", type="primary", 
                              disabled=st.session_state.adversarial_running or not st.session_state.protocol_text.strip(),
                              use_container_width=True)
    stop_button = col2.button("‚èπÔ∏è Stop Adversarial Testing", 
                             disabled=not st.session_state.adversarial_running,
                             use_container_width=True)
    
    # Progress and status section
    if st.session_state.adversarial_running or st.session_state.adversarial_status_message:
        status_container = st.container()
        with status_container:
            # Enhanced status display
            if st.session_state.adversarial_status_message:
                # Use different colors based on status message content
                status_msg = st.session_state.adversarial_status_message
                if "Success" in status_msg or "‚úÖ" in status_msg:
                    st.success(status_msg)
                elif "Error" in status_msg or "üí•" in status_msg or "‚ö†Ô∏è" in status_msg:
                    st.error(status_msg)
                elif "Stop" in status_msg or "‚èπÔ∏è" in status_msg:
                    st.warning(status_msg)
                else:
                    st.info(status_msg)
            
            # Enhanced progress tracking
            if st.session_state.adversarial_running:
                # Progress bar with iteration info
                current_iter = len(st.session_state.get("adversarial_confidence_history", []))
                max_iter = st.session_state.adversarial_max_iter
                progress = min(current_iter / max(1, max_iter), 1.0)
                
                # Progress bar with percentage
                st.progress(progress, text=f"Iteration {current_iter}/{max_iter} ({int(progress*100)}%)")
                
                # Real-time metrics
                if st.session_state.get("adversarial_confidence_history"):
                    current_confidence = st.session_state.adversarial_confidence_history[-1]
                    col1, col2, col3, col4 = st.columns(4)
                    col1.metric("üìä Current Confidence", f"{current_confidence:.1f}%")
                    col2.metric("üí∞ Est. Cost (USD)", f"${st.session_state.adversarial_cost_estimate_usd:.4f}")
                    col3.metric("üî§ Prompt Tokens", f"{st.session_state.adversarial_total_tokens_prompt:,}")
                    col4.metric("üìù Completion Tokens", f"{st.session_state.adversarial_total_tokens_completion:,}")
                
                # Enhanced logs with auto-scroll
                with st.expander("üîç Real-time Logs", expanded=True):
                    if st.session_state.adversarial_log:
                        # Show last 50 entries instead of 20 for better visibility
                        log_content = "\n".join(st.session_state.adversarial_log[-50:])
                        st.text_area("Activity Log", value=log_content, height=300, 
                                    key="adversarial_log_display", 
                                    help="Auto-updating log of adversarial testing activities")
                    else:
                        st.info("‚è≥ Waiting for adversarial testing to start...")

    # If adversarial testing has results, show them with enhanced visualization
            if st.session_state.adversarial_results and not st.session_state.adversarial_running:
                with st.expander("üèÜ Adversarial Testing Results", expanded=True):
                    results = st.session_state.adversarial_results
                    
                    # Enhanced metrics dashboard with better organization
                    st.markdown("### üìä Performance Summary")
                    
                    # Main metrics in cards with enhanced styling
                    col1, col2, col3, col4 = st.columns(4)
                    col1.metric("‚úÖ Final Approval Rate", f"{results.get('final_approval_rate', 0):.1f}%")
                    col2.metric("üîÑ Iterations Completed", len(results.get('iterations', [])))
                    col3.metric("üí∞ Total Cost (USD)", f"${results.get('cost_estimate_usd', 0):.4f}")
                    col4.metric("üî§ Total Tokens", f"{results.get('tokens', {}).get('prompt', 0) + results.get('tokens', {}).get('completion', 0):,}")
                    
                    # Add a visual summary dashboard
                    st.markdown("### üìä Visual Summary")
                    # Create a summary visualization
                    summary_col1, summary_col2 = st.columns(2)
                    with summary_col1:
                        st.markdown('<div class="chart-container">', unsafe_allow_html=True)
                        st.markdown('<div class="chart-title">Approval Rate Progress</div>', unsafe_allow_html=True)
                        if results.get('iterations'):
                            confidence_history = [iter.get("approval_check", {}).get("approval_rate", 0) 
                                                for iter in results.get('iterations', [])]
                            st.line_chart(confidence_history)
                        st.markdown('</div>', unsafe_allow_html=True)
                    
                    with summary_col2:
                        st.markdown('<div class="chart-container">', unsafe_allow_html=True)
                        st.markdown('<div class="chart-title">Issue Severity Distribution</div>', unsafe_allow_html=True)
                        if results.get('iterations'):
                            # Aggregate severity data
                            severity_counts = {"low": 0, "medium": 0, "high": 0, "critical": 0}
                            for iteration in results.get('iterations', []):
                                critiques = iteration.get("critiques", [])
                                for critique in critiques:
                                    critique_json = critique.get("critique_json", {})
                                    issues = critique_json.get("issues", [])
                                    for issue in issues:
                                        severity = issue.get("severity", "low").lower()
                                        if severity in severity_counts:
                                            severity_counts[severity] += 1
                            
                            # Create a simple bar chart representation
                            max_count = max(severity_counts.values()) if severity_counts.values() else 1
                            for severity, count in severity_counts.items():
                                if count > 0:
                                    bar_length = int((count / max_count) * 20)
                                    emoji = {"low": "üü¢", "medium": "üü°", "high": "üü†", "critical": "üî¥"}[severity]
                                    st.write(f"{emoji} {severity.capitalize()}: {'‚ñà' * bar_length} ({count})")
                        st.markdown('</div>', unsafe_allow_html=True)
                    
                    # Advanced analytics dashboard
                    analytics = generate_advanced_analytics(results)
                    st.markdown("### üìä Advanced Analytics")
                    
                    # Create a dashboard card layout
                    dashboard_col1, dashboard_col2, dashboard_col3, dashboard_col4 = st.columns(4)
                    with dashboard_col1:
                        st.markdown('<div class="dashboard-card">', unsafe_allow_html=True)
                        st.markdown("### üîí Security")
                        st.metric("Strength", f"{analytics.get('security_strength', 0):.1f}%")
                        st.markdown('</div>', unsafe_allow_html=True)
                    with dashboard_col2:
                        st.markdown('<div class="dashboard-card">', unsafe_allow_html=True)
                        st.markdown("### üìã Compliance")
                        st.metric("Coverage", f"{analytics.get('compliance_coverage', 0):.1f}%")
                        st.markdown('</div>', unsafe_allow_html=True)
                    with dashboard_col3:
                        st.markdown('<div class="dashboard-card">', unsafe_allow_html=True)
                        st.markdown("### üßæ Clarity")
                        st.metric("Score", f"{analytics.get('clarity_score', 0):.1f}%")
                        st.markdown('</div>', unsafe_allow_html=True)
                    with dashboard_col4:
                        st.markdown('<div class="dashboard-card">', unsafe_allow_html=True)
                        st.markdown("### ‚úÖ Completeness")
                        st.metric("Score", f"{analytics.get('completeness_score', 0):.1f}%")
                        st.markdown('</div>', unsafe_allow_html=True)
                    
                    # Efficiency and resolution metrics
                    efficiency_col1, efficiency_col2 = st.columns(2)
                    with efficiency_col1:
                        st.markdown('<div class="dashboard-card">', unsafe_allow_html=True)
                        st.markdown("### ‚ö° Efficiency")
                        st.metric("Score", f"{analytics.get('efficiency_score', 0):.1f}%")
                        st.markdown('</div>', unsafe_allow_html=True)
                    with efficiency_col2:
                        st.markdown('<div class="dashboard-card">', unsafe_allow_html=True)
                        st.markdown("### üéØ Resolution")
                        st.metric("Rate", f"{analytics.get('issue_resolution_rate', 0):.1f}%")
                        st.markdown('</div>', unsafe_allow_html=True)
                    
                    # Detailed metrics tabs
                    metrics_tab1, metrics_tab2, metrics_tab3 = st.tabs(["üìà Confidence Trend", "üèÜ Model Performance", "üßÆ Issue Analysis"])
                    
                    with metrics_tab1:
                        # Confidence trend chart
                        if results.get('iterations'):
                            confidence_history = [iter.get("approval_check", {}).get("approval_rate", 0) 
                                                for iter in results.get('iterations', [])]
                            if confidence_history:
                                # Enhanced visualization
                                max_confidence = max(confidence_history)
                                min_confidence = min(confidence_history)
                                avg_confidence = sum(confidence_history) / len(confidence_history)
                                
                                st.line_chart(confidence_history)
                                col1, col2, col3, col4 = st.columns(4)
                                col1.metric("üìà Peak Confidence", f"{max_confidence:.1f}%")
                                col2.metric("üìâ Lowest Confidence", f"{min_confidence:.1f}%")
                                col3.metric("üìä Average Confidence", f"{avg_confidence:.1f}%")
                                col4.metric("üìä Final Confidence", f"{confidence_history[-1]:.1f}%")
                                
                                # Confidence improvement
                                if len(confidence_history) > 1:
                                    improvement = confidence_history[-1] - confidence_history[0]
                                    if improvement > 0:
                                        st.success(f"üöÄ Confidence improved by {improvement:.1f}%")
                                    elif improvement < 0:
                                        st.warning(f"‚ö†Ô∏è Confidence decreased by {abs(improvement):.1f}%")
                                    else:
                                        st.info("‚û°Ô∏è Confidence remained stable")
                    
                    with metrics_tab2:
                        # Model performance analysis
                        if st.session_state.get("adversarial_model_performance"):
                            model_performance = st.session_state.adversarial_model_performance
                            st.markdown("### üèÜ Top Performing Models")
                            
                            # Sort models by score
                            sorted_models = sorted(model_performance.items(), key=lambda x: x[1].get("score", 0), reverse=True)
                            
                            # Display top 5 models with enhanced visualization
                            for i, (model_id, perf) in enumerate(sorted_models[:5]):
                                score = perf.get("score", 0)
                                issues = perf.get("issues_found", 0)
                                st.progress(min(score / 100, 1.0), 
                                           text=f"#{i+1} {model_id} - Score: {score}, Issues Found: {issues}")
                        else:
                            st.info("No model performance data available.")
                    
                    with metrics_tab3:
                        # Issue analysis
                        if results.get('iterations'):
                            # Aggregate issue data
                            total_issues = 0
                            severity_counts = {"low": 0, "medium": 0, "high": 0, "critical": 0}
                            category_counts = {}
                            
                            for iteration in results.get('iterations', []):
                                critiques = iteration.get("critiques", [])
                                for critique in critiques:
                                    critique_json = critique.get("critique_json", {})
                                    issues = critique_json.get("issues", [])
                                    total_issues += len(issues)
                                    
                                    for issue in issues:
                                        # Count by severity
                                        severity = issue.get("severity", "low").lower()
                                        if severity in severity_counts:
                                            severity_counts[severity] += 1
                                        
                                        # Count by category
                                        category = issue.get("category", "uncategorized")
                                        category_counts[category] = category_counts.get(category, 0) + 1
                            
                            col1, col2 = st.columns(2)
                            with col1:
                                st.markdown("### üéØ Issue Severity Distribution")
                                for severity, count in severity_counts.items():
                                    if count > 0:
                                        emoji = {"low": "üü¢", "medium": "üü°", "high": "üü†", "critical": "üî¥"}[severity]
                                        st.write(f"{emoji} {severity.capitalize()}: {count}")
                            with col2:
                                st.markdown("### üìö Issue Categories")
                                # Show top 5 categories
                                sorted_categories = sorted(category_counts.items(), key=lambda x: x[1], reverse=True)
                                for category, count in sorted_categories[:5]:
                                    st.write(f"üè∑Ô∏è {category}: {count}")
                            
                            st.metric("üîç Total Issues Found", total_issues)
                    
                    # Protocol comparison and analysis
                    st.markdown("### üìÑ Protocol Analysis")
                    final_sop = results.get('final_sop', '')
                    original_sop = st.session_state.protocol_text
                    
                    if final_sop and original_sop:
                        # Tabs for different views
                        protocol_tab1, protocol_tab2, protocol_tab3, protocol_tab4 = st.tabs(["üîÑ Comparison", "üìÑ Final Protocol", "üîç Structure Analysis", "ü§ñ AI Insights"])
                        
                        with protocol_tab1:
                            st.markdown("### üîÑ Protocol Evolution")
                            # Simple comparison metrics
                            original_complexity = calculate_protocol_complexity(original_sop)
                            final_complexity = calculate_protocol_complexity(final_sop)
                            
                            col1, col2, col3 = st.columns(3)
                            col1.metric("üìè Original Length", f"{original_complexity['word_count']} words")
                            col2.metric("üìè Final Length", f"{final_complexity['word_count']} words")
                            col3.metric("üìä Length Change", 
                                       f"{final_complexity['word_count'] - original_complexity['word_count']} words", 
                                       f"{((final_complexity['word_count'] / max(1, original_complexity['word_count'])) - 1) * 100:.1f}%")
                            
                            # Show both protocols side by side
                            col1, col2 = st.columns(2)
                            with col1:
                                st.markdown("**Original Protocol**")
                                st.text_area("Original", value=original_sop, height=300, key="original_protocol_display")
                            with col2:
                                st.markdown("**Hardened Protocol**")
                                st.text_area("Final", value=final_sop, height=300, key="final_protocol_display")
                            
                            # Add detailed comparison
                            st.markdown("### üìä Detailed Comparison")
                            st.markdown(render_protocol_comparison(original_sop, final_sop, "Original Protocol", "Hardened Protocol"), unsafe_allow_html=True)
                        
                        with protocol_tab2:
                            st.markdown("### üìÑ Final Hardened Protocol")
                            st.code(final_sop, language="markdown")
                            # Add download button
                            st.download_button(
                                label="üì• Download Final Protocol",
                                data=final_sop,
                                file_name="hardened_protocol.md",
                                mime="text/markdown"
                            )
                        
                        with protocol_tab3:
                            st.markdown("### üîç Protocol Structure Analysis")
                            # Add protocol analysis
                            complexity = calculate_protocol_complexity(final_sop)
                            structure = extract_protocol_structure(final_sop)
                            
                            col1, col2 = st.columns(2)
                            with col1:
                                st.markdown("**üìè Complexity Metrics**")
                                st.metric("Words", complexity["word_count"])
                                st.metric("Sentences", complexity["sentence_count"])
                                st.metric("Paragraphs", complexity["paragraph_count"])
                                st.metric("Complexity Score", complexity["complexity_score"])
                                st.metric("Unique Words", complexity["unique_words"])
                            
                            with col2:
                                st.markdown("**üß© Structure Analysis**")
                                st.write("Numbered Steps:", "‚úÖ" if structure["has_numbered_steps"] else "‚ùå")
                                st.write("Bullet Points:", "‚úÖ" if structure["has_bullet_points"] else "‚ùå")
                                st.write("Headers:", "‚úÖ" if structure["has_headers"] else "‚ùå")
                                st.write("Preconditions:", "‚úÖ" if structure["has_preconditions"] else "‚ùå")
                                st.write("Postconditions:", "‚úÖ" if structure["has_postconditions"] else "‚ùå")
                                st.write("Error Handling:", "‚úÖ" if structure["has_error_handling"] else "‚ùå")
                                st.metric("SectionsIn Protocol", structure["section_count"])
                        
                        with protocol_tab4:
                            st.markdown("### ü§ñ AI Insights Dashboard")
                            st.markdown(render_ai_insights_dashboard(final_sop), unsafe_allow_html=True)

with tab2:
    render_adversarial_testing_tab()

# Add start/stop functionality for adversarial testing
if start_button:
    # Validate inputs before starting
    errors = []
    
    if not st.session_state.protocol_text.strip():
        errors.append("üìÑ Please enter a protocol before starting adversarial testing.")
    
    if not st.session_state.openrouter_key:
        errors.append("üîë Please enter your OpenRouter API key.")
    
    if not st.session_state.red_team_models:
        errors.append("üî¥ Please select at least one red team model.")
    
    if not st.session_state.blue_team_models:
        errors.append("üîµ Please select at least one blue team model.")
    
    if st.session_state.adversarial_min_iter > st.session_state.adversarial_max_iter:
        errors.append("üîÑ Min iterations cannot be greater than max iterations.")
    
    if st.session_state.adversarial_confidence < 50 or st.session_state.adversarial_confidence > 100:
        errors.append("üéØ Confidence threshold should be between 50% and 100%.")
    
    # Enhanced validation for model selection
    if len(st.session_state.red_team_models) > 10:
        errors.append("üî¥ Please select 10 or fewer red team models to avoid excessive costs.")
    
    if len(st.session_state.blue_team_models) > 10:
        errors.append("üîµ Please select 10 or fewer blue team models to avoid excessive costs.")
    
    # Show all errors at once
    if errors:
        for error in errors:
            st.error(error)
        st.info("üí° Tip: Check the Quick Start Wizard above for guided setup.")
    else:
        # Thread safety check to prevent multiple concurrent adversarial testing threads
        if st.session_state.adversarial_running:
            st.warning("Adversarial testing is already running. Please wait for it to complete or stop it first.")
        else:
            # Enhanced confirmation dialog with detailed cost estimation
            protocol_length = len(st.session_state.protocol_text.split())
            estimate = estimate_testing_time_and_cost(
                st.session_state.red_team_models,
                st.session_state.blue_team_models,
                st.session_state.adversarial_max_iter,
                protocol_length
            )
            
            estimated_cost = estimate["estimated_cost_usd"]
            estimated_time = estimate["estimated_time_minutes"]
            
            if estimated_cost > 0.1:  # If estimated cost is over $0.10
                st.warning(f"üí∞ Estimated cost: ${estimated_cost:.4f} | ‚è∞ Estimated time: {estimated_time} minutes | üîÑ Operations: {estimate['total_operations']:,}")
                st.info("üí° Tip: Consider reducing the number of models or iterations to control costs.")
                if not st.checkbox("‚úÖ I understand the cost and time estimate and want to proceed", key="cost_confirmation"):
                    st.info("‚ÑπÔ∏è Please confirm you understand the cost and time estimate to proceed.")
                    st.stop()
            
            # Start adversarial testing with enhanced initialization
            st.session_state.adversarial_running = True
            st.session_state.adversarial_status_message = "üöÄ Initializing adversarial testing..."
            
            # Show initialization message
            with st.spinner("Starting adversarial testing process..."):
                # Initialize progress tracking
                st.session_state.adversarial_confidence_history = []
                st.session_state.adversarial_total_tokens_prompt = 0
                st.session_state.adversarial_total_tokens_completion = 0
                st.session_state.adversarial_cost_estimate_usd = 0.0
                
                # Start the testing process in a separate thread
                threading.Thread(target=run_adversarial_testing, daemon=True).start()
            st.rerun()

if stop_button:
    if st.session_state.adversarial_running:
        st.session_state.adversarial_stop_flag = True
        st.warning("‚èπÔ∏è Stop signal sent. Adversarial testing will stop after the current iteration.")
        st.session_state.adversarial_status_message = "‚èπÔ∏è Stopping adversarial testing..."
        # No rerun here, let the loop handle the UI update
    else:
        st.info("‚ÑπÔ∏è Adversarial testing is not currently running.")

# ------------------------------------------------------------------
# 8. Run logic for Evolution tab (Self-Contained Implementation)
# ------------------------------------------------------------------

def _request_openai_compatible_chat(
    api_key: str, base_url: str, model: str, messages: List, extra_headers: Dict,
    temperature: float, top_p: float, frequency_penalty: float, presence_penalty: float,
    max_tokens: int, seed: Optional[int], req_timeout: int = 60, max_retries: int = 5,
    provider: str = "OpenAI"
) -> str:
    url = base_url.rstrip('/') + "/chat/completions"
    headers = {"Content-Type": "application/json", **extra_headers}
    if api_key: headers["Authorization"] = f"Bearer {api_key}"

    payload: Dict[str, Any] = {
        "model": model, "messages": messages, "temperature": temperature, "max_tokens": max_tokens,
        "top_p": top_p, "frequency_penalty": frequency_penalty, "presence_penalty": presence_penalty,
    }
    if PROVIDERS.get(provider, {}).get("omit_model_in_payload"):
        payload.pop("model", None)
    if seed is not None:
        payload["seed"] = int(seed)

    last_err = None
    for attempt in range(max_retries):
        try:
            r = requests.post(url, headers=headers, json=payload, timeout=req_timeout)
            
            # Handle rate limiting and server errors with retry
            if r.status_code in {429, 500, 502, 503, 504}:
                sleep_s = (2 ** attempt) + _rand_jitter_ms()
                time.sleep(sleep_s)
                last_err = Exception(f"HTTP {r.status_code}: {r.text}")
                continue
                
            r.raise_for_status()
            
            # Handle non-JSON responses
            try:
                data = r.json()
            except json.JSONDecodeError as e:
                last_err = Exception(f"Invalid JSON response: {e} - Response: {r.text[:200]}...")
                time.sleep((2 ** attempt) + _rand_jitter_ms())
                continue
                
            # Safely access the response structure
            if not isinstance(data, dict):
                last_err = Exception(f"Unexpected response format: {type(data)} - Response: {data}")
                time.sleep((2 ** attempt) + _rand_jitter_ms())
                continue
                
            # Handle API-specific error formats
            if "error" in data:
                error_msg = data["error"]
                if isinstance(error_msg, dict):
                    error_msg = error_msg.get("message", str(error_msg))
                last_err = Exception(f"API error: {error_msg}")
                time.sleep((2 ** attempt) + _rand_jitter_ms())
                continue
                
            choices = data.get("choices", [])
            if choices:
                choice = choices[0]
                content = choice.get("message", {}).get("content", "")
                if content is not None:
                    return content
                else:
                    last_err = Exception("Empty content in response choice")
            else:
                last_err = Exception("No choices in response")
                
            # If we get here, there was an issue with the response structure
            time.sleep((2 ** attempt) + _rand_jitter_ms())
                
        except requests.exceptions.ConnectionError as e:
            last_err = Exception(f"Connection error: {e}")
            time.sleep((2 ** attempt) + _rand_jitter_ms())
        except requests.exceptions.Timeout as e:
            last_err = Exception(f"Request timeout: {e}")
            time.sleep((2 ** attempt) + _rand_jitter_ms())
        except requests.exceptions.RequestException as e:
            last_err = Exception(f"Request failed: {e}")
            time.sleep((2 ** attempt) + _rand_jitter_ms())
        except Exception as e:
            last_err = e
            time.sleep((2 ** attempt) + _rand_jitter_ms())
            
    raise RuntimeError(f"Request failed after {max_retries} attempts for model {model}: {last_err}")

def run_evolution_internal():
    try:
        with st.session_state.thread_lock:
            st.session_state.evolution_log = []
            st.session_state.evolution_stop_flag = False
        current_protocol = st.session_state.protocol_text
        st.session_state.evolution_current_best = current_protocol

        def log_msg(msg):
            with st.session_state.thread_lock:
                st.session_state.evolution_log.append(f"[{time.strftime('%H:%M:%S')}] {msg}")

        log_msg(f"üöÄ Starting evolution process with {st.session_state.provider}/{st.session_state.model}...")
        try:
            extra_hdrs = json.loads(st.session_state.extra_headers or "{}")
            if not isinstance(extra_hdrs, dict): raise json.JSONDecodeError("JSON is not a dictionary.", "", 0)
        except (ValueError, TypeError):
            log_msg("‚ö†Ô∏è Invalid Extra Headers JSON. Must be a dictionary. Using empty dict.")
            extra_hdrs = {}

        seed_str = str(st.session_state.seed or "").strip()
        seed = None
        if seed_str:
            try:
                seed = int(float(seed_str))  # Handle floats by truncating to int
            except (ValueError, TypeError):
                pass  # Invalid input, keep seed as None

        api_key_to_use = st.session_state.api_key
        provider_info = PROVIDERS.get(st.session_state.provider, {})
        if not api_key_to_use and (env_var := provider_info.get("env")):
            api_key_to_use = os.environ.get(env_var, "")
            if api_key_to_use: log_msg(f"Using API key from env var {env_var}.")

        if not api_key_to_use and provider_info.get("env"):
             log_msg(f"‚ö†Ô∏è No API key provided in UI or in env var {provider_info.get('env')}. Requests may fail.")

        request_functions = {
            "Anthropic": _request_anthropic_chat,
            "Google (Gemini)": _request_google_gemini_chat,
            "Cohere": _request_cohere_chat,
        }
        request_function = request_functions.get(st.session_state.provider, _request_openai_compatible_chat)

        consecutive_failures = 0
        max_consecutive_failures = 3  # Stop after 3 consecutive failures
        
        # Initialize metrics tracking
        iteration_metrics = []
        
        for i in range(st.session_state.max_iterations):
            if st.session_state.evolution_stop_flag:
                log_msg("‚èπÔ∏è Evolution stopped by user.")
                break
            log_msg(f"üîÑ --- Iteration {i+1}/{st.session_state.max_iterations} ---")
            try:
                messages = _compose_messages(
                    st.session_state.system_prompt,
                    f"Current draft:\n\n---\n{current_protocol}\n---\n\nImprove it based on your instructions."
                )
                
                kwargs = {
                    "api_key": api_key_to_use, 
                    "base_url": st.session_state.base_url, 
                    "model": st.session_state.model,
                    "messages": messages, 
                    "extra_headers": extra_hdrs, 
                    "temperature": st.session_state.temperature,
                    "top_p": st.session_state.top_p, 
                    "max_tokens": st.session_state.max_tokens, 
                    "seed": seed,
                }
                if st.session_state.provider not in request_functions:
                    kwargs["frequency_penalty"] = st.session_state.frequency_penalty
                    kwargs["presence_penalty"] = st.session_state.presence_penalty
                    kwargs["provider"] = st.session_state.provider

                improved_protocol = request_function(**kwargs)

                # Enhanced validation with quality metrics
                if improved_protocol and len(improved_protocol.strip()) > len(current_protocol) * 0.7:
                    # Calculate quality metrics
                    original_complexity = calculate_protocol_complexity(current_protocol)
                    improved_complexity = calculate_protocol_complexity(improved_protocol)
                    
                    # Quality improvement check
                    quality_improvement = (
                        improved_complexity["complexity_score"] >= original_complexity["complexity_score"] and
                        improved_complexity["unique_words"] >= original_complexity["unique_words"] * 0.9
                    )
                    
                    if quality_improvement or len(improved_protocol.strip()) > len(current_protocol) * 1.1:
                        log_msg(f"‚úÖ Iteration {i+1} successful. Length: {len(improved_protocol.strip())}")
                        current_protocol = improved_protocol.strip()
                        st.session_state.evolution_current_best = current_protocol
                        consecutive_failures = 0  # Reset failure counter on success
                        
                        # Track metrics
                        iteration_metrics.append({
                            "iteration": i+1,
                            "length": len(current_protocol),
                            "complexity": improved_complexity["complexity_score"],
                            "unique_words": improved_complexity["unique_words"],
                            "improvement": len(current_protocol) - len(st.session_state.protocol_text)
                        })
                    else:
                        log_msg(f"‚ö†Ô∏è Iteration {i+1} result rejected (no significant quality improvement). Length: {len(improved_protocol.strip() if improved_protocol else '')}")
                        consecutive_failures += 1
                else:
                    log_msg(f"‚ö†Ô∏è Iteration {i+1} result rejected (too short or empty). Length: {len(improved_protocol.strip() if improved_protocol else '')}")
                    consecutive_failures += 1
                if (i + 1) % st.session_state.checkpoint_interval == 0:
                    log_msg(f"üíæ Checkpoint at iteration {i+1}.")
            except Exception as e:
                log_msg(f"‚ùå ERROR in iteration {i+1}: {e}")
                consecutive_failures += 1
                time.sleep(2)
            
            # Stop if too many consecutive failures
            if consecutive_failures >= max_consecutive_failures:
                log_msg(f"üõë Stopping evolution due to {consecutive_failures} consecutive failures.")
                break

        log_msg("üèÅ Evolution finished.")
        
        # Log final metrics
        if iteration_metrics:
            final_length = iteration_metrics[-1]["length"]
            initial_length = len(st.session_state.protocol_text)
            total_improvement = final_length - initial_length
            log_msg(f"üìä Total improvement: {total_improvement} characters ({(total_improvement/initial_length)*100:.1f}%)")
        
        st.session_state.protocol_text = current_protocol  # Update the main protocol text
    except Exception as e:
        import traceback
        tb_str = traceback.format_exc()
        log_msg(f"üí• A critical error occurred in the evolution thread: {e}\n{tb_str}")
    finally:
        st.session_state.evolution_running = False

if stop_button:
    st.session_state.evolution_stop_flag = True
    st.warning("Stop signal sent. Evolution will stop after the current iteration.")
    # No rerun here, let the loop handle the UI update

if run_button:
    if st.session_state.protocol_text.strip():
        # Thread safety check to prevent multiple concurrent evolution threads
        if st.session_state.evolution_running:
            st.warning("Evolution is already running. Please wait for it to complete or stop it first.")
        else:
            # Validation
            errors = []
            if not st.session_state.api_key and not os.environ.get(PROVIDERS.get(st.session_state.provider, {}).get("env", ""), ""):
                errors.append("üîë Please enter an API key or set the appropriate environment variable.")
            
            if errors:
                for error in errors:
                    st.error(error)
            else:
                st.session_state.evolution_running = True
                with st.spinner("Starting evolution process..."):
                    threading.Thread(target=run_evolution_internal, daemon=True).start()
                st.rerun()
    else:
        st.warning("Please paste a protocol before starting evolution.")
