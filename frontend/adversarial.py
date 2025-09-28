# This file implements the adversarial generation functionality of the OpenEvolve frontend.
# The purpose of this module is to facilitate AI-driven testing and refinement of ideas,
# code, and other content. It operates on the principle of "AI peer review," where
# different AI agents are assigned to "red team" (critique) and "blue team" (improve)
# roles. An "evaluator" AI then assesses the quality of the improvements.
#
# This process is designed for constructive, iterative improvement and is NOT intended
# for generating malicious prompts, code, or other harmful content. The goal is to
# identify weaknesses and enhance the quality of the content in a controlled and
# ethical manner.

import streamlit as st
import requests
import json
import time
import threading
import traceback
import uuid
import random
import tempfile
import os
from typing import List, Dict, Any, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed

import docx
from fpdf import FPDF

# Import OpenEvolve modules for backend integration

try:
    from openevolve.api import run_evolution as openevolve_run_evolution
    from openevolve.config import Config, LLMModelConfig

    OPENEVOLVE_AVAILABLE = True
except ImportError:
    OPENEVOLVE_AVAILABLE = False
    print("OpenEvolve backend not available - using API-based adversarial testing only")

from openevolve_integration import (
    create_language_specific_evaluator,
    create_specialized_evaluator,
)

from session_utils import (
    _clamp,
    _rand_jitter_ms,
    _approx_tokens,
    _cost_estimate,
    safe_int,
    safe_float,
    _safe_list,
    _extract_json_block,
    _compose_messages,
    APPROVAL_PROMPT,
    RED_TEAM_CRITIQUE_PROMPT,
    BLUE_TEAM_PATCH_PROMPT,
    CODE_REVIEW_RED_TEAM_PROMPT,
    CODE_REVIEW_BLUE_TEAM_PROMPT,
    PLAN_REVIEW_RED_TEAM_PROMPT,
    PLAN_REVIEW_BLUE_TEAM_PROMPT,
    _hash_text,
)

MODEL_META_BY_ID: Dict[str, Dict[str, Any]] = {}
MODEL_META_LOCK = threading.Lock()


@st.cache_data(ttl=600)
def get_openrouter_models(api_key: str) -> List[Dict]:
    """Fetch available models from OpenRouter (cached)."""
    print(f"Using API Key: {api_key}")
    if not api_key:
        return []
    try:
        response = requests.get(
            "https://openrouter.ai/api/v1/models",
            headers={"Authorization": f"Bearer {api_key}"},
            timeout=15,
        )
        response.raise_for_status()
        data = response.json()
        models = data.get("data", []) if isinstance(data, dict) else []
        return models
    except Exception as e:
        st.warning(f"Could not fetch OpenRouter models: {e}")
        return []


def _request_openrouter_chat(
    api_key: str,
    model_id: str,
    messages: List[Dict[str, str]],
    temperature: float,
    top_p: float,
    frequency_penalty: float,
    presence_penalty: float,
    max_tokens: int,
    force_json: bool = False,
    seed: Optional[int] = None,
    req_timeout: int = 60,
    max_retries: int = 5,
) -> Tuple[str, int, int, float]:
    """
    Robust OpenRouter chat call with exponential backoff, jitter, and cost/token estimation.
    Returns: (content, prompt_tokens, completion_tokens, cost)
    """
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://github.com/google/gemini-pro-builder",  # Recommended by OpenRouter
        "X-Title": "OpenEvolve Protocol Improver",  # Recommended by OpenRouter
    }
    payload: Dict[str, Any] = {
        "model": model_id,
        "messages": messages,
        "temperature": _clamp(temperature, 0.0, 2.0),
        "max_tokens": max_tokens,
        "top_p": _clamp(top_p, 0.0, 1.0),
        "frequency_penalty": _clamp(frequency_penalty, -2.0, 2.0),
        "presence_penalty": _clamp(presence_penalty, -2.0, 2.0),
    }
    if seed is not None:
        payload["seed"] = int(seed)
    if force_json:
        payload["response_format"] = {"type": "json_object"}

    last_err = None
    for attempt in range(max_retries):
        try:
            r = requests.post(url, headers=headers, json=payload, timeout=req_timeout)
            if r.status_code == 400:
                # HTTP 400 Bad Request - client error, don't retry
                last_err = Exception(f"HTTP 400 Bad Request: {r.text[:200]}...")
                break  # Break out of retry loop for client errors
            if r.status_code in {429, 500, 502, 503, 504}:
                sleep_s = (2**attempt) + _rand_jitter_ms()
                time.sleep(sleep_s)
                last_err = Exception(f"Transient error {r.status_code}: Retrying...")
                continue
            r.raise_for_status()
            data = r.json()

            # Safely access the first choice to prevent IndexError if "choices" is an empty list.
            choices = data.get("choices", [])
            if choices:
                choice = choices[0]
                content = choice.get("message", {}).get("content", "")
            else:
                content = ""

            usage = data.get("usage", {})
            p_tok = safe_int(
                usage.get("prompt_tokens"), _approx_tokens(json.dumps(messages))
            )
            c_tok = safe_int(
                usage.get("completion_tokens"), _approx_tokens(content or "")
            )
            cost = _cost_estimate(
                p_tok, c_tok, None, None
            )  # Simplified cost calculation
            return content or "", p_tok, c_tok, cost
        except Exception as e:
            last_err = e
            sleep_s = (2**attempt) + _rand_jitter_ms()
            time.sleep(sleep_s)
    raise RuntimeError(
        f"Request failed for {model_id} after {max_retries} attempts: {last_err}"
    )


def _run_adversarial_testing_with_openevolve_backend(
    current_content: str,
    content_type: str,
    red_team_models: List[str],
    blue_team_models: List[str],
    api_key: str,
    base_url: str,
    max_iterations: int,
    confidence_threshold: float,
    max_tokens: int,
    temperature: float,
    top_p: float,
    frequency_penalty: float,
    presence_penalty: float,
    seed: Optional[int],
    max_workers: int,
    rotation_strategy: str,
    red_team_sample_size: int,
    blue_team_sample_size: int,
    custom_requirements: str = "",
    evaluator_system_prompt: str = APPROVAL_PROMPT,
    red_team_prompt: str = RED_TEAM_CRITIQUE_PROMPT,
    blue_team_prompt: str = BLUE_TEAM_PATCH_PROMPT,
    compliance_rules: Optional[List[str]] = None,
    red_team_prompt_enhancement: str = "",
    blue_team_prompt_enhancement: str = "",
    feature_dimensions: Optional[List[str]] = None,
    feature_bins: Optional[int] = None,
    enable_data_augmentation: bool = False,
    augmentation_model_id: str = None,
    augmentation_temperature: float = 0.7,
    enable_human_feedback: bool = False,
    current_iteration: int = 0,
) -> Dict[str, Any]:
    """
    Run adversarial testing using OpenEvolve backend for code content.

    Args:
        current_content: The content to test adversarially
        content_type: Type of content being tested
        red_team_models: List of red team models
        blue_team_models: List of blue team models
        api_key: API key for the LLM provider
        base_url: Base URL for the API
        max_iterations: Maximum number of iterations
        confidence_threshold: Confidence threshold for stopping
        max_tokens: Maximum tokens to generate
        temperature: Temperature for generation
        top_p: Top-p sampling parameter
        frequency_penalty: Frequency penalty
        presence_penalty: Presence penalty
        seed: Random seed
        max_workers: Maximum number of parallel workers
        rotation_strategy: Model rotation strategy
        red_team_sample_size: Number of red team models to sample
        blue_team_sample_size: Number of blue team models to sample
        custom_requirements: Custom requirements for testing
        evaluator_system_prompt: System prompt for evaluation
        red_team_prompt: Prompt for red team
        blue_team_prompt: Prompt for blue team

    Returns:
        Dict[str, Any]: Adversarial testing results
    """
    if not OPENEVOLVE_AVAILABLE:
        st.error("OpenEvolve backend not available for adversarial testing")
        return {"success": False, "error": "OpenEvolve backend not available"}

    try:
        # Create OpenEvolve configuration
        config = Config()

        # Configure LLM models
        llm_models = []

        # Add red team models
        for model_id in red_team_models[:red_team_sample_size]:
            llm_config = LLMModelConfig(
                name=model_id,
                api_key=api_key,
                api_base=base_url if base_url else "https://api.openai.com/v1",
                temperature=temperature,
                top_p=top_p,
                frequency_penalty=frequency_penalty,
                presence_penalty=presence_penalty,
                max_tokens=max_tokens,
                seed=seed,
            )
            llm_models.append(llm_config)

        # Add blue team models
        for model_id in blue_team_models[:blue_team_sample_size]:
            llm_config = LLMModelConfig(
                name=model_id,
                api_key=api_key,
                api_base=base_url if base_url else "https://api.openai.com/v1",
                temperature=temperature,
                top_p=top_p,
                frequency_penalty=frequency_penalty,
                presence_penalty=presence_penalty,
                max_tokens=max_tokens,
                seed=seed,
            )
            llm_models.append(llm_config)

        config.llm.models = llm_models
        config.evolution.max_iterations = max_iterations
        config.evolution.population_size = max_workers
        config.evolution.num_islands = st.session_state.num_islands
        config.evolution.elite_ratio = st.session_state.elite_ratio
        config.evolution.exploration_ratio = st.session_state.exploration_ratio
        config.evolution.exploitation_ratio = st.session_state.exploitation_ratio
        config.evolution.archive_size = st.session_state.archive_size
        config.evolution.checkpoint_interval = st.session_state.checkpoint_interval

        # Configure database settings for multi-objective evolution
        if feature_dimensions is not None:
            config.database.feature_dimensions = feature_dimensions
        if feature_bins is not None:
            config.database.feature_bins = feature_bins

        # Perform data augmentation if enabled
        if enable_data_augmentation and augmentation_model_id:
            _update_adv_log_and_status(
                f"🧪 Augmenting content using {augmentation_model_id}..."
            )
            current_content = generate_adversarial_data_augmentation(
                content=current_content,
                content_type=content_type,
                api_key=api_key,
                model_id=augmentation_model_id,
                temperature=augmentation_temperature,
                max_tokens=max_tokens,  # Use same max_tokens as main evolution
                seed=seed,
            )
            _update_adv_log_and_status("✅ Content augmentation complete.")

        # Create evaluator function based on content_type
        if content_type.startswith("code_"):
            evaluator_instance = create_specialized_evaluator(
                content_type, custom_requirements, compliance_rules
            )
        else:
            evaluator_instance = create_language_specific_evaluator(
                content_type, custom_requirements, compliance_rules
            )
        # Create temporary file for the content with proper evolution markers
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", delete=False
        ) as temp_file:
            # Add evolution markers to the content
            content_with_markers = f"""# EVOLVE-BLOCK-START
{current_content}
# EVOLVE-BLOCK-END"""
            temp_file.write(content_with_markers)
            temp_file_path = temp_file.name

        try:
            # Run adversarial testing using OpenEvolve API
            result = openevolve_run_evolution(
                initial_program=temp_file_path,
                evaluator=evaluator_instance.evaluate,
                config=config,
                iterations=max_iterations,
                output_dir=None,  # Use temporary directory
                cleanup=True,
            )

            # Process results
            if result.best_program and result.best_code:
                # Remove evolution markers from the final result
                best_code = result.best_code
                if "# EVOLVE-BLOCK-START" in best_code:
                    start_idx = best_code.find("# EVOLVE-BLOCK-START") + len(
                        "# EVOLVE-BLOCK-START"
                    )
                    end_idx = best_code.find("# EVOLVE-BLOCK-END")
                    if end_idx != -1:
                        best_code = best_code[start_idx:end_idx].strip()

                # Simulate human feedback capture
                if enable_human_feedback:
                    # For now, a dummy score and comments
                    dummy_score = random.uniform(0.5, 1.0)  # Simulate a score
                    dummy_comments = "Human reviewed and provided general feedback on clarity and relevance."
                    capture_human_feedback(
                        adversarial_example={
                            "id": str(uuid.uuid4()),
                            "content": best_code,
                            "content_type": content_type,
                            "iteration": current_iteration,  # Now using the passed iteration
                        },
                        human_score=dummy_score,
                        human_comments=dummy_comments,
                    )

                return {
                    "success": True,
                    "best_program": result.best_program,
                    "best_score": result.best_score,
                    "best_code": best_code,
                    "metrics": result.metrics,
                    "output_dir": result.output_dir,
                }
            else:
                return {
                    "success": False,
                    "message": "Adversarial testing completed with no improvement.",
                }

        finally:
            # Clean up the temporary file
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)

    except Exception as e:
        st.error(f"Error running adversarial testing with OpenEvolve backend: {e}")
        print(f"Adversarial testing error: {e}")
        import traceback

        traceback.print_exc()
        return {"success": False, "error": str(e)}


def analyze_with_model(
    api_key: str,
    model_id: str,
    sop: str,
    config: Dict,
    system_prompt: str,
    user_suffix: str = "",
    force_json: bool = False,
    seed: Optional[int] = None,
    compliance_requirements: str = "",
    prompt_enhancement: str = "",
) -> Dict[str, Any]:
    """
    Analyzes an SOP with a specific model, handling context limits and returning structured results.
    """
    try:
        if compliance_requirements:
            system_prompt = system_prompt.format(
                compliance_requirements=compliance_requirements
            )
            print(f"Using compliance requirements: {compliance_requirements}")
        if prompt_enhancement:
            system_prompt += f"\n{prompt_enhancement}"
            print(f"Using prompt enhancement: {prompt_enhancement}")
        if st.session_state.get("adversarial_critique_depth"):
            system_prompt += (
                f"\nCritique Depth: {st.session_state.adversarial_critique_depth}"
            )
        if st.session_state.get("adversarial_patch_quality"):
            system_prompt += (
                f"\nPatch Quality: {st.session_state.adversarial_patch_quality}"
            )
        max_tokens = safe_int(config.get("max_tokens"), 8000)
        user_prompt = f"Here is the Standard Operating Procedure (SOP):\n\n---\n\n{sop}\n\n---\n\n{user_suffix}"
        full_prompt_text = system_prompt + user_prompt

        # Simplified context length estimation
        context_len = 8192
        prompt_toks_est = _approx_tokens(full_prompt_text)

        if prompt_toks_est + max_tokens >= context_len:
            err_msg = (
                f"ERROR[{model_id}]: Estimated prompt tokens ({prompt_toks_est}) + max_tokens ({max_tokens}) "
                f"exceeds context window ({context_len}). Skipping."
            )
            return {
                "ok": False,
                "text": err_msg,
                "json": None,
                "ptoks": 0,
                "ctoks": 0,
                "cost": 0.0,
                "model_id": model_id,
            }

        content, p_tok, c_tok, cost = _request_openrouter_chat(
            api_key=api_key,
            model_id=model_id,
            messages=_compose_messages(system_prompt, user_prompt),
            temperature=safe_float(config.get("temperature"), 0.7),
            top_p=safe_float(config.get("top_p"), 1.0),
            frequency_penalty=safe_float(config.get("frequency_penalty"), 0.0),
            presence_penalty=safe_float(config.get("presence_penalty"), 0.0),
            max_tokens=max_tokens,
            force_json=force_json,
            seed=seed,
        )
        json_content = _extract_json_block(content)
        return {
            "ok": True,
            "text": content,
            "json": json_content,
            "ptoks": p_tok,
            "ctoks": c_tok,
            "cost": cost,
            "model_id": model_id,
        }
    except Exception as e:
        return {
            "ok": False,
            "text": f"ERROR[{model_id}]: {e}",
            "json": None,
            "ptoks": 0,
            "ctoks": 0,
            "cost": 0.0,
            "model_id": model_id,
        }


def determine_review_type(content: str) -> str:
    """Determine the appropriate review type based on content analysis.

    Args:
        content (str): The content to analyze

    Returns:
        str: Review type ('general', 'code', 'plan')
    """
    if not content:
        return "general"

    # Convert to lowercase for analysis
    lower_content = content.lower()

    # Enhanced language detection for code review
    # Programming language indicators
    python_indicators = [
        "def ",
        "import ",
        "class ",
        "if __name__ ==",
        "print(",
        "for ",
        "while ",
    ]
    js_indicators = [
        "function ",
        "const ",
        "let ",
        "var ",
        "import ",
        "require(",
        "console.",
        "=>",
    ]
    java_indicators = [
        "public class ",
        "private ",
        "protected ",
        "static ",
        "void ",
        "int ",
        "String ",
        "new ",
    ]
    cpp_indicators = [
        "#include",
        "using namespace",
        "std::",
        "cout <<",
        "cin >>",
        "int main",
    ]
    csharp_indicators = [
        "using System",
        "namespace ",
        "public class",
        "static void Main",
        "Console.WriteLine",
    ]
    go_indicators = ["package ", "import ", "func ", "fmt.", "var "]
    rust_indicators = ["fn ", "let ", "mut ", "use ", "println!", "struct "]

    # Count matches for each language
    lang_counts = {
        "python": sum(
            1 for indicator in python_indicators if indicator in lower_content
        ),
        "javascript": sum(
            1 for indicator in js_indicators if indicator in lower_content
        ),
        "java": sum(1 for indicator in java_indicators if indicator in lower_content),
        "cpp": sum(1 for indicator in cpp_indicators if indicator in lower_content),
        "csharp": sum(
            1 for indicator in csharp_indicators if indicator in lower_content
        ),
        "go": sum(1 for indicator in go_indicators if indicator in lower_content),
        "rust": sum(1 for indicator in rust_indicators if indicator in lower_content),
    }

    # Find the dominant programming language
    dominant_lang = max(lang_counts, key=lang_counts.get)
    dominant_count = lang_counts[dominant_lang]

    # If we have significant code indicators, use code review
    if dominant_count > 2:
        return "code"

    # Check for plan indicators
    plan_indicators = [
        "objective",
        "goal",
        "milestone",
        "deliverable",
        "resource",
        "budget",
        "timeline",
        "schedule",
        "risk",
        "dependency",
        "assumption",
        "stakeholder",
        "communication",
        "review",
        "approval",
    ]

    # Count matches
    code_matches = sum(lang_counts.values())
    plan_matches = sum(1 for indicator in plan_indicators if indicator in lower_content)

    # Determine review type
    if code_matches > plan_matches and code_matches > 2:
        return "code"
    elif plan_matches > code_matches and plan_matches > 2:
        return "plan"
    else:
        return "general"


def get_appropriate_prompts(review_type: str) -> Tuple[str, str]:
    """Get the appropriate prompts based on review type.

    Args:
        review_type (str): Type of review ('general', 'code', 'plan')

    Returns:
        Tuple[str, str]: Red team and blue team prompts
    """
    if review_type == "code":
        return CODE_REVIEW_RED_TEAM_PROMPT, CODE_REVIEW_BLUE_TEAM_PROMPT
    elif review_type == "plan":
        return PLAN_REVIEW_RED_TEAM_PROMPT, PLAN_REVIEW_BLUE_TEAM_PROMPT
    else:
        return RED_TEAM_CRITIQUE_PROMPT, BLUE_TEAM_PATCH_PROMPT


def _severity_rank(sev: str) -> int:
    order = {"low": 0, "medium": 1, "high": 2, "critical": 3}
    return order.get(str(sev).lower(), 0)


def _merge_consensus_sop(
    base_sop: str, blue_patches: List[dict], critiques: List[dict]
) -> Tuple[str, dict]:
    """
    Selects the best patch from the blue team based on coverage, resolution, and quality.
    """
    valid_patches = [
        p
        for p in blue_patches
        if p and (p.get("patch_json") or {}).get("sop", "").strip()
    ]
    if not valid_patches:
        return base_sop, {
            "reason": "no_valid_patches_received",
            "score": -1,
            "resolution_by_severity": {},
            "resolution_by_category": {},
        }

    # Create a lookup for issue severity and category
    issue_details = {}
    for critique in critiques:
        if critique and critique.get("critique_json"):
            for issue in _safe_list(critique["critique_json"], "issues"):
                issue_details[issue.get("title")] = {
                    "severity": issue.get("severity", "low"),
                    "category": issue.get("category", "uncategorized"),
                }

    scored = []
    for patch in valid_patches:
        patch_json = patch.get("patch_json", {})
        sop_text = patch_json.get("sop", "").strip()

        mm = _safe_list(patch_json, "mitigation_matrix")
        residual = _safe_list(patch_json, "residual_risks")

        resolved = sum(1 for r in mm if str(r.get("status", "")).lower() == "resolved")
        mitigated = sum(
            1 for r in mm if str(r.get("status", "")).lower() == "mitigated"
        )

        # Score based on resolved issues, then mitigated, penalize for residuals, and use length as tie-breaker
        coverage_score = (resolved * 2) + mitigated
        final_score = coverage_score - (len(residual) * 2)

        # Track resolution by severity and category
        resolution_by_severity = {}
        resolution_by_category = {}
        for r in mm:
            issue_title = r.get("issue")
            if issue_title in issue_details:
                details = issue_details[issue_title]
                severity = details["severity"]
                category = details["category"]
                status = str(r.get("status", "")).lower()

                if status in ["resolved", "mitigated"]:
                    resolution_by_severity[severity] = (
                        resolution_by_severity.get(severity, 0) + 1
                    )
                    resolution_by_category[category] = (
                        resolution_by_category.get(category, 0) + 1
                    )

        scored.append(
            (
                final_score,
                resolved,
                len(sop_text),
                sop_text,
                patch.get("model"),
                resolution_by_severity,
                resolution_by_category,
            )
        )

    if not scored:
        return base_sop, {
            "reason": "all_patches_were_empty_or_invalid",
            "score": -1,
            "resolution_by_severity": {},
            "resolution_by_category": {},
        }

    # Sort by score, then resolved count, then SOP length
    scored.sort(key=lambda x: (-x[0], -x[1], x[2]))
    best_score, best_resolved, _, best_sop, best_model, best_res_sev, best_res_cat = (
        scored[0]
    )
    diagnostics = {
        "reason": "best_patch_selected",
        "score": best_score,
        "resolved": best_resolved,
        "model": best_model,
        "resolution_by_severity": best_res_sev,
        "resolution_by_category": best_res_cat,
    }
    return best_sop, diagnostics


def _aggregate_red_risk(critiques: List[dict]) -> Dict[str, Any]:
    """Computes an aggregate risk score from all red-team critiques."""
    sev_weight = {"low": 1, "medium": 3, "high": 6, "critical": 12}
    total_weight, issue_count = 0, 0
    categories = {}
    severities = {}

    valid_critiques = [
        c.get("critique_json") for c in critiques if c and c.get("critique_json")
    ]

    for critique in valid_critiques:
        for issue in _safe_list(critique, "issues"):
            sev = str(issue.get("severity", "low")).lower()
            weight = sev_weight.get(sev, 1)
            total_weight += weight
            issue_count += 1
            cat = str(issue.get("category", "uncategorized")).lower()
            categories[cat] = categories.get(cat, 0) + weight
            severities[sev] = severities.get(sev, 0) + 1

    avg_weight = (total_weight / max(1, issue_count)) if issue_count > 0 else 0
    return {
        "total_weight": total_weight,
        "avg_issue_weight": avg_weight,
        "categories": categories,
        "severities": severities,
        "count": issue_count,
    }


def _update_model_performance(critiques: List[dict]):
    """Updates the performance scores of models based on the critiques they generated."""
    with st.session_state.thread_lock:
        if "adversarial_model_performance" not in st.session_state:
            st.session_state.adversarial_model_performance = {}

        sev_weight = {"low": 1, "medium": 3, "high": 6, "critical": 12}
        for critique in critiques:
            model_id = critique.get("model")
            if not model_id:
                continue

            if model_id not in st.session_state.adversarial_model_performance:
                st.session_state.adversarial_model_performance[model_id] = {
                    "score": 0,
                    "issues_found": 0,
                }

            critique_json = critique.get("critique_json")
            if critique_json and isinstance(critique_json.get("issues"), list):
                for issue in critique_json["issues"]:
                    sev = str(issue.get("severity", "low")).lower()
                    st.session_state.adversarial_model_performance[model_id][
                        "score"
                    ] += sev_weight.get(sev, 1)
                    st.session_state.adversarial_model_performance[model_id][
                        "issues_found"
                    ] += 1


def _collect_model_configs(
    model_ids: List[str], max_tokens: int
) -> Dict[str, Dict[str, Any]]:
    return {
        model_id: {
            "temperature": st.session_state.get(f"temp_{model_id}", 0.7),
            "top_p": st.session_state.get(f"topp_{model_id}", 1.0),
            "frequency_penalty": st.session_state.get(f"freqpen_{model_id}", 0.0),
            "presence_penalty": st.session_state.get(f"prespen_{model_id}", 0.0),
            "max_tokens": max_tokens,
        }
        for model_id in model_ids
    }


def _update_adv_log_and_status(msg: str):
    """Thread-safe way to update logs and status message."""
    with st.session_state.thread_lock:
        st.session_state.adversarial_log.append(f"[{time.strftime('%H:%M:%S')}] {msg}")
        st.session_state.adversarial_status_message = msg


def _update_adv_counters(ptoks: int, ctoks: int, cost: float):
    """Thread-safe way to update token and cost counters."""
    with st.session_state.thread_lock:
        st.session_state.adversarial_total_tokens_prompt += ptoks
        st.session_state.adversarial_total_tokens_completion += ctoks
        st.session_state.adversarial_cost_estimate_usd += cost
        print(f"Updated cost: {st.session_state.adversarial_cost_estimate_usd}")


def check_approval_rate(
    api_key: str,
    red_team_models: List[str],
    sop_markdown: str,
    model_configs: Dict,
    seed: Optional[int],
    max_workers: int,
    approval_prompt: str = APPROVAL_PROMPT,
) -> Dict[str, Any]:
    """Asks all red-team models for a final verdict on the SOP."""
    votes, scores, approved = [], [], 0
    total_ptoks, total_ctoks, total_cost = 0, 0, 0.0

    with ThreadPoolExecutor(max_workers=max_workers) as ex:
        future_to_model = {
            ex.submit(
                analyze_with_model,
                api_key,
                model_id,
                sop_markdown,
                model_configs.get(model_id, {}),
                approval_prompt,
                force_json=True,
                seed=seed,
            ): model_id
            for model_id in red_team_models
        }
        for future in as_completed(future_to_model):
            model_id = future_to_model[future]
            res = future.result()
            total_ptoks += res["ptoks"]
            total_ctoks += res["ctoks"]
            total_cost += res["cost"]
            if res.get("ok") and res.get("json"):
                j = res["json"]
                verdict = str(j.get("verdict", "REJECTED")).upper()
                score = _clamp(safe_int(j.get("score"), 0), 0, 100)
                if verdict == "APPROVED":
                    approved += 1
                scores.append(score)
                votes.append(
                    {
                        "model": model_id,
                        "verdict": verdict,
                        "score": score,
                        "reasons": _safe_list(j, "reasons"),
                    }
                )
            else:
                votes.append(
                    {
                        "model": model_id,
                        "verdict": "ERROR",
                        "score": 0,
                        "reasons": [res.get("text")],
                    }
                )

    rate = (approved / max(1, len(red_team_models))) * 100.0
    avg_score = (sum(scores) / max(1, len(scores))) if scores else 0

    # Calculate agreement
    if not votes:
        agreement = 0.0
    else:
        verdicts = [v["verdict"] for v in votes]
        approved_count = verdicts.count("APPROVED")
        rejected_count = verdicts.count("REJECTED")
        agreement = max(approved_count, rejected_count) / len(verdicts) * 100.0

    return {
        "approval_rate": rate,
        "avg_score": avg_score,
        "votes": votes,
        "prompt_tokens": total_ptoks,
        "completion_tokens": total_ctoks,
        "cost": total_cost,
        "agreement": agreement,
    }


def generate_docx_report(results: dict) -> bytes:
    """Generates a DOCX report from the adversarial testing results."""
    document = docx.Document()
    document.add_heading("Adversarial Testing Report", 0)

    document.add_heading("Summary", level=1)
    document.add_paragraph(
        f"Final Approval Rate: {results.get('final_approval_rate', 0.0):.1f}%\n"
        f"Total Iterations: {len(results.get('iterations', []))}\n"
        f"Total Cost (USD): ${results.get('cost_estimate_usd', 0.0):,.4f}\n"
        f"Total Prompt Tokens: {results.get('tokens', {}).get('prompt', 0):,}\n"
        f"Total Completion Tokens: {results.get('tokens', {}).get('completion', 0):,}"
    )

    document.add_heading("Final Hardened SOP", level=1)
    document.add_paragraph(results.get("final_sop", ""))

    document.add_heading("Issues Found", level=1)
    for i, iteration in enumerate(results.get("iterations", [])):
        document.add_heading(f"Iteration {i + 1}", level=2)
        for critique in iteration.get("critiques", []):
            if critique.get("critique_json"):
                for issue in _safe_list(critique["critique_json"], "issues"):
                    document.add_paragraph(
                        f"- {issue.get('title')} ({issue.get('severity')})",
                        style="List Bullet",
                    )

    document.add_heading("Final Votes", level=1)
    if results.get("iterations"):
        for vote in (
            results["iterations"][-1].get("approval_check", {}).get("votes", [])
        ):
            document.add_paragraph(
                f"- {vote.get('model')}: {vote.get('verdict')} ({vote.get('score')})",
                style="List Bullet",
            )

    document.add_heading("Audit Trail", level=1)
    for log_entry in results.get("log", []):
        document.add_paragraph(log_entry)

    from io import BytesIO

    bio = BytesIO()
    document.save(bio)
    return bio.getvalue()


def generate_pdf_report(
    results: dict, watermark: str = None, custom_style: dict = None
) -> bytes:
    """Generates a PDF report from the adversarial testing results."""

    # Set default styling if not provided
    if custom_style is None:
        custom_style = {
            "font_face": "Arial",
            "font_size": 12,
            "header_font_size": 14,
            "title_font_size": 16,
            "primary_color": (42, 82, 152),  # RGB for #2a5298
            "secondary_color": (0, 0, 0),  # RGB for black
            "background_color": (255, 255, 255),  # RGB for white
        }

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font(custom_style["font_face"], size=custom_style["font_size"])

    # Set background color
    pdf.set_fill_color(*custom_style["background_color"])
    pdf.rect(0, 0, 210, 297, "F")  # A4 size in mm

    # Add watermark if provided
    if watermark:
        pdf.set_font(custom_style["font_face"], "B", 50)
        pdf.set_text_color(220, 220, 220)
        pdf.rotate(45)
        pdf.text(60, 150, watermark)
        pdf.rotate(0)
        pdf.set_text_color(0, 0, 0)

    # Set primary color for header text
    pdf.set_text_color(*custom_style["primary_color"])
    pdf.set_font(custom_style["font_face"], "B", custom_style["title_font_size"])
    pdf.cell(200, 10, txt="Adversarial Testing Report", ln=True, align="C")

    # Add a line separator
    pdf.set_draw_color(*custom_style["primary_color"])
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(10)

    # Reset to secondary color for content
    pdf.set_text_color(*custom_style["secondary_color"])

    pdf.set_font(custom_style["font_face"], "B", custom_style["header_font_size"])
    pdf.cell(200, 10, txt="Summary", ln=True)
    pdf.set_font(custom_style["font_face"], size=custom_style["font_size"])
    pdf.multi_cell(
        0,
        10,
        f"Final Approval Rate: {results.get('final_approval_rate', 0.0):.1f}%\n"
        f"Total Iterations: {len(results.get('iterations', []))}\n"
        f"Total Cost (USD): ${results.get('cost_estimate_usd', 0.0):,.4f}\n"
        f"Total Prompt Tokens: {results.get('tokens', {}).get('prompt', 0):,}\n"
        f"Total Completion Tokens: {results.get('tokens', {}).get('completion', 0):,}",
    )

    pdf.ln(10)

    pdf.set_font(custom_style["font_face"], "B", custom_style["header_font_size"])
    pdf.cell(200, 10, txt="Final Hardened SOP", ln=True)
    pdf.set_font(custom_style["font_face"], size=custom_style["font_size"])
    pdf.multi_cell(0, 10, results.get("final_sop", ""))

    pdf.ln(10)

    pdf.set_font(custom_style["font_face"], "B", custom_style["header_font_size"])
    pdf.cell(200, 10, txt="Issues Found", ln=True)
    pdf.set_font(custom_style["font_face"], size=custom_style["font_size"] - 2)
    for i, iteration in enumerate(results.get("iterations", [])):
        pdf.set_font(custom_style["font_face"], "B", custom_style["font_size"] - 1)
        pdf.cell(200, 10, txt=f"Iteration {i + 1}", ln=True)
        pdf.set_font(custom_style["font_face"], size=custom_style["font_size"] - 2)
        for critique in iteration.get("critiques", []):
            if critique.get("critique_json"):
                for issue in _safe_list(critique["critique_json"], "issues"):
                    pdf.multi_cell(
                        0, 10, f"- {issue.get('title')} ({issue.get('severity')})"
                    )

    pdf.ln(10)

    pdf.set_font(custom_style["font_face"], "B", custom_style["header_font_size"])
    pdf.cell(200, 10, txt="Final Votes", ln=True)
    pdf.set_font(custom_style["font_face"], size=custom_style["font_size"] - 2)
    if results.get("iterations"):
        for vote in (
            results["iterations"][-1].get("approval_check", {}).get("votes", [])
        ):
            pdf.multi_cell(
                0,
                10,
                f"- {vote.get('model')}: {vote.get('verdict')} ({vote.get('score')})",
            )

    pdf.ln(10)

    pdf.set_font(custom_style["font_face"], "B", custom_style["header_font_size"])
    pdf.cell(200, 10, txt="Audit Trail", ln=True)
    pdf.set_font(custom_style["font_face"], size=custom_style["font_size"] - 4)
    for log_entry in results.get("log", []):
        pdf.multi_cell(0, 5, log_entry)

    return pdf.output(dest="S").encode("latin-1")


def generate_html_report(
    results: dict, custom_css: str = "", custom_style: dict = None
) -> str:
    """Generates an HTML report from the adversarial testing results."""

    # Set default styling if not provided
    if custom_style is None:
        custom_style = {
            "font_family": "Arial, sans-serif",
            "primary_color": "#4a6fa5",
            "secondary_color": "#2a5298",
            "background_color": "#f8f9fa",
            "card_background": "#ffffff",
            "header_size": "2em",
            "subheader_size": "1.5em",
            "text_size": "1em",
            "border_radius": "8px",
            "box_shadow": "0 4px 6px rgba(0, 0, 0, 0.1)",
        }

    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Adversarial Testing Report</title>
        <style>
            body {{ 
                font-family: {custom_style["font_family"]}; 
                margin: 40px; 
                background-color: {custom_style["background_color"]}; 
                color: {custom_style["secondary_color"]};
            }}
            h1, h2, h3 {{ 
                color: {custom_style["primary_color"]}; 
                font-size: {custom_style["header_size"]};
            }}
            h2 {{ 
                font-size: {custom_style["subheader_size"]};
            }}
            p {{ 
                font-size: {custom_style["text_size"]};
            }}
            .summary {{ 
                background-color: {custom_style["card_background"]}; 
                padding: 15px; 
                border-radius: {custom_style["border_radius"]}; 
                box-shadow: {custom_style["box_shadow"]}; 
                margin-bottom: 20px; 
            }}
            .section {{ 
                margin: 20px 0; 
                background-color: {custom_style["card_background"]}; 
                padding: 20px; 
                border-radius: {custom_style["border_radius"]}; 
                box-shadow: {custom_style["box_shadow"]}; 
            }}
            .log {{ 
                font-family: monospace; 
                font-size: 0.9em; 
                background-color: #f9f9f9; 
                padding: 10px; 
                border-radius: 4px; 
            }}
            table {{ 
                border-collapse: collapse; 
                width: 100%; 
            }}
            th, td {{ 
                border: 1px solid #ddd; 
                padding: 8px; 
                text-align: left; 
            }}
            th {{ 
                background-color: {custom_style["primary_color"]}; 
                color: white; 
            }}
            .metric {{ 
                text-align: center; 
                padding: 10px; 
                background-color: #e9ecef; 
                border-radius: 4px; 
                margin: 5px; 
            }}
            .improvement {{ 
                color: #4caf50; 
                font-weight: bold; 
            }}
            .regression {{ 
                color: #f44336; 
                font-weight: bold; 
            }}
            {custom_css}
        </style>
    </head>
    <body>
        <h1>Adversarial Testing Report</h1>

        <div class="summary">
            <h2>Summary</h2>
            <p><strong>Final Approval Rate:</strong> {results.get("final_approval_rate", 0.0):.1f}%</p>
            <p><strong>Total Iterations:</strong> {len(results.get("iterations", []))}</p>
            <p><strong>Total Cost (USD):</strong> ${results.get("cost_estimate_usd", 0.0):,.4f}</p>
            <p><strong>Total Prompt Tokens:</strong> {results.get("tokens", {}).get("prompt", 0):,}</p>
            <p><strong>Total Completion Tokens:</strong> {results.get("tokens", {}).get("completion", 0):,}</p>
        </div>

        <div class="section">
            <h2>Final Hardened SOP</h2>
            <pre style="white-space: pre-wrap; background-color: #f9f9f9; padding: 15px; border-radius: 4px;">{results.get("final_sop", "")}</pre>
        </div>
    """

    if results.get("iterations"):
        html += """
        <div class="section">
            <h2>Issues Found</h2>
        """
        for i, iteration in enumerate(results.get("iterations", [])):
            html += f"<h3>Iteration {i + 1}</h3><ul>"
            for critique in iteration.get("critiques", []):
                if critique.get("critique_json"):
                    for issue in _safe_list(critique["critique_json"], "issues"):
                        severity = issue.get("severity", "low")
                        severity_color = {
                            "low": "#4caf50",
                            "medium": "#ff9800",
                            "high": "#f44336",
                            "critical": "#9c27b0",
                        }.get(severity, "#000000")
                        html += f"<li><span style='color: {severity_color}; font-weight: bold;'>{severity.upper()}</span>: {issue.get('title')}</li>"
            html += "</ul>"
        html += "</div>"

        html += """
        <div class="section">
            <h2>Final Votes</h2>
            <table>
                <tr><th>Model</th><th>Verdict</th><th>Score</th></tr>
        """
        for vote in (
            results["iterations"][-1].get("approval_check", {}).get("votes", [])
        ):
            verdict = vote.get("verdict", "")
            verdict_color = "#4caf50" if verdict.upper() == "APPROVED" else "#f44336"
            html += f"<tr><td>{vote.get('model')}</td><td style='color: {verdict_color}; font-weight: bold;'>{verdict}</td><td>{vote.get('score')}</td></tr>"
        html += "</table></div>"

        # Add performance chart data
        html += """
        <div class="section">
            <h2>Performance Metrics</h2>
            <div style="display: flex; justify-content: space-around; flex-wrap: wrap;">
        """

        # Approval rate chart
        approval_rates = [
            iter["approval_check"].get("approval_rate", 0)
            for iter in results.get("iterations", [])
        ]
        if approval_rates:
            avg_approval = sum(approval_rates) / len(approval_rates)
            html += f"""
            <div class="metric">
                <div>Avg Approval Rate</div>
                <div style="font-size: 24px;">{avg_approval:.1f}%</div>
            </div>
            """

        # Issue count chart
        issue_counts = [
            iter["agg_risk"].get("count", 0) for iter in results.get("iterations", [])
        ]
        if issue_counts:
            total_issues = sum(issue_counts)
            html += f"""
            <div class="metric">
                <div>Total Issues Found</div>
                <div style="font-size: 24px;">{total_issues}</div>
            </div>
            """

        html += "</div></div>"

    html += """
        <div class="section">
            <h2>Audit Trail</h2>
            <div class="log">
    """
    for log_entry in results.get("log", []):
        html += f"<div>{log_entry}</div>"
    html += """
            </div>
        </div>
    </body>
    </html>
    """

    return html


def generate_latex_report(results: dict) -> str:
    """Generates a LaTeX report from the adversarial testing results."""
    latex = "\\documentclass{article}\n"
    latex += "\\usepackage[utf8]{inputenc}\n"
    latex += "\\usepackage{geometry}\n"
    latex += "\\geometry{a4paper, margin=1in}\n"
    latex += "\\title{Adversarial Testing Report}\n"
    latex += "\\author{OpenEvolve}\n"
    latex += "\\date{\\today}\n"
    latex += "\\begin{document}\n"
    latex += "\\maketitle\n"

    latex += "\\section{Summary}\n"
    latex += f"Final Approval Rate: {results.get('final_approval_rate', 0.0):.1f}\%\\n"
    latex += f"Total Iterations: {len(results.get('iterations', []))}\\n"
    latex += f"Total Cost (USD): ${results.get('cost_estimate_usd', 0.0):,.4f}\\n"
    latex += f"Total Prompt Tokens: {results.get('tokens', {}).get('prompt', 0):,}\\n"
    latex += f"Total Completion Tokens: {results.get('tokens', {}).get('completion', 0):,}\\n"

    latex += "\\section{Final Hardened SOP}\n"
    latex += f"\\begin{{verbatim}}\n{results.get('final_sop', '')}\\end{{verbatim}}\n"

    latex += "\\section{Issues Found}\n"
    for i, iteration in enumerate(results.get("iterations", [])):
        latex += f"\\subsection*{{Iteration {i + 1}}}\n"
        latex += "\\begin{itemize}\n"
        for critique in iteration.get("critiques", []):
            if critique.get("critique_json"):
                for issue in _safe_list(critique["critique_json"], "issues"):
                    latex += f"\\item {issue.get('title')} ({issue.get('severity')})\n"
        latex += "\\end{itemize}\n"

    latex += "\\section{Final Votes}\n"
    latex += "\\begin{itemize}\n"
    if results.get("iterations"):
        for vote in (
            results["iterations"][-1].get("approval_check", {}).get("votes", [])
        ):
            latex += f"\\item {vote.get('model')}: {vote.get('verdict')} ({vote.get('score')})\n"
    latex += "\\end{itemize}\n"

    latex += "\\section{Audit Trail}\n"
    latex += "\\begin{verbatim}\n"
    for log_entry in results.get("log", []):
        latex += f"{log_entry}\\n"
    latex += "\\end{verbatim}\n"

    latex += "\\end{document}\n"

    return latex


def generate_compliance_report(results: dict, compliance_requirements: str) -> str:
    """Generates a compliance report from the adversarial testing results."""
    # This is a simplified compliance report. A real implementation would require a more sophisticated analysis.
    report = "# Compliance Report\n\n"
    report += f"## Compliance Requirements\n\n{compliance_requirements}\n\n"

    report += "## Compliance Status\n\n"
    if results.get("final_approval_rate", 0) >= 90:
        report += "✅ The protocol is compliant with the specified requirements.\n"
    else:
        report += "❌ The protocol is not compliant with the specified requirements.\n"

    return report


def suggest_performance_improvements(current_config: Dict) -> List[str]:
    """Suggest performance improvements for the current configuration.

    Args:
        current_config (Dict): Current adversarial testing configuration

    Returns:
        List[str]: List of suggested improvements
    """
    suggestions = []

    red_models = current_config.get("red_team_models", [])
    blue_models = current_config.get("blue_team_models", [])
    iterations = current_config.get("adversarial_max_iter", 10)
    protocol_text = current_config.get("protocol_text", "")

    # Check for common performance issues
    if len(red_models) > 5:
        suggestions.append(
            "🔴 Reduce red team models to 3-5 for better performance and cost control"
        )

    if len(blue_models) > 5:
        suggestions.append(
            "🔵 Reduce blue team models to 3-5 for better performance and cost control"
        )

    if iterations > 20:
        suggestions.append(
            "🔄 Consider reducing max iterations to 15-20 for faster results"
        )

    if len(protocol_text.split()) > 5000:
        suggestions.append(
            "📄 Your protocol is quite long (>5000 words). Consider breaking it into smaller sections"
        )

    # Check for model diversity
    all_models = red_models + blue_models
    if len(set(all_models)) < len(all_models) * 0.7:
        suggestions.append(
            "🔀 Increase model diversity by selecting models from different providers"
        )

    # Check for expensive model combinations
    expensive_models = [m for m in all_models if "gpt-4" in m or "claude-3-opus" in m]
    if len(expensive_models) > 3:
        suggestions.append(
            "💰 You're using many expensive models. Consider mixing in some cost-effective models"
        )

    # If no suggestions, provide positive feedback
    if not suggestions:
        suggestions.append(
            "✅ Your configuration looks well-balanced for optimal performance!"
        )

    return suggestions


def optimize_model_selection(
    red_team_models: List[str],
    blue_team_models: List[str],
    protocol_complexity: int,
    budget_limit: float = 0.0,
) -> Dict[str, List[str]]:
    """Optimize model selection based on protocol complexity and budget.

    Args:
        red_team_models (List[str]): Available red team models
        blue_team_models (List[str]): Available blue team models
        protocol_complexity (int): Complexity score of the protocol (0-100)
        budget_limit (float): Maximum budget in USD (0 = no limit)

    Returns:
        Dict[str, List[str]]: Optimized model selections
    """
    # Enhanced optimization logic
    optimized = {"red_team": [], "blue_team": []}

    # For complex protocols, use more capable models
    if protocol_complexity > 70:
        # Use premium models for complex protocols
        optimized["red_team"] = [
            m
            for m in red_team_models
            if "gpt-4" in m or "claude-3-opus" in m or "gemini-1.5-pro" in m
        ][:3]
        optimized["blue_team"] = [
            m
            for m in blue_team_models
            if "gpt-4" in m or "claude-3-sonnet" in m or "gemini-1.5-pro" in m
        ][:3]
    elif protocol_complexity > 40:
        # Use balanced models for medium complexity
        optimized["red_team"] = [
            m
            for m in red_team_models
            if "gpt-4" in m or "claude-3-haiku" in m or "gemini-1.5-flash" in m
        ][:3]
        optimized["blue_team"] = [
            m
            for m in blue_team_models
            if "gpt-4" in m or "claude-3-sonnet" in m or "gemini-1.5-flash" in m
        ][:3]
    else:
        # Use cost-effective models for simple protocols
        optimized["red_team"] = [
            m
            for m in red_team_models
            if "gpt-4o-mini" in m or "claude-3-haiku" in m or "llama-3-8b" in m
        ][:3]
        optimized["blue_team"] = [
            m
            for m in blue_team_models
            if "gpt-4o" in m or "claude-3-sonnet" in m or "llama-3-70b" in m
        ][:3]

    # Budget-based optimization
    if budget_limit > 0:
        # Filter out expensive models if budget is constrained
        expensive_models = ["gpt-4", "claude-3-opus", "gemini-1.5-pro"]
        if budget_limit < 0.1:  # Very low budget
            optimized["red_team"] = [
                m
                for m in optimized["red_team"]
                if not any(exp in m for exp in expensive_models)
            ]
            optimized["blue_team"] = [
                m
                for m in optimized["blue_team"]
                if not any(exp in m for exp in expensive_models)
            ]
        elif budget_limit < 0.5:  # Moderate budget
            # Reduce count of expensive models
            expensive_red = [
                m
                for m in optimized["red_team"]
                if any(exp in m for exp in expensive_models)
            ]
            if len(expensive_red) > 1:
                optimized["red_team"] = [
                    m for m in optimized["red_team"] if m not in expensive_red[1:]
                ]

            expensive_blue = [
                m
                for m in optimized["blue_team"]
                if any(exp in m for exp in expensive_models)
            ]
            if len(expensive_blue) > 1:
                optimized["blue_team"] = [
                    m for m in optimized["blue_team"] if m not in expensive_blue[1:]
                ]

    # If no models matched criteria, use defaults
    if not optimized["red_team"]:
        optimized["red_team"] = red_team_models[: min(3, len(red_team_models))]
    if not optimized["blue_team"]:
        optimized["blue_team"] = blue_team_models[: min(3, len(blue_team_models))]

    # Ensure diversity in model selection
    if len(optimized["red_team"]) < 3 and len(red_team_models) >= 3:
        # Add models from different providers
        providers = set()
        selected_models = []
        for model in red_team_models:
            provider = model.split("/")[0] if "/" in model else model.split("-")[0]
            if provider not in providers or len(selected_models) < 3:
                selected_models.append(model)
                providers.add(provider)
        optimized["red_team"] = selected_models[:3]

    if len(optimized["blue_team"]) < 3 and len(blue_team_models) >= 3:
        # Add models from different providers
        providers = set()
        selected_models = []
        for model in blue_team_models:
            provider = model.split("/")[0] if "/" in model else model.split("-")[0]
            if provider not in providers or len(selected_models) < 3:
                selected_models.append(model)
                providers.add(provider)
        optimized["blue_team"] = selected_models[:3]

    return optimized


def estimate_testing_time_and_cost(
    red_team_models: List[str],
    blue_team_models: List[str],
    iterations: int,
    protocol_length: int,
) -> Dict[str, Any]:
    """Estimate testing time and cost based on configuration.

    Args:
        red_team_models (List[str]): Selected red team models
        blue_team_models (List[str]): Selected blue team models
        iterations (int): Number of iterations
        protocol_length (int): Length of protocol in words

    Returns:
        Dict[str, Any]: Time and cost estimates
    """
    # Simplified estimation logic
    # Base estimates per model per iteration
    avg_response_time = 5  # seconds
    avg_cost_per_1000_tokens = 0.002  # USD

    # Calculate total operations
    total_red_operations = len(red_team_models) * iterations
    total_blue_operations = len(blue_team_models) * iterations

    # Estimate time (parallel processing assumed)
    max_parallel_workers = min(6, len(red_team_models) + len(blue_team_models))
    estimated_time_seconds = (
        (total_red_operations + total_blue_operations) / max_parallel_workers
    ) * avg_response_time

    # Estimate cost (simplified token estimation)
    avg_tokens_per_operation = protocol_length * 3  # Rough estimate
    total_tokens = (
        total_red_operations + total_blue_operations
    ) * avg_tokens_per_operation
    estimated_cost = (total_tokens / 1000) * avg_cost_per_1000_tokens

    return {
        "estimated_time_minutes": round(estimated_time_seconds / 60, 1),
        "estimated_cost_usd": round(estimated_cost, 4),
        "total_operations": total_red_operations + total_blue_operations,
        "total_tokens_estimated": total_tokens,
    }


def run_adversarial_testing():
    """Run adversarial testing with content-type-aware routing to OpenEvolve backend."""
    print("run_adversarial_testing function called")
    # Load model performance data for continuous learning
    try:
        with open("model_performance.json", "r") as f:
            st.session_state.adversarial_model_performance = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        st.session_state.adversarial_model_performance = {}

    try:
        # --- Initialization ---
        api_key = st.session_state.openrouter_key
        red_team_base = list(st.session_state.red_team_models or [])
        blue_team_base = list(st.session_state.blue_team_models or [])
        min_iter, max_iter = (
            st.session_state.adversarial_min_iter,
            st.session_state.adversarial_max_iter,
        )
        print(f"Min iterations: {min_iter}, Max iterations: {max_iter}")
        confidence = st.session_state.adversarial_confidence
        max_tokens = st.session_state.adversarial_max_tokens
        json_mode = st.session_state.adversarial_force_json
        max_workers = st.session_state.adversarial_max_workers
        rotation_strategy = st.session_state.adversarial_rotation_strategy
        seed_str = str(st.session_state.adversarial_seed or "").strip()
        seed = None
        if seed_str:
            try:
                seed = int(float(seed_str))  # Handle floats by truncating to int
            except (ValueError, TypeError):
                pass  # Invalid input, keep seed as None

        # Validation
        if not api_key:
            _update_adv_log_and_status(
                "❌ Error: OpenRouter API key is required for adversarial testing."
            )
            with st.session_state.thread_lock:
                st.session_state.adversarial_running = False
            return

        if not red_team_base or not blue_team_base:
            _update_adv_log_and_status(
                "❌ Error: Please select at least one model for both red and blue teams."
            )
            with st.session_state.thread_lock:
                st.session_state.adversarial_running = False
            return

        if not st.session_state.protocol_text.strip():
            _update_adv_log_and_status("❌ Error: Please enter a protocol to test.")
            with st.session_state.thread_lock:
                st.session_state.adversarial_running = False
            return

        with st.session_state.thread_lock:
            st.session_state.adversarial_log = []
            st.session_state.adversarial_stop_flag = False
            st.session_state.adversarial_total_tokens_prompt = 0
            st.session_state.adversarial_total_tokens_completion = 0
            st.session_state.adversarial_cost_estimate_usd = 0.0

        model_configs = _collect_model_configs(
            red_team_base + blue_team_base, max_tokens
        )
        current_sop = st.session_state.protocol_text
        base_hash = _hash_text(current_sop)
        iteration = 0

        # Determine review type and get appropriate prompts
        content_type = "general"
        if st.session_state.get("adversarial_custom_mode", False):
            # Use custom prompts when custom mode is enabled
            red_team_prompt = st.session_state.get(
                "adversarial_custom_red_prompt", RED_TEAM_CRITIQUE_PROMPT
            )
            blue_team_prompt = st.session_state.get(
                "adversarial_custom_blue_prompt", BLUE_TEAM_PATCH_PROMPT
            )
            approval_prompt = st.session_state.get(
                "adversarial_custom_approval_prompt", APPROVAL_PROMPT
            )
            review_type = "custom"
            print(f"Using custom red team prompt: {red_team_prompt}")
            print(f"Using custom blue team prompt: {blue_team_prompt}")
            print(f"Using custom approval prompt: {approval_prompt}")
        else:
            # Use standard prompts based on review type
            if st.session_state.adversarial_review_type == "Auto-Detect":
                review_type = determine_review_type(current_sop)
            elif st.session_state.adversarial_review_type == "Code Review":
                review_type = "code"
                content_type = st.session_state.get(
                    "code_language_type", "code_python"
                )  # Assuming a UI selection for code language
            elif st.session_state.adversarial_review_type == "Plan Review":
                review_type = "plan"
                content_type = (
                    "document_general"  # Or a more specific plan document type
                )
            elif st.session_state.adversarial_review_type == "Legal Document":
                review_type = "document"
                content_type = "document_legal"
            elif st.session_state.adversarial_review_type == "Medical Document":
                review_type = "document"
                content_type = "document_medical"
            elif st.session_state.adversarial_review_type == "Technical Document":
                review_type = "document"
                content_type = "document_technical"
            else:
                review_type = "general"
                content_type = "document_general"

            red_team_prompt, blue_team_prompt = get_appropriate_prompts(review_type)
            print(f"Using review type: {review_type}, content_type: {content_type}")

        # Generate dynamic prompt enhancements
        red_team_prompt_enhancement = ""
        blue_team_prompt_enhancement = ""

        if content_type.startswith("code_"):
            red_team_prompt_enhancement += "\nAs a red teamer, focus on security vulnerabilities, code quality, and performance issues in the code. Look for potential bugs, inefficient algorithms, and non-idiomatic code."
            blue_team_prompt_enhancement += "\nAs a blue teamer, focus on fixing all identified issues, improving code robustness, and optimizing performance. Ensure the code is clean, secure, and follows best practices."
        elif content_type == "document_legal":
            red_team_prompt_enhancement += "\nAs a red teamer, scrutinize the legal document for ambiguities, loopholes, non-compliance with legal standards (e.g., GDPR, CCPA), and potential liabilities."
            blue_team_prompt_enhancement += "\nAs a blue teamer, refine the legal document for clarity, enforce compliance with relevant regulations, and mitigate all identified legal risks."
        elif content_type == "document_medical":
            red_team_prompt_enhancement += "\nAs a red teamer, analyze the medical document for factual inaccuracies, patient privacy violations (e.g., HIPAA), ethical concerns, and clarity for medical professionals."
            blue_team_prompt_enhancement += "\nAs a blue teamer, ensure the medical document is factually accurate, compliant with patient privacy laws, ethically sound, and clearly understandable by medical staff."
        elif content_type == "document_technical":
            red_team_prompt_enhancement += "\nAs a red teamer, review the technical document for technical inaccuracies, outdated information, unclear instructions, and potential security implications of described systems."
            blue_team_prompt_enhancement += "\nAs a blue teamer, update the technical document for accuracy, clarity, and completeness. Ensure all technical details are correct and instructions are easy to follow."
        elif review_type == "plan":
            red_team_prompt_enhancement += "\nAs a red teamer, critique the plan for feasibility, resource allocation, risk assessment, and alignment with strategic objectives. Identify any hidden dependencies or unrealistic timelines."
            blue_team_prompt_enhancement += "\nAs a blue teamer, refine the plan to address all identified risks, optimize resource allocation, and ensure feasibility and strategic alignment."
        else:  # General SOP
            red_team_prompt_enhancement += "\nAs a red teamer, identify any weaknesses, inefficiencies, or potential misinterpretations in the general SOP. Focus on clarity, completeness, and robustness."
            blue_team_prompt_enhancement += "\nAs a blue teamer, improve the general SOP by enhancing clarity, ensuring completeness, and making it more robust against misinterpretation."

        if st.session_state.get("compliance_requirements"):
            red_team_prompt_enhancement += f"\nAlso, specifically check for compliance with the following requirements: {st.session_state.compliance_requirements}"
            blue_team_prompt_enhancement += f"\nEnsure the final output strictly adheres to the following compliance requirements: {st.session_state.compliance_requirements}"

        _update_adv_log_and_status(
            f"🚀 Start: {len(red_team_base)} red / {len(blue_team_base)} blue | seed={seed} | base_hash={base_hash} | rotation={rotation_strategy} | review_type={review_type}"
        )

        # Check if we should use OpenEvolve backend for code-specific features
        if OPENEVOLVE_AVAILABLE and (
            review_type == "code" or review_type == "document"
        ):
            _update_adv_log_and_status(
                "🚀 Using OpenEvolve backend for adversarial testing..."
            )
            _run_adversarial_testing_with_openevolve_backend(
                current_sop,
                content_type,  # Pass the determined content_type
                red_team_base,
                blue_team_base,
                api_key,
                st.session_state.openrouter_base_url,
                max_iter,
                confidence,
                max_tokens,
                st.session_state.get("adversarial_temperature", 0.7),
                st.session_state.get("adversarial_top_p", 0.95),
                st.session_state.get("adversarial_frequency_penalty", 0.0),
                st.session_state.get("adversarial_presence_penalty", 0.0),
                seed,
                max_workers,
                rotation_strategy,
                st.session_state.get("red_team_sample_size", len(red_team_base)),
                st.session_state.get("blue_team_sample_size", len(blue_team_base)),
                st.session_state.get("custom_requirements", ""),
                evaluator_system_prompt=approval_prompt,
                red_team_prompt=red_team_prompt,
                blue_team_prompt=blue_team_prompt,
                compliance_rules=st.session_state.get("compliance_rules", None),
                red_team_prompt_enhancement=red_team_prompt_enhancement,
                blue_team_prompt_enhancement=blue_team_prompt_enhancement,
                feature_dimensions=st.session_state.get(
                    "adversarial_feature_dimensions", None
                ),
                feature_bins=st.session_state.get("adversarial_feature_bins", None),
                enable_data_augmentation=st.session_state.get(
                    "adversarial_enable_data_augmentation", False
                ),
                augmentation_model_id=st.session_state.get(
                    "adversarial_augmentation_model_id", None
                ),
                augmentation_temperature=st.session_state.get(
                    "adversarial_augmentation_temperature", 0.7
                ),
                enable_human_feedback=st.session_state.get(
                    "adversarial_enable_human_feedback", False
                ),
                current_iteration=iteration,
            )
        else:
            _update_adv_log_and_status("🚀 Using API-based adversarial testing...")

    except Exception as e:
        tb_str = traceback.format_exc()
        error_message = f"💥 A critical error occurred: {e}\n{tb_str}"
        _update_adv_log_and_status(error_message)
        with st.session_state.thread_lock:
            st.session_state.adversarial_running = False
            if (
                "adversarial_results" not in st.session_state
                or not st.session_state.adversarial_results
            ):
                st.session_state.adversarial_results = {}
            st.session_state.adversarial_results["critical_error"] = error_message


def _run_adversarial_testing_with_api_backend(
    current_sop,
    api_key,
    red_team_base,
    blue_team_base,
    min_iter,
    max_iter,
    confidence,
    max_tokens,
    json_mode,
    max_workers,
    rotation_strategy,
    seed,
    model_configs,
    red_team_prompt,
    blue_team_prompt,
    approval_prompt,
):
    _update_adv_log_and_status(
        "❌ Error: OpenEvolve backend not available and API-based adversarial testing is not yet implemented."
    )


def capture_human_feedback(
    adversarial_example: Dict[str, Any], human_score: float, human_comments: str
):
    """
    Simulates capturing human feedback on an adversarial example.
    In a real system, this would store feedback in a database or persistent storage.
    """
    feedback_entry = {
        "timestamp": time.time(),
        "adversarial_example_id": adversarial_example.get("id"),
        "human_score": human_score,
        "human_comments": human_comments,
    }
    if "human_feedback_log" not in st.session_state:
        st.session_state.human_feedback_log = []
    st.session_state.human_feedback_log.append(feedback_entry)
    _update_adv_log_and_status(
        f"📝 Captured human feedback for adversarial example {adversarial_example.get('id')}"
    )


def generate_adversarial_data_augmentation(
    content: str,
    content_type: str,
    api_key: str,
    model_id: str,
    temperature: float,
    max_tokens: int,
    seed: Optional[int] = None,
    augmentation_strategy: str = "rephrase",
) -> str:
    """
    Generates an augmented version of the content using an LLM.

    Args:
        content: The original content to augment.
        content_type: The type of content (e.g., "code_python", "document_legal").
        api_key: API key for the LLM provider.
        model_id: Name of the LLM model to use for augmentation.
        temperature: Temperature for generation.
        max_tokens: Maximum tokens to generate.
        seed: Random seed for reproducibility.
        augmentation_strategy: The strategy to use for augmentation (e.g., "rephrase", "add_noise").

    Returns:
        The augmented content string.
    """
    try:
        system_prompt = f"You are an expert content rephraser. Your task is to rephrase the provided {content_type} content to make it more complex, ambiguous, or subtly flawed, without changing its core functionality or intent. Focus on introducing nuances that might challenge an AI reviewer."
        user_prompt = f"Rephrase the following {content_type} content:\n\n---\n\n{content}\n\n---\n\nRephrased content:"

        # Use _request_openrouter_chat for augmentation
        augmented_content, _, _, _ = _request_openrouter_chat(
            api_key=api_key,
            model_id=model_id,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=temperature,
            top_p=1.0,  # Keep top_p high for more diverse rephrasing
            frequency_penalty=0.0,
            presence_penalty=0.0,
            max_tokens=max_tokens,
            seed=seed,
        )
        return augmented_content
    except Exception as e:
        _update_adv_log_and_status(f"Error during adversarial data augmentation: {e}")
        return content  # Return original content on error
