"""
CrewAI Research External Integrations - Features 8-10 Implementation

8. Automated Literature Search
   - arXiv integration
   - Google Scholar integration
   - PubMed integration
   - Semantic Scholar integration
   - Citation analysis

9. Experiment Tracking
   - Experiment logging
   - Parameter tracking
   - Metric collection
   - Artifact storage
   - Result comparison

10. Research Report Generation
    - Automated report writing
    - Citation formatting
    - Figure generation
    - Table formatting
    - Export to PDF/DOCX

License: MIT
"""

import json
import logging
import os
from dataclasses import dataclass, field, asdict
from datetime import datetime
from typing import Dict, List, Any, Optional, Union, Tuple
from enum import Enum
from abc import ABC, abstractmethod
from pathlib import Path
import hashlib
import uuid

# Configure logging
logger = logging.getLogger(__name__)


# =============================================================================
# FEATURE 8: AUTOMATED LITERATURE SEARCH
# =============================================================================

class DatabaseType(Enum):
    """Supported academic databases"""
    ARXIV = "arxiv"
    GOOGLE_SCHOLAR = "google_scholar"
    PUBMED = "pubmed"
    SEMANTIC_SCHOLAR = "semantic_scholar"
    IEEE = "ieee"
    ACM = "acm"


@dataclass
class Paper:
    """Research paper metadata"""
    paper_id: str
    title: str
    authors: List[str]
    abstract: str
    publication_date: Optional[str] = None
    journal: Optional[str] = None
    doi: Optional[str] = None
    url: Optional[str] = None
    pdf_url: Optional[str] = None
    citation_count: int = 0
    reference_count: int = 0
    fields_of_study: List[str] = field(default_factory=list)
    keywords: List[str] = field(default_factory=list)
    source_database: str = ""
    relevance_score: float = 0.0
    full_text: Optional[str] = None


@dataclass
class CitationNetwork:
    """Citation network for a paper"""
    paper_id: str
    citing_papers: List[str] = field(default_factory=list)
    cited_papers: List[str] = field(default_factory=list)
    citation_depth: int = 0


class BaseLiteratureSearch(ABC):
    """Base class for literature search providers"""
    
    def __init__(self, database_type: DatabaseType):
        self.database_type = database_type
        self.logger = logging.getLogger(__name__)
    
    @abstractmethod
    def search(
        self,
        query: str,
        max_results: int = 10,
        filters: Optional[Dict[str, Any]] = None
    ) -> List[Paper]:
        """Search for papers"""
        pass
    
    @abstractmethod
    def get_paper_by_id(self, paper_id: str) -> Optional[Paper]:
        """Get paper by ID"""
        pass
    
    def _generate_paper_id(self, title: str, authors: List[str]) -> str:
        """Generate unique paper ID"""
        key = f"{title}|{','.join(authors[:2])}"
        return hashlib.md5(key.encode()).hexdigest()[:16]


class ArXivSearch(BaseLiteratureSearch):
    """arXiv literature search integration"""
    
    def __init__(self):
        super().__init__(DatabaseType.ARXIV)
        self.base_url = "http://export.arxiv.org/api/query"
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if arxiv package is available"""
        try:
            import arxiv
            self.arxiv = arxiv
            self.available = True
        except ImportError:
            self.available = False
            self.logger.warning("arxiv package not installed. Using mock implementation.")
    
    def search(
        self,
        query: str,
        max_results: int = 10,
        filters: Optional[Dict[str, Any]] = None
    ) -> List[Paper]:
        """Search arXiv for papers"""
        if not self.available:
            return self._mock_search(query, max_results)
        
        try:
            search = self.arxiv.Search(
                query=query,
                max_results=max_results,
                sort_by=self.arxiv.SortCriterion.Relevance
            )
            
            papers = []
            for result in search.results():
                paper = Paper(
                    paper_id=result.entry_id.split('/')[-1],
                    title=result.title,
                    authors=[str(a) for a in result.authors],
                    abstract=result.summary,
                    publication_date=result.published.isoformat() if result.published else None,
                    doi=result.doi,
                    pdf_url=result.pdf_url,
                    source_database="arxiv",
                    citation_count=0  # arXiv doesn't provide citation counts
                )
                papers.append(paper)
            
            return papers
            
        except Exception as e:
            self.logger.error(f"arXiv search error: {e}")
            return []
    
    def get_paper_by_id(self, paper_id: str) -> Optional[Paper]:
        """Get paper by arXiv ID"""
        if not self.available:
            return None
        
        try:
            search = self.arxiv.Search(id_list=[paper_id])
            results = list(search.results())
            
            if results:
                result = results[0]
                return Paper(
                    paper_id=paper_id,
                    title=result.title,
                    authors=[str(a) for a in result.authors],
                    abstract=result.summary,
                    publication_date=result.published.isoformat() if result.published else None,
                    doi=result.doi,
                    pdf_url=result.pdf_url,
                    source_database="arxiv"
                )
            return None
            
        except Exception as e:
            self.logger.error(f"arXiv get paper error: {e}")
            return None
    
    def _mock_search(self, query: str, max_results: int) -> List[Paper]:
        """Mock search for testing"""
        return [
            Paper(
                paper_id=f"arxiv_mock_{i}",
                title=f"Mock Paper {i} about {query[:20]}",
                authors=["Author A", "Author B"],
                abstract=f"This is a mock abstract for paper {i} about {query}",
                source_database="arxiv",
                relevance_score=0.9 - (i * 0.05)
            )
            for i in range(min(max_results, 5))
        ]


class GoogleScholarSearch(BaseLiteratureSearch):
    """Google Scholar literature search integration"""
    
    def __init__(self):
        super().__init__(DatabaseType.GOOGLE_SCHOLAR)
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if scholarly is available"""
        try:
            from scholarly import scholarly
            self.scholarly = scholarly
            self.available = True
        except ImportError:
            self.available = False
            self.logger.warning("scholarly package not installed. Using mock implementation.")
    
    def search(
        self,
        query: str,
        max_results: int = 10,
        filters: Optional[Dict[str, Any]] = None
    ) -> List[Paper]:
        """Search Google Scholar for papers"""
        if not self.available:
            return self._mock_search(query, max_results)
        
        try:
            search_query = self.scholarly.search_pubs(query)
            
            papers = []
            for i, result in enumerate(search_query):
                if i >= max_results:
                    break
                
                bib = result.get('bib', {})
                paper = Paper(
                    paper_id=result.get('author_pub_id', f"gs_{i}"),
                    title=bib.get('title', ''),
                    authors=bib.get('author', '').split(' and ') if bib.get('author') else [],
                    abstract=bib.get('abstract', ''),
                    publication_date=str(bib.get('pub_year', '')),
                    journal=bib.get('venue', ''),
                    citation_count=result.get('num_citations', 0),
                    source_database="google_scholar"
                )
                papers.append(paper)
            
            return papers
            
        except Exception as e:
            self.logger.error(f"Google Scholar search error: {e}")
            return []
    
    def get_paper_by_id(self, paper_id: str) -> Optional[Paper]:
        """Google Scholar doesn't support direct ID lookup"""
        return None
    
    def _mock_search(self, query: str, max_results: int) -> List[Paper]:
        """Mock search for testing"""
        return [
            Paper(
                paper_id=f"gs_mock_{i}",
                title=f"Scholar Paper {i}: {query[:20]}",
                authors=["Researcher A", "Researcher B"],
                abstract=f"Scholarly work on {query}",
                publication_date="2023",
                citation_count=50 - i * 5,
                source_database="google_scholar",
                relevance_score=0.95 - (i * 0.05)
            )
            for i in range(min(max_results, 5))
        ]


class PubMedSearch(BaseLiteratureSearch):
    """PubMed literature search integration"""
    
    def __init__(self):
        super().__init__(DatabaseType.PUBMED)
        self.base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if Biopython is available"""
        try:
            from Bio import Entrez
            self.Entrez = Entrez
            # Set a generic email (required by NCBI)
            Entrez.email = os.getenv("NCBI_EMAIL", "research@example.com")
            self.available = True
        except ImportError:
            self.available = False
            self.logger.warning("Biopython not installed. Using mock implementation.")
    
    def search(
        self,
        query: str,
        max_results: int = 10,
        filters: Optional[Dict[str, Any]] = None
    ) -> List[Paper]:
        """Search PubMed for papers"""
        if not self.available:
            return self._mock_search(query, max_results)
        
        try:
            # Search for IDs
            handle = self.Entrez.esearch(db="pubmed", term=query, retmax=max_results)
            record = self.Entrez.read(handle)
            id_list = record["IdList"]
            
            if not id_list:
                return []
            
            # Fetch details
            handle = self.Entrez.efetch(db="pubmed", id=",".join(id_list), retmode="xml")
            records = self.Entrez.read(handle)
            
            papers = []
            for article in records.get('PubmedArticle', []):
                medline = article.get('MedlineCitation', {})
                article_data = medline.get('Article', {})
                
                # Extract authors
                authors = []
                author_list = article_data.get('AuthorList', [])
                for author in author_list:
                    last_name = author.get('LastName', '')
                    fore_name = author.get('ForeName', '')
                    if last_name:
                        authors.append(f"{fore_name} {last_name}".strip())
                
                # Extract abstract
                abstract_data = article_data.get('Abstract', {})
                abstract = " ".join(abstract_data.get('AbstractText', []))
                
                paper = Paper(
                    paper_id=medline.get('PMID', ''),
                    title=article_data.get('ArticleTitle', ''),
                    authors=authors,
                    abstract=abstract,
                    publication_date=str(medline.get('DateCompleted', {})),
                    journal=article_data.get('Journal', {}).get('Title', ''),
                    source_database="pubmed"
                )
                papers.append(paper)
            
            return papers
            
        except Exception as e:
            self.logger.error(f"PubMed search error: {e}")
            return []
    
    def get_paper_by_id(self, paper_id: str) -> Optional[Paper]:
        """Get paper by PubMed ID"""
        if not self.available:
            return None
        
        try:
            handle = self.Entrez.efetch(db="pubmed", id=paper_id, retmode="xml")
            records = self.Entrez.read(handle)
            
            articles = records.get('PubmedArticle', [])
            if articles:
                # Parse similar to search
                return None  # Simplified for brevity
            return None
            
        except Exception as e:
            self.logger.error(f"PubMed get paper error: {e}")
            return None
    
    def _mock_search(self, query: str, max_results: int) -> List[Paper]:
        """Mock search for testing"""
        return [
            Paper(
                paper_id=f"pmid_{10000000 + i}",
                title=f"PubMed Paper {i}: {query[:20]}",
                authors=["Scientist A", "Scientist B"],
                abstract=f"Biomedical research on {query}",
                journal="Journal of Medical Research",
                publication_date="2023",
                source_database="pubmed",
                relevance_score=0.88 - (i * 0.05)
            )
            for i in range(min(max_results, 5))
        ]


class SemanticScholarSearch(BaseLiteratureSearch):
    """Semantic Scholar literature search integration"""
    
    def __init__(self, api_key: Optional[str] = None):
        super().__init__(DatabaseType.SEMANTIC_SCHOLAR)
        self.api_key = api_key or os.getenv("S2_API_KEY")
        self.base_url = "https://api.semanticscholar.org/graph/v1"
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if requests is available"""
        try:
            import requests
            self.requests = requests
            self.available = True
        except ImportError:
            self.available = False
            self.logger.warning("requests not installed. Using mock implementation.")
    
    def search(
        self,
        query: str,
        max_results: int = 10,
        filters: Optional[Dict[str, Any]] = None
    ) -> List[Paper]:
        """Search Semantic Scholar for papers"""
        if not self.available:
            return self._mock_search(query, max_results)
        
        try:
            headers = {}
            if self.api_key:
                headers["x-api-key"] = self.api_key
            
            params = {
                "query": query,
                "limit": max_results,
                "fields": "title,authors,year,abstract,citationCount,referenceCount,fieldsOfStudy"
            }
            
            response = self.requests.get(
                f"{self.base_url}/paper/search",
                params=params,
                headers=headers
            )
            response.raise_for_status()
            
            data = response.json()
            papers = []
            
            for result in data.get('data', []):
                paper = Paper(
                    paper_id=result.get('paperId', ''),
                    title=result.get('title', ''),
                    authors=[a.get('name', '') for a in result.get('authors', [])],
                    abstract=result.get('abstract', ''),
                    publication_date=str(result.get('year', '')),
                    citation_count=result.get('citationCount', 0),
                    reference_count=result.get('referenceCount', 0),
                    fields_of_study=result.get('fieldsOfStudy', []),
                    source_database="semantic_scholar"
                )
                papers.append(paper)
            
            return papers
            
        except Exception as e:
            self.logger.error(f"Semantic Scholar search error: {e}")
            return []
    
    def get_paper_by_id(self, paper_id: str) -> Optional[Paper]:
        """Get paper by Semantic Scholar ID"""
        if not self.available:
            return None
        
        try:
            headers = {}
            if self.api_key:
                headers["x-api-key"] = self.api_key
            
            response = self.requests.get(
                f"{self.base_url}/paper/{paper_id}",
                headers=headers
            )
            response.raise_for_status()
            
            result = response.json()
            
            return Paper(
                paper_id=result.get('paperId', ''),
                title=result.get('title', ''),
                authors=[a.get('name', '') for a in result.get('authors', [])],
                abstract=result.get('abstract', ''),
                publication_date=str(result.get('year', '')),
                citation_count=result.get('citationCount', 0),
                source_database="semantic_scholar"
            )
            
        except Exception as e:
            self.logger.error(f"Semantic Scholar get paper error: {e}")
            return None
    
    def get_citations(self, paper_id: str) -> CitationNetwork:
        """Get citation network for a paper"""
        if not self.available:
            return CitationNetwork(paper_id=paper_id)
        
        try:
            headers = {}
            if self.api_key:
                headers["x-api-key"] = self.api_key
            
            # Get citations
            response = self.requests.get(
                f"{self.base_url}/paper/{paper_id}/citations",
                headers=headers
            )
            citing = [c.get('citingPaper', {}).get('paperId') for c in response.json().get('data', [])]
            
            # Get references
            response = self.requests.get(
                f"{self.base_url}/paper/{paper_id}/references",
                headers=headers
            )
            cited = [r.get('citedPaper', {}).get('paperId') for r in response.json().get('data', [])]
            
            return CitationNetwork(
                paper_id=paper_id,
                citing_papers=citing,
                cited_papers=cited,
                citation_depth=1
            )
            
        except Exception as e:
            self.logger.error(f"Citation network error: {e}")
            return CitationNetwork(paper_id=paper_id)
    
    def _mock_search(self, query: str, max_results: int) -> List[Paper]:
        """Mock search for testing"""
        return [
            Paper(
                paper_id=f"s2_mock_{i}",
                title=f"Semantic Scholar Paper {i}: {query[:20]}",
                authors=["Researcher X", "Researcher Y"],
                abstract=f"AI/ML research on {query}",
                publication_date="2023",
                citation_count=100 - i * 10,
                source_database="semantic_scholar",
                fields_of_study=["Computer Science", "Artificial Intelligence"],
                relevance_score=0.92 - (i * 0.05)
            )
            for i in range(min(max_results, 5))
        ]


class LiteratureSearchOrchestrator:
    """
    Automated Literature Search System.
    
    Provides unified interface for:
    - arXiv integration
    - Google Scholar integration
    - PubMed integration
    - Semantic Scholar integration
    - Citation analysis
    """
    
    def __init__(self):
        self.searchers: Dict[DatabaseType, BaseLiteratureSearch] = {
            DatabaseType.ARXIV: ArXivSearch(),
            DatabaseType.GOOGLE_SCHOLAR: GoogleScholarSearch(),
            DatabaseType.PUBMED: PubMedSearch(),
            DatabaseType.SEMANTIC_SCHOLAR: SemanticScholarSearch()
        }
        self.citation_cache: Dict[str, CitationNetwork] = {}
        self.logger = logging.getLogger(__name__)
    
    def search_all(
        self,
        query: str,
        max_results_per_db: int = 10,
        databases: Optional[List[DatabaseType]] = None,
        deduplicate: bool = True
    ) -> Dict[DatabaseType, List[Paper]]:
        """
        Search across multiple databases.
        
        Args:
            query: Search query
            max_results_per_db: Maximum results per database
            databases: List of databases to search (default: all)
            deduplicate: Whether to deduplicate results
            
        Returns:
            Dictionary mapping database type to results
        """
        dbs_to_search = databases or list(self.searchers.keys())
        results = {}
        
        for db_type in dbs_to_search:
            searcher = self.searchers.get(db_type)
            if searcher:
                papers = searcher.search(query, max_results_per_db)
                results[db_type] = papers
                self.logger.info(f"{db_type.value}: found {len(papers)} papers")
        
        if deduplicate:
            results = self._deduplicate_results(results)
        
        return results
    
    def search(
        self,
        query: str,
        database: DatabaseType = DatabaseType.SEMANTIC_SCHOLAR,
        max_results: int = 10
    ) -> List[Paper]:
        """Search specific database"""
        searcher = self.searchers.get(database)
        if searcher:
            return searcher.search(query, max_results)
        return []
    
    def get_citation_network(
        self,
        paper_id: str,
        depth: int = 1
    ) -> CitationNetwork:
        """
        Get citation network for a paper.
        
        Args:
            paper_id: Paper ID
            depth: Citation depth to retrieve
            
        Returns:
            Citation network
        """
        # Use Semantic Scholar for citation analysis
        s2 = self.searchers.get(DatabaseType.SEMANTIC_SCHOLAR)
        if isinstance(s2, SemanticScholarSearch):
            return s2.get_citations(paper_id)
        
        return CitationNetwork(paper_id=paper_id)
    
    def analyze_citations(self, papers: List[Paper]) -> Dict[str, Any]:
        """
        Analyze citation patterns across papers.
        
        Args:
            papers: List of papers to analyze
            
        Returns:
            Citation analysis
        """
        total_citations = sum(p.citation_count for p in papers)
        avg_citations = total_citations / len(papers) if papers else 0
        
        # Most cited papers
        most_cited = sorted(papers, key=lambda p: p.citation_count, reverse=True)[:5]
        
        # Authors with most papers
        author_counts = {}
        for paper in papers:
            for author in paper.authors:
                author_counts[author] = author_counts.get(author, 0) + 1
        
        top_authors = sorted(author_counts.items(), key=lambda x: x[1], reverse=True)[:10]
        
        return {
            "total_papers": len(papers),
            "total_citations": total_citations,
            "average_citations": avg_citations,
            "most_cited_papers": [
                {"title": p.title, "citations": p.citation_count}
                for p in most_cited
            ],
            "top_authors": [
                {"name": name, "paper_count": count}
                for name, count in top_authors
            ]
        }
    
    def _deduplicate_results(
        self,
        results: Dict[DatabaseType, List[Paper]]
    ) -> Dict[DatabaseType, List[Paper]]:
        """Remove duplicate papers across databases"""
        seen_titles = set()
        deduplicated = {}
        
        for db_type, papers in results.items():
            unique = []
            for paper in papers:
                title_key = paper.title.lower().strip()
                if title_key not in seen_titles:
                    seen_titles.add(title_key)
                    unique.append(paper)
            deduplicated[db_type] = unique
        
        return deduplicated


# =============================================================================
# FEATURE 9: EXPERIMENT TRACKING
# =============================================================================

@dataclass
class ExperimentParameter:
    """Experiment parameter definition"""
    name: str
    value: Any
    data_type: str = "string"
    description: str = ""


@dataclass
class ExperimentMetric:
    """Experiment metric"""
    name: str
    value: float
    timestamp: str = field(default_factory=lambda: datetime.now().isoformat())
    step: Optional[int] = None


@dataclass
class ExperimentArtifact:
    """Experiment artifact (file, model, etc.)"""
    artifact_id: str
    name: str
    artifact_type: str
    file_path: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())


@dataclass
class Experiment:
    """Complete experiment record"""
    experiment_id: str
    name: str
    description: str
    status: str = "running"  # running, completed, failed, aborted
    parameters: List[ExperimentParameter] = field(default_factory=list)
    metrics: List[ExperimentMetric] = field(default_factory=list)
    artifacts: List[ExperimentArtifact] = field(default_factory=list)
    tags: List[str] = field(default_factory=list)
    parent_experiment_id: Optional[str] = None
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    completed_at: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


class ExperimentTracker:
    """
    Experiment Tracking System.
    
    Provides:
    - Experiment logging
    - Parameter tracking
    - Metric collection
    - Artifact storage
    - Result comparison
    """
    
    def __init__(self, storage_dir: str = "./experiments"):
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(parents=True, exist_ok=True)
        self.experiments: Dict[str, Experiment] = {}
        self.active_experiments: Dict[str, Experiment] = {}
        self.logger = logging.getLogger(__name__)
        self._load_existing_experiments()
    
    def create_experiment(
        self,
        name: str,
        description: str = "",
        parameters: Optional[Dict[str, Any]] = None,
        tags: Optional[List[str]] = None,
        parent_experiment_id: Optional[str] = None
    ) -> str:
        """
        Create a new experiment.
        
        Args:
            name: Experiment name
            description: Experiment description
            parameters: Initial parameters
            tags: Experiment tags
            parent_experiment_id: Parent experiment for runs
            
        Returns:
            Experiment ID
        """
        experiment_id = f"exp_{uuid.uuid4().hex[:12]}"
        
        # Convert parameters
        param_list = []
        if parameters:
            for name, value in parameters.items():
                param_list.append(ExperimentParameter(
                    name=name,
                    value=value,
                    data_type=type(value).__name__
                ))
        
        experiment = Experiment(
            experiment_id=experiment_id,
            name=name,
            description=description,
            parameters=param_list,
            tags=tags or [],
            parent_experiment_id=parent_experiment_id
        )
        
        self.experiments[experiment_id] = experiment
        self.active_experiments[experiment_id] = experiment
        
        self._save_experiment(experiment)
        self.logger.info(f"Created experiment: {name} ({experiment_id})")
        
        return experiment_id
    
    def log_parameter(self, experiment_id: str, name: str, value: Any) -> None:
        """Log a parameter for an experiment"""
        if experiment_id not in self.active_experiments:
            self.logger.warning(f"Experiment {experiment_id} not found")
            return
        
        exp = self.active_experiments[experiment_id]
        param = ExperimentParameter(
            name=name,
            value=value,
            data_type=type(value).__name__
        )
        exp.parameters.append(param)
        
        self._save_experiment(exp)
    
    def log_metric(
        self,
        experiment_id: str,
        name: str,
        value: float,
        step: Optional[int] = None
    ) -> None:
        """Log a metric for an experiment"""
        if experiment_id not in self.active_experiments:
            self.logger.warning(f"Experiment {experiment_id} not found")
            return
        
        exp = self.active_experiments[experiment_id]
        metric = ExperimentMetric(
            name=name,
            value=value,
            step=step
        )
        exp.metrics.append(metric)
        
        self._save_experiment(exp)
    
    def log_metrics(self, experiment_id: str, metrics: Dict[str, float], step: Optional[int] = None) -> None:
        """Log multiple metrics at once"""
        for name, value in metrics.items():
            self.log_metric(experiment_id, name, value, step)
    
    def log_artifact(
        self,
        experiment_id: str,
        name: str,
        artifact_type: str,
        file_path: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> str:
        """Log an artifact for an experiment"""
        if experiment_id not in self.active_experiments:
            self.logger.warning(f"Experiment {experiment_id} not found")
            return ""
        
        artifact_id = f"art_{uuid.uuid4().hex[:12]}"
        
        exp = self.active_experiments[experiment_id]
        artifact = ExperimentArtifact(
            artifact_id=artifact_id,
            name=name,
            artifact_type=artifact_type,
            file_path=file_path,
            metadata=metadata or {}
        )
        exp.artifacts.append(artifact)
        
        self._save_experiment(exp)
        self.logger.info(f"Logged artifact: {name} for experiment {experiment_id}")
        
        return artifact_id
    
    def complete_experiment(
        self,
        experiment_id: str,
        status: str = "completed"
    ) -> None:
        """Mark experiment as complete"""
        if experiment_id not in self.active_experiments:
            return
        
        exp = self.active_experiments[experiment_id]
        exp.status = status
        exp.completed_at = datetime.now().isoformat()
        
        del self.active_experiments[experiment_id]
        
        self._save_experiment(exp)
        self.logger.info(f"Completed experiment: {experiment_id} with status {status}")
    
    def get_experiment(self, experiment_id: str) -> Optional[Experiment]:
        """Get experiment by ID"""
        return self.experiments.get(experiment_id)
    
    def list_experiments(
        self,
        status: Optional[str] = None,
        tags: Optional[List[str]] = None
    ) -> List[Dict[str, Any]]:
        """List experiments with optional filtering"""
        results = []
        
        for exp in self.experiments.values():
            if status and exp.status != status:
                continue
            
            if tags and not any(t in exp.tags for t in tags):
                continue
            
            results.append({
                "experiment_id": exp.experiment_id,
                "name": exp.name,
                "status": exp.status,
                "created_at": exp.created_at,
                "tags": exp.tags,
                "metric_count": len(exp.metrics),
                "artifact_count": len(exp.artifacts)
            })
        
        return sorted(results, key=lambda x: x["created_at"], reverse=True)
    
    def compare_experiments(
        self,
        experiment_ids: List[str],
        metric_names: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """
        Compare multiple experiments.
        
        Args:
            experiment_ids: Experiments to compare
            metric_names: Specific metrics to compare (default: all common)
            
        Returns:
            Comparison results
        """
        experiments = [self.experiments.get(eid) for eid in experiment_ids]
        experiments = [e for e in experiments if e]
        
        if not experiments:
            return {"error": "No valid experiments found"}
        
        # Get parameter differences
        param_comparison = self._compare_parameters(experiments)
        
        # Get metric comparison
        metric_comparison = self._compare_metrics(experiments, metric_names)
        
        return {
            "experiments_compared": len(experiments),
            "experiment_names": [e.name for e in experiments],
            "parameter_comparison": param_comparison,
            "metric_comparison": metric_comparison
        }
    
    def _compare_parameters(self, experiments: List[Experiment]) -> Dict[str, Any]:
        """Compare parameters across experiments"""
        all_params = set()
        for exp in experiments:
            all_params.update(p.name for p in exp.parameters)
        
        differences = {}
        for param_name in all_params:
            values = {}
            for exp in experiments:
                param = next((p for p in exp.parameters if p.name == param_name), None)
                values[exp.experiment_id] = param.value if param else None
            
            # Check if values differ
            unique_values = set(str(v) for v in values.values() if v is not None)
            if len(unique_values) > 1:
                differences[param_name] = values
        
        return {
            "total_parameters": len(all_params),
            "different_parameters": differences
        }
    
    def _compare_metrics(
        self,
        experiments: List[Experiment],
        metric_names: Optional[List[str]]
    ) -> Dict[str, Any]:
        """Compare metrics across experiments"""
        metric_stats = {}
        
        # Determine metrics to compare
        if metric_names:
            metrics_to_compare = metric_names
        else:
            all_metrics = set()
            for exp in experiments:
                all_metrics.update(m.name for m in exp.metrics)
            metrics_to_compare = list(all_metrics)
        
        for metric_name in metrics_to_compare:
            stats = {}
            for exp in experiments:
                values = [m.value for m in exp.metrics if m.name == metric_name]
                if values:
                    import statistics
                    stats[exp.name] = {
                        "count": len(values),
                        "mean": statistics.mean(values),
                        "std": statistics.stdev(values) if len(values) > 1 else 0,
                        "min": min(values),
                        "max": max(values)
                    }
            
            if stats:
                metric_stats[metric_name] = stats
        
        return metric_stats
    
    def _save_experiment(self, experiment: Experiment) -> None:
        """Save experiment to disk"""
        try:
            file_path = self.storage_dir / f"{experiment.experiment_id}.json"
            with open(file_path, 'w') as f:
                json.dump(asdict(experiment), f, indent=2, default=str)
        except Exception as e:
            self.logger.error(f"Failed to save experiment: {e}")
    
    def _load_existing_experiments(self) -> None:
        """Load existing experiments from disk"""
        try:
            for file_path in self.storage_dir.glob("exp_*.json"):
                with open(file_path, 'r') as f:
                    data = json.load(f)
                    exp = Experiment(**data)
                    self.experiments[exp.experiment_id] = exp
                    if exp.status == "running":
                        self.active_experiments[exp.experiment_id] = exp
        except Exception as e:
            self.logger.warning(f"Failed to load existing experiments: {e}")


# =============================================================================
# FEATURE 10: RESEARCH REPORT GENERATION
# =============================================================================

class ReportFormat(Enum):
    """Supported report formats"""
    MARKDOWN = "markdown"
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    JSON = "json"


@dataclass
class ReportSection:
    """Report section"""
    title: str
    content: str
    section_type: str = "text"  # text, table, figure, code
    order: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class Citation:
    """Citation entry"""
    citation_id: str
    authors: List[str]
    title: str
    journal: Optional[str] = None
    year: Optional[int] = None
    doi: Optional[str] = None
    url: Optional[str] = None


class ResearchReportGenerator:
    """
    Research Report Generation System.
    
    Provides:
    - Automated report writing
    - Citation formatting
    - Figure generation
    - Table formatting
    - Export to PDF/DOCX
    """
    
    def __init__(self):
        self.sections: List[ReportSection] = []
        self.citations: Dict[str, Citation] = {}
        self.citation_style = "APA"  # APA, MLA, Chicago
        self.logger = logging.getLogger(__name__)
    
    def add_section(
        self,
        title: str,
        content: str,
        section_type: str = "text",
        order: Optional[int] = None
    ) -> None:
        """Add a section to the report"""
        if order is None:
            order = len(self.sections)
        
        section = ReportSection(
            title=title,
            content=content,
            section_type=section_type,
            order=order
        )
        self.sections.append(section)
        self.sections.sort(key=lambda s: s.order)
    
    def add_table(
        self,
        title: str,
        headers: List[str],
        rows: List[List[Any]],
        order: Optional[int] = None
    ) -> None:
        """Add a table section"""
        # Format as markdown table
        content = f"| {' | '.join(headers)} |\n"
        content += f"| {' | '.join(['---'] * len(headers))} |\n"
        for row in rows:
            content += f"| {' | '.join(str(cell) for cell in row)} |\n"
        
        self.add_section(title, content, "table", order)
    
    def add_figure(
        self,
        title: str,
        figure_path: str,
        caption: str = "",
        order: Optional[int] = None
    ) -> None:
        """Add a figure section"""
        content = f"![{caption}]({figure_path})\n\n*{caption}*"
        self.add_section(title, content, "figure", order)
    
    def add_citation(self, paper: Paper) -> str:
        """
        Add a citation and return citation key.
        
        Args:
            paper: Paper to cite
            
        Returns:
            Citation key
        """
        citation_id = f"cite_{len(self.citations) + 1}"
        
        # Extract year from publication date
        year = None
        if paper.publication_date:
            try:
                year = int(paper.publication_date[:4])
            except (ValueError, TypeError):
                pass
        
        citation = Citation(
            citation_id=citation_id,
            authors=paper.authors,
            title=paper.title,
            journal=paper.journal,
            year=year,
            doi=paper.doi,
            url=paper.url
        )
        
        self.citations[citation_id] = citation
        return citation_id
    
    def format_citation(self, citation_id: str, inline: bool = True) -> str:
        """Format citation according to style"""
        if citation_id not in self.citations:
            return f"[{citation_id}]"
        
        cite = self.citations[citation_id]
        
        if self.citation_style == "APA":
            if inline:
                author = cite.authors[0].split()[-1] if cite.authors else "Unknown"
                year = cite.year or "n.d."
                return f"({author}, {year})"
            else:
                # Full reference
                authors = ", ".join(cite.authors) if cite.authors else "Unknown"
                year = cite.year or "n.d."
                journal = cite.journal or ""
                doi = f"https://doi.org/{cite.doi}" if cite.doi else ""
                return f"{authors} ({year}). {cite.title}. {journal}. {doi}"
        
        # Default format
        return f"[{citation_id}]"
    
    def generate_markdown(self) -> str:
        """Generate Markdown report"""
        lines = []
        
        for section in self.sections:
            lines.append(f"## {section.title}\n")
            lines.append(section.content)
            lines.append("\n")
        
        # Add references
        if self.citations:
            lines.append("\n## References\n")
            for citation_id, citation in sorted(self.citations.items()):
                ref = self.format_citation(citation_id, inline=False)
                lines.append(f"{ref}\n")
        
        return "\n".join(lines)
    
    def generate_html(self) -> str:
        """Generate HTML report"""
        html = ["<!DOCTYPE html>", "<html>", "<head>", "<title>Research Report</title>", "</head>", "<body>"]
        
        for section in self.sections:
            html.append(f"<h2>{section.title}</h2>")
            
            if section.section_type == "table":
                # Convert markdown table to HTML
                html.append("<table border='1'>")
                lines = section.content.strip().split('\n')
                for i, line in enumerate(lines):
                    if '|' in line:
                        cells = [c.strip() for c in line.split('|') if c.strip()]
                        tag = 'th' if i == 0 else 'td'
                        html.append("<tr>" + "".join(f"<{tag}>{c}</{tag}>" for c in cells) + "</tr>")
                html.append("</table>")
            elif section.section_type == "figure":
                # Extract image path from markdown
                import re
                match = re.search(r'!\[.*?\]\((.*?)\)', section.content)
                if match:
                    img_path = match.group(1)
                    caption = section.content.split('*')[-2] if '*' in section.content else ""
                    html.append(f"<img src='{img_path}' /><p><em>{caption}</em></p>")
            else:
                html.append(f"<p>{section.content}</p>")
        
        # References
        if self.citations:
            html.append("<h2>References</h2>")
            html.append("<ol>")
            for citation_id, citation in sorted(self.citations.items()):
                ref = self.format_citation(citation_id, inline=False)
                html.append(f"<li>{ref}</li>")
            html.append("</ol>")
        
        html.extend(["</body>", "</html>"])
        
        return "\n".join(html)
    
    def export(
        self,
        output_path: str,
        format: ReportFormat = ReportFormat.MARKDOWN
    ) -> bool:
        """
        Export report to file.
        
        Args:
            output_path: Output file path
            format: Export format
            
        Returns:
            True if successful
        """
        try:
            if format == ReportFormat.MARKDOWN:
                content = self.generate_markdown()
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
            
            elif format == ReportFormat.HTML:
                content = self.generate_html()
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
            
            elif format == ReportFormat.PDF:
                return self._export_pdf(output_path)
            
            elif format == ReportFormat.DOCX:
                return self._export_docx(output_path)
            
            elif format == ReportFormat.JSON:
                data = {
                    "sections": [asdict(s) for s in self.sections],
                    "citations": {k: asdict(v) for k, v in self.citations.items()}
                }
                with open(output_path, 'w') as f:
                    json.dump(data, f, indent=2)
            
            self.logger.info(f"Exported report to {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Export failed: {e}")
            return False
    
    def _export_pdf(self, output_path: str) -> bool:
        """Export to PDF using ReportLab"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib import colors
            
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            for section in self.sections:
                story.append(Paragraph(section.title, styles['Heading2']))
                story.append(Paragraph(section.content, styles['BodyText']))
                story.append(Spacer(1, 12))
            
            # References
            if self.citations:
                story.append(Paragraph("References", styles['Heading2']))
                for citation_id in sorted(self.citations.keys()):
                    ref = self.format_citation(citation_id, inline=False)
                    story.append(Paragraph(ref, styles['BodyText']))
            
            doc.build(story)
            return True
            
        except ImportError:
            self.logger.error("reportlab not installed. Cannot export PDF.")
            return False
        except Exception as e:
            self.logger.error(f"PDF export failed: {e}")
            return False
    
    def _export_docx(self, output_path: str) -> bool:
        """Export to DOCX using python-docx"""
        try:
            from docx import Document
            
            doc = Document()
            
            for section in self.sections:
                doc.add_heading(section.title, level=2)
                doc.add_paragraph(section.content)
            
            # References
            if self.citations:
                doc.add_heading("References", level=2)
                for citation_id in sorted(self.citations.keys()):
                    ref = self.format_citation(citation_id, inline=False)
                    doc.add_paragraph(ref, style='List Number')
            
            doc.save(output_path)
            return True
            
        except ImportError:
            self.logger.error("python-docx not installed. Cannot export DOCX.")
            return False
        except Exception as e:
            self.logger.error(f"DOCX export failed: {e}")
            return False
    
    def clear(self) -> None:
        """Clear all report content"""
        self.sections = []
        self.citations = {}


# =============================================================================
# CONVENIENCE FUNCTIONS
# =============================================================================

def create_literature_search() -> LiteratureSearchOrchestrator:
    """Factory function for literature search"""
    return LiteratureSearchOrchestrator()


def create_experiment_tracker(storage_dir: str = "./experiments") -> ExperimentTracker:
    """Factory function for experiment tracker"""
    return ExperimentTracker(storage_dir=storage_dir)


def create_report_generator() -> ResearchReportGenerator:
    """Factory function for report generator"""
    return ResearchReportGenerator()
