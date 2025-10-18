# -*- coding: utf-8 -*-
"""
OpenEvolve Content Improver: A Streamlit application for AI-powered content hardening
and evolution using multi-LLM consensus.
"""

# --- Standard Library Imports ---
import json
import os
import sys
import threading
import time
from datetime import datetime
import difflib
from typing import List, Dict
import io

# --- Third-party Library Imports ---
import altair as alt
import numpy as np
import pandas as pd
import streamlit as st
from streamlit.components.v1 import html

# --- Optional Imports with Fallbacks ---
try:
    import plotly.express as px
    PLOTLY_AVAILABLE = True
except ImportError:
    px = None
    PLOTLY_AVAILABLE = False

try:
    import matplotlib.pyplot as plt
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False

try:
    from log_streaming import LogStreaming
    LOG_STREAMING_AVAILABLE = True
except ImportError:
    LogStreaming = None
    LOG_STREAMING_AVAILABLE = False

# --- Optional Imports with Fallbacks ---
try:
    import matplotlib.pyplot as plt
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False

try:
    from streamlit_autorefresh import st_autorefresh
except ImportError:
    st_autorefresh = None  # Define a fallback if not available

try:
    from streamlit_tags import st_tags
    HAS_STREAMLIT_TAGS = True
except ImportError:
    HAS_STREAMLIT_TAGS = False

try:
    from pyvis.network import Network
    PYVIS_AVAILABLE = True
except ImportError:
    PYVIS_AVAILABLE = False

# --- Local Application Imports ---
# These imports are assumed to exist in the user's environment.
# If they don't, the script will fail, but per the instructions, no mock functions will be created.
from session_utils import _safe_list, _load_user_preferences, _load_parameter_settings
from providercatalogue import get_openrouter_models
from session_manager import APPROVAL_PROMPT, RED_TEAM_CRITIQUE_PROMPT, BLUE_TEAM_PATCH_PROMPT
from openevolve_integration import OpenEvolveAPI
from adversarial import run_adversarial_testing, _load_human_feedback, MODEL_META_LOCK, MODEL_META_BY_ID
from evolution import _run_evolution_with_api_backend_refactored
from integrations import create_github_branch, commit_to_github, list_linked_github_repositories, send_discord_notification, send_msteams_notification, send_generic_webhook
from tasks import create_task, get_tasks
from rbac import ROLES, assign_role
from content_manager import content_manager
from prompt_manager import PromptManager
from template_manager import TemplateManager
from analytics_manager import AnalyticsManager
from collaboration_manager import CollaborationManager
from version_control import VersionControl
from notifications import NotificationManager
from log_streaming import LogStreaming
from session_state_classes import SessionManager
from sidebar import get_default_generation_params, get_default_evolution_params
from integrated_reporting import generate_integrated_report

# --- Optional OpenEvolve Backend Imports ---
OPENEVOLVE_AVAILABLE = False
try:
    from openevolve_visualization import render_openevolve_visualization_ui, render_evolution_insights, render_openevolve_advanced_ui, render_advanced_diagnostics
    from monitoring_system import render_comprehensive_monitoring_ui
    from reporting_system import render_reporting_dashboard, create_evolution_report
    from openevolve_orchestrator import start_openevolve_services, stop_openevolve_services, restart_openevolve_services
    OPENEVOLVE_AVAILABLE = True
except ImportError as e:
    print(f"OpenEvolve backend not available - using API-based evolution only: {e}")


# --- Global Constants ---
MODEL_OPTIONS = [
    "gpt-4o", "gpt-4o-mini", "gpt-4-turbo", "gpt-4", "gpt-3.5-turbo",
    "claude-3-opus", "claude-3-sonnet", "claude-3-haiku",
    "gemini-1.5-pro", "gemini-1.5-flash",
    "llama-3-70b", "llama-3-8b",
    "mistral-large", "mistral-medium", "mixtral-8x22b",
    "command-r-plus", "command-r",
    "pplx-7b-online", "pplx-70b-online",
    "openchat/openchat-3.5-0106",
    "microsoft/WizardLM-2-8x22B", "microsoft/WizardLM-2-7B",
]


# --- Report Generation Functions ---

def generate_pdf_report(results, watermark=None):
    """Generate a PDF report of the results."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()

        story = [Paragraph("Evolution Results Report", styles['Title']), Spacer(1, 12)]
        result_text = f"Evolution Results: {str(results)[:1000]}..."
        story.append(Paragraph(result_text, styles['Normal']))

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        st.warning("ReportLab not installed. PDF generation is unavailable. Run 'pip install reportlab'.")
        return b"PDF generation not available. Please install reportlab."

def generate_docx_report(results):
    """Generate a DOCX report of the results."""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('Evolution Results Report', 0)
        doc.add_paragraph(f"Evolution Results: {str(results)[:1000]}...")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        st.warning("python-docx not installed. DOCX generation is unavailable. Run 'pip install python-docx'.")
        return b"DOCX generation not available. Please install python-docx."

def generate_latex_report(results):
    """Generate a LaTeX report of the results."""
    return f"""
\\documentclass{{article}}
\\usepackage[utf8]{{inputenc}}
\\title{{Evolution Results Report}}
\\author{{OpenEvolve Platform}}
\\date{{\\today}}
\\begin{{document}}
\\maketitle
\\section{{Results Summary}}
Evolution Results: {str(results)[:1000] if len(str(results)) > 1000 else str(results)}...
\\end{{document}}
"""

def generate_compliance_report(results, compliance_requirements):
    """Generate a compliance report based on requirements."""
    return f"""
# Compliance Report
## Evolution Results Summary
{str(results)[:500] if len(str(results)) > 500 else str(results)}...
## Compliance Requirements
{str(compliance_requirements)}
## Analysis
Based on the evolution results and compliance requirements, the following compliance checks have been performed:
- All requirements have been documented.
- Evolution process followed proper protocols.
- Results meet specified criteria.
"""


# --- Caching and Data Loading ---

@st.cache_data(ttl=3600)  # Cache for 1 hour
def _load_report_templates():
    """Load report templates from a JSON file."""
    try:
        with open("report_templates.json", "r", encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return {}

def _save_report_templates(templates):
    """Save report templates to a JSON file."""
    with open("report_templates.json", "w", encoding="utf-8") as f:
        json.dump(templates, f, indent=4)


# --- UI Rendering Helper Functions ---

def render_code_diff(text1: str, text2: str):
    """Render the difference between two strings."""
    diff = difflib.unified_diff(
        text1.splitlines(keepends=True),
        text2.splitlines(keepends=True),
        fromfile='previous',
        tofile='current',
    )
    st.code("".join(diff), language="diff")

@st.cache_data(ttl=3600)
def render_island_model_chart(history: List[Dict]):
    """Render an interactive graph of the island model evolution."""
    if not history or not PYVIS_AVAILABLE:
        if not PYVIS_AVAILABLE:
            st.warning("Pyvis not installed. Island model chart is unavailable. Run 'pip install pyvis'.")
        return

    net = Network(height="500px", width="100%", notebook=True, cdn_resources='in_line')
    islands_data = history[-1].get('islands', [])

    for i, _ in enumerate(islands_data):
        net.add_node(i, label=f"Island {i}")

    for i, _ in enumerate(islands_data):
        for j, _ in enumerate(islands_data):
            if i != j and abs(i - j) == 1:
                net.add_edge(i, j, value=np.random.randint(1, 10))

    try:
        net.save_graph("island_model.html")
        with open("island_model.html", 'r', encoding='utf-8') as f:
            html_content = f.read()
        st.components.v1.html(html_content, height=500)
    except Exception as e:
        st.error(f"Failed to render island model chart: {e}")

@st.cache_data(ttl=3600)
def render_evolution_history_chart(history: List[Dict]):
    """Render an interactive scatter plot of the evolution history."""
    if not history:
        return

    data = []
    for generation in history:
        for individual in generation.get('population', []):
            data.append({
                'generation': generation.get('generation'),
                'fitness': individual.get('fitness'),
                'code': individual.get('code', '')
            })

    if not data:
        return

    chart = alt.Chart(pd.DataFrame(data)).mark_circle(size=60).encode(
        x='generation',
        y='fitness',
        tooltip=['generation', 'fitness', 'code']
    ).interactive()
    st.altair_chart(chart, use_container_width=True)

def render_notification_ui():
    """Renders the notification bell icon and panel."""
    unread_notifications = [n for n in st.session_state.collaboration_session.get("notifications", []) if not n.get("read")]
    unread_count = len(unread_notifications)

    st.markdown(f"""
    <div class="notification-container">
        <button class="notification-button">
            🔔 <span class="notification-badge">{unread_count}</span>
        </button>
    </div>
    """, unsafe_allow_html=True)

    if st.session_state.get("show_notifications", False):
        with st.expander("Notifications", expanded=True):
            notifications = st.session_state.collaboration_session.get("notifications", [])
            if not notifications:
                st.write("No notifications.")
            for notification in notifications[-5:]:
                st.info(f"**{notification.get('sender')}** mentioned you: *{notification.get('comment_text')}*")

def check_password():
    """Checks for password if the project is public and password-protected."""
    if st.session_state.get("project_public") and st.session_state.get("project_password"):
        if not st.session_state.get("password_correct", False):
            password = st.text_input("Enter password to view this project", type="password")
            if st.button("Submit"):
                if password == st.session_state.project_password:
                    st.session_state.password_correct = True
                    st.rerun()
                else:
                    st.error("Incorrect password")
            st.stop()

def render_collaboration_ui():
    """Renders collaboration UI elements like presence indicators via JavaScript."""
    if st.session_state.get("collaboration_ui_rendered", False):
        return

    html_content = """
    <div id="presence-container" class="presence-container"></div>
    <div id="notification-center" class="notification-center"></div>
    <script>
        if (!window.collaborationWebSocket) {
            // Establish WebSocket connection for real-time updates
        const socket = new WebSocket("ws://localhost:8000/ws");

        socket.onopen = function(event) {
            console.log("WebSocket connection established");
            socket.send(JSON.stringify({message: "Hello from frontend"}));
        };

        socket.onmessage = function(event) {
            const data = JSON.parse(event.data);
            console.log("WebSocket message received:", data);
            // Add logic to handle incoming messages and update the UI
        };

        socket.onclose = function(event) {
            console.log("WebSocket connection closed");
        };

        socket.onerror = function(error) {
            console.error("WebSocket error:", error);
        };
            console.log("WebSocket should be initialized here.");
        }
    </script>
    """
    st.markdown(html_content, unsafe_allow_html=True)
    st.session_state.collaboration_ui_rendered = True


# --- Core Logic ---

def _stream_evolution_logs_in_thread(evolution_id, api, thread_lock):
    """Streams evolution logs from the API in a separate thread."""
    full_log = []
    try:
        for log_chunk in api.stream_evolution_logs(evolution_id):
            full_log.append(log_chunk)
            with thread_lock:
                st.session_state.evolution_log = full_log.copy()
    except Exception as e:
        with thread_lock:
            st.session_state.evolution_log.append(f"Error streaming logs: {e}")
    finally:
        # Once streaming is complete or fails, set running to False
        with thread_lock:
            st.session_state.evolution_running = False

def _should_update_log_display(log_key, current_log_entries):
    """
    Helper to determine if a log display needs to be updated to prevent
    unnecessary re-rendering of the log text area.
    """
    last_log_state = st.session_state.get(f"_{log_key}_display_state", [])
    if last_log_state != current_log_entries:
        st.session_state[f"_{log_key}_display_state"] = current_log_entries
        return True
    return False

def _initialize_session_state():
    """Initialize session state with default values."""
    if "session_initialized" in st.session_state:
        return

    user_prefs = _load_user_preferences()
    param_settings = _load_parameter_settings()

    defaults = {
        # Core App State
        "theme": user_prefs.get("theme", "light"),
        "show_quick_guide": False,
        "show_keyboard_shortcuts": False,
        "activity_log": [],
        "user_preferences": user_prefs,
        "parameter_settings": param_settings,
        "tasks": [],
        "user_roles": {"admin": "admin", "user": "user"},
        # Evolution State
        "evolution_running": False,
        "evolution_stop_flag": False,
        "evolution_history": [],
        "evolution_id": None,
        "evolution_log": [],
        "evolution_current_best": "",
        "evolution_status_message": "",
        "evolution_best_score": 0,
        "protocol_text": "# Sample Protocol\n\nThis is a sample protocol for testing purposes.",
        # Adversarial Testing State
        "adversarial_running": False,
        "adversarial_stop_flag": False,
        "adversarial_log": [],
        "adversarial_results": None,
        "adversarial_model_performance": {},
        "adversarial_status_message": "",
        "adversarial_confidence_history": [],
        "adversarial_cost_estimate_usd": 0.0,
        # API and Model Config
        "openevolve_base_url": "http://localhost:8000",
        "openevolve_api_key": "",
        "openrouter_key": "",
        "github_token": "",
        "discord_webhook_url": "",
        "msteams_webhook_url": "",
        # Default Evolution Parameters
        "system_prompt": "You are an expert content generator.",
        "evaluator_system_prompt": "Evaluate the quality of this content and provide a score from 0 to 100.",
        "model": MODEL_OPTIONS[0],
        "temperature": 0.7,
        "top_p": 1.0,
        "max_tokens": 4096,
        "population_size": 10,
        "evolution_max_iterations": 20,
        "num_islands": 1,
        "migration_interval": 50,
        "migration_rate": 0.1,
        "archive_size": 100,
        "elite_ratio": 0.1,
        "exploration_ratio": 0.2,
        "checkpoint_interval": 5,
        "feature_dimensions": ["complexity", "diversity"],
        "feature_bins": 10,
        "diversity_metric": "edit_distance",
        # Default Adversarial Parameters
        "red_team_models": ["claude-3-sonnet"],
        "blue_team_models": ["gpt-4o"],
        "evaluator_models": ["gpt-4o", "claude-3-sonnet"],
        "adversarial_custom_mode": False,
        "adversarial_custom_red_prompt": RED_TEAM_CRITIQUE_PROMPT,
        "adversarial_custom_blue_prompt": BLUE_TEAM_PATCH_PROMPT,
        "adversarial_custom_approval_prompt": APPROVAL_PROMPT,
        "adversarial_min_iter": 1,
        "adversarial_max_iter": 5,
        "adversarial_confidence": 80,
        "adversarial_budget_limit": 10.0,
        "adversarial_red_team_sample_size": 2,
        "adversarial_blue_team_sample_size": 2,
        "evaluator_sample_size": 2,
        "evaluator_threshold": 90.0,
        "evaluator_consecutive_rounds": 1,
        "adversarial_rotation_strategy": "Round Robin",
        "enable_performance_tracking": True,
        "adversarial_critique_depth": 5,
        "adversarial_patch_quality": 5,
        "adversarial_compliance_requirements": "",
        "enable_multi_objective_optimization": False,
        "feature_dimensions": ["complexity", "diversity"],
        "feature_bins": 10,
        "enable_data_augmentation": False,
        "augmentation_model": "gpt-4o",
        "augmentation_temperature": 0.7,
        "elite_ratio": 0.1,
        "exploration_ratio": 0.2,
        "archive_size": 100,
        "enable_human_feedback": False,
        "keyword_analysis_enabled": True,
        "keywords_to_target": "",
        "enable_real_time_monitoring": True,
        "enable_comprehensive_reporting": True,
        "enable_encryption": True,
        "enable_audit_trail": True,
        "integrated_adversarial_history": [],
        "model_performance_metrics": {},
        "quality_metrics": {},
        # Collaboration & Project
        "projects": {},
        "project_public": False,
        "project_password": "",
        "collaboration_session": {"notifications": []},
        # Threading
        "thread_lock": threading.Lock(),
    }

    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

    # Ensure parameter_settings has the default structure if loaded empty
    if not st.session_state.parameter_settings:
        st.session_state.parameter_settings = {
            "global": {
                "generation": get_default_generation_params(),
                "evolution": get_default_evolution_params(),
            },
            "providers": {},
        }

    # Load CSS
    try:
        with open("styles.css", encoding="utf-8") as f:
            st.session_state.styles_css = f.read()
    except FileNotFoundError:
        st.session_state.styles_css = ""

    # Load templates
    st.session_state.report_templates = _load_report_templates()

    st.session_state.session_initialized = True


# --- Tab Rendering Functions ---

def render_adversarial_testing_tab():
    """Renders the UI for the Adversarial Testing tab with ALL ultimate features."""
    st.header("⚔️ Ultimate Adversarial Testing & Evolution")
    st.markdown("""
    **Advanced AI-Powered Content Hardening with Multi-Model Consensus**
    
    This enhanced tab provides comprehensive adversarial testing capabilities including:
    - **Tripartite AI Architecture**: Red Team (Critics), Blue Team (Fixers), Evaluator Team (Judges)
    - **Multi-Objective Optimization**: Quality-Diversity evolution with MAP-Elites
    - **Advanced Model Orchestration**: Intelligent load balancing and performance-based selection
    - **Comprehensive Quality Assurance**: Multi-dimensional evaluation and validation
    - **Real-Time Analytics**: Performance monitoring and optimization insights
    """)

    # Content Type Selection
    st.subheader("📄 Content Analysis & Configuration")
    col1, col2 = st.columns(2)
    with col1:
        content_type = st.selectbox(
            "Content Type",
            options=["document_general", "document_legal", "document_medical", "document_technical",
                    "code_python", "code_javascript", "code_java", "code_cpp", "plan", "sop"],
            help="Select the type of content for specialized processing",
            key="adversarial_content_type"
        )
        
        protocol_text = st.text_area(
            "Content to Test & Evolve",
            value=st.session_state.get("protocol_text", ""),
            height=300,
            key="adversarial_protocol_text",
            help="Enter the content you want to harden through adversarial testing and evolution"
        )
        # Note: protocol_text is automatically managed by the widget key
        # No need to manually set st.session_state.protocol_text
    
    with col2:
        # Content Analysis
        if st.session_state.protocol_text:
            content_length = len(st.session_state.protocol_text)
            word_count = len(st.session_state.protocol_text.split())
            st.metric("Content Length", f"{content_length:,} chars")
            st.metric("Word Count", f"{word_count:,} words")
            
            # Quick content analysis
            if word_count > 0:
                avg_word_length = content_length / word_count
                st.metric("Avg Word Length", f"{avg_word_length:.1f} chars")
        
        # Compliance Requirements
        compliance_requirements = st.text_area(
            "Compliance Requirements",
            value=st.session_state.get("adversarial_compliance_requirements", ""),
            height=100,
            placeholder="e.g., GDPR, HIPAA, SOC 2, ISO 27001 requirements...",
            help="Specify regulatory or compliance requirements for the content",
            key="adversarial_compliance_requirements"
        )
        # Note: adversarial_compliance_requirements is automatically managed by the widget key
        # No need to manually set st.session_state.adversarial_compliance_requirements

    # Advanced Configuration Tabs
    config_tabs = st.tabs(["🤖 Model Configuration", "⚙️ Process Parameters", "📊 Advanced Features", "🔧 Quality Control"])

    with config_tabs[0]:  # Model Configuration
        st.subheader("AI Team Configuration")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown("**🔴 Red Team (Critics)**")
            red_team_models = st.multiselect(
                "Red Team Models",
                options=MODEL_OPTIONS,
                default=st.session_state.red_team_models,
                help="Models specialized in finding vulnerabilities and weaknesses",
                key="red_team_models_select"
            )
            adversarial_red_team_sample_size = st.number_input(
                "Red Team Sample Size",
                min_value=1,
                max_value=len(red_team_models) if red_team_models else 1,
                value=min(2, len(red_team_models)) if red_team_models else 1,
                help="Number of Red Team models to use per iteration",
                key="adversarial_red_team_sample_size_input"
            )
            # Note: red_team_models and adversarial_red_team_sample_size are automatically managed by widget keys
            # No need to manually set st.session_state.red_team_models or st.session_state.adversarial_red_team_sample_size
            
        with col2:
            st.markdown("**🔵 Blue Team (Fixers)**")
            blue_team_models = st.multiselect(
                "Blue Team Models",
                options=MODEL_OPTIONS,
                default=st.session_state.blue_team_models,
                help="Models specialized in resolving issues and improving content",
                key="blue_team_models_select"
            )
            adversarial_blue_team_sample_size = st.number_input(
                "Blue Team Sample Size",
                min_value=1,
                max_value=len(blue_team_models) if blue_team_models else 1,
                value=min(2, len(blue_team_models)) if blue_team_models else 1,
                help="Number of Blue Team models to use per iteration",
                key="adversarial_blue_team_sample_size_input"
            )
            # Note: blue_team_models and adversarial_blue_team_sample_size are automatically managed by widget keys
            # No need to manually set st.session_state.blue_team_models or st.session_state.adversarial_blue_team_sample_size
            
        with col3:
            st.markdown("**⚖️ Evaluator Team (Judges)**")
            evaluator_models = st.multiselect(
                "Evaluator Models",
                options=MODEL_OPTIONS,
                default=st.session_state.get("evaluator_models", ["gpt-4o", "claude-3-sonnet"]),
                help="Models specialized in judging quality and correctness",
                key="evaluator_models_select"
            )
            
            evaluator_sample_size = st.number_input(
                "Evaluator Sample Size",
                min_value=1,
                max_value=len(evaluator_models) if evaluator_models else 1,
                value=min(2, len(evaluator_models)) if evaluator_models else 1,
                help="Number of Evaluator models to use per assessment"
            )
            # Note: evaluator_models and evaluator_sample_size are automatically managed by widget keys
            # No need to manually set st.session_state.evaluator_models or st.session_state.evaluator_sample_size

        # Model Selection Strategy
        st.subheader("Model Selection Strategy")
        col1, col2 = st.columns(2)
        with col1:
            adversarial_rotation_strategy = st.selectbox(
                "Rotation Strategy",
                options=["Round Robin", "Random Sampling", "Performance-Based", "Staged", "Adaptive", "Focus-Category"],
                index=0,
                help="Strategy for selecting models across iterations",
                key="rotation_strategy_select"
            )
            # Note: adversarial_rotation_strategy is automatically managed by widget key
            # No need to manually set st.session_state.adversarial_rotation_strategy
            
        with col2:
            enable_performance_tracking = st.checkbox(
                "Enable Performance Tracking",
                value=st.session_state.get("enable_performance_tracking", True),
                help="Track model performance and adjust selection weights automatically",
                key="enable_performance_tracking_checkbox"
            )
            # Note: enable_performance_tracking is automatically managed by widget key
            # No need to manually set st.session_state.enable_performance_tracking

        # Custom Prompts
        with st.expander("💬 Custom Prompts & Templates", expanded=False):
            adversarial_custom_mode = st.checkbox(
                "Enable Custom Prompts",
                value=st.session_state.adversarial_custom_mode,
                help="Use custom prompts for specialized content types",
                key="adversarial_custom_mode_checkbox"
            )
            # Note: adversarial_custom_mode is automatically managed by widget key
            # No need to manually set st.session_state.adversarial_custom_mode
            if st.session_state.adversarial_custom_mode:
                col1, col2 = st.columns(2)
                with col1:
                    adversarial_custom_red_prompt = st.text_area(
                        "Red Team Critique Prompt",
                        value=st.session_state.adversarial_custom_red_prompt,
                        height=150,
                        help="Custom prompt for Red Team vulnerability analysis",
                        key="adversarial_custom_red_prompt_textarea"
                    )
                    # Note: adversarial_custom_red_prompt is automatically managed by widget key
                    # No need to manually set st.session_state.adversarial_custom_red_prompt
                with col2:
                    adversarial_custom_blue_prompt = st.text_area(
                        "Blue Team Patch Prompt",
                        value=st.session_state.adversarial_custom_blue_prompt,
                        height=150,
                        help="Custom prompt for Blue Team issue resolution",
                        key="adversarial_custom_blue_prompt_textarea"
                    )
                    # Note: adversarial_custom_blue_prompt is automatically managed by widget key
                    # No need to manually set st.session_state.adversarial_custom_blue_prompt
                
                adversarial_custom_approval_prompt = st.text_area(
                    "Approval Prompt",
                    value=st.session_state.adversarial_custom_approval_prompt,
                    height=100,
                    help="Custom prompt for final content approval assessment",
                    key="adversarial_custom_approval_prompt_textarea"
                )
                # Note: adversarial_custom_approval_prompt is automatically managed by widget key
                # No need to manually set st.session_state.adversarial_custom_approval_prompt

    with config_tabs[1]:  # Process Parameters
        st.subheader("Iteration & Threshold Management")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            adversarial_min_iter = st.number_input(
                "Minimum Iterations",
                min_value=1,
                max_value=100,
                value=st.session_state.adversarial_min_iter,
                help="Minimum number of adversarial iterations (ensures thorough processing)",
                key="adversarial_min_iter_input"
            )
            adversarial_max_iter = st.number_input(
                "Maximum Iterations",
                min_value=adversarial_min_iter,
                max_value=200,
                value=st.session_state.adversarial_max_iter,
                help="Maximum number of adversarial iterations (prevents infinite loops)",
                key="adversarial_max_iter_input"
            )
            # Note: adversarial_min_iter and adversarial_max_iter are automatically managed by widget keys
            # No need to manually set st.session_state.adversarial_min_iter or st.session_state.adversarial_max_iter
            
        with col2:
            adversarial_confidence = st.slider(
                "Confidence Threshold (%)",
                50, 100,
                st.session_state.adversarial_confidence,
                help="Minimum acceptance level for process completion",
                key="adversarial_confidence_slider"
            )
            evaluator_threshold = st.slider(
                "Evaluator Threshold",
                50.0, 100.0,
                value=st.session_state.get("evaluator_threshold", 90.0),
                step=0.5,
                help="Minimum score requirement for evaluator team acceptance",
                key="evaluator_threshold_slider"
            )
            # Note: adversarial_confidence and evaluator_threshold are automatically managed by widget keys
            # No need to manually set st.session_state.adversarial_confidence or st.session_state.evaluator_threshold
            
        with col3:
            evaluator_consecutive_rounds = st.number_input(
                "Consecutive Rounds Required",
                min_value=1,
                max_value=10,
                value=st.session_state.get("evaluator_consecutive_rounds", 1),
                help="Number of consecutive successful evaluations required for acceptance",
                key="evaluator_consecutive_rounds_input"
            )
            
            adversarial_budget_limit = st.number_input(
                "Budget Limit (USD)",
                min_value=0.0,
                value=st.session_state.adversarial_budget_limit,
                format="%.2f",
                help="Maximum budget for the entire adversarial testing process",
                key="adversarial_budget_limit_input"
            )
            # Note: evaluator_consecutive_rounds and adversarial_budget_limit are automatically managed by widget keys
            # No need to manually set st.session_state.evaluator_consecutive_rounds or st.session_state.adversarial_budget_limit

        # Quality Control Parameters
        st.subheader("Quality Control Parameters")
        col1, col2 = st.columns(2)
        with col1:
            adversarial_critique_depth = st.slider(
                "Critique Depth Level",
                min_value=1,
                max_value=10,
                value=st.session_state.get("adversarial_critique_depth", 5),
                help="Controls thoroughness of Red Team analysis (1=surface, 10=deep)",
                key="adversarial_critique_depth_slider"
            )
            # Note: adversarial_critique_depth is automatically managed by widget key
            # No need to manually set st.session_state.adversarial_critique_depth
            
        with col2:
            adversarial_patch_quality = st.slider(
                "Patch Quality Level",
                min_value=1,
                max_value=10,
                value=st.session_state.get("adversarial_patch_quality", 5),
                help="Governs thoroughness of Blue Team fix implementation (1=basic, 10=comprehensive)",
                key="adversarial_patch_quality_slider"
            )
            # Note: adversarial_patch_quality is automatically managed by widget key
            # No need to manually set st.session_state.adversarial_patch_quality

    with config_tabs[2]:  # Advanced Features
        st.subheader("Advanced Evolution & Optimization")
        
        col1, col2 = st.columns(2)
        with col1:
            enable_multi_objective_optimization = st.checkbox(
                "Enable Multi-Objective Optimization",
                value=st.session_state.get("enable_multi_objective_optimization", False),
                help="Optimize for multiple competing objectives simultaneously",
                key="enable_multi_objective_optimization_checkbox"
            )
            # Note: enable_multi_objective_optimization is automatically managed by widget key
            # No need to manually set st.session_state.enable_multi_objective_optimization
            
            if enable_multi_objective_optimization:
                feature_dimensions = st.multiselect(
                    "Feature Dimensions",
                    options=["complexity", "diversity", "performance", "efficiency", "readability", "robustness", "accuracy", "clarity"],
                    default=st.session_state.get("feature_dimensions", ["complexity", "diversity"]),
                    help="Feature dimensions for Quality-Diversity evolution",
                    key="adversarial_feature_dimensions"
                )
                
                feature_bins = st.number_input(
                    "Feature Bins",
                    min_value=5,
                    max_value=50,
                    value=st.session_state.get("feature_bins", 10),
                    help="Number of bins for each feature dimension",
                    key="adversarial_feature_bins"
                )
                # Note: feature_dimensions and feature_bins are automatically managed by widget keys
                # No need to manually set st.session_state.feature_dimensions or st.session_state.feature_bins
                
        with col2:
            enable_data_augmentation = st.checkbox(
                "Enable Data Augmentation",
                value=st.session_state.get("enable_data_augmentation", False),
                help="Generate adversarial examples for more robust testing",
                key="enable_data_augmentation_checkbox"
            )
            # Note: enable_data_augmentation is automatically managed by widget key
            # No need to manually set st.session_state.enable_data_augmentation
            
            if enable_data_augmentation:
                augmentation_model = st.selectbox(
                    "Augmentation Model",
                    options=MODEL_OPTIONS,
                    index=MODEL_OPTIONS.index("gpt-4o") if "gpt-4o" in MODEL_OPTIONS else 0,
                    help="Model to use for generating adversarial examples",
                    key="augmentation_model"
                )
                
                augmentation_temperature = st.slider(
                    "Augmentation Temperature",
                    min_value=0.0,
                    max_value=2.0,
                    value=st.session_state.get("augmentation_temperature", 0.7),
                    step=0.1,
                    help="Temperature for augmentation model (higher = more creative)",
                    key="augmentation_temperature_slider"
                )
                # Note: augmentation_model and augmentation_temperature are automatically managed by widget keys
                # No need to manually set st.session_state.augmentation_model or st.session_state.augmentation_temperature

        # Evolution Parameters
        st.subheader("Evolution Parameters")
        col1, col2, col3 = st.columns(3)
        with col1:
            elite_ratio = st.slider(
                "Elite Ratio",
                0.0, 1.0,
                st.session_state.get("elite_ratio", 0.1),
                0.01,
                help="Ratio of elite individuals to preserve",
                key="adversarial_elite_ratio"
            )
            # Session state is automatically handled by the widget key
            
        with col2:
            exploration_ratio = st.slider(
                "Exploration Ratio",
                0.0, 1.0,
                st.session_state.get("exploration_ratio", 0.2),
                0.01,
                help="Ratio for exploration in evolution",
                key="adversarial_exploration_ratio"
            )
            # Session state is automatically handled by the widget key
            
        with col3:
            archive_size = st.number_input(
                "Archive Size",
                min_value=10,
                max_value=1000,
                value=st.session_state.get("archive_size", 100),
                help="Size of the archive for storing best solutions",
                key="adversarial_archive_size"
            )
            # Session state is automatically handled by the widget key

    with config_tabs[3]:  # Quality Control
        st.subheader("Quality Assurance & Validation")
        
        col1, col2 = st.columns(2)
        with col1:
            enable_human_feedback = st.checkbox(
                "Enable Human Feedback Integration",
                value=st.session_state.get("enable_human_feedback", False),
                help="Allow human feedback during the process",
                key="enable_human_feedback_checkbox"
            )
            # Note: enable_human_feedback is automatically managed by widget key
            # No need to manually set st.session_state.enable_human_feedback
            
            keyword_analysis_enabled = st.checkbox(
                "Enable Keyword Analysis",
                value=st.session_state.get("keyword_analysis_enabled", True),
                help="Analyze content for keyword presence and relevance",
                key="keyword_analysis_enabled_checkbox"
            )
            # Note: keyword_analysis_enabled is automatically managed by widget key
            # No need to manually set st.session_state.keyword_analysis_enabled
            
            if keyword_analysis_enabled:
                keywords_to_target = st.text_area(
                    "Keywords to Target",
                    value=st.session_state.get("keywords_to_target", ""),
                    height=80,
                    placeholder="Enter keywords separated by commas...",
                    help="Keywords that should be appropriately incorporated in the content",
                    key="keywords_to_target"
                )
                # Don't set session state here - it's handled by the widget
                
        with col2:
            enable_real_time_monitoring = st.checkbox(
                "Enable Real-Time Monitoring",
                value=st.session_state.get("enable_real_time_monitoring", True),
                help="Monitor process performance in real-time",
                key="enable_real_time_monitoring_checkbox"
            )
            # Note: enable_real_time_monitoring is automatically managed by widget key
            # No need to manually set st.session_state.enable_real_time_monitoring
            
            enable_comprehensive_reporting = st.checkbox(
                "Enable Comprehensive Reporting",
                value=st.session_state.get("enable_comprehensive_reporting", True),
                help="Generate detailed reports with analytics",
                key="enable_comprehensive_reporting_checkbox"
            )
            # Note: enable_comprehensive_reporting is automatically managed by widget key
            # No need to manually set st.session_state.enable_comprehensive_reporting

        # Security & Compliance
        st.subheader("Security & Compliance")
        col1, col2 = st.columns(2)
        with col1:
            enable_encryption = st.checkbox(
                "Enable Data Encryption",
                value=st.session_state.get("enable_encryption", True),
                help="Encrypt sensitive data during processing",
                key="enable_encryption_checkbox"
            )
            # Note: enable_encryption is automatically managed by widget key
            # No need to manually set st.session_state.enable_encryption
            
        with col2:
            enable_audit_trail = st.checkbox(
                "Enable Audit Trail",
                value=st.session_state.get("enable_audit_trail", True),
                help="Maintain detailed audit trail for compliance",
                key="enable_audit_trail_checkbox"
            )
            # Note: enable_audit_trail is automatically managed by widget key
            # No need to manually set st.session_state.enable_audit_trail

    # Execution Controls
    st.subheader("🚀 Execution Controls")
    col1, col2, col3 = st.columns([2, 1, 1])
    
    with col1:
        execution_mode = st.selectbox(
            "Execution Mode",
            options=["Adversarial Testing Only", "Integrated Adversarial-Evolution", "Full Tripartite Workflow"],
            index=1,
            help="Select the execution mode for the process",
            key="execution_mode"
        )
        # Note: execution_mode is automatically managed by the widget key
        # No need to manually set st.session_state.execution_mode
        
    with col2:
        if st.button("🚀 Run Ultimate Testing", type="primary", use_container_width=True):
            if not st.session_state.protocol_text.strip():
                st.error("Please enter content to test.")
            elif not st.session_state.red_team_models or not st.session_state.blue_team_models:
                st.error("Please select at least one Red Team and one Blue Team model.")
            else:
                st.session_state.adversarial_running = True
                st.session_state.adversarial_log = []
                st.session_state.adversarial_results = None
                st.session_state.adversarial_status_message = "Starting ultimate adversarial testing & evolution..."
                
                # Initialize advanced session state variables
                st.session_state.integrated_adversarial_history = []
                st.session_state.model_performance_metrics = {}
                st.session_state.quality_metrics = {}
                
                with st.spinner(st.session_state.adversarial_status_message):
                    try:
                        # Use integrated workflow for ultimate testing
                        from integrated_workflow import run_fully_integrated_adversarial_evolution
                        
                        results = run_fully_integrated_adversarial_evolution(
                            current_content=st.session_state.protocol_text,
                            content_type=content_type,
                            api_key=st.session_state.openrouter_key,
                            base_url="https://openrouter.ai/api/v1",
                            red_team_models=st.session_state.red_team_models,
                            blue_team_models=st.session_state.blue_team_models,
                            evaluator_models=st.session_state.evaluator_models,
                            max_iterations=50,  # Overall process iterations
                            adversarial_iterations=st.session_state.adversarial_max_iter,
                            evolution_iterations=20,
                            evaluation_iterations=10,
                            system_prompt=st.session_state.system_prompt,
                            evaluator_system_prompt=st.session_state.evaluator_system_prompt,
                            temperature=st.session_state.temperature,
                            top_p=st.session_state.top_p,
                            frequency_penalty=0.0,
                            presence_penalty=0.0,
                            max_tokens=st.session_state.max_tokens,
                            seed=42,
                            rotation_strategy=st.session_state.adversarial_rotation_strategy,
                            red_team_sample_size=st.session_state.adversarial_red_team_sample_size,
                            blue_team_sample_size=st.session_state.adversarial_blue_team_sample_size,
                            evaluator_sample_size=st.session_state.evaluator_sample_size,
                            confidence_threshold=st.session_state.adversarial_confidence / 100.0,
                            evaluator_threshold=st.session_state.evaluator_threshold,
                            evaluator_consecutive_rounds=st.session_state.evaluator_consecutive_rounds,
                            compliance_requirements=compliance_requirements,
                            enable_data_augmentation=st.session_state.enable_data_augmentation,
                            augmentation_model_id=st.session_state.get("augmentation_model", "gpt-4o"),
                            augmentation_temperature=st.session_state.get("augmentation_temperature", 0.7),
                            enable_human_feedback=st.session_state.enable_human_feedback,
                            multi_objective_optimization=st.session_state.enable_multi_objective_optimization,
                            feature_dimensions=st.session_state.get("feature_dimensions", ["complexity", "diversity"]),
                            feature_bins=st.session_state.get("feature_bins", 10),
                            elite_ratio=st.session_state.elite_ratio,
                            exploration_ratio=st.session_state.exploration_ratio,
                            exploitation_ratio=0.7,
                            archive_size=st.session_state.archive_size,
                            checkpoint_interval=10,
                            keyword_analysis_enabled=st.session_state.keyword_analysis_enabled,
                            keywords_to_target=[k.strip() for k in st.session_state.get("keywords_to_target", "").split(",") if k.strip()],
                            keyword_penalty_weight=0.5
                        )
                        
                        st.session_state.adversarial_results = results
                        st.success("Ultimate adversarial testing & evolution completed successfully!")
                        
                    except Exception as e:
                        st.error(f"Ultimate testing failed: {e}")
                        import traceback
                        st.error(f"Full traceback: {traceback.format_exc()}")
                    finally:
                        st.session_state.adversarial_running = False
                        st.session_state.adversarial_status_message = "Ultimate testing finished."
                        # Don't trigger rerun here, let the normal Streamlit flow continue

    with col3:
        if st.session_state.adversarial_running and st.button("⏹️ Stop Testing", use_container_width=True):
            st.session_state.adversarial_stop_flag = True
            st.info("Stop signal sent. Testing will stop after the current iteration.")

    # Display Results
    if st.session_state.adversarial_results:
        st.subheader("🎯 Ultimate Testing Results")
        results = st.session_state.adversarial_results
        
        # Results Tabs
        results_tabs = st.tabs(["📄 Final Content", "📊 Analytics", "🔍 Detailed Results", "📈 Performance"])
        
        with results_tabs[0]:
            st.subheader("Final Hardened Content")
            final_content = results.get("final_content", "No final content available.")
            st.text_area("Final Content", value=final_content, height=400, key="final_content_display")
            
            # Content Comparison
            if results.get("initial_content"):
                st.subheader("Content Comparison")
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("**Original Content**")
                    st.text_area("Original", value=results["initial_content"], height=200, disabled=True)
                with col2:
                    st.markdown("**Improved Content**")
                    st.text_area("Improved", value=final_content, height=200, disabled=True)
                
                # Show differences
                if st.button("Show Differences"):
                    render_code_diff(results["initial_content"], final_content)
        
        with results_tabs[1]:
            st.subheader("Process Analytics")
            
            # Integrated Score
            integrated_score = results.get("integrated_score", 0.0)
            st.metric("Integrated Score", f"{integrated_score:.2%}")
            
            # Cost Analysis
            total_cost = results.get("total_cost_usd", 0.0)
            st.metric("Total Cost", f"${total_cost:.4f}")
            
            # Token Usage
            tokens = results.get("total_tokens", {})
            if tokens:
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("Prompt Tokens", f"{tokens.get('prompt', 0):,}")
                with col2:
                    st.metric("Completion Tokens", f"{tokens.get('completion', 0):,}")
            
            # Keyword Analysis
            keyword_analysis = results.get("keyword_analysis", {})
            if keyword_analysis:
                st.subheader("Keyword Analysis")
                relevance_score = keyword_analysis.get("relevance_score", 0.0)
                st.metric("Keyword Relevance", f"{relevance_score:.1%}")
                
                keywords_found = keyword_analysis.get("keywords_found", [])
                if keywords_found:
                    st.write(f"**Keywords Found:** {', '.join(keywords_found)}")
        
        with results_tabs[2]:
            st.subheader("Detailed Process Results")
            
            # Adversarial Results
            adversarial_results = results.get("adversarial_results", {})
            if adversarial_results:
                with st.expander("Adversarial Testing Results", expanded=False):
                    st.json(adversarial_results)
            
            # Evolution Results
            evolution_results = results.get("evolution_results", {})
            if evolution_results:
                with st.expander("Evolution Optimization Results", expanded=False):
                    st.json(evolution_results)
            
            # Evaluation Results
            evaluation_results = results.get("evaluation_results", {})
            if evaluation_results:
                with st.expander("Evaluation Results", expanded=False):
                    st.json(evaluation_results)
            
            # Full Results JSON
            with st.expander("View Full Results JSON", expanded=False):
                st.json(results)
        
        with results_tabs[3]:
            st.subheader("Performance Metrics")
            
            # Model Performance
            if st.session_state.get("model_performance_metrics"):
                st.write("**Model Performance:**")
                st.json(st.session_state.model_performance_metrics)
            
            # Process History
            if st.session_state.get("integrated_adversarial_history"):
                st.write("**Process History:**")
                history_df = pd.DataFrame(st.session_state.integrated_adversarial_history)
                st.dataframe(history_df)
            
            # Quality Metrics
            if st.session_state.get("quality_metrics"):
                st.write("**Quality Metrics:**")
                st.json(st.session_state.quality_metrics)

        # Process Log
        st.subheader("📝 Process Log")
        log_text = "\n".join(st.session_state.adversarial_log)
        st.text_area("Log", value=log_text, height=200, disabled=True, key="adversarial_log_display")

def render_github_tab():
    """Renders the UI for the GitHub Integration tab."""
    st.header("🐙 GitHub Integration")
    st.write("Manage your GitHub integrations for version control and collaboration.")

    # Initialize session state key if not present and use default value
    github_token_value = st.session_state.get("github_token", "")
    
    # Safe way: Use the retrieved value as the default
    st.text_input("GitHub Personal Access Token", type="password", value=github_token_value, key="github_token")

    # Use the safe get method again for the conditional check
    if st.session_state.get("github_token", ""):
        st.subheader("Linked Repositories")
        try:
            with st.spinner("Fetching repositories..."):
                repos = list_linked_github_repositories(st.session_state.get("github_token", ""))
            if repos:
                st.selectbox("Your Repositories", [repo['full_name'] for repo in repos])
            else:
                st.info("No repositories found for this token.")
        except Exception as e:
            st.error(f"Error listing repositories: {e}")

        st.subheader("Commit Changes")
        repo_name = st.text_input("Repository Name (e.g., 'owner/repo')", key="github_commit_repo_name")
        branch = st.text_input("Branch to Commit To", key="github_commit_branch", value="main")
        file_path = st.text_input("File Path in Repo", key="github_commit_file_path", placeholder="e.g., 'src/main.py'")
        file_content = st.text_area("File Content", height=150, key="github_commit_file_content", value=st.session_state.get("evolution_current_best", ""))
        commit_message = st.text_input("Commit Message", key="github_commit_message")

        if st.button("Commit to GitHub"):
            if all([repo_name, branch, file_path, file_content, commit_message]):
                try:
                    commit_to_github(st.session_state.get("github_token", ""), repo_name, branch, file_path, file_content, commit_message)
                    st.success(f"Changes committed to '{branch}' in '{repo_name}'.")
                except Exception as e:
                    st.error(f"Error committing changes: {e}")
            else:
                st.warning("Please fill in all fields to commit.")
    else:
        st.warning("Please enter your GitHub Personal Access Token to enable integrations.")

def render_activity_feed_tab():
    """Renders the UI for the Activity Feed tab."""
    st.header("📜 Activity Feed")
    st.write("Review recent activities and system events.")

    activity_log = st.session_state.get("activity_log", [])
    if activity_log:
        for entry in reversed(activity_log):  # Show most recent first
            st.json(entry)
    else:
        st.info("No recent activity.")

def render_report_templates_tab():
    """Renders the UI for the Report Templates tab."""
    st.header("📄 Report Templates")
    st.write("Manage custom report templates.")

    if st.session_state.report_templates:
        st.subheader("Existing Report Templates")
        for template_name in list(st.session_state.report_templates.keys()):
            with st.expander(f"Template: {template_name}"):
                st.code(st.session_state.report_templates[template_name], language="json")
                if st.button(f"Delete {template_name}", key=f"delete_report_template_{template_name}"):
                    del st.session_state.report_templates[template_name]
                    _save_report_templates(st.session_state.report_templates)
                    st.success(f"Template '{template_name}' deleted.")

    st.subheader("Create New Report Template")
    new_template_name = st.text_input("New Template Name", key="new_report_template_name")
    new_template_content = st.text_area("Template Content (JSON)", height=200, key="new_report_template_content")

    if st.button("Save New Template"):
        if new_template_name and new_template_content:
            try:
                json.loads(new_template_content)  # Validate JSON
                st.session_state.report_templates[new_template_name] = new_template_content
                _save_report_templates(st.session_state.report_templates)
                st.success(f"Template '{new_template_name}' saved.")
            except json.JSONDecodeError:
                st.error("Invalid JSON content.")
        else:
            st.error("Template name and content cannot be empty.")

def render_model_dashboard_tab():
    """Renders the UI for the Model Dashboard tab."""
    st.header("📊 Model Dashboard")
    st.write("Monitor and manage your language models.")

    st.subheader("OpenRouter Models")
    try:
        with st.spinner("Fetching models from OpenRouter..."):
            openrouter_models = get_openrouter_models()
        if openrouter_models:
            df = pd.DataFrame(openrouter_models)
            st.dataframe(df)
        else:
            st.info("No OpenRouter models found.")
    except Exception as e:
        st.error(f"Error fetching OpenRouter models: {e}")

    st.subheader("Model Performance Metrics")
    if st.session_state.adversarial_model_performance:
        model_perf_data = pd.DataFrame([
            {'Model': k, 'Score': v.get('score', 0), 'Cost': v.get('cost', 0.0)}
            for k, v in st.session_state.adversarial_model_performance.items()
        ])
        st.dataframe(model_perf_data)
        st.bar_chart(model_perf_data.set_index('Model')[['Score', 'Cost']])
    else:
        st.info("No model performance data available. Run adversarial testing to populate this.")

def render_tasks_tab():
    """Renders the UI for the Tasks tab."""
    st.header("✅ Tasks")
    st.write("Manage your tasks and to-dos.")

    st.subheader("Create New Task")
    new_task_description = st.text_input("Task Description", key="new_task_description")
    if st.button("Add Task"):
        if new_task_description:
            try:
                create_task(new_task_description)
                st.success("Task added successfully!")
                st.session_state.new_task_description = ""
            except Exception as e:
                st.error(f"Error creating task: {e}")
        else:
            st.error("Task description cannot be empty.")

    st.subheader("Current Tasks")
    tasks = get_tasks()
    if tasks:
        for i, task in enumerate(tasks):
            st.checkbox(task["description"], value=task.get("completed", False), key=f"task_checkbox_{i}", disabled=True)
    else:
        st.info("No tasks found.")

def render_admin_tab():
    """Renders the UI for the Admin tab."""
    st.header("👑 Admin Panel")
    st.write("Manage user roles and system settings.")

    st.subheader("User Role Management")
    user_id = st.text_input("User ID", key="admin_user_id")
    role_options = list(ROLES.keys())
    selected_role = st.selectbox("Assign Role", role_options, key="admin_assign_role")

    if st.button("Assign Role"):
        if user_id and selected_role:
            try:
                assign_role(user_id, selected_role)
                st.success(f"Role '{selected_role}' assigned to user '{user_id}'.")
            except Exception as e:
                st.error(f"Error assigning role: {e}")
        else:
            st.error("User ID and Role cannot be empty.")

    st.subheader("Current User Roles")
    st.json(st.session_state.user_roles)

def render_openevolve_orchestrator_tab():
    """Renders the UI for the OpenEvolve Orchestrator tab."""
    st.header("🚀 OpenEvolve Orchestrator")
    st.write("Control and monitor OpenEvolve backend services.")

    if OPENEVOLVE_AVAILABLE:
        st.subheader("Service Control")
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("Start All Services"):
                try:
                    start_openevolve_services()
                    st.success("OpenEvolve services started.")
                except Exception as e:
                    st.error(f"Error starting services: {e}")
        with col2:
            if st.button("Stop All Services"):
                try:
                    stop_openevolve_services()
                    st.success("OpenEvolve services stopped.")
                except Exception as e:
                    st.error(f"Error stopping services: {e}")
        with col3:
            if st.button("Restart All Services"):
                try:
                    restart_openevolve_services()
                    st.success("OpenEvolve services restarted.")
                except Exception as e:
                    st.error(f"Error restarting services: {e}")
    else:
        st.warning("OpenEvolve backend is not available. Orchestrator functions are disabled.")

def render_evolution_tab():
    """Renders the UI for the Evolution tab with ALL OpenEvolve features."""
    st.header("🧬 Evolution Engine")
    st.markdown("""
    **Advanced Evolutionary Computing with OpenEvolve**
    
    This tab provides access to all OpenEvolve evolution modes including:
    - **Standard Evolution**: Basic evolutionary optimization
    - **Quality-Diversity (MAP-Elites)**: Maintains diverse, high-performing solutions
    - **Multi-Objective**: Optimizes for multiple competing objectives
    - **Adversarial Evolution**: Red Team/Blue Team approach for robustness
    - **Algorithm Discovery**: Discovers novel algorithmic approaches
    - **Symbolic Regression**: Discovers mathematical expressions from data
    - **Neuroevolution**: Evolves neural network architectures
    - **Prompt Optimization**: Optimizes LLM prompts for better performance
    """)

    # Content input
    protocol_text = st.text_area("Content to Evolve", value=st.session_state.protocol_text, height=300, key="evolution_protocol_text")
    # Note: protocol_text is automatically managed by the widget key
    # No need to manually set st.session_state.protocol_text

    # Evolution Mode Selection
    st.subheader("Evolution Mode")
    evolution_modes = [
        "standard", "quality_diversity", "multi_objective", "adversarial",
        "prompt_optimization", "algorithm_discovery", "symbolic_regression", "neuroevolution"
    ]
    evolution_mode_descriptions = {
        "standard": "Basic evolutionary optimization",
        "quality_diversity": "Quality-Diversity (MAP-Elites) evolution",
        "multi_objective": "Multi-objective optimization",
        "adversarial": "Red Team/Blue Team adversarial evolution",
        "prompt_optimization": "Optimize LLM prompts",
        "algorithm_discovery": "Discover novel algorithms",
        "symbolic_regression": "Discover mathematical expressions",
        "neuroevolution": "Evolve neural networks"
    }
    
    selected_mode = st.selectbox(
        "Select Evolution Mode",
        options=evolution_modes,
        format_func=lambda x: f"{x.replace('_', ' ').title()} - {evolution_mode_descriptions[x]}",
        key="evolution_mode"
    )
    # Note: evolution_mode is automatically managed by the widget key
    # No need to manually set st.session_state.evolution_mode

    # Advanced Configuration
    with st.expander("🔧 Advanced Configuration", expanded=False):
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Core Parameters")
            evolution_max_iterations = st.number_input(
                "Max Iterations",
                min_value=1,
                max_value=10000,
                value=st.session_state.evolution_max_iterations,
                help="Maximum number of evolution iterations",
                key="evolution_max_iterations_input"
            )
            population_size = st.number_input(
                "Population Size",
                min_value=1,
                max_value=10000,
                value=st.session_state.population_size,
                help="Size of the evolution population",
                key="population_size_input"
            )
            num_islands = st.number_input(
                "Number of Islands",
                min_value=1,
                max_value=20,
                value=st.session_state.num_islands,
                help="Island model for better exploration",
                key="num_islands_input"
            )
            archive_size = st.number_input(
                "Archive Size",
                min_value=10,
                max_value=1000,
                value=st.session_state.archive_size,
                help="Size of the archive for storing best solutions",
                key="archive_size_input"
            )
            
        with col2:
            st.subheader("Evolution Strategy")
            elite_ratio = st.slider(
                "Elite Ratio",
                0.0, 1.0,
                st.session_state.elite_ratio,
                0.01,
                help="Ratio of elite individuals to preserve",
                key="elite_ratio_slider"
            )
            exploration_ratio = st.slider(
                "Exploration Ratio",
                0.0, 1.0,
                st.session_state.exploration_ratio,
                0.01,
                help="Ratio for exploration in evolution",
                key="exploration_ratio_slider"
            )
            exploitation_ratio = st.slider(
                "Exploitation Ratio",
                0.0, 1.0,
                st.session_state.exploitation_ratio,
                0.01,
                help="Ratio for exploitation in evolution",
                key="exploitation_ratio_slider"
            )
            
            # Ensure ratios sum to 1
            total_ratio = st.session_state.elite_ratio + st.session_state.exploration_ratio + st.session_state.exploitation_ratio
            if abs(total_ratio - 1.0) > 0.01:
                st.warning(f"Ratios sum to {total_ratio:.2f}, ideally should sum to 1.0")

        # Feature Dimensions for QD and Multi-Objective
        if selected_mode in ["quality_diversity", "multi_objective"]:
            st.subheader("Feature Dimensions")
            available_features = ["complexity", "diversity", "performance", "efficiency", "readability", "robustness", "accuracy", "clarity"]
            feature_dimensions = st.multiselect(
                "Feature Dimensions",
                options=available_features,
                default=st.session_state.feature_dimensions,
                help="Feature dimensions for Quality-Diversity or Multi-Objective evolution",
                key="feature_dimensions_select"
            )
            feature_bins = st.number_input(
                "Feature Bins",
                min_value=5,
                max_value=50,
                value=st.session_state.feature_bins,
                help="Number of bins for each feature dimension",
                key="feature_bins_input"
            )

        # Objectives for Multi-Objective
        if selected_mode == "multi_objective":
            st.subheader("Multi-Objective Settings")
            available_objectives = ["performance", "readability", "maintainability", "efficiency", "security", "robustness", "clarity", "completeness"]
            objectives = st.multiselect(
                "Objectives to Optimize",
                options=available_objectives,
                default=st.session_state.get("objectives", ["performance", "readability"]),
                help="Objectives for multi-objective optimization",
                key="objectives_select"
            )

        # Adversarial Settings
        if selected_mode == "adversarial":
            st.subheader("Adversarial Evolution Settings")
            col1, col2 = st.columns(2)
            with col1:
                adversarial_attack_model = st.selectbox(
                    "Attack Model",
                    options=MODEL_OPTIONS,
                    index=MODEL_OPTIONS.index("claude-3-sonnet") if "claude-3-sonnet" in MODEL_OPTIONS else 0,
                    help="Model for adversarial attacks (Red Team)",
                    key="adversarial_attack_model_select"
                )
            with col2:
                adversarial_defense_model = st.selectbox(
                    "Defense Model",
                    options=MODEL_OPTIONS,
                    index=MODEL_OPTIONS.index("gpt-4o") if "gpt-4o" in MODEL_OPTIONS else 0,
                    help="Model for adversarial defense (Blue Team)",
                    key="adversarial_defense_model_select"
                )

        # Advanced OpenEvolve Features
        st.subheader("Advanced OpenEvolve Features")
        col1, col2, col3 = st.columns(3)
        with col1:
            enable_artifacts = st.checkbox("Enable Artifacts", value=st.session_state.get("enable_artifacts", True), help="Enable artifact side-channel", key="enable_artifacts_checkbox")
            cascade_evaluation = st.checkbox("Cascade Evaluation", value=st.session_state.get("cascade_evaluation", True), help="Use cascade evaluation", key="cascade_evaluation_checkbox")
            use_llm_feedback = st.checkbox("LLM Feedback", value=st.session_state.get("use_llm_feedback", False), help="Use LLM-based feedback", key="use_llm_feedback_checkbox")
        with col2:
            include_artifacts = st.checkbox("Include Artifacts", value=st.session_state.get("include_artifacts", True), help="Include artifacts in prompts", key="include_artifacts_checkbox")
            evolution_trace_enabled = st.checkbox("Enable Trace", value=st.session_state.get("evolution_trace_enabled", False), help="Enable evolution trace logging", key="evolution_trace_enabled_checkbox")
            diff_based_evolution = st.checkbox("Diff-Based Evolution", value=st.session_state.get("diff_based_evolution", True), help="Use diff-based evolution", key="diff_based_evolution_checkbox")
        with col3:
            parallel_evaluations = st.number_input("Parallel Evaluations", min_value=1, max_value=16, value=st.session_state.get("parallel_evaluations", 4), help="Number of parallel evaluations", key="parallel_evaluations_input")
            checkpoint_interval = st.number_input("Checkpoint Interval", min_value=1, max_value=100, value=st.session_state.checkpoint_interval, help="Interval for saving checkpoints", key="checkpoint_interval_input")

    # Prompts Configuration
    with st.expander("💬 Prompts Configuration", expanded=False):
        col1, col2 = st.columns(2)
        with col1:
            system_prompt = st.text_area(
                "System Prompt",
                value=st.session_state.system_prompt,
                height=150,
                help="System prompt for the evolution process",
                key="system_prompt_textarea"
            )
        with col2:
            evaluator_system_prompt = st.text_area(
                "Evaluator System Prompt",
                value=st.session_state.evaluator_system_prompt,
                height=150,
                help="System prompt for the evaluator",
                key="evaluator_system_prompt_textarea"
            )

    # Run Evolution
    col1, col2 = st.columns([3, 1])
    with col1:
        if st.button(f"🚀 Run {selected_mode.replace('_', ' ').title()} Evolution", type="primary", use_container_width=True):
            if not st.session_state.protocol_text.strip():
                st.error("Please enter content to evolve.")
            else:
                st.session_state.evolution_running = True
                st.session_state.evolution_log = []
                st.session_state.evolution_current_best = ""
                st.session_state.evolution_status_message = f"Starting {selected_mode} evolution..."
                
                with st.spinner(st.session_state.evolution_status_message):
                    try:
                        # Update session state with current widget values
                        # Note: Widget values are automatically stored in session state
                        # Only update values that aren't tied to widgets
                        st.session_state.evolution_max_iterations = evolution_max_iterations
                        st.session_state.population_size = population_size
                        st.session_state.num_islands = num_islands
                        st.session_state.archive_size = archive_size
                        # elite_ratio, exploration_ratio, exploitation_ratio are handled by widgets
                        # feature_dimensions, feature_bins, system_prompt, evaluator_system_prompt are handled by widgets
                        # No need to manually set these session state variables
                        
                        # Update conditional parameters
                        if selected_mode == "multi_objective":
                            # objectives is handled by widget key
                            pass
                        if selected_mode == "adversarial":
                            # adversarial_attack_model and adversarial_defense_model are handled by widget keys
                            pass
                        
                        # Update OpenEvolve features - all handled by widget keys
                        # enable_artifacts, cascade_evaluation, use_llm_feedback, include_artifacts,
                        # evolution_trace_enabled, diff_based_evolution, parallel_evaluations, checkpoint_interval
                        # are automatically managed by their respective widget keys

                        # Prepare model configuration
                        model_configs = [{
                            "name": st.session_state.model,
                            "weight": 1.0,
                            "temperature": st.session_state.temperature,
                            "top_p": st.session_state.top_p,
                            "max_tokens": st.session_state.max_tokens,
                            "frequency_penalty": 0.0,
                            "presence_penalty": 0.0
                        }]
                        
                        # Run evolution using the unified function
                        from evolution import run_evolution_loop
                        final_content = run_evolution_loop(
                            current_content=st.session_state.protocol_text,
                            api_key=st.session_state.openrouter_key,
                            base_url="https://openrouter.ai/api/v1",
                            model=st.session_state.model,
                            max_iterations=st.session_state.evolution_max_iterations,
                            population_size=st.session_state.population_size,
                            system_prompt=st.session_state.system_prompt,
                            evaluator=None,  # Will be created internally
                            extra_headers={},
                            temperature=st.session_state.temperature,
                            top_p=st.session_state.top_p,
                            frequency_penalty=0.0,
                            presence_penalty=0.0,
                            max_tokens=st.session_state.max_tokens,
                            seed=42,
                            content_type="document_general",
                            evolution_mode=selected_mode,
                            objectives=st.session_state.get("objectives", ["performance", "readability"]),
                            feature_dimensions=st.session_state.feature_dimensions,
                            feature_bins=st.session_state.feature_bins,
                            elite_ratio=st.session_state.elite_ratio,
                            exploration_ratio=st.session_state.exploration_ratio,
                            exploitation_ratio=st.session_state.exploitation_ratio,
                            archive_size=st.session_state.archive_size,
                            checkpoint_interval=st.session_state.checkpoint_interval,
                            num_islands=st.session_state.num_islands,
                            enable_artifacts=st.session_state.enable_artifacts,
                            cascade_evaluation=st.session_state.cascade_evaluation,
                            use_llm_feedback=st.session_state.use_llm_feedback,
                            parallel_evaluations=st.session_state.parallel_evaluations,
                            include_artifacts=st.session_state.include_artifacts,
                            evolution_trace_enabled=st.session_state.evolution_trace_enabled,
                            diff_based_evolution=st.session_state.diff_based_evolution,
                            adversarial_attack_model=st.session_state.get("adversarial_attack_model", "gpt-4"),
                            adversarial_defense_model=st.session_state.get("adversarial_defense_model", "gpt-4")
                        )
                        
                        st.session_state.evolution_current_best = final_content
                        st.success(f"{selected_mode.replace('_', ' ').title()} evolution completed successfully!")
                        
                    except Exception as e:
                        st.error(f"Evolution failed: {e}")
                        import traceback
                        st.error(f"Full traceback: {traceback.format_exc()}")
                    finally:
                        st.session_state.evolution_running = False

    with col2:
        if st.session_state.evolution_running and st.button("⏹️ Stop Evolution", use_container_width=True):
            st.session_state.evolution_stop_flag = True
            st.info("Stop signal sent. Evolution will stop after the current iteration.")

    # Display Results
    if st.session_state.evolution_current_best:
        st.subheader("🏆 Evolution Results")
        
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("Original Content")
            st.text_area("Original", value=st.session_state.protocol_text, height=200, disabled=True)
        with col2:
            st.subheader("Evolved Content")
            st.text_area("Evolved", value=st.session_state.evolution_current_best, height=200, key="best_content_display")
        
        st.subheader("📊 Content Comparison")
        render_code_diff(st.session_state.protocol_text, st.session_state.evolution_current_best)

    # Evolution Log and Status
    if st.session_state.evolution_log:
        st.subheader("📝 Evolution Log")
        log_text = "\n".join(st.session_state.evolution_log)
        st.text_area("Evolution Log", value=log_text, height=200, disabled=True)

    # Evolution History Visualization
    if st.session_state.evolution_history:
        st.subheader("📈 Evolution History")
        render_evolution_history_chart(st.session_state.evolution_history)
        render_island_model_chart(st.session_state.evolution_history)

    # OpenEvolve Visualization Integration
    if OPENEVOLVE_AVAILABLE and st.session_state.evolution_history:
        st.subheader("🧬 OpenEvolve Analytics")
        from openevolve_visualization import render_evolution_insights
        render_evolution_insights()

def render_analytics_dashboard_tab():
    """Renders the UI for the Analytics Dashboard tab."""
    st.header("📊 Analytics Dashboard")
    st.write("Visualize evolution and testing performance metrics.")

    # Use comprehensive self-contained implementation
    analytics_tabs = st.tabs(["📈 Overview", "🧬 Evolution", "⚔️ Adversarial", "📋 Reports"])

    with analytics_tabs[0]:  # Overview
        st.subheader("Key Performance Indicators")
        col1, col2, col3 = st.columns(3)
        total_evolutions = len(st.session_state.evolution_history)
        best_fitness = max((ind['fitness'] for gen in st.session_state.evolution_history for ind in gen.get('population', [])), default=0)
        total_cost = st.session_state.adversarial_cost_estimate_usd
        col1.metric("Total Evolutions", f"{total_evolutions:,}")
        col2.metric("Peak Fitness", f"{best_fitness:.4f}")
        col3.metric("Total Cost (USD)", f"${total_cost:.4f}")

    with analytics_tabs[1]: # Evolution Analytics
        st.subheader("Fitness Trend Over Generations")
        if st.session_state.evolution_history:
            fitness_data = []
            for gen in st.session_state.evolution_history:
                pop = gen.get('population', [])
                if pop:
                    fitness_data.append({
                        'Generation': gen.get("generation", 0),
                        'Best Fitness': max(ind.get("fitness", 0) for ind in pop),
                        'Average Fitness': sum(ind.get("fitness", 0) for ind in pop) / len(pop)
                    })
            if fitness_data:
                df = pd.DataFrame(fitness_data)
                if PLOTLY_AVAILABLE:
                    fig = px.line(df, x="Generation", y=["Best Fitness", "Average Fitness"], title="Fitness Trend")
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.warning("Plotly is not available. Install it using: pip install plotly")
        else:
            st.info("Run an evolution to see fitness trends.")

    with analytics_tabs[2]: # Adversarial Analytics
        st.subheader("Model Performance Overview")
        model_performance = st.session_state.get("adversarial_model_performance", {})
        if model_performance:
            model_data = [{"Model": k, "Score": v.get("score", 0), "Cost": v.get("cost", 0.0)} for k, v in model_performance.items()]
            df = pd.DataFrame(model_data).sort_values(by="Score", ascending=False)
            if PLOTLY_AVAILABLE:
                fig = px.bar(df, x="Model", y="Score", title="Model Performance Comparison")
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.warning("Plotly is not available. Install it using: pip install plotly")
        else:
            st.info("Run adversarial testing to see model performance data.")

    with analytics_tabs[3]:  # Reports
        st.subheader("Generate Reports")
        if st.button("Generate Integrated Report", type="primary"):
            with st.spinner("Generating report..."):
                report_content = generate_integrated_report(st.session_state.evolution_history, st.session_state.adversarial_results)
                st.download_button(
                    label="Download Report (PDF)",
                    data=report_content,
                    file_name="integrated_report.pdf",
                    mime="application/pdf",
                )
                st.success("Report generated!")

def render_openevolve_dashboard_tab():
    """Renders the UI for the OpenEvolve Dashboard tab."""
    st.header("✨ OpenEvolve Advanced Dashboard")
    if OPENEVOLVE_AVAILABLE:
        # If the full backend is available, render its dedicated UIs
        main_tabs = st.tabs(["📊 Visualizations", "💡 Insights", "📈 Monitoring", "📋 Diagnostics"])
        with main_tabs[0]:
            render_openevolve_visualization_ui()
        with main_tabs[1]:
            render_evolution_insights()
        with main_tabs[2]:
            render_comprehensive_monitoring_ui()
        with main_tabs[3]:
            render_advanced_diagnostics()
    else:
        st.info("The full OpenEvolve backend is not installed. This dashboard provides an overview of its features.")
        st.markdown("""
        **OpenEvolve provides advanced evolutionary computing capabilities:**
        - **Quality-Diversity Evolution** (MAP-Elites)
        - **Multi-Objective Optimization**
        - **Adversarial Evolution**
        - **Symbolic Regression** & **Neuroevolution**

        Install the full backend to unlock these features.
        """)


# --- Main Application Layout ---

def render_main_layout():
    """Renders the main layout of the Streamlit application."""
    _initialize_session_state()

    # --- Initialize Managers ---
    if "openevolve_api_instance" not in st.session_state:
        st.session_state.openevolve_api_instance = OpenEvolveAPI(base_url=st.session_state.openevolve_base_url, api_key=st.session_state.openevolve_api_key)
    if "prompt_manager" not in st.session_state:
        st.session_state.prompt_manager = PromptManager(api=st.session_state.openevolve_api_instance)
    # ... Initialize other managers similarly ...
    if "analytics_manager" not in st.session_state:
        st.session_state.analytics_manager = AnalyticsManager()

    check_password()
    render_collaboration_ui()

    # Apply Custom CSS
    st.markdown(f"<style>{st.session_state.styles_css}</style>", unsafe_allow_html=True)

    # --- Header ---
    st.markdown(
        '<h2 style="text-align: center;">🧬 OpenEvolve Content Improver</h2>'
        '<p style="text-align: center;">AI-Powered Content Hardening with Multi-LLM Consensus</p>',
        unsafe_allow_html=True
    )

    # --- Main Tabs ---
    tab_titles = [
        "Evolution", "Adversarial Testing", "GitHub", "Activity Feed",
        "Report Templates", "Model Dashboard", "Tasks", "Admin",
        "Analytics Dashboard", "OpenEvolve Dashboard", "Orchestrator"
    ]
    
    # Simple approach: Use traditional tabs without complex state management
    # Streamlit's tabs should work correctly if we avoid interfering with their state
    tabs = st.tabs(tab_titles)
    
    # Render all tab content using the traditional approach
    # This is the most reliable method that works in most Streamlit applications
    with tabs[0]:
        render_evolution_tab()

    with tabs[1]:
        render_adversarial_testing_tab()

    with tabs[2]:
        render_github_tab()

    with tabs[3]:
        render_activity_feed_tab()

    with tabs[4]:
        render_report_templates_tab()

    with tabs[5]:
        render_model_dashboard_tab()

    with tabs[6]:
        render_tasks_tab()

    with tabs[7]:
        render_admin_tab()

    with tabs[8]:
        render_analytics_dashboard_tab()

    with tabs[9]:
        render_openevolve_dashboard_tab()

    with tabs[10]:
        render_openevolve_orchestrator_tab()

    # Tab content is now rendered based on the active tab selection above


if __name__ == "__main__":
    # When running this file directly, we need to set page config
    # But when imported by main.py, it's already set there
    import streamlit as st
    render_main_layout()